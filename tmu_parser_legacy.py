
import re
import warnings
import io
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import numpy as np
import pandas as pd


PARSER_BUILD_ID_V55 = "v55-safe-datetime-pandas3"
PARSER_BUILD_ID_V59 = "v59-choke-zero-test-breaks-20260621"
PARSER_BUILD_ID_V58 = "v58-persistent-units-test-separation-20260621"
PARSER_BUILD_ID_V57 = PARSER_BUILD_ID_V58
PARSER_BUILD_ID_V56 = PARSER_BUILD_ID_V57

PARSER_COMPAT_BUILD_ID_V72 = "v72-expro-column-alignment-20260626"
KEYWORDS = [
    "date", "time", "well", "choke", "whp", "w.h.p", "wellhead", "sep", "separator",
    "gas", "formation", "gross", "oil", "water", "bsw", "bs&w", "wc", "salinity",
    "h2s", "co2", "pump", "pumping", "pressure", "n2", "nitrogen", "ct", "rate",
]

UNIT_KEYWORDS = [
    "psi", "psig", "psia", "bbl", "bbl/d", "stb", "stb/d", "bpd", "mmscf",
    "mmscf/d", "scf", "ppm", "nacl", "mole", "deg api", "api", "%",
    "c", "f", "hh:mm", "d/mm", "gpm", "bpm",
]

PARENT_HEADERS = {
    "gas", "oil", "water", "gross rate", "production", "liquid rate",
    "liquid metering", "gas metering", "pressure and temperature measurements",
    "mpfm parameters", "fluid properties", "esp parameters", "well head",
    "separator", "oil or cond.", "water reading", "ratio", "tmu",
}

BASE_NON_PLOT_COLS = {
    "source", "sheet", "well", "date", "time", "time_text", "datetime", "note", "test_unit"
}

COLUMN_LABELS = {
    "choke_pct": "Choke Opening (%)",
    "whp_psi": "WHP (psi)",
    "flp_psi": "FLP (psi)",
    "flt_c": "FLT (°C)",
    "sep_p_psi": "Separator Pressure (psi)",
    "gas_temp_c": "Gas Temp (°C)",
    "gas_temp_f": "Gas Temp (°F)",
    "gas_sg": "Gas Specific Gravity",
    "orifice_size_in": "Orifice Size (in)",
    "gas_rate_mmscfd": "Total Gas Rate (MMSCF/D)",
    "gas_formation_mmscfd": "Formation Gas Rate (MMSCF/D)",
    "gor_scf_bbl": "GOR (scf/bbl)",
    "h2s_ppm": "H2S (ppm)",
    "co2_mole_pct": "CO2 (mole %)",
    "oil_temp_c": "Oil Temp (°C)",
    "oil_temp_f": "Oil Temp (°F)",
    "oil_kf": "Oil K.F Factor",
    "oil_api": "Oil Gravity (API)",
    "oil_meter_increment_bbl": "Oil Meter Increment (bbl)",
    "oil_cmf": "Oil CMF",
    "oil_rate_stbd": "Oil Rate (STB/D)",
    "water_rate_bpd": "Water Rate (BBL/D)",
    "bsw_pct": "BS&W (%)",
    "salinity_kppm": "Salinity (K ppm NaCl)",
    "gross_rate_bpd": "Gross Rate (BBL/D)",
    "water_cum_bbl": "Water Cum (bbl)",
    "pumping_pressure_psi": "Pumping Pressure (psi)",
    "n2_rate_scfm": "N2 Rate (scfm)",
    "ct_pressure_psi": "CT Pressure (psi)",
    "ct_depth_m": "CT Depth (m)",
    "ct_running_speed_ftmin": "CT Running Speed (ft/min)",
    "ct_pipe_weight_lbf": "CT Pipe Weight (lbf)",
    "u2_pass_side_pump_pressure_psi": "U2 Pass-side Pump Pressure (psi)",
    "choke_size_64": "Choke Size (/64 in)",
    "choke_ambiguous": "Choke (unit not stated)",
    "choke_unified": "Unified Choke",
    "flow_press_psi": "Flow Pressure (psi)",
    "flow_temp_c": "Flow Temp (°C)",
    "wellhead_temp_c": "Wellhead Temp (°C)",
    "wellhead_temp_f": "Wellhead Temp (°F)",
    "sep_temp_c": "Separator Temp (°C)",
    "sep_temp_f": "Separator Temp (°F)",
    "mpfm_press_psig": "MPFM Pressure (psig)",
    "mpfm_temp_f": "MPFM Temp (°F)",
    "dp_mbar": "DP (mbar)",
    "qoil_s_stbd": "QOil(S) (STB/D)",
    "qwat_s_bpd": "QWat(S) (BBL/D)",
    "qgas_s_mmscfd": "QGas(S) (MMSCF/D)",
    "qoil_a_bpd": "QOil(A) (BBL/D)",
    "qwat_a_bpd": "QWat(A) (BBL/D)",
    "qgas_a_mmcfd": "QGas(A) (MMCF/D)",
    "wlr_s_pct": "WLR(S) (%)",
    "qgross_s_bpd": "QGross(S) (BBL/D)",
    "gor_s_scf_stb": "GOR(S) (SCF/STB)",
    "gvf_a_pct": "GVF(A) (%)",
    "flt_f": "FLT (°F)",
    "flow_temp_f": "Flow Temp (°F)",
    "flow_temp_c": "Flow Temp (°C)",
    "wellhead_temp_c": "Wellhead Temp (°C)",
    "wellhead_temp_f": "Wellhead Temp (°F)",
    "sep_temp_c": "Separator Temp (°C)",
    "sep_temp_f": "Separator Temp (°F)",
    "sep_dp_inh2o": "Separator DP (inH2O)",
    "gas_dp_inh2o": "Gas DP (inH2O)",
    "oil_cum_bbl": "Oil Cum (bbl)",
    "wat_cum_bbl": "Water Cum (bbl)",
    "gas_cum_mscf": "Gas Cum (MSCF)",
    "cgr_bbl_mmscf": "CGR (BBL/MMSCF)",
    "oil_sg": "Oil SG",
    "water_sg": "Water SG",
    "water_ph": "Water pH",
    "pump_freq_hz": "Pump Frequency (Hz)",
    "pump_intake_pressure_psi": "Pi / Intake Pressure (psi)",
    "pump_discharge_pressure_psi": "Pd / Discharge Pressure (psi)",
    "motor_current_amp": "Motor Current (A)",
    "motor_ama_amp": "AMA / Motor Current (A)",
    "intake_temp_c": "Intake Temperature (°C)",
    "intake_temp_f": "Intake Temperature (°F)",
    "motor_temp_c": "Motor Temperature (°C)",
    "motor_temp_f": "Motor Temperature (°F)",
    "motor_load_pct": "Motor Load (%)",
    "vibration_x": "Vibration X",
    "vibration_y": "Vibration Y",
    "vibration_z": "Vibration Z",
    "drive_freq_hz": "Drive Frequency (Hz)",
    "us_press_psi": "Upstream Pressure (psi)",
    "us_temp_c": "Upstream Temp (°C)",
    "ds_press_psi": "Downstream Pressure (psi)",
    "fw_pct": "FW (%)",
    "twc_pct": "TWC (%)",
    "liquid_volume_bbl": "Liquid Volume (bbl)",
    "gor_mmscf_bbl": "GOR (MMSCF/BBL)",
}


def normalize_text(x: object) -> str:
    if pd.isna(x):
        return ""
    s = str(x).strip().lower()
    s = s.replace("\n", " ").replace("\r", " ")
    s = re.sub(r"[\t_]+", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def safe_text(x: object) -> str:
    """Return a safe display/header string for any Excel/PDF cell value."""
    try:
        if pd.isna(x):
            return ""
    except Exception:
        pass
    try:
        s = str(x).strip()
    except Exception:
        s = ""
    if s.lower() in ["nan", "nat", "none"]:
        return ""
    return s


def safe_join(items, sep: str = " ") -> str:
    """Join mixed objects safely; prevents float-in-header join crashes."""
    return sep.join(safe_text(i) for i in items if safe_text(i))


def clean_header(s: object) -> str:
    s = normalize_text(s)
    s = s.replace("&", " and ")
    s = re.sub(r"[^a-z0-9/%.\- =]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def canonical_key(s: object) -> str:
    """Normalize a user/company header into a stable alias key.

    This is used by the Streamlit column-mapping UI.  Example:
    "W.H.P (psig)" and "WHP psig" become comparable keys.
    """
    c = clean_header(s)
    c = c.replace(" and ", " ")
    c = re.sub(r"[^a-z0-9]+", "_", c).strip("_")
    return c


def standard_column_options(include_meta: bool = False) -> Dict[str, str]:
    """Return canonical parser fields for the user-facing mapping UI."""
    labels = dict(COLUMN_LABELS)
    if include_meta:
        labels.update({
            "well": "Well Name",
            "date": "Date",
            "time": "Time",
            "datetime": "Date & Time",
            "note": "Note / Event",
        })
    return dict(sorted(labels.items(), key=lambda kv: kv[1].lower()))


def is_datetime_like(x: object) -> bool:
    import datetime as _dt
    return isinstance(x, (_dt.datetime, _dt.date, _dt.time, pd.Timestamp))


def is_numeric_like(x: object) -> bool:
    if pd.isna(x):
        return False
    if isinstance(x, (int, float, np.number)) and not isinstance(x, bool):
        return True
    s = str(x).strip().replace(",", "")
    if not s:
        return False
    return bool(re.fullmatch(r"[-+]?\d+(\.\d+)?", s))


def row_keywords_score(row: pd.Series) -> int:
    txt = safe_join([normalize_text(v) for v in row.tolist()])
    return sum(1 for kw in KEYWORDS if kw in txt)

def row_looks_like_data(row: pd.Series) -> bool:
    """Detect actual data/event rows so they are not swallowed as header rows."""
    vals = [v for v in row.tolist() if normalize_text(v)]
    if not vals:
        return False

    first = vals[0]
    first_txt = normalize_text(first)

    # Dates/timestamps in first column are almost always data rows, not header rows.
    if is_datetime_like(first):
        return True

    if re.match(r"^\d{1,2}[:.]\d{2}(:\d{2})?$", first_txt):
        # EXPRO/PDF rows that start with time.
        return True

    if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}(\s+\d{1,2}[:.]\d{2})?", first_txt):
        return True

    return False


def row_looks_like_units(row: pd.Series) -> bool:
    vals = [v for v in row.tolist() if normalize_text(v)]
    if not vals:
        return False
    txt = safe_join([normalize_text(v) for v in vals])
    unit_count = sum(1 for kw in UNIT_KEYWORDS if kw in txt)
    text_count = sum(not is_numeric_like(v) for v in vals)
    return unit_count >= 2 and text_count >= 2


def header_score(row: pd.Series) -> float:
    if row_looks_like_data(row):
        return -100.0

    kw = row_keywords_score(row)
    non_empty = int(row.notna().sum())
    textish = sum((not is_numeric_like(v)) and normalize_text(v) != "" for v in row.tolist())
    units = 6 if row_looks_like_units(row) else 0

    # Strongly favor rows that contain "date/time" plus operational keywords.
    row_txt = safe_join(normalize_text(v) for v in row.tolist())
    time_bonus = 12 if ("date" in row_txt and "time" in row_txt) or "hh:mm" in row_txt else 0

    return kw * 4 + min(textish, 25) + min(non_empty, 25) * 0.25 + units + time_bonus


def detect_header_row(raw_df: pd.DataFrame, max_scan_rows: int = 80) -> int:
    """Find the first row of a multi-row header block.

    Many well-test templates use broad parent rows above the actual column labels:
    e.g. PRODUCTION / GAS RATE / RATIO, or TMU / LAB DATA / Gas Metering.
    This scores candidate header blocks instead of a single row.
    """
    if raw_df.empty:
        return 0

    scan_rows = min(max_scan_rows, len(raw_df))
    best_idx = 0
    best_score = -1.0

    for i in range(scan_rows):
        if row_looks_like_data(raw_df.iloc[i]):
            continue

        # Do not start a header block from a single metadata/reference cell such as
        # "14.73 Psi and 60 °F Gas line Meter Run Size...". Parent merged headers
        # like "Gas" or "Oil" are still allowed if they match PARENT_HEADERS.
        row_vals_norm = [normalize_text(v) for v in raw_df.iloc[i].tolist() if normalize_text(v)]
        row_txt = safe_join(row_vals_norm)
        has_parent_cell = any(v in PARENT_HEADERS for v in row_vals_norm)
        has_date_time_marker = ("date" in row_txt and "time" in row_txt) or "hh:mm" in row_txt
        if len(row_vals_norm) < 2 and not has_parent_cell and not has_date_time_marker:
            continue

        row_score = header_score(raw_df.iloc[i])
        if row_score < 0:
            continue

        group_score = row_score
        row_joined_for_group = safe_join([normalize_text(v) for v in raw_df.iloc[i].tolist()])
        group_has_date_time = "date" in row_joined_for_group or "hh:mm" in row_joined_for_group
        header_rows = 1

        # Look ahead and include subheader/unit rows until a data row starts.
        for r in range(i + 1, min(len(raw_df), i + 6)):
            if row_looks_like_data(raw_df.iloc[r]):
                break
            if is_header_like_subrow(raw_df.iloc[r]) or row_looks_like_units(raw_df.iloc[r]) or header_score(raw_df.iloc[r]) > 8:
                group_score += max(header_score(raw_df.iloc[r]), 0) * 0.95
                header_rows += 1
                row_txt = safe_join([normalize_text(v) for v in raw_df.iloc[r].tolist()])
                if "date" in row_txt or "hh:mm" in row_txt:
                    group_has_date_time = True
            else:
                # allow one blank row inside metadata, but not inside a header block
                break

        if group_has_date_time:
            group_score += 15
        if header_rows >= 2:
            group_score += 8

        if group_score > best_score:
            best_score = group_score
            best_idx = i

    return best_idx if best_score >= 10 else 0


def is_header_like_subrow(row: pd.Series) -> bool:
    vals = [v for v in row.tolist() if normalize_text(v)]
    if not vals:
        return False

    if row_looks_like_data(row):
        return False

    txt = safe_join([normalize_text(v) for v in vals])
    kw_count = sum(1 for kw in KEYWORDS if kw in txt)
    unit_count = sum(1 for kw in UNIT_KEYWORDS if kw in txt)
    numeric_count = sum(is_numeric_like(v) and not is_datetime_like(v) for v in vals)
    text_count = sum(not is_numeric_like(v) for v in vals)

    return (kw_count + unit_count >= 2 and text_count >= 2 and numeric_count <= max(3, text_count + 1))


def fill_parent_groups(row_values: Iterable[object]) -> List[str]:
    """Forward-fill merged parent/group headers across blank cells.

    This handles rows such as:
      PRESSURE AND TEMPERATURE MEASUREMENTS | ... | PRODUCTION | GAS RATE
    where Excel merged cells appear as a value followed by blanks.
    """
    vals = [str(v).strip() if normalize_text(v) else "" for v in row_values]
    non_empty = sum(1 for v in vals if v)
    total = len(vals)

    normalized_vals = [normalize_text(v) for v in vals if normalize_text(v)]
    has_parent_cell = any(v in PARENT_HEADERS for v in normalized_vals)
    short_group_row = all(len(v) <= 60 for v in normalized_vals)

    should_fill = non_empty > 0 and short_group_row and (
        non_empty <= max(3, int(total * 0.55))
        or has_parent_cell
    )

    if not should_fill:
        return vals

    out = []
    current = ""
    for v in vals:
        if v:
            current = v
            out.append(v)
        else:
            out.append(current)
    return out


def build_combined_headers(raw_df: pd.DataFrame, header_row: int, max_header_rows: int = 6):
    header_rows = [header_row]

    for r in range(header_row + 1, min(len(raw_df), header_row + max_header_rows)):
        if row_looks_like_data(raw_df.iloc[r]):
            break
        if is_header_like_subrow(raw_df.iloc[r]) or row_looks_like_units(raw_df.iloc[r]) or header_score(raw_df.iloc[r]) > 8:
            header_rows.append(r)
        else:
            break

    header_matrix = []
    for r in header_rows:
        vals = fill_parent_groups(raw_df.iloc[r].tolist())
        header_matrix.append(vals)

    headers = []
    for c in range(raw_df.shape[1]):
        parts = []
        for row_vals in header_matrix:
            raw = row_vals[c] if c < len(row_vals) else ""
            val = normalize_text(raw)
            if not val or val == "nan" or is_datetime_like(raw):
                continue
            part = str(raw).strip()
            if part and part.lower() != "nan" and part not in parts:
                parts.append(part)

        # Safety: some Excel templates may pass numeric/float objects into
        # the combined header parts through merged or unit rows. Convert again
        # before joining to avoid: sequence item X: expected str instance, float found.
        parts = [str(p).strip() for p in parts if str(p).strip() and str(p).strip().lower() != "nan"]
        header = safe_join(parts).strip() or f"Column_{c + 1}"
        # compact repeated whitespace and obvious duplicate fragments
        header = re.sub(r"\s+", " ", str(header))
        headers.append(str(header))

    return make_unique(headers), header_rows




def make_unique(names: Iterable[object]) -> List[str]:
    seen: Dict[str, int] = {}
    out: List[str] = []

    for n in names:
        # Safety against float/numeric headers from Excel merged/unit rows.
        name = str(n).strip() if str(n).strip() else "unnamed"
        name = re.sub(r"\s+", " ", name)
        if name.lower() == "nan":
            name = "unnamed"
        if name in seen:
            seen[name] += 1
            name = f"{name}.{seen[name]}"
        else:
            seen[name] = 0
        out.append(str(name))

    return out




def table_from_raw(raw_df: pd.DataFrame) -> pd.DataFrame:
    raw_df = raw_df.dropna(how="all").dropna(axis=1, how="all")
    if raw_df.empty:
        return raw_df

    hdr = detect_header_row(raw_df)
    headers, header_rows = build_combined_headers(raw_df, hdr)
    data_start = max(header_rows) + 1

    df = raw_df.iloc[data_start:].copy()
    df.columns = [str(h) for h in headers[: df.shape[1]]]
    df = df.dropna(how="all")
    return df



def canonical_candidate_score(canon: str, column_name: str) -> int:
    c = clean_header(column_name)
    score = 0

    # Prefer exact operational rate/pressure fields over broader parent labels.
    if canon == "bsw_pct":
        if "bsw" in c or "bs and w" in c:
            score += 20
        if "wh" in c:
            score -= 8
    if canon == "oil_rate_stbd" and ("oil rate" in c or "oil q" in c or "oil or cond" in c):
        score += 20
    if canon == "water_rate_bpd" and ("water rate" in c or "water q" in c or "water reading" in c):
        score += 20
    if canon == "gross_rate_bpd" and ("gross" in c or "liq q" in c or "liquid rate" in c):
        score += 20
    if canon == "gas_rate_mmscfd" and ("gas rate" in c or "mmscf/d" in c):
        score += 20
    if canon == "gas_formation_mmscfd" and ("formation gas" in c or "gas formation" in c):
        score += 30
    if canon == "pumping_pressure_psi" and ("pump p" in c or "pumping p" in c or "pumping.p" in c or "pump.p" in c or "pump pressure" in c or "pumping pressure" in c):
        score += 25
    if canon == "gas_formation_mmscfd" and ("formation gas" in c or "gas formation" in c):
        score += 30
    if canon == "pumping_pressure_psi" and ("pump p" in c or "pumping p" in c or "pumping.p" in c or "pump.p" in c or "pump pressure" in c or "pumping pressure" in c):
        score += 25
    if canon == "sep_p_psi" and ("sep" in c or "separator" in c):
        score += 20
    if canon == "whp_psi" and ("well head" in c or "fthp" in c or "whp" in c or "u/s" in c):
        score += 20
    if canon == "datetime" and ("date" in c and "time" in c):
        score += 20

    # Penalize total/cumulative/event columns when looking for rates.
    if canon in {"oil_rate_stbd", "water_rate_bpd", "gross_rate_bpd", "gas_rate_mmscfd"}:
        if "cum" in c or "volume" in c:
            score -= 12

    return score



def canonical_candidate_score(canon: str, column_name: str) -> int:
    c = clean_header(column_name)
    score = 0

    if canon == "bsw_pct":
        if "bsw" in c or "bs and w" in c:
            score += 25
        if "wh" in c:
            score -= 8
    if canon == "oil_rate_stbd" and ("oil rate" in c or "oil q" in c or ("oil or condensate" in c and "rate" in c)):
        score += 25
    if canon == "water_rate_bpd" and ("water rate" in c or "water q" in c or ("water reading" in c and "rate" in c)):
        score += 25
    if canon == "gross_rate_bpd" and ("gross" in c or "liq q" in c or "liquid rate" in c):
        score += 20
    if canon == "gas_rate_mmscfd" and ("gas rate" in c or "mmscf/d" in c):
        score += 20
    if canon == "sep_p_psi" and ("sep" in c or "separator" in c or "gasp" in c):
        score += 20
    if canon == "whp_psi" and ("well head" in c or "fthp" in c or "whp" in c or "u/s" in c):
        score += 20
    if canon == "datetime" and ("date" in c and "time" in c):
        score += 20
    if canon == "pump_intake_pressure_psi" and re.search(r"\b(pi|intake)\b", c):
        score += 30
    if canon == "pump_discharge_pressure_psi" and re.search(r"\b(pd|discharge|disch)\b", c):
        score += 30
    if canon == "motor_current_amp" and ("amp" in c or "current" in c or re.fullmatch(r"cur\.?", c)):
        score += 30
    if canon == "pump_freq_hz" and ("freq" in c or "hz" in c):
        score += 30
    if canon == "salinity_kppm":
        if "salinity" in c:
            score += 35
        if "nacl" in c:
            score += 35
        if re.search(r"\bppm\b", c):
            score += 10
        if "water" in c:
            score += 5
        if any(term in c for term in ["api", "kf", "factor", "psig", "psi", "bbl/d", "bpd", "stb/d", "mmscf", "scf/stb", "h2o", "rate"]):
            score -= 40

    if canon in {"oil_rate_stbd", "water_rate_bpd", "gross_rate_bpd", "gas_rate_mmscfd"}:
        if "cum" in c or "volume" in c:
            score -= 14

    return score


def best_canonical_name(column_name: str) -> Optional[str]:
    c = clean_header(column_name)
    original = normalize_text(column_name)
    cpad = f" {c} "

    if c.startswith("calcul") or c in {"factor", "fb", "ftf", "fg", "fpv", "y2"}:
        return None

    # Date/time. Avoid broad fragments like "d m" because they match "liquid metering".
    if ("date" in c and "time" in c) or "hh:mm" in c or re.search(r"dd[/\-]?\s*mon[/\-]?\s*yy", c) or "[dd-mm-yy]" in original:
        return "datetime"
    if re.fullmatch(r"(date|test date|start date|dd/mm/yy|ddmmyy)", c) or (re.search(r"\bdate\b", c) and "time" not in c and "last update" not in c):
        return "date"
    if re.fullmatch(r"(time|test time|hour|hh:mm:ss?|hh mm)", c) or ("time" in c and ("hh:mm" in c or "hh mm" in c)):
        return "time"
    if c in {"hhmm", "hh mm", "time hh mm", "time hh:mm"}:
        return "time"

    if re.fullmatch(r"well( name| no\.?)?", c) or c in {"well_name", "wellname"}:
        return "well"
    if "event" in c or "remark" in c or "comment" in c or re.search(r"\bnote\b", c):
        return "note"

    # Choke.
    if "choke" in c:
        if "size" in c or "/64" in c or '64"' in c:
            return "choke_size_64"
        return "choke_pct"

    # Temperatures before pressures, because parent headers often contain "well head" or "separator".
    if re.search(r"\bflt\b", c) or "flow line temp" in c or "flowline temp" in c:
        return "flt_f" if ("deg f" in c or "°f" in c or re.search(r"\bf\b", c)) else "flt_c"
    has_temp_child = bool(re.search(r"\btemp\.?\b", c)) or "deg f" in c or "deg. f" in c or "deg c" in c or "deg. c" in c or "°f" in c or "°c" in c
    if "flow temp" in c or "d/s temp" in c or "ds temp" in c or "downstream temp" in c:
        return "flow_temp_f" if ("deg f" in c or "°f" in c) else "flow_temp_c"
    if "u/s temp" in c or "upstream temp" in c:
        return "us_temp_c"
    if "gast" in c or re.search(r"\bgas\s+(gas\s+)?t\b", c) or ("gas" in c and ("gas temp" in c or "gas temperature" in c)):
        if re.search(r"\bgas\s+(gas\s+)?t\s+c\b", c) or re.search(r"\bgas\s+temp\s+c\b", c):
            return "gas_temp_c"
        if re.search(r"\bgas\s+(gas\s+)?t\s+f\b", c) or re.search(r"\bgas\s+temp\s+f\b", c):
            return "gas_temp_f"
        return "gas_temp_f" if ("deg f" in c or "°f" in c or re.search(r"\bf\b", c)) else "gas_temp_c"
    if "oilt" in c or re.search(r"\boil\s+(oil\s+)?t\b", c) or ("oil" in c and ("oil temp" in c or "oil temperature" in c)):
        if re.search(r"\boil\s+(oil\s+)?t\s+c\b", c) or re.search(r"\boil\s+temp\s+c\b", c):
            return "oil_temp_c"
        if re.search(r"\boil\s+(oil\s+)?t\s+f\b", c) or re.search(r"\boil\s+temp\s+f\b", c):
            return "oil_temp_f"
        return "oil_temp_f" if ("deg f" in c or "°f" in c or re.search(r"\bf\b", c)) else "oil_temp_c"
    if "oil or condensate" in c and ("deg f" in c or "deg. f" in c or "deg c" in c or "deg. c" in c):
        return "oil_temp_f" if ("deg f" in c or "deg. f" in c) else "oil_temp_c"
    if "well head" in c and has_temp_child and "psig" not in c and "psi-g" not in c and "psi" not in c:
        return "wellhead_temp_f" if ("deg f" in c or "deg. f" in c or "°f" in c) else "wellhead_temp_c"
    if ("separator" in c or "sep" in c) and has_temp_child and "psig" not in c and "psi-g" not in c and "psi" not in c and "inh2o" not in c:
        return "sep_temp_f" if ("deg f" in c or "deg. f" in c or "°f" in c) else "sep_temp_c"

    # Gross rate can be polluted by merged parent rows that contain "salinity".
    if "gross rate" in c and ("bbl/d" in c or "bpd" in c):
        return "gross_rate_bpd"

    # Fluid properties before rates because parent words can include "oil".
    if "bs and w" in c or "bs&w" in original or "bsw" in c or "water cut" in c or "watercut" in c or re.fullmatch(r"wc %?", c):
        return "bsw_pct"
    if "fw" == c or re.fullmatch(r"fw %?", c) or " lab data fw " in f" {c} ":
        return "fw_pct"
    if "twc" in c:
        return "twc_pct"
    if ("salinity" in c or "nacl" in c or "salt" in c) and "gross rate" not in c:
        # Avoid false salinity mapping caused by merged-header forward-fill.
        # Example bad headers from wide TMU sheets: "api Salinity BBL/D",
        # "[psig] Salinity BBL/D", "Factor Salinity BBL/D". Those are
        # helper/pressure/rate columns that inherited the word Salinity from a
        # neighboring column; they must not be plotted as salinity.
        salinity_conflicts = [
            "api", "kf", "factor", "psig", "psi", "psia", "bbl/d", "bpd",
            "stb/d", "mmscf", "scf/stb", "h2o", "gross rate", "oil t",
            "air =1", "chart", "orifice", "gor", "rate",
        ]
        has_conflict = any(term in c for term in salinity_conflicts)
        has_strong_salinity_unit = ("nacl" in c) or re.search(r"\bsalinity\b.*\bppm\b", c)
        if has_conflict and not has_strong_salinity_unit:
            return None
        return "salinity_kppm"
    if ("oil" in c and "sg" in c):
        return "oil_sg"
    if ("water" in c and "sg" in c):
        return "water_sg"
    if ("water" in c and "ph" in c) or re.fullmatch(r"ph", c):
        return "water_ph"
    if ("gravity" in c and "api" in c) or "oil api" in c or "grav. deg api" in c:
        return "oil_api"
    if ("gravity" in c and "air" in c) or "gas sg" in c or "gas gravity" in c or "air = 1" in c:
        return "gas_sg"
    if "sp.gr" in c or "sp gr" in c or "specific gravity" in c:
        return "gas_sg"
    if re.search(r"\bh2s\b", c):
        return "h2s_ppm"
    if re.search(r"\bco2\b|co₂", c):
        return "co2_mole_pct"

    # MPFM/EXPRO explicit columns.
    if "qoil" in c and "(s" in c:
        return "qoil_s_stbd"
    if "qwat" in c and "(s" in c:
        return "qwat_s_bpd"
    if "qgas" in c and "(s" in c:
        return "qgas_s_mmscfd"
    if "qoil" in c and "(a" in c:
        return "qoil_a_bpd"
    if "qwat" in c and "(a" in c:
        return "qwat_a_bpd"
    if "qgas" in c and "(a" in c:
        return "qgas_a_mmcfd"
    if "qgross" in c:
        return "qgross_s_bpd"
    if re.search(r"\bwlr\b", c):
        return "wlr_s_pct"
    if "gvf" in c:
        return "gvf_a_pct"

    # Rates / cumulative volumes. Order matters.
    if "oil cum" in c or ("oil" in c and "cum" in c) or re.search(r"\bcum\.?\s*stb\b", c):
        return "oil_cum_bbl"
    if "wat cum" in c or "water cum" in c or "water cumulative" in c:
        return "wat_cum_bbl"
    if "gas cum" in c:
        return "gas_cum_mscf"
    if "liquid volume" in c or "lquid volume" in c:
        return "liquid_volume_bbl"

    # Formation gas must be checked before total gas because headers can contain
    # "formation gas rate", which also includes the phrase "gas rate".
    if "gas formation" in c or "formation gas" in c:
        return "gas_formation_mmscfd"

    # Gas rate first because some templates have polluted headers such as
    # "GAS WATER READING RATE MMSCF/D" after merged-cell extraction.
    if ("total gas rate" in c) or ("gas rate" in c) or ("mmscf/d" in c and "gas" in c) or re.search(r"\brate mmscf/d\b", c):
        return "gas_rate_mmscfd"

    if "oil rate" in c or "oil q" in c or ("oil or condensate" in c and "rate" in c) or (re.search(r"\brate\b", c) and "stb/d" in c and "water" not in c and "gas" not in c and "gross" not in c):
        return "oil_rate_stbd"
    if "water rate" in c or "water q" in c or "bwpd" in c or ("water reading" in c and "rate" in c and "mmscf" not in c):
        return "water_rate_bpd"
    if "gross rate" in c or ("gross" in c and ("bbl" in c or "bpd" in c)):
        return "gross_rate_bpd"
    if "liq q" in c:
        return "gross_rate_bpd"
    if "liquid rate" in c:
        if "oil" in c:
            return "oil_rate_stbd"
        if "water" in c:
            return "water_rate_bpd"
        return "gross_rate_bpd"
    if re.search(r"\brate\b", c) and "bbl/d" in c and "oil" not in c and "gas" not in c and "gross" not in c:
        return "water_rate_bpd"

    # ESP / artificial-lift short aliases.  These are intentionally handled
    # before generic pressure/temperature rules because headers like Pi/Pd/Amp/Freq
    # are short and may otherwise remain raw fallback columns.  The strict fullmatch
    # rules avoid misclassifying broad text.
    if re.fullmatch(r"(pi|p/i|p int|pint|pump intake p|pump intake pressure|intake pressure|intake p|pin|pi psi|pi psig)", c):
        return "pump_intake_pressure_psi"
    if re.fullmatch(r"(pd|p/d|p dis|pdis|pump discharge p|pump discharge pressure|discharge pressure|disch pressure|discharge p|pd psi|pd psig)", c):
        return "pump_discharge_pressure_psi"
    if re.fullmatch(r"(amp|amps|ampere|amperage|current|motor current|pump current|run current|cur|cur\.?|i)", c) or ("amp" in c and "temp" not in c):
        return "motor_current_amp"
    if re.fullmatch(r"(ama|a m a)", c):
        return "motor_ama_amp"
    if re.fullmatch(r"(freq|frequency|hz|run freq|run frequency|operating freq|operating frequency|pump freq|pump frequency)", c):
        return "pump_freq_hz"
    if re.fullmatch(r"(drive freq|drive frequency|vsd freq|vfd freq|vfd frequency|speed hz)", c):
        return "drive_freq_hz"
    if re.fullmatch(r"(ti|t/i|intake temp|intake temperature|pump intake temp|pump intake temperature|ti c|ti deg c)", c):
        return "intake_temp_c"
    if re.fullmatch(r"(ti f|ti deg f|intake temp f|intake temperature f)", c):
        return "intake_temp_f"
    if re.fullmatch(r"(tm|t/m|motor temp|motor temperature|motor winding temp|tm c|tm deg c)", c):
        return "motor_temp_c"
    if re.fullmatch(r"(tm f|tm deg f|motor temp f|motor temperature f)", c):
        return "motor_temp_f"
    if re.fullmatch(r"(motor load|load pct|load %|motor load %)", c):
        return "motor_load_pct"
    if re.fullmatch(r"(vx|vib x|vibration x|x vibration)", c):
        return "vibration_x"
    if re.fullmatch(r"(vy|vib y|vibration y|y vibration)", c):
        return "vibration_y"
    if re.fullmatch(r"(vz|vib z|vibration z|z vibration)", c):
        return "vibration_z"

    # Pressures / DP.
    if re.search(r"\bflp\b|flow line pressure|flowline pressure", c):
        return "flp_psi"
    if "flow press" in c or "d/s press" in c or "ds press" in c or "downstream" in c:
        return "flow_press_psi"
    if re.search(r"\bfthp\b|\bwhp\b|w\.?h\.?p|wellhead pressure|well head pressure", c):
        return "whp_psi"
    if ("well head" in c or "upstream" in c or "u/s" in c) and ("press" in c or "psi" in c):
        return "whp_psi"
    if ("sep" in c or "separator" in c or "gasp" in c or "gas p" in c) and ("dp" in c or "diff" in c or "inh2o" in c):
        return "sep_dp_inh2o"
    if "gasdp" in c:
        return "gas_dp_inh2o"
    if ("sep" in c or "separator" in c or "gasp" in c or "gas p" in c) and ("p" in c or "pressure" in c or "psig" in c or "psi" in c):
        return "sep_p_psi"
    if " dp " in f" {c} " or (re.search(r"\bdiff\b", c) and ("h2o" in c or "mbar" in c)):
        return "dp_mbar" if "mbar" in c else "sep_dp_inh2o"
    if re.search(r"\bpump\s*[.\-/ ]?\s*p\b|\bpumping\s*[.\-/ ]?\s*p\b|pumping[. ]*pressure|pump[. ]*pressure|circulation pressure", c):
        return "pumping_pressure_psi"
    if re.search(r"\bct pressure\b|ct press", c):
        return "ct_pressure_psi"

    # Ratios.
    if "cgr" in c:
        return "cgr_bbl_mmscf"
    if re.search(r"\bgor\b", c):
        if "mmscf/bbl" in c:
            return "gor_mmscf_bbl"
        return "gor_scf_bbl"

    # Other operational parameters.
    if "orifice" in c and "size" in c:
        return "orifice_size_in"
    if "mtr inc" in c or "meter inc" in c or "meter increment" in c:
        return "oil_meter_increment_bbl"
    if "cmf" in c and "oil" in c:
        return "oil_cmf"
    if "k .f" in c or c == "kf":
        return "oil_kf"
    if "n2" in c or "nitrogen" in c:
        return "n2_rate_scfm"
    if "ct depth" in c or ("depth" in c and "ct" in c):
        return "ct_depth_m"
    if "running speed" in c:
        return "ct_running_speed_ftmin"
    if "pipe weight" in c:
        return "ct_pipe_weight_lbf"
    if "u2 pass" in c and "pressure" in c:
        return "u2_pass_side_pump_pressure_psi"
    if "pump freq" in c or ("freq" in c and ("hz" in c or "pump" in c)):
        return "pump_freq_hz"

    return None

def extract_number(value: object) -> float:
    """Strict numeric extraction.

    This intentionally reads only numbers at the start of the value.
    It avoids pulling 0.750 from long event text such as
    "Installed orifice plate size 0.750".
    """
    import datetime as _dt

    if value is None or (isinstance(value, float) and np.isnan(value)):
        return np.nan
    if isinstance(value, _dt.time):
        return np.nan
    if isinstance(value, (_dt.datetime, _dt.date, pd.Timestamp)):
        return np.nan
    if isinstance(value, (int, float, np.number)) and not isinstance(value, bool):
        return float(value)

    s = str(value).strip().replace(",", "")
    if not s or re.fullmatch(r"n/?a|nan|null|-", s, flags=re.I):
        return np.nan

    # Accept ordinary decimals and scientific notation.  XLSX XML may store
    # small calculated values as text such as ``9.827356E-2``.  The old regex
    # stopped before ``E-2`` and incorrectly returned 9.827356 instead of
    # 0.09827356.
    m = re.match(
        r"^\s*[-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][-+]?\d+)?",
        s,
    )
    if not m:
        return np.nan

    try:
        return float(m.group(0))
    except ValueError:
        return np.nan


def clean_numeric_series(series: pd.Series, canonical_name: str) -> pd.Series:
    s = series.map(extract_number).astype(float)

    # Excel often stores choke opening as a fraction (0.32 = 32%).
    if canonical_name == "choke_pct":
        valid = s.dropna()
        if not valid.empty and valid.median() <= 1.5:
            s = s.where(s > 1.5, s * 100.0)
        s = s.where((s >= 0) & (s <= 100))

    # Many field templates label the column as in/64 but store 30/64 as
    # Excel 30% (raw XML value 0.30). Restore the displayed numerator.
    if canonical_name == "choke_size_64":
        valid = s.dropna()
        if not valid.empty and valid.max() < 1.0:
            s = s * 100.0
        s = s.where((s >= 0) & (s <= 256))

    # BS&W can be stored either as percent (1 = 1%) or fraction (0.006 = 0.6%).
    # Convert only clear fraction-style values, but do not turn 1% into 100%.
    if canonical_name == "bsw_pct":
        valid = s.dropna()
        if not valid.empty and valid.max() <= 1.0 and valid.median() < 0.25:
            s = s * 100.0

    # Store salinity as K ppm because most plots/report templates use K ppm NaCl.
    if canonical_name == "salinity_kppm":
        med = s.dropna().median() if s.notna().any() else np.nan
        if pd.notna(med) and med > 1000:
            s = s / 1000.0

    return s


def _datetime_year_is_safe_v55(value, min_year: int = 1900, max_year: int = 2100) -> bool:
    try:
        if pd.isna(value):
            return False
    except Exception:
        pass
    try:
        year = int(pd.Timestamp(value).year)
        return min_year <= year <= max_year
    except Exception:
        return False


def _safe_datetime_scalar_v55(value, *, dayfirst: bool = True, allow_excel_serial: bool = False):
    """Parse one date/datetime without allowing numeric measurements to become years.

    Pandas 3 on Python 3.14 may parse a plain value such as 1423 or 2613 as a
    real year using second-resolution timestamps. Assigning that result into a
    nanosecond datetime Series then raises OutOfBoundsDatetime or an internal
    AssertionError. TMU files also contain many four-digit pressures/rates, so
    plain numeric values must not be treated as calendar years.
    """
    import datetime as _dt

    try:
        if value is None or pd.isna(value):
            return pd.NaT
    except Exception:
        if value is None:
            return pd.NaT

    if isinstance(value, pd.Timestamp):
        return value if _datetime_year_is_safe_v55(value) else pd.NaT
    if isinstance(value, _dt.datetime):
        ts = pd.Timestamp(value)
        return ts if _datetime_year_is_safe_v55(ts) else pd.NaT
    if isinstance(value, _dt.date):
        ts = pd.Timestamp(value)
        return ts if _datetime_year_is_safe_v55(ts) else pd.NaT

    if isinstance(value, (int, float, np.number)) and not isinstance(value, bool):
        try:
            number = float(value)
        except Exception:
            return pd.NaT
        # Normal Excel serial dates for the supported 1900-2100 window are
        # approximately 1..73415. Restrict further to modern field data to avoid
        # interpreting pressures/rates such as 1423 or 2613 as dates.
        if allow_excel_serial and 20000.0 <= number <= 80000.0:
            try:
                ts = pd.Timestamp("1899-12-30") + pd.to_timedelta(number, unit="D")
                return ts if _datetime_year_is_safe_v55(ts) else pd.NaT
            except Exception:
                return pd.NaT
        return pd.NaT

    text = str(value).strip()
    if not text or text.lower() in {"nan", "nat", "none", "null", "-"}:
        return pd.NaT

    # Reject plain numeric strings. In these workbooks they are overwhelmingly
    # rates, pressures, depths, or counts—not dates. Permit YYYYMMDD explicitly.
    compact = text.replace(",", "")
    if re.fullmatch(r"[-+]?\d+(?:\.\d+)?", compact):
        if re.fullmatch(r"\d{8}", compact):
            try:
                ts = pd.to_datetime(compact, format="%Y%m%d", errors="coerce")
                return ts if _datetime_year_is_safe_v55(ts) else pd.NaT
            except Exception:
                return pd.NaT
        if allow_excel_serial:
            try:
                number = float(compact)
                if 20000.0 <= number <= 80000.0:
                    ts = pd.Timestamp("1899-12-30") + pd.to_timedelta(number, unit="D")
                    return ts if _datetime_year_is_safe_v55(ts) else pd.NaT
            except Exception:
                pass
        return pd.NaT

    # Require a recognisable date marker before calling pandas/dateutil. This
    # prevents arbitrary operational text or isolated values from becoming dates.
    date_like = bool(
        re.search(r"\d{1,4}[/-]\d{1,2}[/-]\d{1,4}", text)
        or re.search(r"\d{1,2}[ -](?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[ -]\d{2,4}", text, re.I)
        or re.search(r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[ -]\d{1,2}[, -]+\d{2,4}", text, re.I)
        or re.match(r"^\d{4}-\d{1,2}-\d{1,2}[T ]", text)
    )
    if not date_like:
        return pd.NaT

    iso = bool(re.match(r"^\d{4}[-/]\d{1,2}[-/]\d{1,2}", text))
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            ts = pd.to_datetime(text, errors="coerce", dayfirst=False if iso else dayfirst)
    except Exception:
        return pd.NaT
    return ts if _datetime_year_is_safe_v55(ts) else pd.NaT


def _safe_datetime_series_v55(series: pd.Series, *, dayfirst: bool = True,
                              allow_excel_serial: bool = False,
                              normalize_date: bool = False) -> pd.Series:
    """Return a nanosecond-safe Series without assigning mixed datetime units."""
    values = []
    for value in series.tolist():
        ts = _safe_datetime_scalar_v55(
            value, dayfirst=dayfirst, allow_excel_serial=allow_excel_serial
        )
        if normalize_date and pd.notna(ts):
            try:
                ts = pd.Timestamp(ts).normalize()
            except Exception:
                ts = pd.NaT
        values.append(ts)

    # All out-of-range values have already been removed, so this construction is
    # safe on pandas 2.x and 3.x and avoids .loc assignment between datetime units.
    try:
        arr = pd.to_datetime(values, errors="coerce")
        return pd.Series(arr, index=series.index, dtype="datetime64[ns]")
    except Exception:
        clean = []
        for value in values:
            clean.append(value if _datetime_year_is_safe_v55(value) else pd.NaT)
        return pd.Series(clean, index=series.index, dtype="datetime64[ns]")


def parse_date_series(series: pd.Series) -> pd.Series:
    # Date columns may contain genuine Excel serial dates, but only within the
    # modern safe range. Four-digit measurements are rejected.
    return _safe_datetime_series_v55(
        series, dayfirst=True, allow_excel_serial=True, normalize_date=True
    )


def parse_time_series(series: pd.Series) -> pd.Series:
    import datetime as _dt

    def one(x):
        try:
            if x is None or pd.isna(x):
                return pd.NaT
        except Exception:
            if x is None:
                return pd.NaT

        if isinstance(x, pd.Timestamp):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, _dt.datetime):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, _dt.time):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, (_dt.timedelta, pd.Timedelta)):
            try:
                seconds = int(round(pd.Timedelta(x).total_seconds())) % 86400
                return pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)
            except Exception:
                return pd.NaT

        if isinstance(x, (int, float, np.number)) and not isinstance(x, bool):
            try:
                xf = float(x)
                # Some TMU templates keep cumulative Excel time values above 1
                # (for example 2.0833 = 02:00 on a later copied date row). Use
                # only the fractional day, while rejecting large measurements.
                if 0 <= xf < 10:
                    seconds = int(round((xf % 1.0) * 24 * 3600)) % 86400
                    return pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)
            except Exception:
                return pd.NaT

        text = str(x).strip()
        if re.fullmatch(r"[-+]?\d+(?:\.\d+)?", text):
            try:
                xf = float(text)
                if 0 <= xf < 10:
                    seconds = int(round((xf % 1.0) * 24 * 3600)) % 86400
                    return pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)
            except Exception:
                return pd.NaT
        m = re.fullmatch(r"\s*(\d{1,2})[:.](\d{2})(?:[:.](\d{2}))?\s*(am|pm)?\s*", text, re.I)
        if m:
            hh = int(m.group(1))
            mm = int(m.group(2))
            ss = int(m.group(3) or 0)
            ampm = (m.group(4) or "").lower()
            if ampm == "pm" and hh < 12:
                hh += 12
            if ampm == "am" and hh == 12:
                hh = 0
            if 0 <= hh <= 23 and 0 <= mm <= 59 and 0 <= ss <= 59:
                return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=hh, minutes=mm, seconds=ss)
        return pd.NaT

    values = [one(x) for x in series.tolist()]
    return pd.Series(values, index=series.index, dtype="datetime64[ns]")


def parse_datetime_series(series: pd.Series) -> pd.Series:
    """Parse mixed datetimes while rejecting numeric measurements as years.

    This implementation deliberately avoids assigning pandas 3 second-resolution
    timestamps into a nanosecond Series, which caused the internal pandas error
    reported on Python 3.14.
    """
    header = clean_header(getattr(series, "name", ""))
    allow_serial = bool(re.search(r"date|datetime|timestamp", header))
    return _safe_datetime_series_v55(
        series, dayfirst=True, allow_excel_serial=allow_serial, normalize_date=False
    )

def combine_date_time(
    date_series: Optional[pd.Series] = None,
    time_series: Optional[pd.Series] = None,
    datetime_series: Optional[pd.Series] = None,
) -> pd.Series:
    """Build a real datetime column from any available date/time inputs.

    Important fix: Excel files often contain monthly/daily production history with a
    Date column only and no separate Time column. Older versions returned all NaT
    when time_series was missing, causing duplicate-removal to collapse the whole
    sheet into one row and the workbook was rejected as "no usable time-series".
    Date-only rows are now valid readings at midnight (00:00).
    """
    idx = None
    for s in [date_series, time_series, datetime_series]:
        if s is not None:
            idx = s.index
            break

    if idx is None:
        return pd.Series(dtype="datetime64[ns]")

    if datetime_series is not None:
        dt = parse_datetime_series(datetime_series)
    else:
        dt = pd.Series(pd.NaT, index=idx, dtype="datetime64[ns]")

    dates = parse_date_series(date_series) if date_series is not None else pd.Series(pd.NaT, index=idx, dtype="datetime64[ns]")
    times = parse_time_series(time_series) if time_series is not None else pd.Series(pd.NaT, index=idx, dtype="datetime64[ns]")

    out = []
    for d, t, existing in zip(dates, times, dt):
        if pd.notna(existing):
            out.append(pd.Timestamp(existing))
        elif pd.notna(d) and pd.notna(t):
            out.append(pd.Timestamp(d.date()) + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second))
        elif pd.notna(d):
            # Date-only production-history sheets are valid time series.
            out.append(pd.Timestamp(d.date()))
        elif pd.notna(t):
            out.append(pd.Timestamp("1900-01-01") + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second))
        else:
            out.append(pd.NaT)

    return pd.Series(out, index=idx, dtype="datetime64[ns]")





def adjust_datetime_rollover(out: pd.DataFrame, parsed_time: Optional[pd.Series]) -> pd.DataFrame:
    """Fix midnight rollover in templates where DATE is copied down but TIME resets from 23:30 to 00:00.

    Example: date column remains 2026-06-05, rows are 23:30 then 00:00.
    The 00:00 row is really 2026-06-06 00:00. We adjust in original row order before sorting.
    """
    if out is None or out.empty or parsed_time is None:
        return out
    if "date" not in out.columns or "datetime" not in out.columns:
        return out
    if parsed_time.notna().sum() < 2:
        return out

    dates = pd.to_datetime(out["date"], errors="coerce")
    times = parsed_time
    adjusted = out["datetime"].copy()

    day_offset = 0
    prev_valid_time = None
    prev_base_date = None

    for idx in out.index:
        d = dates.loc[idx]
        t = times.loc[idx]
        if pd.isna(d) or pd.isna(t):
            continue

        base_date = pd.Timestamp(d.date())

        # If the sheet explicitly advances the DATE column, reset offset.
        if prev_base_date is not None and base_date > prev_base_date:
            day_offset = 0

        # If DATE did not advance but TIME moved backwards, assume midnight rollover.
        if (
            prev_valid_time is not None
            and prev_base_date is not None
            and base_date == prev_base_date
            and (t.hour, t.minute, t.second) < (prev_valid_time.hour, prev_valid_time.minute, prev_valid_time.second)
        ):
            day_offset += 1

        adjusted.loc[idx] = base_date + pd.Timedelta(days=day_offset) + pd.Timedelta(
            hours=t.hour, minutes=t.minute, seconds=t.second
        )

        prev_valid_time = t
        prev_base_date = base_date

    out = out.copy()
    out["datetime"] = adjusted
    out["date"] = pd.to_datetime(out["datetime"], errors="coerce").dt.floor("D")
    out["time_text"] = pd.to_datetime(out["datetime"], errors="coerce").dt.strftime("%H:%M")
    return out


def normalize_well_name(value: object) -> Optional[str]:
    if value is None or pd.isna(value):
        return None

    s = str(value).strip()
    if not s:
        return None

    # Remove obvious label prefixes.
    s = re.sub(r"(?i)\bwell\s*(name|no\.?)?\s*[:=#-]*", "", s).strip()
    s = s.replace("#", "-")
    s = re.sub(r"\s*-\s*", "-", s)
    s = re.sub(r"\s+", " ", s).strip(" :-_")
    if not s:
        return None

    # Reject plain dates/numbers.
    if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", s):
        return None
    if re.fullmatch(r"\d+(\.\d+)?", s):
        return None

    # Normalize common forms.
    up = s.upper()

    # BAHGA 11 -> BAHGA-11
    up = re.sub(r"\b(BAHGA|ASSIL|KARAM|MAGD|ELMAGD|BED|SITRA|BAPETCO)\s+(\d)", r"\1-\2", up)

    # MAGD C 86-4 / ELMAGD 86-4 kept readable.
    up = re.sub(r"\bMAGD-C-", "MAGD C-", up)
    up = re.sub(r"\bELMAGD-", "ELMAGD-", up)

    # Remove duplicate spaces and spaces around hyphens.
    up = re.sub(r"\s*-\s*", "-", up)
    up = re.sub(r"\s+", " ", up).strip()

    # Keep not-too-long candidates only.
    if len(up) > 30:
        return None

    return up


def guess_well_from_name(name: str) -> Optional[str]:
    s = str(name)

    # Remove date-like parenthetical blocks to avoid returning 03-07 or 8-12 as a well.
    s_clean = re.sub(r"\([^)]*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}[^)]*\)", " ", s)
    s_clean = re.sub(r"\b\d{4}[-_]\d{2}[-_]\d{2}\b", " ", s_clean)

    patterns = [
        r"\b(BAHGA\s*#?\s*\d+[A-Z]?)\b",
        r"\b(MAGD\s*C\s*\d+\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b(ELMAGD\s*\d+\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b(BED\s*\d+[A-Z]?\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b(B\d+[A-Z]*\d*\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b([A-Z][A-Z0-9]*\s*C?\s*\d+\s*[-_]\s*\d+[A-Z]?)\b",
    ]

    for pat in patterns:
        m = re.search(pat, s_clean, flags=re.I)
        if m:
            well = normalize_well_name(m.group(1))
            # Avoid false matches created by spreadsheet column labels plus dates, e.g. "Y2 06-06".
            if well and re.fullmatch(r"[XYZ]\d+[- ]\d{1,2}[-_/]\d{1,2}", well):
                continue
            if well:
                return well

    return None


def extract_well_from_raw(raw_df: pd.DataFrame, source_name: str = "", sheet_name: str = "") -> Optional[str]:
    """Extract well name from workbook metadata rows before falling back to filename."""
    try:
        max_r = min(len(raw_df), 40)
        max_c = min(raw_df.shape[1], 30)

        for r in range(max_r):
            row_vals = [raw_df.iloc[r, c] for c in range(max_c)]
            for c, val in enumerate(row_vals):
                txt = str(val).strip() if not pd.isna(val) else ""
                if not txt:
                    continue

                # Cell contains the whole value: "WELL: Bahga#9"
                m = re.search(r"(?i)\bwell\s*(?:name|no\.?)?\s*[:=#-]+\s*([A-Za-z0-9# \-_]+)", txt)
                if m:
                    well = normalize_well_name(m.group(1))
                    if well:
                        return well

                # Label in this cell, value in nearby cells to the right.
                if re.search(r"(?i)^\s*well\s*(name|no\.?)?\s*:?\s*$", txt):
                    for k in range(c + 1, min(c + 5, max_c)):
                        cand = normalize_well_name(raw_df.iloc[r, k])
                        if cand:
                            return cand

                # Some templates use "Well  :" in one cell and the next cell has value.
                if re.search(r"(?i)\bwell\b", txt) and len(txt) <= 20:
                    for k in range(c + 1, min(c + 5, max_c)):
                        cand = normalize_well_name(raw_df.iloc[r, k])
                        if cand and not cand.startswith("TEST"):
                            return cand

        # As a last metadata attempt, look for known well patterns in all early text.
        text_blob = safe_join(
            [raw_df.iloc[r, c] for r in range(max_r) for c in range(max_c)]
        )
        guessed = guess_well_from_name(text_blob)
        if guessed:
            return guessed

    except Exception:
        pass

    return guess_well_from_name(f"{source_name} {sheet_name}")



def raw_output_column_name(column_name: object, existing_columns: Iterable[str]) -> str:
    """Create a stable fallback name for numeric columns that do not match aliases yet."""
    label = str(column_name).strip() if str(column_name).strip() else "Column"
    safe = clean_header(label)
    safe = re.sub(r"[^a-z0-9]+", "_", safe).strip("_")
    if not safe:
        safe = "column"
    safe = safe[:80]
    base = f"raw__{safe}"
    name = base
    i = 2
    existing = set(existing_columns)
    while name in existing:
        name = f"{base}_{i}"
        i += 1
    return name


def is_useful_raw_numeric_column(column_name: object, series: pd.Series) -> bool:
    """Decide whether an unmapped column should still appear as a user-selectable raw column."""
    header = clean_header(column_name)
    if not header or header.startswith("unnamed") or header.startswith("column_"):
        return False
    if header in {"date", "time", "datetime", "well", "note", "event", "comments", "remarks"}:
        return False
    if "event" in header or "remark" in header or "comment" in header:
        return False

    # Avoid hidden calculation/helper columns that may be present far to the right
    # of operational Excel templates.
    helper_terms = [
        "calcul", "calc", "factor", "ftf", "fg", "fb", "fpv", "y2",
        "h2o", "psia", "base condition", "intermediate", "correction",
    ]
    if any(term in header for term in helper_terms):
        return False

    nums = series.map(extract_number).astype(float)
    return nums.notna().any()



def _choke_header_kind_v56(header: object) -> str:
    """Return pct, size64, or ambiguous without mixing the two units."""
    raw = str(header or "").lower()
    h = clean_header(header)
    if re.search(r"/\s*64|64\s*(?:th|ths)|in\s*/\s*64|size", raw + " " + h, flags=re.I):
        return "size64"
    if "%" in raw or re.search(r"\bpercent(?:age)?\b|\bopening\b", h):
        return "pct"
    return "ambiguous"


def _split_choke_series_v56(series: pd.Series, header: object):
    """Split one source choke column without guessing missing units.

    Returns percentage, /64-size, ambiguous-raw, and a confidence score.
    Explicit text/header units are authoritative. Bare values under a plain
    "Choke" header are preserved in ``choke_ambiguous`` so the user can choose
    percentage or /64 interpretation in the dashboard.
    """
    idx = series.index
    pct = pd.Series(np.nan, index=idx, dtype=float)
    size = pd.Series(np.nan, index=idx, dtype=float)
    ambiguous = pd.Series(np.nan, index=idx, dtype=float)
    kind = _choke_header_kind_v56(header)
    text = series.astype(str).str.strip()
    explicit_fraction = text.str.contains(r"[-+]?\d+(?:\.\d+)?\s*/\s*64\b", regex=True, case=False, na=False)
    explicit_pct = text.str.contains(r"%|\bpercent(?:age)?\b", regex=True, case=False, na=False)
    nums = series.map(extract_number).astype(float)
    column_has_fraction = bool(explicit_fraction.any())

    if explicit_fraction.any():
        size.loc[explicit_fraction] = nums.loc[explicit_fraction]
    if explicit_pct.any():
        vals = nums.loc[explicit_pct]
        pct.loc[explicit_pct] = vals.where(vals > 1.0, vals * 100.0)

    remaining = ~(explicit_fraction | explicit_pct) & nums.notna()
    vals = nums.loc[remaining]
    if kind == "size64":
        # Several TMU templates store displayed 30/64 as raw XML 0.30.
        size.loc[remaining] = vals.where(vals >= 1.0, vals * 100.0)
    elif kind == "pct":
        pct.loc[remaining] = vals.where(vals > 1.0, vals * 100.0)
    elif column_has_fraction:
        # Once the same column explicitly contains /64 entries, other bare
        # values above 1 are most consistently interpreted as /64 numerators.
        # Values <=1 are typical fractional percentage entries.
        frac_open = vals <= 1.0
        pct.loc[vals.index[frac_open]] = vals.loc[frac_open] * 100.0
        as_size = (vals > 1.0) & (vals <= 256.0)
        size.loc[vals.index[as_size]] = vals.loc[as_size]
    else:
        # No unit evidence: preserve the raw number and let the user decide.
        ambiguous.loc[remaining] = vals

    pct = pct.where((pct >= 0) & (pct <= 100)).round(3)
    size = size.where((size >= 0) & (size <= 256)).round(3)
    ambiguous = ambiguous.where((ambiguous >= 0) & (ambiguous <= 256)).round(3)
    explicit_score = 120 if kind in {"pct", "size64"} else 70
    return pct, size, ambiguous, explicit_score


def _collect_choke_columns_v56(df: pd.DataFrame):
    pct_candidates = []
    size_candidates = []
    ambiguous_candidates = []
    for col in df.columns:
        header = str(col)
        if "choke" not in clean_header(header):
            continue
        pct, size, ambiguous, score = _split_choke_series_v56(df[col], header)
        kind = _choke_header_kind_v56(header)
        if pct.notna().any():
            pct_candidates.append((score + (20 if kind == "pct" else 0), pct))
        if size.notna().any():
            size_candidates.append((score + (20 if kind == "size64" else 0), size))
        if ambiguous.notna().any():
            ambiguous_candidates.append((score, ambiguous))

    def combine(candidates):
        if not candidates:
            return None
        result = pd.Series(np.nan, index=df.index, dtype=float)
        for _, candidate in sorted(candidates, key=lambda x: x[0], reverse=True):
            result = result.combine_first(candidate)
        return result

    return combine(pct_candidates), combine(size_candidates), combine(ambiguous_candidates)


def standardize_dataframe(
    df: pd.DataFrame,
    source_name: str = "",
    sheet_name: str = "",
    default_well: Optional[str] = None,
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    # Drop fully empty rows, but do NOT blindly drop fully empty columns here.
    # Some templates have a real field column whose values are all N/A, such as
    # Salinity = N/A. If we drop that exact Salinity column first, a later
    # helper/calculation column with a polluted merged header like
    # "api Salinity BBL/D" can be selected by mistake and plotted as salinity.
    df = df.copy().dropna(how="all")
    if df.empty:
        return pd.DataFrame()

    keep_cols = []
    for col in df.columns:
        if not df[col].isna().all() or best_canonical_name(str(col)) is not None:
            keep_cols.append(col)
    df = df.loc[:, keep_cols]
    if df.empty:
        return pd.DataFrame()

    # Choose the best source column for each canonical field.
    best_map: Dict[str, tuple] = {}
    for col in df.columns:
        canon = best_canonical_name(str(col))
        if not canon:
            continue
        score = canonical_candidate_score(canon, str(col))
        if canon not in best_map or score > best_map[canon][1]:
            best_map[canon] = (col, score)

    mapping = [(canon, col_score[0]) for canon, col_score in best_map.items()]

    out = pd.DataFrame(index=df.index)
    out["source"] = source_name
    out["sheet"] = sheet_name

    mapped_source_cols = set()
    for canon, col in mapping:
        mapped_source_cols.add(col)
        if canon in ["well", "date", "time", "datetime", "note"]:
            out[canon] = df[col]
        else:
            out[canon] = clean_numeric_series(df[col], canon)

    # Choke needs content-aware unit handling. A single workbook family may use
    # percentage, /64-inch size, or even both in separate columns. Rebuild the
    # two canonical series from the original source columns and never combine
    # unlike units into one plotted curve.
    choke_pct_v56, choke_size_v56, choke_ambiguous_v57 = _collect_choke_columns_v56(df)
    mapped_choke_header = str(best_map.get("choke_pct", ("", 0))[0]) if "choke_pct" in best_map else ""
    mapped_choke_kind = _choke_header_kind_v56(mapped_choke_header) if mapped_choke_header else "ambiguous"
    if choke_pct_v56 is not None and choke_pct_v56.notna().any():
        out["choke_pct"] = choke_pct_v56
    elif "choke_pct" in out.columns and mapped_choke_kind == "pct":
        out["choke_pct"] = pd.to_numeric(out["choke_pct"], errors="coerce").where(lambda x: (x >= 0) & (x <= 100))
    else:
        out.drop(columns=["choke_pct"], inplace=True, errors="ignore")
    if choke_size_v56 is not None and choke_size_v56.notna().any():
        out["choke_size_64"] = choke_size_v56
    elif "choke_size_64" in out.columns:
        out["choke_size_64"] = pd.to_numeric(out["choke_size_64"], errors="coerce").where(lambda x: (x >= 0) & (x <= 256))
    if choke_ambiguous_v57 is not None and choke_ambiguous_v57.notna().any():
        out["choke_ambiguous"] = choke_ambiguous_v57

    # Fallback: keep numeric columns not yet recognized by aliases only when the
    # table is mostly unknown. If a TMU template already produced many canonical
    # fields (WHP, FLP, Sep P, Gas rate, Oil rate, etc.), keeping every extra
    # helper/calculation/unit column creates confusing labels such as
    # "Raw: Column" or "Raw: Psig" and can even cause the Excel loader to choose
    # a lower-quality unit-row parse. Unknown templates still get raw columns.
    mapped_numeric_count = sum(
        1 for canon, _ in mapping
        if canon not in {"well", "date", "time", "datetime", "note"}
    )
    keep_raw_fallback = mapped_numeric_count < 6

    for col in df.columns:
        if col in mapped_source_cols:
            continue
        if not keep_raw_fallback:
            continue
        if not is_useful_raw_numeric_column(col, df[col]):
            continue
        raw_name = raw_output_column_name(col, out.columns)
        out[raw_name] = df[col].map(extract_number).astype(float)

    # Parse date/time/datetime robustly.
    if "datetime" in out.columns:
        out["datetime"] = parse_datetime_series(out["datetime"])

    if "date" in out.columns:
        out["date"] = parse_date_series(out["date"])
        out["date"] = out["date"].ffill().bfill()

    parsed_time = None
    if "time" in out.columns:
        parsed_time = parse_time_series(out["time"])
        out["time_text"] = parsed_time.dt.strftime("%H:%M")
        out.loc[parsed_time.isna(), "time_text"] = ""

    if ("datetime" not in out.columns or out["datetime"].isna().all()) and ("date" in out.columns or "time" in out.columns):
        out["datetime"] = combine_date_time(out.get("date"), out.get("time"), out.get("datetime"))

    # Derive date/time_text from datetime when a combined Date & Time column exists.
    if "datetime" in out.columns and out["datetime"].notna().any():
        if "date" not in out.columns:
            out["date"] = out["datetime"].dt.floor("D")
        if "time_text" not in out.columns:
            out["time_text"] = out["datetime"].dt.strftime("%H:%M")

    # Correct midnight rollover before sorting.
    if parsed_time is not None:
        out = adjust_datetime_rollover(out, parsed_time)

    # Well name: source metadata > sheet/filename fallback.
    if "well" in out.columns:
        out["well"] = out["well"].map(lambda x: normalize_well_name(x) or "")
    else:
        out["well"] = ""

    if out["well"].astype(str).str.strip().eq("").all():
        well = default_well or guess_well_from_name(f"{source_name} {sheet_name}") or "Unknown"
        out["well"] = well
    else:
        fallback = default_well or guess_well_from_name(f"{source_name} {sheet_name}") or "Unknown"
        out["well"] = out["well"].replace("", np.nan).ffill().bfill().fillna(fallback)

    numeric_cols = [c for c in out.columns if c not in BASE_NON_PLOT_COLS]
    canonical_numeric_cols = [c for c in numeric_cols if not str(c).startswith("raw__")]
    row_filter_cols = canonical_numeric_cols if canonical_numeric_cols else numeric_cols

    if row_filter_cols:
        useful = out[row_filter_cols].notna().any(axis=1)

        # For tables with a separate TIME column, require a valid time.
        # This removes event rows and rows such as Final Average that contain numbers but are not readings.
        if parsed_time is not None and parsed_time.notna().sum() >= 2:
            useful = useful & parsed_time.notna()
        elif "datetime" in out.columns and out["datetime"].notna().sum() >= 2:
            useful = useful & out["datetime"].notna()

        # Drop pure average/summary/total rows even if the date column was forward-filled.
        raw_text_for_row = df.apply(lambda r: safe_join(list(r.values), " | "), axis=1).str.lower()
        useful = useful & ~raw_text_for_row.str.contains(
            r"final\s+average|average|avg|summary|grand\s+total|^total",
            regex=True,
            na=False,
        )

        out = out.loc[useful].copy()

    if out.empty:
        return pd.DataFrame()

    # v67: repair isolated wrong Date cells and midnight rollover while rows
    # are still in the original worksheet order.  Sorting first would move an
    # isolated bad date to the end and make it impossible to recognize.
    out = _repair_measurement_datetime_sequence_v67(out)

    # Remove duplicate readings at the same well/datetime. This also cleans old dashboard-exported CSVs
    # that may contain a duplicated Final Average row with a copied timestamp.
    if "datetime" in out.columns and "well" in out.columns and out["datetime"].notna().any():
        # Do not let all-NaT datetimes collapse a valid table to one row.
        out = out.drop_duplicates(subset=["well", "datetime"], keep="first")

    out["well"] = out["well"].astype(str).str.strip().replace("", "Unknown")
    sort_cols = ["well"] + (["datetime"] if "datetime" in out.columns else [])
    out = out.sort_values(sort_cols, na_position="last").reset_index(drop=True)

    return out


def dataframe_quality_score(df: pd.DataFrame) -> int:
    """Score parsed tables so the Excel loader keeps the real operational table.

    The old scoring over-valued the number of numeric columns. In wide TMU
    spreadsheets, a bad unit-row parse can create dozens of raw columns such as
    raw__psig/raw__column and beat the proper multi-row header parse. This scorer
    now strongly prefers:
      - recognized canonical TMU columns,
      - real calendar datetimes (not 1900 time-only placeholders),
      - date + time together,
      - fewer raw fallback columns.
    """
    if df is None or df.empty:
        return -10_000

    numeric_cols = available_numeric_columns(df)
    if not numeric_cols:
        return -10_000

    canonical_cols = [c for c in numeric_cols if not str(c).startswith("raw__")]
    raw_cols = [c for c in numeric_cols if str(c).startswith("raw__")]

    score = 0
    score += min(len(df), 500) * 4
    score += len(canonical_cols) * 300
    score += len(raw_cols) * 5

    # Penalize raw-only parses. They are useful as a fallback, but should not win
    # over a table where columns have been identified as WHP, FLP, gas rate, etc.
    if not canonical_cols:
        score -= 700
    if raw_cols and len(raw_cols) > max(5, len(canonical_cols) * 2):
        score -= (len(raw_cols) - max(5, len(canonical_cols) * 2)) * 20

    if "datetime" in df.columns:
        dt = pd.to_datetime(df["datetime"], errors="coerce")
        valid_dt = dt.dropna()
        score += int(valid_dt.size) * 12
        real_dt_count = int((valid_dt.dt.year > 1970).sum()) if not valid_dt.empty else 0
        score += real_dt_count * 45
        # Time-only fallback creates dates around 1900. It can be plotted, but it
        # must never beat a parse with real file dates.
        if valid_dt.size and real_dt_count == 0:
            score -= 900

    if "date" in df.columns:
        dates = pd.to_datetime(df["date"], errors="coerce").dropna()
        real_date_count = int((dates.dt.year > 1970).sum()) if not dates.empty else 0
        score += real_date_count * 25

    if "time_text" in df.columns:
        score += int(df["time_text"].astype(str).str.strip().ne("").sum()) * 5

    # Prefer the common field-test columns that users expect to see by name.
    priority_cols = {
        "choke_pct", "choke_ambiguous", "whp_psi", "flp_psi", "flt_c", "sep_p_psi",
        "gas_rate_mmscfd", "oil_rate_stbd", "water_rate_bpd", "gross_rate_bpd",
        "bsw_pct", "salinity_kppm",
    }
    score += sum(1 for c in priority_cols if c in df.columns) * 150

    return int(score)

def dataframe_key(df: pd.DataFrame) -> tuple:
    """Return a loose de-duplication key for repeated parsing attempts."""
    if df is None or df.empty:
        return (0, "", "")
    sheet = str(df["sheet"].iloc[0]) if "sheet" in df.columns and len(df) else ""
    well = str(df["well"].iloc[0]) if "well" in df.columns and len(df) else ""
    dt_min = ""
    dt_max = ""
    if "datetime" in df.columns:
        dt = pd.to_datetime(df["datetime"], errors="coerce").dropna()
        if not dt.empty:
            dt_min = str(dt.min())
            dt_max = str(dt.max())
    return (len(df), sheet, well, dt_min, dt_max)


def parse_datetime_value_count(series: pd.Series) -> int:
    """Count values that contain both date and time information."""
    parsed = parse_datetime_series(series)
    raw = series.astype(str).str.strip()
    raw_has_date_and_time = raw.str.contains(r"\d{1,4}[-/]\d{1,2}[-/]\d{1,4}.*\d{1,2}[:.]\d{2}|\d{1,2}[:.]\d{2}.*\d{1,4}[-/]\d{1,2}[-/]\d{1,4}", regex=True, na=False)

    import datetime as _dt
    object_has_time = series.map(
        lambda x: isinstance(x, (pd.Timestamp, _dt.datetime)) and (x.hour != 0 or x.minute != 0 or x.second != 0)
        if not pd.isna(x) else False
    )
    return int((parsed.notna() & (raw_has_date_and_time | object_has_time)).sum())


def parse_date_value_count(series: pd.Series) -> int:
    return int(parse_date_series(series).notna().sum())


def parse_time_value_count(series: pd.Series) -> int:
    parsed = parse_time_series(series)
    # Treat pure dates parsed as midnight as weak time evidence unless the raw value had ':' or was an Excel time fraction.
    raw = series.astype(str).str.strip()
    raw_time_like = raw.str.contains(r"\d{1,2}[:.]\d{2}", regex=True, na=False)
    excel_fraction_time = series.map(lambda x: isinstance(x, (int, float, np.number)) and not isinstance(x, bool) and 0 <= float(x) < 1 if not pd.isna(x) else False)
    return int((parsed.notna() & (raw_time_like | excel_fraction_time)).sum())


def best_value_column(raw_df: pd.DataFrame, kind: str, exclude: Optional[set] = None) -> Optional[object]:
    """Detect date/time/datetime columns from cell values, not only from headers."""
    exclude = exclude or set()
    best_col = None
    best_score = 0
    for col in raw_df.columns:
        if col in exclude:
            continue
        s = raw_df[col].dropna()
        if s.empty:
            continue
        sample = s.head(80)
        if kind == "datetime":
            score = parse_datetime_value_count(sample)
        elif kind == "date":
            score = parse_date_value_count(sample)
        else:
            score = parse_time_value_count(sample)
        # Require at least two useful values so metadata cells are not mistaken for a series.
        if score >= 2 and score > best_score:
            best_score = score
            best_col = col
    return best_col


def standardize_loose_timeseries(
    raw_df: pd.DataFrame,
    source_name: str = "",
    sheet_name: str = "",
    default_well: Optional[str] = None,
) -> pd.DataFrame:
    """Last-resort Excel parser for sheets whose headers are blank/merged/not recognized.

    It detects date/time columns from the actual values and keeps every numeric column as
    a raw plotting series. This prevents valid field Excel sheets from being rejected just
    because their header text does not match the alias list yet.
    """
    if raw_df is None or raw_df.empty:
        return pd.DataFrame()

    df = raw_df.copy().dropna(how="all").dropna(axis=1, how="all")
    if df.empty:
        return pd.DataFrame()

    # If the first non-empty row is mostly text, use it as labels for nicer raw column names.
    first_row = df.iloc[0]
    text_cells = sum((not is_numeric_like(v)) and bool(normalize_text(v)) for v in first_row.tolist())
    non_empty = sum(bool(normalize_text(v)) for v in first_row.tolist())
    first_row_is_data = row_looks_like_data(first_row) or (
        parse_date_value_count(pd.Series(first_row.tolist())) >= 1
        and parse_time_value_count(pd.Series(first_row.tolist())) >= 1
        and sum(is_numeric_like(v) for v in first_row.tolist()) >= 1
    )
    if not first_row_is_data and non_empty >= 2 and text_cells >= max(2, non_empty // 2):
        labels = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(first_row.tolist())])
        df = df.iloc[1:].copy()
        df.columns = labels[: df.shape[1]]
    else:
        df.columns = [f"Column_{i + 1}" for i in range(df.shape[1])]

    df = df.dropna(how="all")
    if df.empty:
        return pd.DataFrame()

    datetime_col = best_value_column(df, "datetime")
    used = set([datetime_col]) if datetime_col is not None else set()
    date_col = None if datetime_col is not None else best_value_column(df, "date", used)
    if date_col is not None:
        used.add(date_col)
    time_col = best_value_column(df, "time", used)
    if time_col is not None:
        used.add(time_col)

    out = pd.DataFrame(index=df.index)
    out["source"] = source_name
    out["sheet"] = sheet_name
    out["well"] = default_well or guess_well_from_name(f"{source_name} {sheet_name}") or "Unknown"

    if datetime_col is not None:
        out["datetime"] = parse_datetime_series(df[datetime_col])
    if date_col is not None:
        out["date"] = parse_date_series(df[date_col]).ffill().bfill()
    if time_col is not None:
        parsed_time = parse_time_series(df[time_col])
        out["time_text"] = parsed_time.dt.strftime("%H:%M")
        out.loc[parsed_time.isna(), "time_text"] = ""
    else:
        parsed_time = None

    if ("datetime" not in out.columns or out["datetime"].isna().all()) and (date_col is not None or time_col is not None):
        out["datetime"] = combine_date_time(out.get("date"), df[time_col] if time_col is not None else None, out.get("datetime"))

    if "datetime" in out.columns and out["datetime"].notna().any():
        if "date" not in out.columns:
            out["date"] = out["datetime"].dt.floor("D")
        if "time_text" not in out.columns:
            out["time_text"] = out["datetime"].dt.strftime("%H:%M")

    # Keep all numeric columns except detected date/time columns. Use header labels when possible.
    for col in df.columns:
        if col in used:
            continue
        nums = df[col].map(extract_number).astype(float)
        if nums.notna().sum() < 2:
            continue
        raw_name = raw_output_column_name(col, out.columns)
        out[raw_name] = nums

    numeric_cols = [c for c in out.columns if c not in BASE_NON_PLOT_COLS]
    if not numeric_cols:
        return pd.DataFrame()

    useful = out[numeric_cols].notna().any(axis=1)
    if "datetime" in out.columns and out["datetime"].notna().sum() >= 2:
        useful &= out["datetime"].notna()
    elif "time_text" in out.columns and out["time_text"].astype(str).str.strip().ne("").sum() >= 2:
        useful &= out["time_text"].astype(str).str.strip().ne("")

    out = out.loc[useful].copy()
    if out.empty:
        return pd.DataFrame()

    if "datetime" in out.columns and "well" in out.columns and out["datetime"].notna().any():
        # Do not let all-NaT datetimes collapse a valid table to one row.
        out = out.drop_duplicates(subset=["well", "datetime"], keep="first")

    sort_cols = ["well"] + (["datetime"] if "datetime" in out.columns else [])
    return out.sort_values(sort_cols, na_position="last").reset_index(drop=True)


def parse_excel_sheet_attempts(raw: pd.DataFrame, source_name: str, sheet_name: str, default_well: Optional[str]) -> List[pd.DataFrame]:
    """Try several Excel interpretations and return valid time-series candidates."""
    attempts: List[pd.DataFrame] = []

    def add_candidate(candidate: pd.DataFrame):
        if candidate is not None and not candidate.empty and is_valid_timeseries(candidate):
            attempts.append(candidate)

    # Existing smart multi-row header parser.
    table = table_from_raw(raw)
    add_candidate(standardize_dataframe(table, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # Common case: Excel already has a simple first-row header. Try it explicitly.
    df0 = raw.copy().dropna(how="all").dropna(axis=1, how="all")
    if not df0.empty and len(df0) >= 2:
        headers = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(df0.iloc[0].tolist())])
        direct = df0.iloc[1:].copy()
        direct.columns = headers[: direct.shape[1]]
        add_candidate(standardize_dataframe(direct, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # Scan possible header rows. This helps when title/metadata rows are above the actual table.
    max_scan = min(60, len(raw))
    for header_row in range(max_scan):
        row = raw.iloc[header_row]
        if row_looks_like_data(row):
            continue
        score = header_score(row)
        if score < 4:
            continue
        candidate_raw = raw.iloc[header_row:].dropna(how="all").dropna(axis=1, how="all")
        if candidate_raw.empty or len(candidate_raw) < 3:
            continue
        headers = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(candidate_raw.iloc[0].tolist())])
        candidate = candidate_raw.iloc[1:].copy()
        candidate.columns = headers[: candidate.shape[1]]
        add_candidate(standardize_dataframe(candidate, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # Last-resort value-based parser. This catches headerless or unrecognized templates.
    add_candidate(standardize_loose_timeseries(raw, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # De-duplicate and keep the strongest parse for this sheet.
    unique: Dict[tuple, pd.DataFrame] = {}
    for cand in attempts:
        key = dataframe_key(cand)
        if key not in unique or dataframe_quality_score(cand) > dataframe_quality_score(unique[key]):
            unique[key] = cand

    return sorted(unique.values(), key=dataframe_quality_score, reverse=True)


def is_valid_timeseries(df: pd.DataFrame) -> bool:
    """Return True only for real plottable time-series tables.

    A table must have at least one numeric reading column and repeated date/time
    evidence.  This prevents non-data text/PDF files from being accepted just
    because a few numbers were found somewhere in the document.
    """
    if df is None or df.empty:
        return False
    numeric_cols = available_numeric_columns(df)
    if len(numeric_cols) == 0:
        return False
    if "datetime" in df.columns and pd.to_datetime(df["datetime"], errors="coerce").notna().sum() >= 2:
        return True
    if "time_text" in df.columns and df["time_text"].astype(str).str.strip().ne("").sum() >= 2:
        return True
    if "date" in df.columns and pd.to_datetime(df["date"], errors="coerce").notna().sum() >= 2:
        return True
    return False


def is_usable_single_message_table(df: pd.DataFrame) -> bool:
    """Allow one-row WhatsApp/TMU messages while rejecting random text/PDF pages.

    WhatsApp reports are often one reading only, so they cannot pass the stricter
    repeated-time-series test.  They still must contain at least one numeric
    operational field and at least one usable date/time marker.
    """
    if df is None or df.empty:
        return False
    if not available_numeric_columns(df):
        return False
    has_dt = "datetime" in df.columns and pd.to_datetime(df["datetime"], errors="coerce").notna().any()
    has_date = "date" in df.columns and pd.to_datetime(df["date"], errors="coerce").notna().any()
    has_time = "time_text" in df.columns and df["time_text"].astype(str).str.strip().ne("").any()
    return bool(has_dt or has_date or has_time)


def filter_usable_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """Final safety filter for all upload types."""
    return [t for t in tables if is_valid_timeseries(t) or is_usable_single_message_table(t)]



def parse_expro_mpfm_text(text: str, source_name: str = "EXPRO_MPFM_PDF") -> pd.DataFrame:
    """Parse EXPRO MPFM PDF Data & Events text rows.

    Keeps only rows where the time is followed by the full numeric MPFM reading set.
    Skips event/comment rows such as 'BS&W is ...', 'bypassed the meter', etc.
    """
    if not text or "Data & Events" not in text or "QOil" not in text or "QWat" not in text:
        return pd.DataFrame()

    m = re.search(r"Well\s+(?:No\s+)?([A-Z0-9]+\s*-\s*\d+)", text, flags=re.I)
    if m:
        well = re.sub(r"\s*-\s*", "-", m.group(1).strip().upper())
    else:
        well = guess_well_from_name(source_name) or "Unknown"

    current_date = pd.NaT
    rows = []
    # EXPRO Data & Events rows contain one choke column followed by 24 MPFM
    # measurements. Older builds inserted an extra ``choke_ambiguous`` slot,
    # shifting every value one column to the right (for example QGross became
    # BS&W). Keep the schema exactly aligned with the report header.
    expro_cols = [
        "choke_size_64", "whp_psi", "flow_press_psi", "flow_temp_c",
        "mpfm_press_psig", "mpfm_temp_f", "dp_mbar",
        "qoil_s_stbd", "qwat_s_bpd", "qgas_s_mmscfd",
        "qoil_a_bpd", "qwat_a_bpd", "qgas_a_mmcfd",
        "wlr_s_pct", "qgross_s_bpd",
        "oil_sg", "water_sg", "water_ph", "salinity_kppm",
        "gas_sg", "co2_mole_pct", "h2s_ppm", "gor_s_scf_stb", "gvf_a_pct", "pump_freq_hz",
    ]

    date_pat = re.compile(r"^\s*(\d{1,2}/\d{1,2}/\d{2,4})\s*$")
    time_pat = re.compile(r"^\s*(\d{1,2}:\d{2}:\d{2})\s+(.*)$")

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        dm = date_pat.match(line)
        if dm:
            current_date = pd.to_datetime(dm.group(1), errors="coerce", dayfirst=True)
            continue

        tm = time_pat.match(line)
        if not tm:
            continue

        time_txt, rest = tm.group(1), tm.group(2).strip()
        if not re.match(r"^[+-]?\d", rest):
            continue

        nums = re.findall(r"[-+]?\d*\.\d+|[-+]?\d+", rest.replace(",", ""))
        if len(nums) < len(expro_cols):
            continue

        vals = [float(x) for x in nums[: len(expro_cols)]]
        row = dict(zip(expro_cols, vals))
        row["source"] = source_name
        row["sheet"] = "EXPRO_MPFM_Data_Events"
        row["well"] = well
        row["date"] = current_date
        row["time"] = time_txt
        row["time_text"] = time_txt[:5]

        t = parse_time_series(pd.Series([time_txt])).iloc[0]
        if pd.notna(current_date) and pd.notna(t):
            row["datetime"] = pd.Timestamp(current_date.date()) + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
        else:
            row["datetime"] = pd.NaT

        row["oil_rate_stbd"] = row["qoil_s_stbd"]
        row["water_rate_bpd"] = row["qwat_s_bpd"]
        row["gas_rate_mmscfd"] = row["qgas_s_mmscfd"]
        row["gross_rate_bpd"] = row["qgross_s_bpd"]
        row["bsw_pct"] = row["wlr_s_pct"]
        row["gor_scf_bbl"] = row["gor_s_scf_stb"]

        if pd.notna(row["salinity_kppm"]) and row["salinity_kppm"] > 1000:
            row["salinity_kppm"] = row["salinity_kppm"] / 1000.0

        rows.append(row)

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    df = df.drop_duplicates(subset=["well", "datetime", "source"], keep="first")
    df = df.sort_values(["well", "datetime"], na_position="last").reset_index(drop=True)
    return df



def filter_preferred_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """When a workbook has both summary/form sheets and detailed time-series sheets,
    keep the detailed time-series tables by default.
    """
    if len(tables) <= 1:
        return tables

    detailed = []
    for df in tables:
        if df is None or df.empty or len(df) < 3:
            continue
        dt_unique = df["datetime"].nunique(dropna=True) if "datetime" in df.columns else 0
        time_unique = df["time_text"].nunique(dropna=True) if "time_text" in df.columns else 0
        sheet_name = str(df["sheet"].iloc[0]).lower() if "sheet" in df.columns and len(df) else ""
        if max(dt_unique, time_unique) >= 3 and sheet_name not in {"form", "summary", "test summary"}:
            detailed.append(df)

    return detailed if detailed else tables

def load_tabular_file_base(uploaded_file) -> List[pd.DataFrame]:
    name = uploaded_file.name
    suffix = name.split(".")[-1].lower()
    tables: List[pd.DataFrame] = []

    if suffix in ["xlsx", "xls"]:
        xls = pd.ExcelFile(uploaded_file)
        for sheet in xls.sheet_names:
            raw = pd.read_excel(xls, sheet_name=sheet, header=None)
            default_well = extract_well_from_raw(raw, source_name=name, sheet_name=sheet)
            sheet_candidates = parse_excel_sheet_attempts(
                raw,
                source_name=name,
                sheet_name=sheet,
                default_well=default_well,
            )
            if sheet_candidates:
                # Keep only the best interpretation for each sheet to avoid duplicate plots.
                tables.append(sheet_candidates[0])

        tables = filter_preferred_tables(tables)

    elif suffix == "csv":
        # First try normal CSV headers. This is important for CSVs exported from this dashboard.
        uploaded_file.seek(0)
        try:
            header_df = pd.read_csv(uploaded_file)
            header_df = header_df.loc[:, ~header_df.columns.astype(str).str.match(r"^Unnamed")]
            known_header_hits = sum(1 for c in header_df.columns if best_canonical_name(str(c)) is not None)
            if known_header_hits >= 3:
                default_well = extract_well_from_raw(header_df, source_name=name, sheet_name="CSV")
                std = standardize_dataframe(header_df, source_name=name, sheet_name="CSV", default_well=default_well)
                if is_valid_timeseries(std):
                    tables.append(std)
        except Exception:
            pass

        if not tables:
            uploaded_file.seek(0)
            raw = pd.read_csv(uploaded_file, header=None)
            default_well = extract_well_from_raw(raw, source_name=name, sheet_name="CSV")
            table = table_from_raw(raw)
            std = standardize_dataframe(table, source_name=name, sheet_name="CSV", default_well=default_well)
            if is_valid_timeseries(std):
                tables.append(std)

    elif suffix == "txt":
        text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        msg_rows = parse_many_tmu_messages(text, source_name=name)
        if is_usable_single_message_table(msg_rows):
            tables.append(msg_rows)

    elif suffix == "docx":
        try:
            from docx import Document

            doc = Document(uploaded_file)
            text = "\n".join(safe_text(p.text) for p in doc.paragraphs)
            msg_rows = parse_many_tmu_messages(text, source_name=name)
            if is_usable_single_message_table(msg_rows):
                tables.append(msg_rows)

            for i, t in enumerate(doc.tables):
                rows = [[cell.text for cell in row.cells] for row in t.rows]
                raw = pd.DataFrame(rows)
                default_well = extract_well_from_raw(raw, source_name=name, sheet_name=f"DOCX_Table_{i + 1}")
                table = table_from_raw(raw)
                std = standardize_dataframe(table, source_name=name, sheet_name=f"DOCX_Table_{i + 1}", default_well=default_well)
                if is_valid_timeseries(std):
                    tables.append(std)
        except Exception as e:
            raise RuntimeError(f"Could not read DOCX file {name}: {e}")

    elif suffix == "pdf":
        try:
            import pdfplumber

            with pdfplumber.open(uploaded_file) as pdf:
                full_text = []
                for pi, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    full_text.append(text)

                    for ti, table in enumerate(page.extract_tables() or []):
                        raw = pd.DataFrame(table)
                        default_well = extract_well_from_raw(raw, source_name=name, sheet_name=f"PDF_Page_{pi + 1}_Table_{ti + 1}")
                        table_df = table_from_raw(raw)
                        std = standardize_dataframe(table_df, source_name=name, sheet_name=f"PDF_Page_{pi + 1}_Table_{ti + 1}", default_well=default_well)
                        if is_valid_timeseries(std):
                            tables.append(std)

                all_text = "\n".join(safe_text(x) for x in full_text)

                # Special support for EXPRO MPFM PDF reports where the Data & Events
                # table is embedded as text rather than a clean extractable table.
                # If this detailed parser succeeds, keep it and discard summary/event-only tables.
                expro_rows = parse_expro_mpfm_text(all_text, source_name=name)
                if not expro_rows.empty:
                    tables = [expro_rows]
                else:
                    msg_rows = parse_many_tmu_messages(all_text, source_name=name)
                    if is_usable_single_message_table(msg_rows):
                        tables.append(msg_rows)
                    tables = filter_preferred_tables(tables)
        except Exception as e:
            raise RuntimeError(f"Could not read PDF file {name}: {e}")

    return filter_usable_tables(tables)


def value_by_patterns(text: str, patterns: List[str]) -> Optional[str]:
    for p in patterns:
        pattern = rf"{p}\s*[:=@-]?\s*([^\n\r]+)"
        m = re.search(pattern, text, flags=re.I)
        if m:
            value = m.group(1).strip()
            value = re.split(
                r"\s{2,}(?=(choke|w\.?h\.?p|sep|gas|gross|oil|water|bs|salinity|h2s|co2|note|pumping)\b)",
                value,
                flags=re.I,
            )[0].strip()
            return value
    return None


def parse_salinity_to_kppm(value: object) -> float:
    if value is None:
        return np.nan

    s = str(value).replace(",", "")
    num = extract_number(s)
    if pd.isna(num):
        return np.nan
    if re.search(r"\bk\b|kppm|k ppm", s, flags=re.I):
        return num
    if num > 1000:
        return num / 1000.0
    return num


def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    text = message.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)

    row: Dict[str, object] = {"source": source_name, "sheet": "WhatsApp_Text"}

    tmu = re.search(r"\b(PICO\s*T\s*MU[-\s]*\d+|PICO\s*TMU[-\s]*\d+|TMU[-\s]*\d+)\b", text, flags=re.I)
    if tmu:
        row["test_unit"] = re.sub(r"\s+", " ", tmu.group(1).upper()).replace("T MU", "TMU")

    date_v = value_by_patterns(text, [r"\bdate\b"])
    time_v = value_by_patterns(text, [r"\btime\s*@?", r"\btime\b"])
    well_v = value_by_patterns(text, [r"\bwell\s*name\b", r"\bwell\b"])

    if well_v:
        row["well"] = well_v.split()[0].strip()
    else:
        row["well"] = guess_well_from_name(text) or "Unknown"

    if date_v:
        row["date"] = pd.to_datetime(date_v, errors="coerce", dayfirst=True)
    if time_v:
        row["time"] = time_v

    fields = {
        "whp_psi": [r"\bw\.?\s*h\.?\s*p\.?\b", r"\bwellhead pressure\b"],
        "sep_p_psi": [r"\bsep\.?\s*p\.?\b", r"\bseparator pressure\b"],
        "gas_rate_mmscfd": [r"\bgas rate\b"],
        "gas_formation_mmscfd": [r"\bgas formation\b", r"\bformation gas\b"],
        "gross_rate_bpd": [r"\bgross rate\b", r"\bgross\b"],
        "oil_rate_stbd": [r"\boil rate\b"],
        "water_rate_bpd": [r"\bwater rate\b"],
        "bsw_pct": [r"\bbs\s*&\s*w\b", r"\bbsw\b", r"\bwc\b", r"\bwater cut\b"],
        "salinity_kppm": [r"\bsalinity\b"],
        "h2s_ppm": [r"\bh2s\b"],
        "co2_mole_pct": [r"\bco2\b", r"\bco₂\b"],
        "water_cum_bbl": [r"\bwater cum\b", r"\bwater cumulative\b"],
        "pumping_pressure_psi": [r"\bpumping\s*[.\-/ ]?\s*p\b", r"\bpumping pressure\b", r"\bpump pressure\b"],
        "n2_rate_scfm": [r"\bn2 standard rate\b", r"\bn2 rate\b", r"\bnitrogen rate\b"],
        "note": [r"\bnote\b"],
    }

    for canon, pats in fields.items():
        raw_val = value_by_patterns(text, pats)
        if raw_val is None:
            continue
        if canon == "note":
            row[canon] = raw_val.strip()
        elif canon == "salinity_kppm":
            row[canon] = parse_salinity_to_kppm(raw_val)
        else:
            row[canon] = extract_number(raw_val)

    # Parse choke separately so 24/64 is never plotted as 24%.
    choke_raw = value_by_patterns(text, [r"\bchoke\b"])
    if choke_raw is not None:
        choke_series = pd.Series([choke_raw])
        pct_s, size_s, ambiguous_s, _ = _split_choke_series_v56(choke_series, "Choke " + choke_raw)
        if pct_s.notna().any():
            row["choke_pct"] = float(pct_s.dropna().iloc[0])
        if size_s.notna().any():
            row["choke_size_64"] = float(size_s.dropna().iloc[0])
        if ambiguous_s.notna().any():
            row["choke_ambiguous"] = float(ambiguous_s.dropna().iloc[0])

    if "date" in row and "time" in row:
        d = pd.to_datetime(row["date"], errors="coerce", dayfirst=True)
        tser = parse_time_series(pd.Series([row["time"]]))
        if pd.notna(d) and pd.notna(tser.iloc[0]):
            t = tser.iloc[0]
            row["datetime"] = pd.Timestamp(d.date()) + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
            row["time_text"] = f"{t.hour:02d}:{t.minute:02d}"

    return row


def split_messages(text: str) -> List[str]:
    markers = list(re.finditer(r"(?=PICO\s*T?\s*MU|TMU[-\s]*\d+|Date\s*:)", text, flags=re.I))
    if len(markers) <= 1:
        return [text]

    chunks = []
    for i, m in enumerate(markers):
        start = m.start()
        end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

    return chunks


def parse_many_tmu_messages(text: str, source_name: str = "WhatsApp_Text") -> pd.DataFrame:
    rows = []
    for chunk in split_messages(text):
        row = parse_tmu_message(chunk, source_name=source_name)
        has_rate = any(
            k in row
            for k in [
                "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi",
                "pumping_pressure_psi", "gas_rate_mmscfd", "bsw_pct",
            ]
        )
        if has_rate:
            rows.append(row)

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    for col in ["source", "sheet", "well", "datetime", "date", "time_text"]:
        if col not in df.columns:
            df[col] = np.nan

    return df


def apply_fill_method(df: pd.DataFrame, numeric_cols: List[str], method: str) -> pd.DataFrame:
    if df.empty or not numeric_cols or method == "No fill":
        return df

    out = df.sort_values(["well", "datetime"], na_position="last").copy()

    if method == "Forward fill":
        out[numeric_cols] = out.groupby("well", dropna=False)[numeric_cols].ffill()

    elif method == "Forward + backward fill":
        out[numeric_cols] = out.groupby("well", dropna=False)[numeric_cols].ffill().bfill()

    elif method == "Linear interpolation by row":
        out[numeric_cols] = out.groupby("well", dropna=False)[numeric_cols].transform(
            lambda x: x.interpolate(limit_direction="both")
        )

    elif method == "Time interpolation":
        pieces = []
        for _, g in out.groupby("well", dropna=False):
            g = g.copy()
            if "datetime" in g.columns and g["datetime"].notna().sum() >= 2:
                g = g.set_index("datetime")
                g[numeric_cols] = g[numeric_cols].interpolate(method="time", limit_direction="both")
                g = g.reset_index()
            else:
                g[numeric_cols] = g[numeric_cols].interpolate(limit_direction="both")
            pieces.append(g)
        out = pd.concat(pieces, ignore_index=True)

    return out


def apply_user_column_mappings(df: pd.DataFrame, mapping: Optional[Dict[str, str]] = None) -> pd.DataFrame:
    """Apply user-confirmed column mappings after auto parsing.

    mapping keys may be actual DataFrame column names (for example raw__amp) or
    normalized alias keys from canonical_key(). Values are canonical parser names
    such as motor_current_amp, pump_intake_pressure_psi, or the special value
    __drop__ to hide a column.

    If target already exists, non-null values from the source fill gaps only; this
    prevents a manual alias from overwriting a stronger auto-detected column.
    """
    if df is None or df.empty or not mapping:
        return df

    out = df.copy()
    options = standard_column_options(include_meta=False)
    drop_cols = []

    for col in list(out.columns):
        key_candidates = [str(col), canonical_key(col)]
        target = None
        for key in key_candidates:
            if key in mapping:
                target = mapping.get(key)
                break
        if not target or target in {"__keep__", "Keep as-is"}:
            continue
        if target == "__drop__":
            drop_cols.append(col)
            continue
        if target not in options:
            continue
        if col in BASE_NON_PLOT_COLS:
            continue

        vals = clean_numeric_series(out[col], target)
        if target in out.columns:
            out[target] = pd.to_numeric(out[target], errors="coerce").combine_first(vals)
        else:
            out[target] = vals
        if col != target:
            drop_cols.append(col)

    if drop_cols:
        out = out.drop(columns=[c for c in drop_cols if c in out.columns], errors="ignore")
    return out


def available_numeric_columns(df: pd.DataFrame) -> List[str]:
    audit_flags = {
        "gas_formation_derived", "n2_rate_derived", "total_gas_derived",
        "review_required", "ocr_approved", "is_event", "is_duplicate",
    }
    columns = []
    seen = set()
    for c in df.columns:
        if c in seen or c in BASE_NON_PLOT_COLS or c in audit_flags:
            continue
        seen.add(c)
        positions = [i for i, name in enumerate(df.columns) if name == c]
        for pos in positions:
            series = df.iloc[:, pos]
            if pd.api.types.is_bool_dtype(series.dtype):
                continue
            if pd.api.types.is_numeric_dtype(series.dtype) and series.notna().any():
                columns.append(c)
                break
    return columns


def column_label(column_name: str) -> str:
    if str(column_name).startswith("raw__"):
        label = str(column_name)[5:].replace("_", " ").strip()
        return f"Raw: {label.title()}"
    return COLUMN_LABELS.get(column_name, column_name)

# -----------------------------------------------------------------------------
# v43 additions: WhatsApp export ZIP + CTU image OCR + safe test segmentation
# -----------------------------------------------------------------------------

# Keep CTU / OCR metadata out of normal numeric detection and plotting unless the
# value columns themselves are selected.
BASE_NON_PLOT_COLS.update({
    "source_type", "ocr_template", "ocr_fields_found", "ocr_status", "ocr_confidence",
    "image_file", "attachment_name", "chat_sender", "chat_datetime", "message_index",
    "test_id", "test_start", "test_end", "test_sequence", "link_status",
    "suggested_well", "suggested_test_id", "suggested_link_reason", "suggested_link_gap_hours",
    "review_required", "caption_text", "whatsapp_message_body", "source_member",
})

COLUMN_LABELS.update({
    "ctu_weight_lbf": "CTU Weight (LBF)",
    "ctu_lt_weight_lbf": "CTU Lt Weight (LBF)",
    "ctu_wellhead_pressure_psi": "CTU Wellhead Pressure (psi)",
    "ctu_circulation_pressure_psi": "CTU Circulation Pressure (psi)",
    "ctu_reel_depth_ft": "CTU Reel Depth (ft)",
    "ctu_reel_speed_ftmin": "CTU Reel Speed (ft/min)",
    "ctu_fluid_rate_bpm": "CTU Fluid Rate (bpm)",
    "ctu_n2_rate_scfm": "CTU N2 Flow (scf/min)",
    "ctu_fluid_total_bbl": "CTU Fluid Total (bbl)",
    "ctu_n2_total_scf": "CTU N2 Total (scf)",
})

for _kw in [
    "circulation", "reel", "lt weight", "fluid total", "n2 total", "flare test",
    "wellhead pressure", "all data", "ctu", "coil tubing", "coiled tubing",
]:
    if _kw not in KEYWORDS:
        KEYWORDS.append(_kw)

IMAGE_SUFFIXES = {"jpg", "jpeg", "png"}
DATA_SUFFIXES = {"xlsx", "xls", "csv", "docx", "pdf", "txt"}


class UploadedBytes(io.BytesIO):
    """BytesIO object that behaves like a Streamlit UploadedFile for recursive parsers."""

    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


def append_note(old_note, new_note):
    old = "" if old_note is None or pd.isna(old_note) else str(old_note).strip()
    new = str(new_note or "").strip()
    if not new:
        return old
    if not old:
        return new
    if new.lower() in old.lower():
        return old
    return f"{old}; {new}"


def parse_datetime_from_filename(name: str):
    """Parse WhatsApp media filenames such as PHOTO-2026-06-12-20-35-05.jpg."""
    name = str(name or "")
    patterns = [
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(\d{2})[-_](\d{2})[-_](20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
    ]
    for p in patterns:
        m = re.search(p, name)
        if not m:
            continue
        nums = list(map(int, m.groups()))
        try:
            if len(nums) == 6 and str(nums[0]).startswith("20"):
                y, mo, d, hh, mm, ss = nums
            elif len(nums) == 5:
                y, mo, d, hh, mm = nums
                ss = 0
            else:
                d, mo, y, hh, mm, ss = nums
            return pd.Timestamp(year=y, month=mo, day=d, hour=hh, minute=mm, second=ss)
        except Exception:
            pass
    return pd.NaT


def parse_attachment_reference(body: str) -> str:
    """Return attached filename mentioned in exported WhatsApp text, if any."""
    text = str(body or "")
    patterns = [
        r"<attached:\s*([^>]+)>",
        r"([^\s<>]+\.(?:jpg|jpeg|png|xlsx|xls|csv|pdf|docx))",
    ]
    for p in patterns:
        m = re.search(p, text, flags=re.I)
        if m:
            return Path(m.group(1).strip()).name
    return ""


def _try_import_ocr_libs():
    try:
        from PIL import Image, ImageOps, ImageEnhance
        import pytesseract
        return Image, ImageOps, ImageEnhance, pytesseract
    except Exception:
        return None, None, None, None


def _numbers_from_ocr_text(txt: str) -> List[str]:
    return re.findall(r"-?\d+(?:\.\d+)?", str(txt or "").replace(" ", ""))


def _best_number_from_ocr_text(txt: str):
    nums = _numbers_from_ocr_text(txt)
    if not nums:
        return np.nan
    # Prefer the longest candidate; this avoids choosing a stray unit digit.
    nums = sorted(nums, key=lambda x: len(x.replace(".", "").replace("-", "")), reverse=True)
    return extract_number(nums[0])


def _ocr_numeric_region_pil(crop_img):
    """OCR a numeric value from a CTU/PICO HMI ROI using multiple simple preprocesses."""
    Image, ImageOps, ImageEnhance, pytesseract = _try_import_ocr_libs()
    if Image is None:
        return np.nan, 0.0, "ocr_dependency_missing"

    best_val = np.nan
    best_score = -1
    best_text = ""

    try:
        gray = crop_img.convert("L")
        gray = ImageOps.autocontrast(gray)
    except Exception:
        return np.nan, 0.0, "ocr_image_error"

    variants = []
    for invert in [False, True]:
        g0 = ImageOps.invert(gray) if invert else gray
        for contrast in [2.0, 3.0, 5.0]:
            g = ImageEnhance.Contrast(g0).enhance(contrast)
            g = ImageEnhance.Sharpness(g).enhance(2.0)
            g = g.resize((max(90, g.width * 4), max(40, g.height * 4)))
            variants.append(g)
            for thr in [110, 140, 170, 200]:
                variants.append(g.point(lambda p, t=thr: 255 if p > t else 0))

    for img_try in variants:
        try:
            txt = pytesseract.image_to_string(
                img_try,
                config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.",
            )
        except Exception:
            continue

        nums = _numbers_from_ocr_text(txt)
        if not nums:
            continue
        val = _best_number_from_ocr_text(txt)
        digit_count = max(len(n.replace(".", "").replace("-", "")) for n in nums)
        score = digit_count + (2 if any("." in n for n in nums) else 0)
        if pd.notna(val) and score > best_score:
            best_score = score
            best_val = val
            best_text = str(txt).strip()

    confidence = min(1.0, max(0.0, best_score / 8.0)) if best_score >= 0 else 0.0
    return best_val, confidence, best_text


# ROIs are value boxes only, relative to the full image. They fit the CTU "ALL DATA"
# screen style in the user's sample; all rows remain review_required unless safely linked.


def _ocr_numeric_region_cv2(crop_img):
    """Optional OpenCV OCR path that improves low-contrast colored HMI digits."""
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return np.nan, 0.0, "cv2_or_tesseract_missing"

    try:
        arr = _np.array(crop_img.convert("RGB"))
        bgr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
        bgr = cv2.resize(bgr, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
        lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        gray = clahe.apply(l)
        variants = [gray, 255 - gray]
        for block in [31, 51, 71]:
            variants.append(cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block, 3))
            variants.append(cv2.adaptiveThreshold(255 - gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block, 3))

        best_val = np.nan
        best_score = -1
        best_text = ""
        for v in variants:
            try:
                txt = pytesseract.image_to_string(
                    v,
                    config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.",
                )
            except Exception:
                continue
            nums = _numbers_from_ocr_text(txt)
            if not nums:
                continue
            val = _best_number_from_ocr_text(txt)
            digit_count = max(len(n.replace(".", "").replace("-", "")) for n in nums)
            score = digit_count + (2 if any("." in n for n in nums) else 0)
            if pd.notna(val) and score > best_score:
                best_score = score
                best_val = val
                best_text = str(txt).strip()
        confidence = min(1.0, max(0.0, best_score / 8.0)) if best_score >= 0 else 0.0
        return best_val, confidence, best_text
    except Exception:
        return np.nan, 0.0, "cv2_ocr_error"


def _ocr_numeric_region(crop_img):
    """Use PIL and OpenCV OCR paths, then choose the higher-confidence numeric value."""
    pil_val, pil_conf, pil_text = _ocr_numeric_region_pil(crop_img)
    cv_val, cv_conf, cv_text = _ocr_numeric_region_cv2(crop_img)
    if pd.notna(cv_val) and (pd.isna(pil_val) or cv_conf >= pil_conf):
        return cv_val, cv_conf, cv_text
    return pil_val, pil_conf, pil_text

CTU_ALL_DATA_ROIS = {
    "ctu_weight_lbf": (0.26, 0.13, 0.55, 0.30),
    "ctu_lt_weight_lbf": (0.66, 0.13, 0.92, 0.30),
    "ctu_wellhead_pressure_psi": (0.31, 0.31, 0.55, 0.47),
    "ctu_circulation_pressure_psi": (0.64, 0.31, 0.92, 0.47),
    "ctu_reel_depth_ft": (0.31, 0.49, 0.55, 0.66),
    "ctu_reel_speed_ftmin": (0.62, 0.49, 0.92, 0.66),
    "ctu_fluid_rate_bpm": (0.31, 0.66, 0.55, 0.80),
    "ctu_n2_rate_scfm": (0.64, 0.66, 0.92, 0.80),
    "ctu_fluid_total_bbl": (0.31, 0.81, 0.55, 0.95),
    "ctu_n2_total_scf": (0.64, 0.81, 0.92, 0.95),
}


def parse_ctu_all_data_screen_image(uploaded_file, source_name="Image_OCR") -> pd.DataFrame:
    """Parse CTU/PICO ALL DATA screen photos as auxiliary OCR rows.

    Safety rule: OCR rows are never treated as confirmed well-test readings by
    themselves. If the image does not explicitly carry a well/date, the row is
    marked review_required and remains Well=Unknown until the app/user approves
    linking. This avoids silently assigning CTU data to the wrong test.
    """
    Image, _, _, _ = _try_import_ocr_libs()
    if Image is None:
        return pd.DataFrame()

    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        img = Image.open(uploaded_file).convert("RGB")
        w, h = img.size
        file_name = getattr(uploaded_file, "name", source_name)
        dt_from_name = parse_datetime_from_filename(file_name)

        row = {
            "source": source_name,
            "sheet": "CTU_Image_OCR",
            "source_type": "ctu_image_ocr",
            "ocr_template": "ctu_all_data",
            "image_file": file_name,
            "well": "Unknown",
            "link_status": "unlinked_needs_review",
            "review_required": True,
        }
        if pd.notna(dt_from_name):
            row["datetime"] = dt_from_name
            row["date"] = pd.Timestamp(dt_from_name.date())
            row["time_text"] = pd.Timestamp(dt_from_name).strftime("%H:%M")

        fields_found = 0
        confidences = []
        for field, (x1, y1, x2, y2) in CTU_ALL_DATA_ROIS.items():
            crop = img.crop((int(x1 * w), int(y1 * h), int(x2 * w), int(y2 * h)))
            val, conf, _raw = _ocr_numeric_region(crop)
            if pd.notna(val):
                row[field] = float(val)
                fields_found += 1
                confidences.append(float(conf))

        row["ocr_fields_found"] = fields_found
        row["ocr_confidence"] = float(np.mean(confidences)) if confidences else 0.0
        row["ocr_status"] = "parsed_review_required" if fields_found >= 2 else "low_confidence_or_not_ctu_screen"

        if fields_found < 1:
            return pd.DataFrame()
        return pd.DataFrame([row])
    except Exception:
        return pd.DataFrame()


WHATSAPP_SYSTEM_PATTERNS = [
    r"messages and calls are end-to-end encrypted",
    r"joined using this group",
    r"changed the subject",
    r"changed the group description",
    r"changed this group's icon",
    r"this message was deleted",
    r"missed voice call",
    r"missed video call",
    r"created group",
]


def is_system_or_noise_message(body: str) -> bool:
    b = normalize_text(body)
    if not b:
        return True
    for p in WHATSAPP_SYSTEM_PATTERNS:
        if re.search(p, b, flags=re.I):
            return True
    # Very short acknowledgement messages are not useful data.
    if b in {"ok", "okay", "thanks", "thank you", "done", "تمام", "اوكي"}:
        return True
    return False


def score_tmu_body(body: str) -> int:
    b = normalize_text(body)
    if is_system_or_noise_message(b):
        return 0
    score = 0
    for kw in [
        "tmu", "well", "date", "time", "choke", "w.h.p", "whp", "sep",
        "gas rate", "gross rate", "oil rate", "water rate", "bsw", "bs&w",
        "salinity", "h2s", "co2", "pumping", "n2", "flare test", "all data",
        "circulation pressure", "reel depth", "wellhead pressure",
    ]:
        if kw in b:
            score += 1
    score += min(5, len(re.findall(r"[-+]?\d+(?:\.\d+)?", body)))
    return score


def parse_whatsapp_export_messages(text: str) -> List[Dict[str, object]]:
    """Parse WhatsApp exported _chat.txt into message dicts.

    Handles common Android/iOS formats:
      12/06/2026, 00:30 - Sender: body
      [12/06/2026, 00:30:05] Sender: body
    Multi-line messages are preserved.
    """
    lines = str(text or "").replace("\ufeff", "").splitlines()
    messages = []
    current = None

    patterns = [
        r"^(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{2,4}),\s*(?P<time>\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\s*-\s*(?P<sender>[^:]+):\s*(?P<body>.*)$",
        r"^\[(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{2,4}),\s*(?P<time>\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\]\s*(?P<sender>[^:]+):\s*(?P<body>.*)$",
        r"^(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\s+(?P<time>\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\s*-\s*(?P<sender>[^:]+):\s*(?P<body>.*)$",
    ]

    for line in lines:
        line = line.rstrip()
        if not line.strip():
            continue

        matched = False
        for p in patterns:
            m = re.match(p, line)
            if m:
                if current:
                    messages.append(current)
                current = m.groupdict()
                current["sender"] = str(current.get("sender", "")).strip()
                current["body"] = str(current.get("body", "")).strip()
                matched = True
                break

        if not matched and current:
            current["body"] = str(current.get("body", "")) + "\n" + line.strip()

    if current:
        messages.append(current)

    for i, m in enumerate(messages):
        m["message_index"] = i
        m["datetime"] = pd.to_datetime(
            f"{m.get('date', '')} {m.get('time', '')}",
            errors="coerce",
            dayfirst=True,
        )
        m["attachment_name"] = parse_attachment_reference(m.get("body", ""))

    return messages


def parse_whatsapp_export_text(text: str, source_name="WhatsApp_Export") -> pd.DataFrame:
    """Parse exported WhatsApp chat text into TMU production rows.

    Important: this parser uses well name and explicit TMU keywords. It does not
    assume one chat equals one test.
    """
    messages = parse_whatsapp_export_messages(text)
    if not messages:
        return pd.DataFrame()

    rows = []
    for m in messages:
        body = str(m.get("body", ""))
        if score_tmu_body(body) < 4:
            continue

        row = parse_tmu_message(body, source_name=source_name)
        # If a message is only a continuation line, parse_tmu_message may not get
        # its own date/time. Use WhatsApp timestamp as fallback only.
        msg_dt = m.get("datetime", pd.NaT)
        if ("datetime" not in row or pd.isna(row.get("datetime", pd.NaT))) and pd.notna(msg_dt):
            row["datetime"] = msg_dt
            row["date"] = pd.Timestamp(msg_dt).floor("D")
            row["time_text"] = pd.Timestamp(msg_dt).strftime("%H:%M")

        row["source_type"] = "whatsapp_export_text"
        row["chat_sender"] = m.get("sender", "")
        row["chat_datetime"] = msg_dt
        row["message_index"] = m.get("message_index", np.nan)
        row["attachment_name"] = m.get("attachment_name", "")
        row["whatsapp_message_body"] = body[:500]
        row["link_status"] = "text_confirmed_by_well" if row.get("well") and str(row.get("well")).lower() != "unknown" else "text_needs_well_review"

        has_useful = any(
            k in row for k in [
                "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi",
                "sep_p_psi", "pumping_pressure_psi", "gas_rate_mmscfd", "bsw_pct",
            ]
        )
        if has_useful:
            rows.append(row)

    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def parse_whatsapp_plain_or_export_text(text: str, source_name="WhatsApp_Text") -> pd.DataFrame:
    """Try exported-chat parsing first; fallback to pasted/block TMU parser."""
    export_df = parse_whatsapp_export_text(text, source_name=source_name)
    if not export_df.empty:
        return export_df
    df = parse_many_tmu_messages(text, source_name=source_name)
    if not df.empty:
        df["source_type"] = "pasted_whatsapp_text"
        df["link_status"] = "text_confirmed_by_well"
    return df


def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    """Assign stable test_id by well name + time gap.

    Rule requested by user:
      - Same well continues the same test until the time gap exceeds gap_hours.
      - A different well is always a different test stream.
      - Unknown/unlinked OCR rows are NOT silently assigned to a well/test.
    """
    if df is None or df.empty:
        return df

    out = df.copy()
    if "well" not in out.columns:
        out["well"] = "Unknown"
    if "datetime" not in out.columns:
        out["datetime"] = pd.NaT

    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["well"] = out["well"].fillna("Unknown").astype(str).str.strip().replace("", "Unknown")
    out["test_id"] = out.get("test_id", pd.Series([np.nan] * len(out), index=out.index))
    out["test_sequence"] = out.get("test_sequence", pd.Series([np.nan] * len(out), index=out.index))

    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}

    for well, g in sortable.groupby("well", dropna=False):
        well_txt = str(well or "Unknown").strip() or "Unknown"
        if well_txt.lower() == "unknown":
            # Keep unknown rows separate/unlinked. They may be reviewed in the app.
            for i in g.index:
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well", np.nan)
            continue

        seq = 0
        last_dt = pd.NaT
        current_id = None
        current_start = pd.NaT
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                # Same well but no time: keep as separate manual-review row.
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, seq)
                continue

            new_test = (
                current_id is None
                or pd.isna(last_dt)
                or (pd.Timestamp(dt) - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            )
            if new_test:
                seq += 1
                current_start = pd.Timestamp(dt)
                current_id = f"{well_txt}_{current_start.strftime('%Y%m%d_%H%M')}"

            assigned[i] = (current_id, seq)
            last_dt = pd.Timestamp(dt)

    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = tid
        out.at[i, "test_sequence"] = seq

    # Add test start/end for confirmed test rows.
    if "test_id" in out.columns and "datetime" in out.columns:
        valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
        if not valid.empty:
            starts = valid.groupby("test_id")["datetime"].min()
            ends = valid.groupby("test_id")["datetime"].max()
            out["test_start"] = out["test_id"].map(starts)
            out["test_end"] = out["test_id"].map(ends)

    return out


def suggest_links_for_ocr_rows(df: pd.DataFrame, max_gap_hours: float = 3.0) -> pd.DataFrame:
    """Suggest, but do not apply, well/test links for CTU image rows.

    A suggestion is created only when there is one nearest confirmed non-OCR row
    within max_gap_hours. The row remains review_required; app/user must approve
    before it becomes part of the test plot.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    for c in ["source_type", "well", "datetime", "test_id", "suggested_well", "suggested_test_id", "suggested_link_reason", "suggested_link_gap_hours"]:
        if c not in out.columns:
            out[c] = np.nan
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")

    anchor_mask = (
        out["datetime"].notna()
        & out["well"].notna()
        & (out["well"].astype(str).str.strip().str.lower() != "unknown")
        & (~out["source_type"].astype(str).str.contains("ocr", case=False, na=False))
    )
    anchors = out[anchor_mask].copy()
    if anchors.empty:
        return out

    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    for i, row in out[ocr_mask].iterrows():
        dt = row.get("datetime", pd.NaT)
        if pd.isna(dt):
            continue
        deltas = (anchors["datetime"] - pd.Timestamp(dt)).abs()
        if deltas.empty:
            continue
        nearest_idx = deltas.idxmin()
        gap_h = float(deltas.loc[nearest_idx].total_seconds() / 3600.0)
        if gap_h <= float(max_gap_hours):
            out.at[i, "suggested_well"] = anchors.at[nearest_idx, "well"]
            out.at[i, "suggested_test_id"] = anchors.at[nearest_idx, "test_id"]
            out.at[i, "suggested_link_gap_hours"] = round(gap_h, 3)
            out.at[i, "suggested_link_reason"] = f"Nearest confirmed text/Excel row within {gap_h:.2f} hr; review before approving."
            out.at[i, "link_status"] = "suggested_needs_user_approval"
            out.at[i, "review_required"] = True
    return out


def approve_suggested_ocr_links(df: pd.DataFrame) -> pd.DataFrame:
    """Apply suggested OCR links after explicit user approval in Streamlit."""
    if df is None or df.empty:
        return df
    out = df.copy()
    required = {"suggested_well", "suggested_test_id", "source_type"}
    if not required.issubset(set(out.columns)):
        return out
    mask = (
        out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
        & out["suggested_well"].notna()
        & out["suggested_test_id"].notna()
    )
    out.loc[mask, "well"] = out.loc[mask, "suggested_well"]
    out.loc[mask, "test_id"] = out.loc[mask, "suggested_test_id"]
    out.loc[mask, "link_status"] = "ocr_link_approved"
    out.loc[mask, "review_required"] = False
    return out


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 100) -> List[pd.DataFrame]:
    """v43 upload router.

    Supports existing Excel/CSV/DOCX/PDF/TXT parsing plus:
      - WhatsApp exported ZIP bundles (_chat.txt + attachments + images)
      - Direct CTU screen image OCR (JPG/JPEG/PNG)

    OCR safety: CTU image rows are parsed as auxiliary rows and left unlinked until
    reviewed/approved in the app. No nearest-well fill is applied automatically.
    """
    name = getattr(uploaded_file, "name", "uploaded")
    suffix = Path(str(name)).suffix.lower().lstrip(".")

    if suffix in IMAGE_SUFFIXES:
        if not parse_images:
            return []
        return filter_usable_tables([parse_ctu_all_data_screen_image(uploaded_file, source_name=name)])

    if suffix == "zip":
        try:
            raw = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
            tables: List[pd.DataFrame] = []
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                members = [m for m in zf.namelist() if not m.endswith("/") and not Path(m).name.startswith("._")]

                # Build chat-message index first so OCR/image rows can get timestamps and suggestions.
                all_messages: List[Dict[str, object]] = []
                chat_names = []
                for member in members:
                    member_name = Path(member).name
                    if Path(member_name).suffix.lower().lstrip(".") == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        text = zf.read(member).decode("utf-8", errors="ignore")
                        chat_names.append(member_name)
                        all_messages.extend(parse_whatsapp_export_messages(text))
                        df = parse_whatsapp_export_text(text, source_name=f"{name}:{member_name}")
                        if not df.empty:
                            df["source_member"] = member
                            tables.append(df)

                attachment_context = {}
                for m in all_messages:
                    att = str(m.get("attachment_name", "") or "").strip()
                    if not att:
                        continue
                    attachment_context[Path(att).name] = m

                # Parse non-chat attachments and OCR images.
                for member in members:
                    member_name = Path(member).name
                    ext = Path(member_name).suffix.lower().lstrip(".")
                    if not member_name or ext not in (DATA_SUFFIXES | IMAGE_SUFFIXES):
                        continue
                    if ext == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        continue

                    sub_file = UploadedBytes(zf.read(member), member_name)
                    if ext in IMAGE_SUFFIXES:
                        if not parse_images or max_ocr_images <= 0:
                            continue
                        max_ocr_images -= 1
                    sub_tables = load_tabular_file(sub_file, parse_images=parse_images, max_ocr_images=max_ocr_images)
                    ctx = attachment_context.get(member_name, {})
                    for t in sub_tables or []:
                        if t is None or t.empty:
                            continue
                        t = t.copy()
                        t["attachment_name"] = member_name
                        t["source_member"] = member
                        if ctx:
                            t["chat_sender"] = ctx.get("sender", "")
                            t["chat_datetime"] = ctx.get("datetime", pd.NaT)
                            t["message_index"] = ctx.get("message_index", np.nan)
                            if "datetime" not in t.columns or pd.to_datetime(t["datetime"], errors="coerce").isna().all():
                                if pd.notna(ctx.get("datetime", pd.NaT)):
                                    t["datetime"] = ctx.get("datetime")
                                    t["date"] = pd.Timestamp(ctx.get("datetime")).floor("D")
                                    t["time_text"] = pd.Timestamp(ctx.get("datetime")).strftime("%H:%M")
                            if ext in IMAGE_SUFFIXES:
                                t["caption_text"] = str(ctx.get("body", ""))[:500]
                        tables.append(t)

            if not tables:
                return []
            merged = pd.concat(tables, ignore_index=True, sort=False)
            merged = assign_test_ids(merged, gap_hours=12.0)
            # No nearest-time OCR linking here. CTU/OCR image rows remain unlinked
            # unless the user manually selects the correct Well/Test in Streamlit.
            return filter_usable_tables([merged])
        except Exception as e:
            raise RuntimeError(f"Could not read WhatsApp ZIP {name}: {e}")

    if suffix == "txt":
        try:
            text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        df = parse_whatsapp_plain_or_export_text(text, source_name=name)
        if not df.empty:
            df = assign_test_ids(df, gap_hours=12.0)
            return filter_usable_tables([df])
        return []

    tables = load_tabular_file_base(uploaded_file)
    out_tables = []
    for t in tables or []:
        if t is not None and not t.empty:
            t = assign_test_ids(t, gap_hours=12.0)
            out_tables.append(t)
    return filter_usable_tables(out_tables)

# v43.1 safety refinements: allow direct OCR rows even if no timestamp; keep them unlinked.
def parse_datetime_from_filename(name: str):
    """Parse common WhatsApp media filenames, including 'WhatsApp Image 2026-06-13 at 15.29.01.jpeg'."""
    name = str(name or "")
    patterns = [
        r"(20\d{2})[-_](\d{2})[-_](\d{2})\s+at\s+(\d{1,2})[.:\-_](\d{2})[.:\-_](\d{2})",
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(\d{2})[-_](\d{2})[-_](20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
    ]
    for p in patterns:
        m = re.search(p, name, flags=re.I)
        if not m:
            continue
        nums = list(map(int, m.groups()))
        try:
            if len(nums) == 6 and str(nums[0]).startswith("20"):
                y, mo, d, hh, mm, ss = nums
            elif len(nums) == 5:
                y, mo, d, hh, mm = nums
                ss = 0
            else:
                d, mo, y, hh, mm, ss = nums
            return pd.Timestamp(year=y, month=mo, day=d, hour=hh, minute=mm, second=ss)
        except Exception:
            pass
    return pd.NaT


def is_usable_ocr_table(df: pd.DataFrame) -> bool:
    if df is None or df.empty:
        return False
    if "source_type" not in df.columns:
        return False
    if not df["source_type"].astype(str).str.contains("ocr", case=False, na=False).any():
        return False
    return any(c in df.columns and pd.to_numeric(df[c], errors="coerce").notna().any() for c in COLUMN_LABELS if str(c).startswith("ctu_"))


def filter_usable_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """Final safety filter for all upload types, including direct CTU OCR rows."""
    return [t for t in tables if is_valid_timeseries(t) or is_usable_single_message_table(t) or is_usable_ocr_table(t)]

# v43.2 optimized OCR overrides: fewer OCR passes for Streamlit performance.
def _ocr_numeric_region_pil(crop_img):
    Image, ImageOps, ImageEnhance, pytesseract = _try_import_ocr_libs()
    if Image is None:
        return np.nan, 0.0, "ocr_dependency_missing"
    try:
        gray = ImageOps.autocontrast(crop_img.convert("L"))
        g = ImageEnhance.Contrast(gray).enhance(3.0)
        g = ImageEnhance.Sharpness(g).enhance(2.0)
        g = g.resize((max(90, g.width * 4), max(40, g.height * 4)))
        variants = [g, ImageOps.invert(g), g.point(lambda p: 255 if p > 150 else 0)]
        best = (-1, np.nan, "")
        for im in variants:
            try:
                txt = pytesseract.image_to_string(im, config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.")
            except Exception:
                continue
            nums = _numbers_from_ocr_text(txt)
            if not nums:
                continue
            val = _best_number_from_ocr_text(txt)
            score = max(len(n.replace(".", "").replace("-", "")) for n in nums) + (2 if any("." in n for n in nums) else 0)
            if pd.notna(val) and score > best[0]:
                best = (score, val, str(txt).strip())
        conf = min(1.0, max(0.0, best[0] / 8.0)) if best[0] >= 0 else 0.0
        return best[1], conf, best[2]
    except Exception:
        return np.nan, 0.0, "ocr_image_error"


def _ocr_numeric_region_cv2(crop_img):
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return np.nan, 0.0, "cv2_or_tesseract_missing"
    try:
        arr = _np.array(crop_img.convert("RGB"))
        bgr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
        bgr = cv2.resize(bgr, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
        lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
        l, _, _ = cv2.split(lab)
        gray = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8)).apply(l)
        variants = [gray, 255 - gray, cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 51, 3)]
        best = (-1, np.nan, "")
        for im in variants:
            try:
                txt = pytesseract.image_to_string(im, config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.")
            except Exception:
                continue
            nums = _numbers_from_ocr_text(txt)
            if not nums:
                continue
            val = _best_number_from_ocr_text(txt)
            score = max(len(n.replace(".", "").replace("-", "")) for n in nums) + (2 if any("." in n for n in nums) else 0)
            if pd.notna(val) and score > best[0]:
                best = (score, val, str(txt).strip())
        conf = min(1.0, max(0.0, best[0] / 8.0)) if best[0] >= 0 else 0.0
        return best[1], conf, best[2]
    except Exception:
        return np.nan, 0.0, "cv2_ocr_error"


def _ocr_numeric_region(crop_img):
    pil_val, pil_conf, pil_text = _ocr_numeric_region_pil(crop_img)
    cv_val, cv_conf, cv_text = _ocr_numeric_region_cv2(crop_img)
    if pd.notna(cv_val) and (pd.isna(pil_val) or cv_conf >= pil_conf):
        return cv_val, cv_conf, cv_text
    return pil_val, pil_conf, pil_text

# v43.3 dtype-safe test ID assignment override.
def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    out = df.copy()
    if "well" not in out.columns:
        out["well"] = "Unknown"
    if "datetime" not in out.columns:
        out["datetime"] = pd.NaT
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["well"] = out["well"].fillna("Unknown").astype(str).str.strip().replace("", "Unknown")
    out["test_id"] = out["test_id"].astype("object") if "test_id" in out.columns else pd.Series([None] * len(out), index=out.index, dtype="object")
    out["test_sequence"] = pd.to_numeric(out["test_sequence"], errors="coerce") if "test_sequence" in out.columns else pd.Series([np.nan] * len(out), index=out.index, dtype="float")
    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}
    for well, g in sortable.groupby("well", dropna=False):
        well_txt = str(well or "Unknown").strip() or "Unknown"
        if well_txt.lower() == "unknown":
            for i in g.index:
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well", np.nan)
            continue
        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            new_test = current_id is None or pd.isna(last_dt) or (pd.Timestamp(dt) - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{pd.Timestamp(dt).strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = pd.Timestamp(dt)
    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = tid
        out.at[i, "test_sequence"] = seq
    if "test_id" in out.columns and "datetime" in out.columns:
        valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
        if not valid.empty:
            starts = valid.groupby("test_id")["datetime"].min()
            ends = valid.groupby("test_id")["datetime"].max()
            out["test_start"] = out["test_id"].map(starts)
            out["test_end"] = out["test_id"].map(ends)
    return out

# v43.4 preserve old TMU parser but add operation-note detection such as Flare test.
_parse_tmu_message_base_v43 = parse_tmu_message

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v43(message, source_name=source_name)
    text = str(message or "")
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row


# -----------------------------------------------------------------------------
# v44 safety patch: no nearest-time CTU linking + robust WhatsApp bold well names
# -----------------------------------------------------------------------------
def clean_well_name_value(value: object) -> str:
    """Clean WhatsApp/Excel well names without inventing a well.

    Examples:
      '*S8-58*' -> 'S8-58'
      '*'       -> 'Unknown'
    """
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-"}:
        return "Unknown"
    # Prefer the first token/candidate containing at least one digit, which is how
    # almost all field well names appear: S8-58, B3C18-7, A-C83-1, BED15-33.
    token_match = re.search(r"[A-Za-z0-9]+(?:[-/][A-Za-z0-9]+)*", s)
    if token_match:
        candidate = token_match.group(0).strip("-_/ .")
        if candidate and re.search(r"[A-Za-z0-9]", candidate):
            return candidate
    return "Unknown"


def suggest_links_for_ocr_rows(df: pd.DataFrame, max_gap_hours: float = 0.0) -> pd.DataFrame:
    """Disabled in v44 by design.

    CTU/OCR rows are not linked or suggested by nearest time because that can
    attach image data to the wrong test. The app asks the user to manually choose
    the Well/Test ID after reviewing the image OCR numbers.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    if "source_type" in out.columns:
        ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
        for c in ["link_status", "review_required"]:
            if c not in out.columns:
                out[c] = pd.Series([None] * len(out), index=out.index, dtype="object")
        out.loc[ocr_mask, "link_status"] = out.loc[ocr_mask, "link_status"].fillna("ocr_manual_link_required")
        out.loc[ocr_mask, "review_required"] = True
    return out


def approve_suggested_ocr_links(df: pd.DataFrame) -> pd.DataFrame:
    """No-op kept for backward compatibility with older app.py builds."""
    return df


_parse_tmu_message_base_v44 = parse_tmu_message

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v44(message, source_name=source_name)
    row["well"] = clean_well_name_value(row.get("well", "Unknown"))
    text = str(message or "")
    if row.get("well", "Unknown") == "Unknown":
        guessed = guess_well_from_name(re.sub(r"[*_`~]+", "", text))
        row["well"] = clean_well_name_value(guessed)
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row


def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    """Assign test IDs by well name and time gap only.

    - Different well names are always separate streams.
    - Same well continues until the inactive gap exceeds gap_hours.
    - OCR rows with Unknown well stay Unlinked until manual review.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    for c in ["source", "sheet", "source_type", "well", "datetime", "test_id", "test_sequence", "link_status", "review_required"]:
        if c not in out.columns:
            default = None if c in {"well", "test_id", "link_status", "source_type", "source", "sheet"} else np.nan
            out[c] = pd.Series([default] * len(out), index=out.index, dtype="object")
    out["well"] = out["well"].apply(clean_well_name_value).astype("object")
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["test_id"] = out["test_id"].astype("object")
    out["link_status"] = out["link_status"].astype("object")
    out["review_required"] = out["review_required"].astype("object")

    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}
    for well, g in sortable.groupby("well", dropna=False):
        well_txt = clean_well_name_value(well)
        if well_txt == "Unknown":
            for i, row in g.iterrows():
                is_ocr = str(row.get("source_type", "")).lower().find("ocr") >= 0
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well" if is_ocr else "Unknown_Well_Unlinked", np.nan)
            continue
        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            dt = pd.Timestamp(dt)
            new_test = current_id is None or pd.isna(last_dt) or (dt - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{dt.strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = dt

    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = str(tid)
        out.at[i, "test_sequence"] = seq

    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    unknown_ocr = ocr_mask & ((out["well"].astype(str).str.lower() == "unknown") | out["test_id"].astype(str).str.startswith("Unlinked"))
    out.loc[unknown_ocr, "link_status"] = out.loc[unknown_ocr, "link_status"].fillna("ocr_manual_link_required")
    out.loc[unknown_ocr, "review_required"] = True

    valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
    if not valid.empty:
        starts = valid.groupby("test_id")["datetime"].min()
        ends = valid.groupby("test_id")["datetime"].max()
        out["test_start"] = out["test_id"].map(starts)
        out["test_end"] = out["test_id"].map(ends)
    return out


# v44.1 extra well-name cleanup for WhatsApp bold/plain first-line well names.
_guess_well_from_name_base_v441 = guess_well_from_name

def guess_well_from_name(name: str) -> Optional[str]:
    base = _guess_well_from_name_base_v441(name)
    if base:
        return clean_well_name_value(base) if 'clean_well_name_value' in globals() else base
    s = re.sub(r"[*_`~]+", "", str(name or ""))
    extra_patterns = [
        r"\b(OB\s*[-_ ]\s*\d+[A-Z]?)\b",          # Obaiyed OB-69
        r"\b(S\d+\s*[-_ ]\s*\d+[A-Z]?)\b",        # S8-58
        r"\b([A-Z]{1,4}\d+[A-Z]*\s*[-_ ]\s*\d+[A-Z]?)\b",
    ]
    for pat in extra_patterns:
        m = re.search(pat, s, flags=re.I)
        if m:
            return clean_well_name_value(m.group(1)) if 'clean_well_name_value' in globals() else normalize_well_name(m.group(1))
    return None

# Override clean_well_name_value to reject non-well words such as 'to' and '*'.
def clean_well_name_value(value: object) -> str:
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-", "to"}:
        return "Unknown"
    # Candidate must contain at least one digit to avoid false wells from normal words.
    m = re.search(r"[A-Za-z0-9]*\d[A-Za-z0-9]*(?:[-/][A-Za-z0-9]+)*", s)
    if m:
        candidate = m.group(0).strip("-_/ .")
        if candidate and re.search(r"\d", candidate):
            return candidate.upper() if re.match(r"^[A-Za-z]{1,6}[-/ ]?\d", candidate) else candidate
    return "Unknown"

# Re-override parser and test-id assignment so the final clean_well_name_value is used.
_parse_tmu_message_base_v441 = _parse_tmu_message_base_v44

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v441(message, source_name=source_name)
    text = str(message or "")
    cleaned_guess = clean_well_name_value(row.get("well", "Unknown"))
    if cleaned_guess == "Unknown":
        guessed = guess_well_from_name(text)
        cleaned_guess = clean_well_name_value(guessed)
    row["well"] = cleaned_guess
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row

# v44.2 final well cleaning: keep prefixes like OB-69, S8-58, B3C18-7, A-C83-1.
def clean_well_name_value(value: object) -> str:
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-", "to"}:
        return "Unknown"
    # Find all candidate well-looking tokens that contain digits and optional letter prefixes.
    pat = r"\b(?:[A-Za-z]{1,12}[-/ ]*)?[A-Za-z0-9]*\d[A-Za-z0-9]*(?:[-/][A-Za-z0-9]+)*\b"
    candidates = []
    for m in re.finditer(pat, s):
        cand = re.sub(r"\s+", "", m.group(0).strip("-_/ ."))
        if not cand or not re.search(r"\d", cand):
            continue
        if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", cand):
            continue
        candidates.append(cand)
    if candidates:
        # Prefer the last candidate because strings like 'Obaiyed OB-69' contain
        # a field name first and the actual well token at the end.
        return candidates[-1].upper()
    return "Unknown"

_parse_tmu_message_base_v442 = _parse_tmu_message_base_v441

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v442(message, source_name=source_name)
    text = str(message or "")
    cleaned_guess = clean_well_name_value(row.get("well", "Unknown"))
    if cleaned_guess == "Unknown":
        guessed = guess_well_from_name(text)
        cleaned_guess = clean_well_name_value(guessed)
    row["well"] = cleaned_guess
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row

# Re-override assign_test_ids to use v44.2 clean_well_name_value.
def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    out = df.copy()
    for c in ["source", "sheet", "source_type", "well", "datetime", "test_id", "test_sequence", "link_status", "review_required"]:
        if c not in out.columns:
            default = None if c in {"well", "test_id", "link_status", "source_type", "source", "sheet"} else np.nan
            out[c] = pd.Series([default] * len(out), index=out.index, dtype="object")
    out["well"] = out["well"].apply(clean_well_name_value).astype("object")
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["test_id"] = out["test_id"].astype("object")
    out["link_status"] = out["link_status"].astype("object")
    out["review_required"] = out["review_required"].astype("object")
    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}
    for well, g in sortable.groupby("well", dropna=False):
        well_txt = clean_well_name_value(well)
        if well_txt == "Unknown":
            for i, row in g.iterrows():
                is_ocr = str(row.get("source_type", "")).lower().find("ocr") >= 0
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well" if is_ocr else "Unknown_Well_Unlinked", np.nan)
            continue
        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            dt = pd.Timestamp(dt)
            new_test = current_id is None or pd.isna(last_dt) or (dt - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{dt.strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = dt
    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = str(tid)
        out.at[i, "test_sequence"] = seq
    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    unknown_ocr = ocr_mask & ((out["well"].astype(str).str.lower() == "unknown") | out["test_id"].astype(str).str.startswith("Unlinked"))
    out.loc[unknown_ocr, "link_status"] = out.loc[unknown_ocr, "link_status"].fillna("ocr_manual_link_required")
    out.loc[unknown_ocr, "review_required"] = True
    valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
    if not valid.empty:
        starts = valid.groupby("test_id")["datetime"].min()
        ends = valid.groupby("test_id")["datetime"].max()
        out["test_start"] = out["test_id"].map(starts)
        out["test_end"] = out["test_id"].map(ends)
    return out


# -----------------------------------------------------------------------------
# v45 FINAL SAFETY OVERRIDES
# -----------------------------------------------------------------------------
# Purpose:
# 1) Completely disable nearest-time CTU/OCR suggestions and any automatic OCR link.
# 2) Prevent pandas dtype errors when text well names such as '*S8-58*' are assigned.
# 3) Interpret max_ocr_images=0 as no limit; image OCR is skipped only when parse_images=False.

PARSER_BUILD = "v45_whatsapp_zip_ctu_ocr_safe_no_nearest_link"


def _v45_series_object(default=None, n=0, index=None):
    return pd.Series([default] * int(n), index=index, dtype="object")


def safe_object_columns(df: pd.DataFrame, columns=None) -> pd.DataFrame:
    """Force mixed text/status columns to object dtype before assignment."""
    if df is None or df.empty:
        return df
    out = df.copy()
    if columns is None:
        columns = [
            "well", "test_id", "source", "sheet", "source_type", "link_status",
            "review_required", "suggested_well", "suggested_test_id",
            "suggested_link_reason", "image_file", "attachment_name", "source_member",
            "chat_sender", "caption_text", "ocr_template", "ocr_status",
        ]
    for c in columns:
        if c not in out.columns:
            out[c] = _v45_series_object(None, len(out), out.index)
        else:
            out[c] = out[c].astype("object")
    return out


def suggest_links_for_ocr_rows(df: pd.DataFrame, max_gap_hours: float = 0.0) -> pd.DataFrame:
    """v45 disabled: no nearest-time suggestion, no automatic CTU/OCR well/test fill."""
    if df is None or df.empty:
        return df
    out = safe_object_columns(df)
    ocr_mask = out.get("source_type", pd.Series([""] * len(out), index=out.index)).astype(str).str.contains("ocr", case=False, na=False)
    out.loc[ocr_mask, "link_status"] = out.loc[ocr_mask, "link_status"].where(
        out.loc[ocr_mask, "link_status"].notna(), "ocr_manual_link_required"
    )
    out.loc[ocr_mask, "review_required"] = True
    # Keep old columns empty if older app versions look for them.
    out.loc[ocr_mask, "suggested_well"] = None
    out.loc[ocr_mask, "suggested_test_id"] = None
    out.loc[ocr_mask, "suggested_link_reason"] = "disabled_no_nearest_time_linking"
    return out


def approve_suggested_ocr_links(df: pd.DataFrame) -> pd.DataFrame:
    """v45 no-op; user must manually select well/test in Streamlit review table."""
    return df


def clean_well_name_value(value: object) -> str:
    """Final robust well-name cleanup for WhatsApp bold/hidden characters."""
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "").replace("\ufeff", "")
    s = re.sub(r"[*_`~]+", "", s)  # WhatsApp markdown bold/italic
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-", "to", "*"}:
        return "Unknown"
    # Avoid dates becoming wells.
    if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", s):
        return "Unknown"
    # Pure numeric values are usually BS&W/salinity/rates from partial WhatsApp messages, not well names.
    if re.fullmatch(r"\d+(?:\.\d+)?", s):
        return "Unknown"
    # Examples: S8-58, OB-69, B3C18-7, A-C83-1, BED15-33.
    pat = r"\b(?:[A-Za-z]{1,12}[-/ ]*)?[A-Za-z0-9]*\d[A-Za-z0-9]*(?:[-/][A-Za-z0-9]+)*\b"
    candidates = []
    for m in re.finditer(pat, s):
        cand = re.sub(r"\s+", "", m.group(0).strip("-_/ ."))
        if not cand or not re.search(r"\d", cand):
            continue
        if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", cand):
            continue
        candidates.append(cand)
    if candidates:
        return candidates[-1].upper()
    return "Unknown"


def guess_well_from_name(name: str) -> Optional[str]:
    """Final well guess used after WhatsApp markdown cleanup."""
    s = re.sub(r"[*_`~]+", "", str(name or ""))
    # Prefer explicit Obaiyed/OB style when present.
    patterns = [
        r"\b(OB\s*[-_ ]\s*\d+[A-Z]?)\b",
        r"\b(S\d+\s*[-_ ]\s*\d+[A-Z]?)\b",
        r"\b([A-Z]{1,4}\d+[A-Z]*\s*[-_ ]\s*\d+[A-Z]?)\b",
        r"\b([A-Z]{1,4}\s*[-_ ]\s*[A-Z]?\d+[A-Z]*\s*[-_ ]\s*\d+[A-Z]?)\b",
    ]
    for pat in patterns:
        m = re.search(pat, s, flags=re.I)
        if m:
            well = clean_well_name_value(m.group(1))
            if well != "Unknown":
                return well
    return None


_parse_tmu_message_before_v45 = parse_tmu_message

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_before_v45(message, source_name=source_name)
    text = str(message or "")
    cleaned = clean_well_name_value(row.get("well", "Unknown"))
    if cleaned == "Unknown":
        cleaned = clean_well_name_value(guess_well_from_name(text))
    row["well"] = cleaned
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row


def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    """Assign test IDs by well name and same-well time gap only.

    CTU/OCR rows with Unknown well remain unlinked. Nothing is filled by nearest time.
    """
    if df is None or df.empty:
        return df
    out = safe_object_columns(df)
    out["well"] = out["well"].apply(clean_well_name_value).astype("object")
    out["datetime"] = pd.to_datetime(out.get("datetime"), errors="coerce")
    if "test_sequence" not in out.columns:
        out["test_sequence"] = np.nan
    out["test_sequence"] = pd.to_numeric(out["test_sequence"], errors="coerce")

    sortable_cols = [c for c in ["well", "datetime", "source"] if c in out.columns]
    sortable = out.sort_values(sortable_cols, na_position="last").copy() if sortable_cols else out.copy()
    assigned = {}

    for well, g in sortable.groupby("well", dropna=False):
        well_txt = clean_well_name_value(well)
        if well_txt == "Unknown":
            for i, row in g.iterrows():
                is_ocr = "ocr" in str(row.get("source_type", "")).lower()
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well" if is_ocr else "Unknown_Well_Unlinked", np.nan)
            continue

        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            dt = pd.Timestamp(dt)
            new_test = current_id is None or pd.isna(last_dt) or (dt - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{dt.strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = dt

    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = str(tid)
        out.at[i, "test_sequence"] = seq

    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    unknown_ocr = ocr_mask & ((out["well"].astype(str).str.lower() == "unknown") | out["test_id"].astype(str).str.startswith("Unlinked"))
    out.loc[unknown_ocr, "link_status"] = "ocr_manual_link_required"
    out.loc[unknown_ocr, "review_required"] = True

    valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
    if not valid.empty:
        starts = valid.groupby("test_id")["datetime"].min()
        ends = valid.groupby("test_id")["datetime"].max()
        out["test_start"] = out["test_id"].map(starts)
        out["test_end"] = out["test_id"].map(ends)
    return out


# Keep a final explicit router so older code paths cannot call the old nearest-time suggester.
def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 300) -> List[pd.DataFrame]:
    name = getattr(uploaded_file, "name", "uploaded")
    suffix = Path(str(name)).suffix.lower().lstrip(".")

    if suffix in IMAGE_SUFFIXES:
        if not parse_images:
            return []
        df = parse_ctu_all_data_screen_image(uploaded_file, source_name=name)
        if df is not None and not df.empty:
            df = safe_object_columns(df)
            df = assign_test_ids(df, gap_hours=12.0)
        return filter_usable_tables([df])

    if suffix == "zip":
        try:
            raw = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
            tables: List[pd.DataFrame] = []
            ocr_count = 0
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                members = [m for m in zf.namelist() if not m.endswith("/") and not Path(m).name.startswith("._")]

                all_messages: List[Dict[str, object]] = []
                for member in members:
                    member_name = Path(member).name
                    ext = Path(member_name).suffix.lower().lstrip(".")
                    if ext == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        text = zf.read(member).decode("utf-8", errors="ignore")
                        msgs = parse_whatsapp_export_messages(text)
                        all_messages.extend(msgs)
                        df = parse_whatsapp_export_text(text, source_name=f"{name}:{member_name}")
                        if df is not None and not df.empty:
                            df = safe_object_columns(df)
                            df["source_member"] = member
                            tables.append(df)

                attachment_context = {}
                for m in all_messages:
                    att = str(m.get("attachment_name", "") or "").strip()
                    if att:
                        attachment_context[Path(att).name] = m

                for member in members:
                    member_name = Path(member).name
                    ext = Path(member_name).suffix.lower().lstrip(".")
                    if not member_name or ext not in (DATA_SUFFIXES | IMAGE_SUFFIXES):
                        continue
                    if ext == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        continue

                    if ext in IMAGE_SUFFIXES:
                        if not parse_images:
                            continue
                        # max_ocr_images=0 means no limit in v45.
                        if int(max_ocr_images or 0) > 0 and ocr_count >= int(max_ocr_images):
                            continue
                        ocr_count += 1

                    sub_file = UploadedBytes(zf.read(member), member_name)
                    sub_tables = load_tabular_file(sub_file, parse_images=parse_images, max_ocr_images=max_ocr_images)
                    ctx = attachment_context.get(member_name, {})
                    for t in sub_tables or []:
                        if t is None or t.empty:
                            continue
                        t = safe_object_columns(t)
                        t["attachment_name"] = member_name
                        t["source_member"] = member
                        if ctx:
                            t["chat_sender"] = ctx.get("sender", "")
                            t["chat_datetime"] = ctx.get("datetime", pd.NaT)
                            t["message_index"] = ctx.get("message_index", np.nan)
                            # Use exact WhatsApp attachment timestamp only when the attachment parser did not provide time.
                            if "datetime" not in t.columns or pd.to_datetime(t["datetime"], errors="coerce").isna().all():
                                if pd.notna(ctx.get("datetime", pd.NaT)):
                                    t["datetime"] = ctx.get("datetime")
                                    t["date"] = pd.Timestamp(ctx.get("datetime")).floor("D")
                                    t["time_text"] = pd.Timestamp(ctx.get("datetime")).strftime("%H:%M")
                            if ext in IMAGE_SUFFIXES:
                                t["caption_text"] = str(ctx.get("body", ""))[:500]
                        tables.append(t)

            if not tables:
                return []
            merged = pd.concat(tables, ignore_index=True, sort=False)
            merged = safe_object_columns(merged)
            merged = assign_test_ids(merged, gap_hours=12.0)
            # Explicitly keep OCR rows unlinked. No nearest-time suggestions.
            merged = suggest_links_for_ocr_rows(merged, max_gap_hours=0.0)
            return filter_usable_tables([merged])
        except Exception as e:
            raise RuntimeError(f"Could not read WhatsApp ZIP {name}: {e}")

    if suffix == "txt":
        try:
            text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        df = parse_whatsapp_plain_or_export_text(text, source_name=name)
        if df is not None and not df.empty:
            df = safe_object_columns(df)
            df = assign_test_ids(df, gap_hours=12.0)
            return filter_usable_tables([df])
        return []

    tables = load_tabular_file_base(uploaded_file)
    out_tables = []
    for t in tables or []:
        if t is not None and not t.empty:
            t = safe_object_columns(t)
            t = assign_test_ids(t, gap_hours=12.0)
            out_tables.append(t)
    return filter_usable_tables(out_tables)

# -----------------------------------------------------------------------------
# v46 final overrides: stricter WhatsApp TMU row acceptance + safer CTU OCR.
# -----------------------------------------------------------------------------
PRODUCTION_NUMERIC_FIELDS_V46 = [
    "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "gas_rate_mmscfd",
    "gas_formation_mmscfd", "whp_psi", "sep_p_psi", "flp_psi", "bsw_pct",
    "salinity_kppm", "h2s_ppm", "co2_mole_pct", "pumping_pressure_psi",
    "n2_rate_scfm", "choke_pct", "choke_size_64", "choke_ambiguous",
]
CORE_RATE_OR_PRESSURE_FIELDS_V46 = [
    "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "gas_rate_mmscfd",
    "whp_psi", "sep_p_psi", "pumping_pressure_psi", "bsw_pct",
]

CTU_PLAUSIBLE_RANGES_V46 = {
    "ctu_weight_lbf": (-200000.0, 200000.0),
    "ctu_lt_weight_lbf": (-200000.0, 200000.0),
    "ctu_wellhead_pressure_psi": (-50.0, 2500.0),
    "ctu_circulation_pressure_psi": (-50.0, 5000.0),
    "ctu_reel_depth_ft": (-1000.0, 40000.0),
    "ctu_reel_speed_ftmin": (-1000.0, 1000.0),
    "ctu_fluid_rate_bpm": (-10.0, 50.0),
    "ctu_n2_rate_scfm": (-100.0, 20000.0),
    "ctu_fluid_total_bbl": (-100.0, 100000.0),
    "ctu_n2_total_scf": (-100.0, 100000000.0),
}

CTU_REQUIRED_FIELDS_V46 = {
    "ctu_weight_lbf", "ctu_wellhead_pressure_psi", "ctu_circulation_pressure_psi",
    "ctu_reel_depth_ft", "ctu_reel_speed_ftmin", "ctu_fluid_rate_bpm", "ctu_n2_rate_scfm",
}


def _row_has_valid_numeric_v46(row: Dict[str, object], fields: List[str]) -> bool:
    for k in fields:
        if k in row:
            try:
                if pd.notna(pd.to_numeric(pd.Series([row.get(k)]), errors="coerce").iloc[0]):
                    return True
            except Exception:
                pass
    return False


def _clean_whatsapp_message_body_v46(body: object) -> str:
    text = str(body or "")
    text = text.replace("\u200e", "").replace("\u200f", "").replace("\ufeff", "")
    text = text.replace("\xa0", " ")
    # WhatsApp formatting only, not data.
    text = re.sub(r"[*_`~]+", "", text)
    # Remove attachment-only tails that can pollute value extraction.
    text = re.sub(r"<attached:[^>]+>", "", text, flags=re.I)
    text = re.sub(r"\bimage omitted\b|\bvideo omitted\b|\baudio omitted\b", "", text, flags=re.I)
    return text.strip()


def _is_probably_full_tmu_reading_v46(body: str) -> bool:
    b = normalize_text(body)
    if is_system_or_noise_message(b):
        return False
    # Accept normal PICO/TMU production reports, not comments like "return oil & water".
    has_identity = bool(re.search(r"\b(pico\s*t?mu|tmu[-\s]*\d+|well\s*name|well\s*:)", b, flags=re.I))
    numeric_line_hits = 0
    for pat in [
        r"\bchoke\b\s*[:=@]", r"\bw\.?\s*h\.?\s*p\.?\b\s*[:=@]", r"\bsep\.?\s*p\.?\b\s*[:=@]",
        r"\bgross\s*rate\b\s*[:=@]", r"\boil\s*rate\b\s*[:=@]", r"\bwater\s*rate\b\s*[:=@]",
        r"\bgas\s*rate\b\s*[:=@]", r"\bbs\s*&\s*w\b\s*[:=@]", r"\bpumping\s*p\b\s*[:=@]",
    ]:
        if re.search(pat, body, flags=re.I):
            numeric_line_hits += 1
    return has_identity and numeric_line_hits >= 2


def parse_whatsapp_export_text(text: str, source_name="WhatsApp_Export") -> pd.DataFrame:
    """v46: parse only real TMU numeric readings from WhatsApp export.

    This intentionally drops chat comments/rubbish and operation notes that do not
    contain actual numeric test readings. Notes can be added manually in the app.
    """
    messages = parse_whatsapp_export_messages(text)
    if not messages:
        return pd.DataFrame()

    rows = []
    for m in messages:
        body_raw = str(m.get("body", ""))
        body = _clean_whatsapp_message_body_v46(body_raw)
        if not _is_probably_full_tmu_reading_v46(body):
            continue
        row = parse_tmu_message(body, source_name=source_name)
        # Must contain at least one real numeric production/pressure value.
        if not _row_has_valid_numeric_v46(row, CORE_RATE_OR_PRESSURE_FIELDS_V46):
            continue
        # Avoid rows that have only choke or metadata but no plotted test reading.
        valid_numeric_count = sum(
            1 for k in PRODUCTION_NUMERIC_FIELDS_V46
            if k in row and pd.notna(pd.to_numeric(pd.Series([row.get(k)]), errors="coerce").iloc[0])
        )
        if valid_numeric_count < 2:
            continue

        msg_dt = m.get("datetime", pd.NaT)
        if ("datetime" not in row or pd.isna(row.get("datetime", pd.NaT))) and pd.notna(msg_dt):
            row["datetime"] = msg_dt
            row["date"] = pd.Timestamp(msg_dt).floor("D")
            row["time_text"] = pd.Timestamp(msg_dt).strftime("%H:%M")

        row["source_type"] = "whatsapp_export_text"
        row["chat_sender"] = m.get("sender", "")
        row["chat_datetime"] = msg_dt
        row["message_index"] = m.get("message_index", np.nan)
        row["attachment_name"] = m.get("attachment_name", "")
        row["whatsapp_message_body"] = body[:500]
        row["link_status"] = "text_confirmed_by_well" if clean_well_name_value(row.get("well")) != "Unknown" else "text_needs_well_review"
        rows.append(row)

    if not rows:
        return pd.DataFrame()
    df = pd.DataFrame(rows)
    # Remove exact duplicates created by repeated quoted/attached messages.
    subset = [c for c in ["well", "datetime", "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi", "sep_p_psi"] if c in df.columns]
    if subset:
        df = df.drop_duplicates(subset=subset, keep="first")
    return df.reset_index(drop=True)


def parse_whatsapp_plain_or_export_text(text: str, source_name="WhatsApp_Text") -> pd.DataFrame:
    export_df = parse_whatsapp_export_text(text, source_name=source_name)
    if export_df is not None and not export_df.empty:
        return export_df
    df = parse_many_tmu_messages(_clean_whatsapp_message_body_v46(text), source_name=source_name)
    if df is not None and not df.empty:
        df["source_type"] = "pasted_whatsapp_text"
        df["link_status"] = "text_confirmed_by_well"
    return df


def _normalize_ctu_ocr_value_v46(field: str, value: object) -> float:
    try:
        v = float(value)
    except Exception:
        return np.nan
    if not np.isfinite(v):
        return np.nan

    lo, hi = CTU_PLAUSIBLE_RANGES_V46.get(field, (-np.inf, np.inf))
    # Tesseract often loses decimal points on HMI screens: 12096 -> 120.96.
    # Try decimal restoration before rejecting.
    candidates = [v]
    for div in [10.0, 100.0, 1000.0]:
        candidates.append(v / div)
    # Preserve negative sign for speed if needed.
    for c in candidates:
        if lo <= c <= hi:
            # For pressure fields, prefer decimal-restored value when raw is impossible.
            if lo <= v <= hi:
                return float(v)
            return float(c)
    return np.nan


_parse_ctu_all_data_screen_image_base_v46 = parse_ctu_all_data_screen_image

def parse_ctu_all_data_screen_image(uploaded_file, source_name="Image_OCR") -> pd.DataFrame:
    """v46: stricter CTU/PICO screen OCR.

    Only returns a row when enough plausible CTU fields are detected. Random chat
    photos are ignored instead of becoming rubbish OCR rows.
    """
    df = _parse_ctu_all_data_screen_image_base_v46(uploaded_file, source_name=source_name)
    if df is None or df.empty:
        return pd.DataFrame()
    out = df.copy()
    found = 0
    confidences = []
    for field in list(CTU_PLAUSIBLE_RANGES_V46.keys()):
        if field in out.columns:
            v = _normalize_ctu_ocr_value_v46(field, out.at[out.index[0], field])
            if pd.notna(v):
                out.at[out.index[0], field] = v
                found += 1
            else:
                out = out.drop(columns=[field], errors="ignore")
    try:
        conf = float(out.get("ocr_confidence", pd.Series([0.0])).iloc[0])
        confidences.append(conf)
    except Exception:
        pass
    key_found = sum(1 for f in CTU_REQUIRED_FIELDS_V46 if f in out.columns and pd.notna(pd.to_numeric(out[f], errors="coerce").iloc[0]))
    # Stricter gate: enough plausible fields, or at least multiple key fields.
    if found < 4 or key_found < 3:
        return pd.DataFrame()
    out["ocr_fields_found"] = found
    out["ocr_status"] = "parsed_review_required"
    out["link_status"] = "ocr_manual_link_required"
    out["review_required"] = True
    return out


# v46 final usable filters: a table with only empty rows/comments is not usable.
def is_usable_single_message_table(df: pd.DataFrame) -> bool:
    if df is None or df.empty:
        return False
    # OCR handled separately.
    if "source_type" in df.columns and df["source_type"].astype(str).str.contains("ocr", case=False, na=False).any():
        return is_usable_ocr_table(df)
    has_dt = "datetime" in df.columns and pd.to_datetime(df["datetime"], errors="coerce").notna().any()
    if not has_dt:
        return False
    return any(c in df.columns and pd.to_numeric(df[c], errors="coerce").notna().any() for c in CORE_RATE_OR_PRESSURE_FIELDS_V46)


def filter_usable_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    out = []
    for t in tables or []:
        if t is None or t.empty:
            continue
        # Drop rows with no useful numeric values, except valid OCR rows.
        tt = t.copy()
        if "source_type" in tt.columns:
            ocr_mask = tt["source_type"].astype(str).str.contains("ocr", case=False, na=False)
        else:
            ocr_mask = pd.Series([False] * len(tt), index=tt.index)
        non_ocr = ~ocr_mask
        useful = pd.Series([False] * len(tt), index=tt.index)
        for c in CORE_RATE_OR_PRESSURE_FIELDS_V46:
            if c in tt.columns:
                useful |= pd.to_numeric(tt[c], errors="coerce").notna()
        ocr_useful = pd.Series([False] * len(tt), index=tt.index)
        for c in CTU_PLAUSIBLE_RANGES_V46.keys():
            if c in tt.columns:
                ocr_useful |= pd.to_numeric(tt[c], errors="coerce").notna()
        keep = (non_ocr & useful) | (ocr_mask & ocr_useful)
        tt = tt.loc[keep].copy()
        if not tt.empty and (is_valid_timeseries(tt) or is_usable_single_message_table(tt) or is_usable_ocr_table(tt)):
            out.append(tt)
    return out

# v46.1 OCR decimal preference override.
def _normalize_ctu_ocr_value_v46(field: str, value: object) -> float:
    try:
        v = float(value)
    except Exception:
        return np.nan
    if not np.isfinite(v):
        return np.nan
    lo, hi = CTU_PLAUSIBLE_RANGES_V46.get(field, (-np.inf, np.inf))
    abs_v = abs(v)
    # When decimals disappear, pressures like 120.96 often become 12096.
    if field in {"ctu_wellhead_pressure_psi", "ctu_circulation_pressure_psi"}:
        if abs_v >= 10000 and lo <= v / 100.0 <= hi:
            return float(v / 100.0)
        if abs_v >= 2500:
            for div in [100.0, 10.0, 1000.0]:
                c = v / div
                if lo <= c <= hi:
                    return float(c)
    if field in {"ctu_reel_speed_ftmin", "ctu_fluid_rate_bpm"}:
        if abs_v > hi:
            for div in [100.0, 10.0, 1000.0]:
                c = v / div
                if lo <= c <= hi:
                    return float(c)
    if field in {"ctu_weight_lbf", "ctu_lt_weight_lbf", "ctu_reel_depth_ft"}:
        if not (lo <= v <= hi):
            for div in [100.0, 10.0, 1000.0]:
                c = v / div
                if lo <= c <= hi:
                    return float(c)
    return float(v) if lo <= v <= hi else np.nan

# v46.2 final router: exact same-message CTU linking from WhatsApp caption only.
def _apply_exact_caption_context_to_ocr_v46(t: pd.DataFrame, ctx: Dict[str, object], ext: str) -> pd.DataFrame:
    if t is None or t.empty:
        return t
    out = safe_object_columns(t)
    if ext not in IMAGE_SUFFIXES:
        return out
    caption = _clean_whatsapp_message_body_v46(ctx.get("body", "")) if ctx else ""
    if caption:
        out["caption_text"] = caption[:1000]
    is_ocr = out.get("source_type", pd.Series([""] * len(out), index=out.index)).astype(str).str.contains("ocr", case=False, na=False)
    if not is_ocr.any():
        return out
    if not caption:
        return out
    try:
        cap_row = parse_tmu_message(caption, source_name="WhatsApp_caption_context")
    except Exception:
        cap_row = {}
    cap_well = clean_well_name_value(cap_row.get("well", "Unknown"))
    cap_dt = pd.to_datetime(cap_row.get("datetime", pd.NaT), errors="coerce")
    if cap_well != "Unknown":
        out.loc[is_ocr, "well"] = cap_well
        out.loc[is_ocr, "link_status"] = "ocr_linked_by_same_whatsapp_message_caption"
        out.loc[is_ocr, "review_required"] = True  # still allow review, but not Unknown
    if pd.notna(cap_dt):
        out.loc[is_ocr, "datetime"] = cap_dt
        out.loc[is_ocr, "date"] = pd.Timestamp(cap_dt).floor("D")
        out.loc[is_ocr, "time_text"] = pd.Timestamp(cap_dt).strftime("%H:%M")
    elif ctx and pd.notna(ctx.get("datetime", pd.NaT)):
        msg_dt = ctx.get("datetime")
        out.loc[is_ocr, "datetime"] = msg_dt
        out.loc[is_ocr, "date"] = pd.Timestamp(msg_dt).floor("D")
        out.loc[is_ocr, "time_text"] = pd.Timestamp(msg_dt).strftime("%H:%M")
    if cap_row.get("test_unit"):
        out.loc[is_ocr, "test_unit"] = cap_row.get("test_unit")
    return out


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    name = getattr(uploaded_file, "name", "uploaded")
    suffix = Path(str(name)).suffix.lower().lstrip(".")

    if suffix in IMAGE_SUFFIXES:
        if not parse_images:
            return []
        df = parse_ctu_all_data_screen_image(uploaded_file, source_name=name)
        if df is not None and not df.empty:
            df = safe_object_columns(df)
            df = assign_test_ids(df, gap_hours=12.0)
        return filter_usable_tables([df])

    if suffix == "zip":
        try:
            raw = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
            tables: List[pd.DataFrame] = []
            ocr_count = 0
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                members = [m for m in zf.namelist() if not m.endswith("/") and not Path(m).name.startswith("._")]
                all_messages: List[Dict[str, object]] = []
                for member in members:
                    member_name = Path(member).name
                    ext = Path(member_name).suffix.lower().lstrip(".")
                    if ext == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        text = zf.read(member).decode("utf-8", errors="ignore")
                        msgs = parse_whatsapp_export_messages(text)
                        all_messages.extend(msgs)
                        df = parse_whatsapp_export_text(text, source_name=f"{name}:{member_name}")
                        if df is not None and not df.empty:
                            df = safe_object_columns(df)
                            df["source_member"] = member
                            tables.append(df)

                attachment_context: Dict[str, Dict[str, object]] = {}
                for m in all_messages:
                    att = str(m.get("attachment_name", "") or "").strip()
                    if att:
                        attachment_context[Path(att).name] = m

                for member in members:
                    member_name = Path(member).name
                    ext = Path(member_name).suffix.lower().lstrip(".")
                    if not member_name or ext not in (DATA_SUFFIXES | IMAGE_SUFFIXES):
                        continue
                    if ext == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        continue
                    if ext in IMAGE_SUFFIXES:
                        if not parse_images:
                            continue
                        if int(max_ocr_images or 0) > 0 and ocr_count >= int(max_ocr_images):
                            continue
                        ocr_count += 1

                    sub_file = UploadedBytes(zf.read(member), member_name)
                    sub_tables = load_tabular_file(sub_file, parse_images=parse_images, max_ocr_images=max_ocr_images)
                    ctx = attachment_context.get(member_name, {})
                    for t in sub_tables or []:
                        if t is None or t.empty:
                            continue
                        t = safe_object_columns(t)
                        t["attachment_name"] = member_name
                        t["source_member"] = member
                        if ctx:
                            t["chat_sender"] = ctx.get("sender", "")
                            t["chat_datetime"] = ctx.get("datetime", pd.NaT)
                            t["message_index"] = ctx.get("message_index", np.nan)
                            if ext in IMAGE_SUFFIXES:
                                t = _apply_exact_caption_context_to_ocr_v46(t, ctx, ext)
                            else:
                                if "datetime" not in t.columns or pd.to_datetime(t["datetime"], errors="coerce").isna().all():
                                    if pd.notna(ctx.get("datetime", pd.NaT)):
                                        t["datetime"] = ctx.get("datetime")
                                        t["date"] = pd.Timestamp(ctx.get("datetime")).floor("D")
                                        t["time_text"] = pd.Timestamp(ctx.get("datetime")).strftime("%H:%M")
                        tables.append(t)
            if not tables:
                return []
            merged = pd.concat(tables, ignore_index=True, sort=False)
            merged = safe_object_columns(merged)
            merged = assign_test_ids(merged, gap_hours=12.0)
            merged = suggest_links_for_ocr_rows(merged, max_gap_hours=0.0)
            return filter_usable_tables([merged])
        except Exception as e:
            raise RuntimeError(f"Could not read WhatsApp ZIP {name}: {e}")

    if suffix == "txt":
        try:
            text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        df = parse_whatsapp_plain_or_export_text(text, source_name=name)
        if df is not None and not df.empty:
            df = safe_object_columns(df)
            df = assign_test_ids(df, gap_hours=12.0)
            return filter_usable_tables([df])
        return []

    tables = load_tabular_file_base(uploaded_file)
    out_tables = []
    for t in tables or []:
        if t is not None and not t.empty:
            t = safe_object_columns(t)
            t = assign_test_ids(t, gap_hours=12.0)
            out_tables.append(t)
    return filter_usable_tables(out_tables)

# -----------------------------------------------------------------------------
# v48 final safety override: never miss Pumping Pressure columns from Excel.
# -----------------------------------------------------------------------------
def _looks_like_pumping_pressure_column_v48(col: object) -> bool:
    c = clean_header(col)
    c2 = canonical_key(col)
    txt = f" {c} {c2} "
    if any(bad in txt for bad in ["frequency", "freq", "hz", "speed", "rate", "n2", "nitrogen", "depth", "temp", "temperature", "total"]):
        # Do not confuse pump frequency/rate/temp with pressure.
        if not re.search(r"\b(pump|pumping)[_\s.\-/]*p\b|pumping[.\s_\-/]*pressure|pump[.\s_\-/]*pressure", txt):
            return False
    return bool(re.search(
        r"\b(pump|pumping)[_\s.\-/]*p\b|pumping[.\s_\-/]*pressure|pump[.\s_\-/]*pressure|circulation[.\s_\-/]*pressure",
        txt,
        flags=re.I,
    ))


def ensure_pumping_pressure_column_v48(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    out = df.copy()
    current = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series([np.nan] * len(out), index=out.index)), errors="coerce")
    if current.notna().sum() > 0:
        out["pumping_pressure_psi"] = current
        return out
    candidates = []
    for col in list(out.columns):
        if str(col) == "pumping_pressure_psi":
            continue
        if _looks_like_pumping_pressure_column_v48(col):
            vals = pd.to_numeric(
                out[col].astype(str)
                .str.replace(",", "", regex=False)
                .str.replace("%", "", regex=False)
                .str.extract(r"([-+]?\d+(?:\.\d+)?)", expand=False),
                errors="coerce",
            )
            if vals.notna().sum() >= max(1, min(3, len(out) // 20)):
                candidates.append((int(vals.notna().sum()), str(col), vals))
    if candidates:
        candidates.sort(reverse=True, key=lambda x: x[0])
        out["pumping_pressure_psi"] = candidates[0][2]
    return out


_load_tabular_file_v47_final = load_tabular_file

def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    tables = _load_tabular_file_v47_final(uploaded_file, parse_images=parse_images, max_ocr_images=max_ocr_images)
    fixed = []
    for t in tables or []:
        if t is not None and not t.empty:
            fixed.append(ensure_pumping_pressure_column_v48(t))
    return fixed


# -----------------------------------------------------------------------------
# v49: robust hidden-column Pumping Pressure extraction from Excel.
# Some TMU Excel files keep the real Pumping Pressure column far to the right
# (for example hidden column EG with header "pumping.p" / unit "psi").  The
# standard table parser may reject or rename hidden helper columns, so this
# final pass scans the raw workbook and merges that pressure series by datetime.
# -----------------------------------------------------------------------------
_load_tabular_file_v48_streamlit_final = load_tabular_file


def _uploaded_bytes_v49(uploaded_file):
    name = getattr(uploaded_file, "name", "uploaded")
    try:
        data = uploaded_file.getvalue()
    except Exception:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        data = uploaded_file.read()
    return data, name


def _parse_time_series_v49(series: pd.Series) -> pd.Series:
    """Improved Excel time parser: supports values like 1.5 as 12:00 next-day fraction."""
    import datetime as _dt
    def one(x):
        if pd.isna(x):
            return pd.NaT
        if isinstance(x, pd.Timestamp):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, _dt.datetime):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, _dt.time):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, (int, float, np.number)) and not isinstance(x, bool):
            xf = float(x)
            if xf >= 0:
                frac = xf % 1.0
                seconds = int(round(frac * 24 * 3600))
                return pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)
        s = str(x)
        m = re.search(r"(\d{1,2})[:.](\d{2})(?:[:.](\d{2}))?", s)
        if m:
            hh = int(m.group(1))
            mm = int(m.group(2))
            ss = int(m.group(3) or 0)
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=hh, minutes=mm, seconds=ss)
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            return pd.to_datetime(s, errors="coerce")
    return series.map(one)


def _combine_date_time_v49(date_series, time_series):
    dates = parse_date_series(date_series)
    times = _parse_time_series_v49(time_series)
    out = []
    for d, t, raw_t in zip(dates, times, time_series):
        if pd.isna(d):
            out.append(pd.NaT)
            continue
        extra_days = 0
        try:
            if isinstance(raw_t, (int, float, np.number)) and not isinstance(raw_t, bool):
                extra_days = int(np.floor(float(raw_t)))
        except Exception:
            extra_days = 0
        if pd.notna(t):
            out.append(pd.Timestamp(d.date()) + pd.Timedelta(days=extra_days, hours=t.hour, minutes=t.minute, seconds=t.second))
        else:
            out.append(pd.Timestamp(d.date()))
    return pd.Series(out, index=date_series.index, dtype="datetime64[ns]")


def _looks_like_pumping_pressure_header_v49(value: object, unit_value: object = "") -> bool:
    c = clean_header(value)
    ck = canonical_key(value)
    u = clean_header(unit_value)
    txt = f" {c} {ck} {u} "
    if not re.search(r"pump|pumping|circulation", txt, flags=re.I):
        return False
    if re.search(r"freq|frequency|hz|speed|rate|n2|nitrogen|temp|temperature|depth|total", txt, flags=re.I):
        if not re.search(r"pump[\s._/-]*p\b|pumping[\s._/-]*p\b|pump[\s._/-]*pressure|pumping[\s._/-]*pressure|circulation[\s._/-]*pressure", txt, flags=re.I):
            return False
    return bool(re.search(r"pump[\s._/-]*p\b|pumping[\s._/-]*p\b|pump[\s._/-]*pressure|pumping[\s._/-]*pressure|pumping\.p|circulation[\s._/-]*pressure", txt, flags=re.I))


def extract_hidden_pumping_pressure_from_excel_v49(data: bytes, name: str) -> pd.DataFrame:
    suffix = str(name).split(".")[-1].lower()
    if suffix not in {"xlsx", "xls"}:
        return pd.DataFrame()
    rows = []
    try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None, dtype=object)
    except Exception:
        return pd.DataFrame()

    for sheet_name, raw in sheets.items():
        if raw is None or raw.empty:
            continue
        default_well = extract_well_from_raw(raw, source_name=name, sheet_name=str(sheet_name)) or "Unknown"
        scan_rows = min(25, len(raw))
        candidate_cols = []
        for r in range(scan_rows):
            for c in range(raw.shape[1]):
                val = raw.iat[r, c] if c < raw.shape[1] else None
                unit = raw.iat[r + 1, c] if r + 1 < len(raw) else ""
                if _looks_like_pumping_pressure_header_v49(val, unit):
                    candidate_cols.append((r, c))
        for header_r, c in candidate_cols:
            value_series = pd.to_numeric(
                pd.Series(raw.iloc[header_r + 1 :, c]).astype(str)
                .str.replace(",", "", regex=False)
                .str.extract(r"([-+]?\d+(?:\.\d+)?)", expand=False),
                errors="coerce",
            )
            if value_series.notna().sum() < 3:
                continue
            idx = value_series.index
            # Prefer first two columns as Date/Time, which is how these TMU sheets are structured.
            dt = _combine_date_time_v49(raw.iloc[idx, 0], raw.iloc[idx, 1])
            valid = value_series.notna() & dt.notna()
            if valid.sum() < 3:
                continue
            for row_idx in value_series[valid].index:
                rows.append({
                    "source": name,
                    "sheet": str(sheet_name),
                    "well": clean_well_name_value(default_well),
                    "datetime": pd.Timestamp(dt.loc[row_idx]),
                    "date": pd.Timestamp(dt.loc[row_idx]).floor("D"),
                    "time_text": pd.Timestamp(dt.loc[row_idx]).strftime("%H:%M"),
                    "pumping_pressure_psi": float(value_series.loc[row_idx]),
                    "source_type": "excel_hidden_pumping_pressure",
                })
    if not rows:
        return pd.DataFrame()
    out = pd.DataFrame(rows).drop_duplicates(subset=["well", "datetime"], keep="last")
    return out


def merge_pumping_pressure_tables_v49(tables, pump_df: pd.DataFrame):
    if pump_df is None or pump_df.empty:
        return tables
    if not tables:
        return [pump_df]
    merged_tables = []
    pump_df = pump_df.copy()
    pump_df["datetime"] = pd.to_datetime(pump_df["datetime"], errors="coerce")
    for df in tables:
        if df is None or df.empty:
            merged_tables.append(df)
            continue
        out = df.copy()
        if "datetime" not in out.columns:
            merged_tables.append(out)
            continue
        out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
        if "well" not in out.columns:
            out["well"] = "Unknown"
        tmp = pump_df[["well", "datetime", "pumping_pressure_psi"]].copy()
        tmp["well"] = tmp["well"].astype(str)
        out["well"] = out["well"].astype(str)
        out = out.merge(tmp, on=["well", "datetime"], how="left", suffixes=("", "__hidden"))
        if "pumping_pressure_psi__hidden" in out.columns:
            base = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce")
            hidden = pd.to_numeric(out["pumping_pressure_psi__hidden"], errors="coerce")
            out["pumping_pressure_psi"] = base.combine_first(hidden)
            out = out.drop(columns=["pumping_pressure_psi__hidden"], errors="ignore")
        merged_tables.append(out)
    return merged_tables


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).split(".")[-1].lower()
    base_tables = _load_tabular_file_v48_streamlit_final(UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images)
    fixed_tables = [ensure_pumping_pressure_column_v48(t) for t in (base_tables or []) if t is not None and not t.empty]
    if suffix in {"xlsx", "xls"}:
        pump_df = extract_hidden_pumping_pressure_from_excel_v49(data, name)
        fixed_tables = merge_pumping_pressure_tables_v49(fixed_tables, pump_df)
        fixed_tables = [ensure_pumping_pressure_column_v48(t) for t in fixed_tables if t is not None and not t.empty]
    return fixed_tables


# -----------------------------------------------------------------------------
# v51: stronger pumping-pressure rescue and cache-busting build id
# -----------------------------------------------------------------------------
PARSER_BUILD_ID = "v52_xml_pump_notes_layout_20260616"

_load_tabular_file_v50_final = load_tabular_file


def _clean_merge_well_v51(x: object) -> str:
    try:
        y = clean_well_name_value(x)
    except Exception:
        y = x
    s = str(y or "").strip().strip("*")
    if not s or s.lower() in {"nan", "none", "unknown", "to", "*"}:
        return "Unknown"
    return s.upper()


def _to_number_series_v51(s: pd.Series) -> pd.Series:
    return pd.to_numeric(
        s.astype(str)
        .str.replace(",", "", regex=False)
        .str.replace("%", "", regex=False)
        .str.extract(r"([-+]?\d+(?:\.\d+)?)", expand=False),
        errors="coerce",
    )


def _looks_like_pumping_pressure_header_v51(value: object, unit_value: object = "") -> bool:
    # Handles hidden Excel headers such as: pumping.p / psi, Pumping P, Pump P,
    # Pump Pressure, Pumping Pressure, Circulation Pressure.
    h = f" {clean_header(value)} {canonical_key(value)} "
    u = f" {clean_header(unit_value)} {canonical_key(unit_value)} "
    txt = h + u
    if not re.search(r"pump|pumping|circulation", txt, flags=re.I):
        return False
    # Avoid pump frequency / speed / rate / temp unless the compact header clearly says P/pressure.
    if re.search(r"freq|frequency|hz|speed|rate|n2|nitrogen|temp|temperature|depth|total", txt, flags=re.I):
        if not re.search(r"pump[\s._/-]*p\b|pumping[\s._/-]*p\b|pump[\s._/-]*pressure|pumping[\s._/-]*pressure|circulation[\s._/-]*pressure|pumping\.p|pump\.p", txt, flags=re.I):
            return False
    return bool(re.search(r"pump[\s._/-]*p\b|pumping[\s._/-]*p\b|pump[\s._/-]*pressure|pumping[\s._/-]*pressure|pumping\.p|pump\.p|circulation[\s._/-]*pressure", txt, flags=re.I))


def _date_time_from_raw_v51(raw: pd.DataFrame, idx) -> pd.Series:
    # First try columns A/B, because the uploaded TMU sheets put Date and Time there.
    try:
        dt = _combine_date_time_v49(raw.iloc[idx, 0], raw.iloc[idx, 1])
        if pd.to_datetime(dt, errors="coerce").notna().sum() >= 3:
            return dt
    except Exception:
        pass

    # Then scan for date/time headers in the first 35 rows.
    date_col = None
    time_col = None
    scan_rows = min(35, len(raw))
    for r in range(scan_rows):
        for c in range(raw.shape[1]):
            v = clean_header(raw.iat[r, c] if c < raw.shape[1] else "")
            if date_col is None and re.search(r"\bdate\b", v):
                date_col = c
            if time_col is None and re.search(r"\btime\b|hh:mm", v):
                time_col = c
        if date_col is not None and time_col is not None:
            break
    if date_col is not None and time_col is not None:
        try:
            dt = _combine_date_time_v49(raw.iloc[idx, date_col], raw.iloc[idx, time_col])
            if pd.to_datetime(dt, errors="coerce").notna().sum() >= 3:
                return dt
        except Exception:
            pass
    return pd.Series([pd.NaT] * len(idx), index=idx)


def _extract_pumping_pressure_from_pandas_raw_v51(data: bytes, name: str) -> pd.DataFrame:
    suffix = str(name).split(".")[-1].lower()
    if suffix not in {"xlsx", "xls"}:
        return pd.DataFrame()
    try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None, dtype=object)
    except Exception:
        return pd.DataFrame()

    rows = []
    for sheet_name, raw in sheets.items():
        if raw is None or raw.empty:
            continue
        # Excel files may have formatting extended to XFD (16,384 columns).
        # Prune all-empty columns before scanning headers, otherwise regex scanning
        # becomes extremely slow.  Hidden data columns such as EG are preserved
        # because they contain real values.
        raw = raw.dropna(axis=1, how="all")
        if raw.shape[1] > 400:
            sample = raw.head(250).astype(str).replace({"nan": "", "NaT": ""})
            keep_cols = [c for c in sample.columns if sample[c].str.strip().ne("").any()]
            raw = raw.loc[:, keep_cols]
        default_well = _clean_merge_well_v51(extract_well_from_raw(raw, source_name=name, sheet_name=str(sheet_name)) or "Unknown")
        candidate_cols = []
        scan_rows = min(50, len(raw))
        for r in range(scan_rows):
            for c in range(raw.shape[1]):
                val = raw.iat[r, c]
                unit_below = raw.iat[r + 1, c] if r + 1 < len(raw) else ""
                unit_above = raw.iat[r - 1, c] if r - 1 >= 0 else ""
                unit_right = raw.iat[r, c + 1] if c + 1 < raw.shape[1] else ""
                if _looks_like_pumping_pressure_header_v51(val, f"{unit_below} {unit_above} {unit_right}"):
                    candidate_cols.append((r, c))
        seen = set()
        for header_r, c in candidate_cols:
            if (header_r, c) in seen:
                continue
            seen.add((header_r, c))
            # Use data below the header, but also allow one or two unit rows directly below it.
            for start_r in [header_r + 1, header_r + 2, header_r + 3]:
                if start_r >= len(raw):
                    continue
                idx = raw.index[start_r:]
                vals = _to_number_series_v51(pd.Series(raw.iloc[start_r:, c], index=idx))
                if vals.notna().sum() < 3:
                    continue
                dt = _date_time_from_raw_v51(raw, idx)
                valid = vals.notna() & pd.to_datetime(dt, errors="coerce").notna()
                if valid.sum() < 3:
                    continue
                for row_idx in vals[valid].index:
                    dtv = pd.Timestamp(dt.loc[row_idx])
                    rows.append({
                        "source": name,
                        "sheet": str(sheet_name),
                        "well": default_well,
                        "well_key_v51": default_well,
                        "datetime": dtv,
                        "date": dtv.floor("D"),
                        "time_text": dtv.strftime("%H:%M"),
                        "pumping_pressure_psi": float(vals.loc[row_idx]),
                        "source_type": "excel_raw_pumping_pressure_v51",
                    })
                break
    if not rows:
        return pd.DataFrame()
    out = pd.DataFrame(rows)
    out = out.drop_duplicates(subset=["sheet", "well_key_v51", "datetime"], keep="last")
    return out


def _extract_pumping_pressure_from_openpyxl_v51(data: bytes, name: str) -> pd.DataFrame:
    # Extra fallback for hidden/far-right columns.  openpyxl sees hidden columns
    # such as EG even when the visual sheet hides them.
    suffix = str(name).split(".")[-1].lower()
    if suffix not in {"xlsx"}:
        return pd.DataFrame()
    try:
        from openpyxl import load_workbook
    except Exception:
        return pd.DataFrame()
    try:
        wb = load_workbook(io.BytesIO(data), read_only=False, data_only=True)
    except Exception:
        return pd.DataFrame()

    rows = []
    for ws in wb.worksheets:
        max_r = min(ws.max_row or 0, 5000)
        max_c = min(ws.max_column or 0, 300)
        if max_r < 3 or max_c < 3:
            continue
        # Build a small raw dataframe for well extraction and datetime parsing.
        raw_vals = []
        for r in range(1, max_r + 1):
            raw_vals.append([ws.cell(r, c).value for c in range(1, max_c + 1)])
        raw = pd.DataFrame(raw_vals)
        default_well = _clean_merge_well_v51(extract_well_from_raw(raw, source_name=name, sheet_name=str(ws.title)) or "Unknown")
        candidate_cols = []
        for r in range(1, min(50, max_r) + 1):
            for c in range(1, max_c + 1):
                val = ws.cell(r, c).value
                if val is None:
                    continue
                unit_below = ws.cell(r + 1, c).value if r + 1 <= max_r else ""
                unit_above = ws.cell(r - 1, c).value if r - 1 >= 1 else ""
                unit_right = ws.cell(r, c + 1).value if c + 1 <= max_c else ""
                if _looks_like_pumping_pressure_header_v51(val, f"{unit_below} {unit_above} {unit_right}"):
                    candidate_cols.append((r - 1, c - 1))
        for header_r0, c0 in candidate_cols:
            for start_r0 in [header_r0 + 1, header_r0 + 2, header_r0 + 3]:
                if start_r0 >= len(raw):
                    continue
                idx = raw.index[start_r0:]
                vals = _to_number_series_v51(pd.Series(raw.iloc[start_r0:, c0], index=idx))
                if vals.notna().sum() < 3:
                    continue
                dt = _date_time_from_raw_v51(raw, idx)
                valid = vals.notna() & pd.to_datetime(dt, errors="coerce").notna()
                if valid.sum() < 3:
                    continue
                for row_idx in vals[valid].index:
                    dtv = pd.Timestamp(dt.loc[row_idx])
                    rows.append({
                        "source": name,
                        "sheet": str(ws.title),
                        "well": default_well,
                        "well_key_v51": default_well,
                        "datetime": dtv,
                        "date": dtv.floor("D"),
                        "time_text": dtv.strftime("%H:%M"),
                        "pumping_pressure_psi": float(vals.loc[row_idx]),
                        "source_type": "excel_openpyxl_pumping_pressure_v51",
                    })
                break
    if not rows:
        return pd.DataFrame()
    out = pd.DataFrame(rows)
    out = out.drop_duplicates(subset=["sheet", "well_key_v51", "datetime"], keep="last")
    return out


def extract_any_pumping_pressure_from_excel_v51(data: bytes, name: str) -> pd.DataFrame:
    frames = []
    for fn in [_extract_pumping_pressure_from_pandas_raw_v51, _extract_pumping_pressure_from_openpyxl_v51, extract_hidden_pumping_pressure_from_excel_v49]:
        try:
            df = fn(data, name)
            if df is not None and not df.empty:
                frames.append(df)
        except Exception:
            pass
    if not frames:
        return pd.DataFrame()
    out = pd.concat(frames, ignore_index=True, sort=False)
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["well"] = out["well"].map(_clean_merge_well_v51)
    out["well_key_v51"] = out["well"].map(_clean_merge_well_v51)
    out["pumping_pressure_psi"] = pd.to_numeric(out["pumping_pressure_psi"], errors="coerce")
    out = out.dropna(subset=["datetime", "pumping_pressure_psi"])
    out = out.drop_duplicates(subset=["well_key_v51", "datetime"], keep="last")
    return out


def merge_pumping_pressure_tables_v51(tables, pump_df: pd.DataFrame):
    if pump_df is None or pump_df.empty:
        return tables
    pump_df = pump_df.copy()
    pump_df["datetime"] = pd.to_datetime(pump_df["datetime"], errors="coerce")
    pump_df["well_key_v51"] = pump_df.get("well", "Unknown").map(_clean_merge_well_v51) if isinstance(pump_df.get("well", None), pd.Series) else "Unknown"
    pump_df = pump_df.dropna(subset=["datetime", "pumping_pressure_psi"])
    if pump_df.empty:
        return tables
    if not tables:
        return [pump_df]

    merged_tables = []
    used_any = False
    for df in tables:
        if df is None or df.empty:
            merged_tables.append(df)
            continue
        out = df.copy()
        if "datetime" not in out.columns:
            merged_tables.append(out)
            continue
        out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
        if "well" not in out.columns:
            out["well"] = "Unknown"
        out["well_key_v51"] = out["well"].map(_clean_merge_well_v51)
        tmp = pump_df[["well_key_v51", "datetime", "pumping_pressure_psi"]].copy()

        # First and safest: clean well + datetime.
        out = out.merge(tmp, on=["well_key_v51", "datetime"], how="left", suffixes=("", "__pump_v51"))
        if "pumping_pressure_psi__pump_v51" in out.columns:
            base = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce")
            hidden = pd.to_numeric(out["pumping_pressure_psi__pump_v51"], errors="coerce")
            out["pumping_pressure_psi"] = base.combine_first(hidden)
            out = out.drop(columns=["pumping_pressure_psi__pump_v51"], errors="ignore")

        # Second: datetime only, but only when there is a single real well in the uploaded table.
        if pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce").notna().sum() == 0:
            real_wells = [w for w in out["well_key_v51"].dropna().unique().tolist() if w and w != "Unknown"]
            if len(real_wells) == 1:
                tmp2 = pump_df[["datetime", "pumping_pressure_psi"]].drop_duplicates("datetime", keep="last")
                out = out.drop(columns=[c for c in out.columns if c.endswith("__pump_v51")], errors="ignore")
                out = out.merge(tmp2, on="datetime", how="left", suffixes=("", "__pump_v51"))
                if "pumping_pressure_psi__pump_v51" in out.columns:
                    base = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce")
                    hidden = pd.to_numeric(out["pumping_pressure_psi__pump_v51"], errors="coerce")
                    out["pumping_pressure_psi"] = base.combine_first(hidden)
                    out = out.drop(columns=["pumping_pressure_psi__pump_v51"], errors="ignore")

        # Third: row-order fallback for the same single-well sheet when datetime parsing differs slightly.
        current_count = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce").notna().sum()
        if current_count == 0:
            real_wells = [w for w in out["well_key_v51"].dropna().unique().tolist() if w and w != "Unknown"]
            if len(real_wells) == 1 and abs(len(out) - len(pump_df)) <= max(3, int(0.15 * max(len(out), len(pump_df)))):
                out_sorted_idx = out.sort_values("datetime").index.tolist()
                pump_sorted = pump_df.sort_values("datetime")["pumping_pressure_psi"].reset_index(drop=True)
                n = min(len(out_sorted_idx), len(pump_sorted))
                out["pumping_pressure_psi"] = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce")
                for pos in range(n):
                    if pd.isna(out.at[out_sorted_idx[pos], "pumping_pressure_psi"]):
                        out.at[out_sorted_idx[pos], "pumping_pressure_psi"] = pump_sorted.iloc[pos]

        final_count = pd.to_numeric(out.get("pumping_pressure_psi", pd.Series(np.nan, index=out.index)), errors="coerce").notna().sum()
        if final_count > 0:
            used_any = True
        out = out.drop(columns=["well_key_v51"], errors="ignore")
        merged_tables.append(out)

    # If no existing table accepted the pumping series, append it so the user can still see it.
    if not used_any:
        merged_tables.append(pump_df.drop(columns=["well_key_v51"], errors="ignore"))
    return merged_tables


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).split(".")[-1].lower()
    base_tables = _load_tabular_file_v50_final(UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images)
    fixed_tables = [ensure_pumping_pressure_column_v48(t) for t in (base_tables or []) if t is not None and not t.empty]
    if suffix in {"xlsx", "xls"}:
        # Speed guard: if the normal/v49 parser already found Pumping Pressure,
        # do not run the heavier raw/openpyxl rescue again.  This also avoids
        # repeated workbook reads on Streamlit Cloud.
        has_pump = False
        for _t in fixed_tables:
            if "pumping_pressure_psi" in _t.columns and pd.to_numeric(_t["pumping_pressure_psi"], errors="coerce").notna().sum() > 0:
                has_pump = True
                break
        if not has_pump:
            pump_df = extract_any_pumping_pressure_from_excel_v51(data, name)
            fixed_tables = merge_pumping_pressure_tables_v51(fixed_tables, pump_df)
            fixed_tables = [ensure_pumping_pressure_column_v48(t) for t in fixed_tables if t is not None and not t.empty]
    return fixed_tables


# -----------------------------------------------------------------------------
# v52: very fast XML-based pumping-pressure rescue for hidden/far-right columns.
# This solves templates such as S8-58 where Pumping Pressure is in hidden/far
# right column EG with header "pumping.p" / unit "psi".  It reads the XLSX XML
# directly instead of relying only on the main table parser, so Excel hidden
# columns, styled empty columns, and max-column=16384 artifacts do not hide the
# field.
# -----------------------------------------------------------------------------
import zipfile as _zipfile_v52
import xml.etree.ElementTree as _ET_v52


def _xlsx_col_to_num_v52(col_letters: str) -> int:
    n = 0
    for ch in str(col_letters).upper():
        if "A" <= ch <= "Z":
            n = n * 26 + (ord(ch) - ord("A") + 1)
    return n


def _xlsx_cell_ref_parts_v52(ref: str):
    m = re.match(r"^([A-Z]+)(\d+)$", str(ref).upper())
    if not m:
        return None, None
    return _xlsx_col_to_num_v52(m.group(1)), int(m.group(2))


def _read_shared_strings_v52(zf) -> list:
    try:
        data = zf.read("xl/sharedStrings.xml")
    except Exception:
        return []
    ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    root = _ET_v52.fromstring(data)
    out = []
    for si in root.findall("m:si", ns):
        parts = []
        for t in si.findall(".//m:t", ns):
            parts.append(t.text or "")
        out.append("".join(parts))
    return out


def _iter_xlsx_sheet_cells_v52(zf, sheet_path: str, shared_strings: list):
    """Yield (row, col, value) for non-empty worksheet cells."""
    ns_uri = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
    try:
        with zf.open(sheet_path) as fh:
            for event, elem in _ET_v52.iterparse(fh, events=("end",)):
                if elem.tag != ns_uri + "c":
                    continue
                ref = elem.attrib.get("r", "")
                col, row = _xlsx_cell_ref_parts_v52(ref)
                if not row or not col:
                    elem.clear()
                    continue
                typ = elem.attrib.get("t")
                value = None
                if typ == "inlineStr":
                    parts = []
                    for tnode in elem.findall(".//" + ns_uri + "t"):
                        parts.append(tnode.text or "")
                    value = "".join(parts)
                else:
                    vnode = elem.find(ns_uri + "v")
                    if vnode is not None and vnode.text is not None:
                        raw = vnode.text
                        if typ == "s":
                            try:
                                value = shared_strings[int(float(raw))]
                            except Exception:
                                value = raw
                        else:
                            value = raw
                if value not in [None, ""]:
                    yield row, col, value
                elem.clear()
    except Exception:
        return


def _workbook_sheet_paths_v52(zf):
    """Return list of (sheet_name, sheet_path) for workbook sheets."""
    ns_main = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main", "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    rel_ns = {"pr": "http://schemas.openxmlformats.org/package/2006/relationships"}
    try:
        wb_root = _ET_v52.fromstring(zf.read("xl/workbook.xml"))
        rel_root = _ET_v52.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
    except Exception:
        return []
    rels = {}
    for rel in rel_root.findall("pr:Relationship", rel_ns):
        rid = rel.attrib.get("Id")
        target = rel.attrib.get("Target", "")
        if rid:
            if not target.startswith("xl/"):
                target = "xl/" + target.lstrip("/")
            rels[rid] = target
    sheets = []
    for sh in wb_root.findall(".//m:sheet", ns_main):
        name = sh.attrib.get("name", "Sheet")
        rid = sh.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        path = rels.get(rid)
        if path:
            sheets.append((name, path))
    return sheets


def _excel_datetime_from_values_v52(date_val, time_val):
    def _to_float(x):
        try:
            if x is None or str(x).strip() == "":
                return np.nan
            return float(str(x).strip())
        except Exception:
            return np.nan

    # Date may be Excel serial, a datetime-like string, or already parsed text.
    d = pd.NaT
    fv = _to_float(date_val)
    if pd.notna(fv) and fv > 1000:
        try:
            d = pd.Timestamp("1899-12-30") + pd.to_timedelta(float(fv), unit="D")
        except Exception:
            d = pd.NaT
    if pd.isna(d):
        d = pd.to_datetime(date_val, errors="coerce", dayfirst=True)

    # Time may be Excel serial fraction, serial datetime, datetime.time string, HH:MM, or text.
    t = None
    tf = _to_float(time_val)
    if pd.notna(tf):
        frac = tf % 1.0
        try:
            seconds = int(round(frac * 24 * 3600))
            t = (pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)).time()
        except Exception:
            t = None
    if t is None:
        ts = pd.to_datetime(str(time_val), errors="coerce")
        if pd.notna(ts):
            t = pd.Timestamp(ts).time()
    if pd.isna(d) or t is None:
        return pd.NaT
    return pd.Timestamp.combine(pd.Timestamp(d).date(), t)


def _well_from_sheet_or_file_v52(sheet_name: str, name: str) -> str:
    source = f"{sheet_name} {name}"
    # Preserve field names like B3 C18-7 / B3C18-7.  Older generic guessing may
    # return only C18-7, which blocks well+datetime merging.
    m = re.search(r"\b(B\d+)\s*[-_ ]?\s*(C\d{1,3}[-_ ]?\d+)\b", source, flags=re.I)
    if m:
        return (m.group(1) + m.group(2)).replace(" ", "").replace("_", "-").upper()
    m = re.search(r"\b(S\d+[-_ ]?\d+)\b", source, flags=re.I)
    if m:
        return m.group(1).replace(" ", "").replace("_", "-").upper()
    return guess_well_from_name(str(sheet_name)) or guess_well_from_name(str(name)) or "Unknown"


def _extract_pumping_pressure_from_xlsx_xml_v52(data: bytes, name: str) -> pd.DataFrame:
    rows = []
    try:
        zf = _zipfile_v52.ZipFile(io.BytesIO(data))
    except Exception:
        return pd.DataFrame()
    with zf:
        shared = _read_shared_strings_v52(zf)
        sheet_paths = _workbook_sheet_paths_v52(zf)
        for sheet_name, sheet_path in sheet_paths:
            sparse = {}
            max_row = 0
            for r, c, v in _iter_xlsx_sheet_cells_v52(zf, sheet_path, shared):
                sparse[(r, c)] = v
                max_row = max(max_row, r)
            if not sparse:
                continue

            # Find all credible pumping-pressure header cells.  Reject operation
            # comments like "start pumping N2" because they do not look like Pump P.
            header_cells = []
            for (r, c), v in sparse.items():
                current_txt = clean_header(v)
                # Only the header cell itself may identify the field.  Neighboring
                # cells are used only for units/context, not to turn a plain "PSI"
                # unit cell into a Pumping Pressure header.
                if not re.search(r"pump|pumping|circulation", current_txt, flags=re.I):
                    continue
                unit_context = " ".join(str(sparse.get((rr, cc), "")) for rr in range(r - 2, r + 3) for cc in range(c - 1, c + 2))
                if _looks_like_pumping_pressure_header_v51(v, unit_context):
                    header_cells.append((r, c, v))
            if not header_cells:
                continue

            # Find date/time columns from header rows.  If template does not label
            # them, default to A=date, B=time, which is used by the field templates.
            date_col, time_col = 1, 2
            for rr in range(1, min(12, max_row) + 1):
                for cc in range(1, 12):
                    txt = clean_header(sparse.get((rr, cc), ""))
                    if txt in {"date", "d mm yyy", "d mmm yyyy"} or txt.startswith("date"):
                        date_col = cc
                    if txt in {"time", "hh mm", "hh:mm"} or txt.startswith("time"):
                        time_col = cc

            well = _well_from_sheet_or_file_v52(str(sheet_name), str(name))
            for hr, pc, hv in header_cells:
                for r in range(hr + 1, max_row + 1):
                    pump_raw = sparse.get((r, pc))
                    pump_val = extract_number(pump_raw)
                    if pd.isna(pump_val):
                        continue
                    dt = _excel_datetime_from_values_v52(sparse.get((r, date_col)), sparse.get((r, time_col)))
                    if pd.isna(dt):
                        continue
                    rows.append({
                        "source": name,
                        "sheet": sheet_name,
                        "well": well,
                        "datetime": dt,
                        "date": pd.Timestamp(dt).normalize(),
                        "time_text": pd.Timestamp(dt).strftime("%H:%M"),
                        "pumping_pressure_psi": float(pump_val),
                        "source_type": "excel_xml_pumping_pressure_v52",
                        "pump_header_cell": f"R{hr}C{pc}",
                    })
    if not rows:
        return pd.DataFrame()
    out = pd.DataFrame(rows)
    out["well"] = out["well"].map(_clean_merge_well_v51)
    out["well_key_v51"] = out["well"].map(_clean_merge_well_v51)
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["pumping_pressure_psi"] = pd.to_numeric(out["pumping_pressure_psi"], errors="coerce")
    out = out.dropna(subset=["datetime", "pumping_pressure_psi"])
    out = out.drop_duplicates(subset=["well_key_v51", "datetime"], keep="last")
    return out


def extract_any_pumping_pressure_from_excel_v52(data: bytes, name: str) -> pd.DataFrame:
    # Fast XML first.  If it succeeds, trust it and skip the slow pandas/openpyxl
    # fallbacks.  This is critical for templates whose used range extends to XFD.
    try:
        xml_df = _extract_pumping_pressure_from_xlsx_xml_v52(data, name)
        if xml_df is not None and not xml_df.empty:
            out = xml_df.copy()
        else:
            raise ValueError("XML pump extraction found no rows")
    except Exception:
        frames = []
        for fn in [_extract_pumping_pressure_from_pandas_raw_v51, _extract_pumping_pressure_from_openpyxl_v51, extract_hidden_pumping_pressure_from_excel_v49]:
            try:
                df = fn(data, name)
                if df is not None and not df.empty:
                    frames.append(df)
            except Exception:
                pass
        if not frames:
            return pd.DataFrame()
        out = pd.concat(frames, ignore_index=True, sort=False)
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["well"] = out.get("well", "Unknown").map(_clean_merge_well_v51) if isinstance(out.get("well", None), pd.Series) else "Unknown"
    out["well_key_v51"] = out["well"].map(_clean_merge_well_v51)
    out["pumping_pressure_psi"] = pd.to_numeric(out["pumping_pressure_psi"], errors="coerce")
    out = out.dropna(subset=["datetime", "pumping_pressure_psi"])
    # Prefer XML/header-specific rows over broad pandas/openpyxl rows when duplicates exist.
    out["_pump_priority_v52"] = out.get("source_type", "").astype(str).str.contains("xml|hidden", case=False, na=False).astype(int)
    out = out.sort_values(["well_key_v51", "datetime", "_pump_priority_v52"]).drop_duplicates(subset=["well_key_v51", "datetime"], keep="last")
    out = out.drop(columns=["_pump_priority_v52"], errors="ignore")
    return out


def force_restore_pumping_pressure_v52(tables, pump_df: pd.DataFrame):
    """Always merge pump rescue; append pump-only table only if merge still fails."""
    if pump_df is None or pump_df.empty:
        return tables
    fixed = merge_pumping_pressure_tables_v51(tables, pump_df)
    # Make sure the canonical column is numeric and visible in every returned table.
    out_tables = []
    any_visible = False
    for t in fixed or []:
        if t is None or t.empty:
            continue
        t = ensure_pumping_pressure_column_v48(t)
        if "pumping_pressure_psi" in t.columns:
            t["pumping_pressure_psi"] = pd.to_numeric(t["pumping_pressure_psi"], errors="coerce")
            if t["pumping_pressure_psi"].notna().sum() > 0:
                any_visible = True
        out_tables.append(t)
    if not any_visible:
        out_tables.append(pump_df.drop(columns=["well_key_v51"], errors="ignore"))
    return out_tables


# Keep previous loader, but force cache-busted XML Pumping Pressure rescue for Excel.
_load_tabular_file_v51_final = load_tabular_file

def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).split(".")[-1].lower()
    # Use v50 base loader to avoid v51 slow openpyxl rescue. v52 XML rescue is applied below.
    base_loader = globals().get("_load_tabular_file_v50_final", _load_tabular_file_v51_final)
    base_tables = base_loader(UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images)
    fixed_tables = [ensure_pumping_pressure_column_v48(t) for t in (base_tables or []) if t is not None and not t.empty]
    if suffix in {"xlsx", "xls"}:
        pump_df = extract_any_pumping_pressure_from_excel_v52(data, name)
        fixed_tables = force_restore_pumping_pressure_v52(fixed_tables, pump_df)
        fixed_tables = [ensure_pumping_pressure_column_v48(t) for t in fixed_tables if t is not None and not t.empty]
    return fixed_tables

# =============================================================================
# v53 — safe XLSX preflight / fast XML loader + cross-file de-duplication
# =============================================================================
# Why this exists:
# Some field workbooks have an accidentally inflated Excel used range such as
# A1:FU1048518 or A1:XFC56.  pandas/openpyxl may allocate or scan the entire
# range, which can exhaust Streamlit Cloud memory and show the generic
# "Oh no. Error running app" page.  v53 detects those workbooks before the
# normal loader and reads only non-empty XML cells in a bounded data region.

import zipfile as _zipfile_v53
import xml.etree.ElementTree as _ET_v53

PARSER_BUILD_ID_V53 = "v54-safe-xlsx-object-dtype-20260621"
PARSER_BUILD_ID_V54 = "v54-safe-xlsx-object-dtype-20260621"

# Internal bookkeeping columns must never appear as plot features.
try:
    BASE_NON_PLOT_COLS.update({
        "_upload_order", "_table_order", "_source_row_order",
        "_well_key_v53", "_datetime_key_v53", "_completeness_v53",
    })
except Exception:
    pass


def _xlsx_col_number_v53(col_letters: str) -> int:
    n = 0
    for ch in str(col_letters or "").upper():
        if "A" <= ch <= "Z":
            n = n * 26 + (ord(ch) - 64)
    return n


def _xlsx_ref_parts_v53(ref: str):
    m = re.match(r"^([A-Za-z]+)(\d+)$", str(ref or ""))
    if not m:
        return 0, 0
    return int(m.group(2)), _xlsx_col_number_v53(m.group(1))


def _xlsx_dimension_parts_v53(ref: str):
    last = str(ref or "A1").split(":")[-1]
    return _xlsx_ref_parts_v53(last)


def xlsx_preflight_v53(data: bytes) -> dict:
    """Inspect XLSX dimensions without loading cells into pandas/openpyxl."""
    result = {
        "is_xlsx": False,
        "suspicious": False,
        "max_declared_row": 0,
        "max_declared_col": 0,
        "sheet_xml_uncompressed": 0,
        "reasons": [],
    }
    try:
        with _zipfile_v53.ZipFile(io.BytesIO(data)) as zf:
            result["is_xlsx"] = True
            for info in zf.infolist():
                if not (info.filename.startswith("xl/worksheets/sheet") and info.filename.endswith(".xml")):
                    continue
                result["sheet_xml_uncompressed"] += int(info.file_size or 0)
                try:
                    with zf.open(info.filename) as fh:
                        head = fh.read(32768)
                    m = re.search(br"<dimension\s+ref=\"([^\"]+)\"", head)
                    if m:
                        r, c = _xlsx_dimension_parts_v53(m.group(1).decode("utf-8", errors="ignore"))
                        result["max_declared_row"] = max(result["max_declared_row"], r)
                        result["max_declared_col"] = max(result["max_declared_col"], c)
                except Exception:
                    pass
    except Exception:
        return result

    # These limits are far above normal TMU sheets but below pathological Excel
    # used ranges.  Actual non-empty cells outside the bounds are exceptionally
    # unlikely for production-test sheets.
    if result["max_declared_row"] > 20000:
        result["reasons"].append("inflated row dimension")
    if result["max_declared_col"] > 512:
        result["reasons"].append("inflated column dimension")
    if result["sheet_xml_uncompressed"] > 25_000_000:
        result["reasons"].append("very large worksheet XML")
    result["suspicious"] = bool(result["reasons"])
    return result


def _read_shared_strings_fast_v53(zf) -> list:
    try:
        root = _ET_v53.fromstring(zf.read("xl/sharedStrings.xml"))
    except Exception:
        return []
    ns_uri = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
    out = []
    for si in root.findall(ns_uri + "si"):
        out.append("".join((t.text or "") for t in si.findall(".//" + ns_uri + "t")))
    return out


def _iter_sheet_cells_limited_v53(zf, sheet_path: str, shared_strings: list,
                                  max_rows: int = 20000, max_cols: int = 512):
    """Yield non-empty cells, stopping before inflated empty worksheet tails."""
    ns_uri = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
    with zf.open(sheet_path) as fh:
        for _, elem in _ET_v53.iterparse(fh, events=("end",)):
            if elem.tag != ns_uri + "c":
                continue
            row, col = _xlsx_ref_parts_v53(elem.attrib.get("r", ""))
            if row <= 0 or col <= 0:
                elem.clear()
                continue
            # Excel serializes cells in row order. Once the bound is exceeded,
            # the remaining million-row formatted tail can be ignored safely.
            if row > max_rows:
                elem.clear()
                break
            if col > max_cols:
                elem.clear()
                continue
            typ = elem.attrib.get("t")
            value = None
            if typ == "inlineStr":
                value = "".join((tn.text or "") for tn in elem.findall(".//" + ns_uri + "t"))
            else:
                vnode = elem.find(ns_uri + "v")
                if vnode is not None and vnode.text is not None:
                    raw = vnode.text
                    if typ == "s":
                        try:
                            value = shared_strings[int(float(raw))]
                        except Exception:
                            value = raw
                    elif typ == "b":
                        value = str(raw).strip() == "1"
                    else:
                        value = raw
            if value not in (None, ""):
                yield row, col, value
            elem.clear()


def _excel_serial_date_v53(value):
    """Convert only plausible modern Excel serial dates.

    Values such as 1234, 1423, 2600, and 2613 are operational measurements in
    these TMU sheets and must never be interpreted as calendar years/dates.
    """
    try:
        f = float(str(value).strip())
    except Exception:
        return value
    if 20000.0 <= f <= 80000.0:
        try:
            ts = pd.Timestamp("1899-12-30") + pd.to_timedelta(f, unit="D")
            return ts if 1900 <= int(ts.year) <= 2100 else value
        except Exception:
            return value
    return value


def _excel_serial_time_v53(value):
    try:
        f = float(str(value).strip())
    except Exception:
        return value
    # Some copied TMU sheets use cumulative Excel day values above 1 or 2 in
    # the Time column. Only the fractional day is the clock time. Since this
    # function is called only for a strongly detected Time column, values below
    # 10 days are safe to normalize and must not be dropped.
    if 0 <= f < 10:
        try:
            seconds = int(round((f % 1.0) * 86400)) % 86400
            return (pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)).time()
        except Exception:
            return value
    return value


def _normalize_excel_date_time_columns_v53(raw: pd.DataFrame) -> pd.DataFrame:
    if raw is None or raw.empty:
        return raw

    # Rebuild from Python lists. This prevents Arrow/string/datetime extension
    # arrays from surviving into the raw worksheet and avoids pandas setitem bugs.
    out = pd.DataFrame(
        [[cell for cell in row] for row in raw.to_numpy(dtype=object).tolist()],
        index=raw.index.copy(),
        columns=raw.columns.copy(),
        dtype=object,
    )

    date_cols = []
    time_cols = []
    header_end_by_col = {}
    for r in range(min(30, len(out))):
        for c in range(out.shape[1]):
            txt = clean_header(out.iat[r, c])
            if not txt:
                continue
            if re.search(r"(^|\b)date($|\b)|d/mm|d-mmm|yyyy", txt):
                date_cols.append(c)
                header_end_by_col[c] = max(header_end_by_col.get(c, 0), r)
            if re.search(r"(^|\b)time($|\b)|hh:mm", txt):
                time_cols.append(c)
                header_end_by_col[c] = max(header_end_by_col.get(c, 0), r)

    # Use A/B fallback only when their values actually look like date/time data.
    if not date_cols and out.shape[1] >= 1:
        sample = out.iloc[: min(len(out), 100), 0].tolist()
        plausible = sum(
            1 for v in sample
            if isinstance(v, (pd.Timestamp,))
            or (isinstance(v, (int, float, np.number)) and not isinstance(v, bool) and 20000 <= float(v) <= 80000)
            or bool(re.search(r"\d{1,4}[/-]\d{1,2}[/-]\d{1,4}", str(v or "")))
        )
        if plausible >= 2:
            date_cols = [0]
    if not time_cols and out.shape[1] >= 2:
        sample = out.iloc[: min(len(out), 100), 1].tolist()
        plausible = sum(
            1 for v in sample
            if hasattr(v, "hour") and hasattr(v, "minute") and not isinstance(v, (str, int, float, np.number))
            or (isinstance(v, (int, float, np.number)) and not isinstance(v, bool) and 0 <= float(v) < 2)
            or bool(re.fullmatch(r"\s*\d{1,2}[:.]\d{2}(?::\d{2})?\s*(?:am|pm)?\s*", str(v or ""), re.I))
        )
        if plausible >= 2:
            time_cols = [1]

    for c in sorted(set(date_cols)):
        start = header_end_by_col.get(c, -1) + 1
        values = out.iloc[:, c].tolist()
        for i in range(max(0, start), len(values)):
            values[i] = _excel_serial_date_v53(values[i])
        out.iloc[:, c] = pd.Series(values, index=out.index, dtype=object).to_numpy(dtype=object)

    for c in sorted(set(time_cols)):
        start = header_end_by_col.get(c, -1) + 1
        values = out.iloc[:, c].tolist()
        for i in range(max(0, start), len(values)):
            values[i] = _excel_serial_time_v53(values[i])
        out.iloc[:, c] = pd.Series(values, index=out.index, dtype=object).to_numpy(dtype=object)

    return pd.DataFrame(out.to_numpy(dtype=object), index=out.index, columns=out.columns, dtype=object)

def _raw_dataframe_from_sheet_xml_v53(zf, sheet_path: str, shared_strings: list,
                                      max_rows: int = 20000, max_cols: int = 512):
    sparse = {}
    max_r = 0
    max_c = 0
    for r, c, v in _iter_sheet_cells_limited_v53(
        zf, sheet_path, shared_strings, max_rows=max_rows, max_cols=max_cols
    ):
        sparse[(r, c)] = v
        max_r = max(max_r, r)
        max_c = max(max_c, c)
    if not sparse or max_r <= 0 or max_c <= 0:
        return pd.DataFrame()
    matrix = [[None] * max_c for _ in range(max_r)]
    for (r, c), v in sparse.items():
        matrix[r - 1][c - 1] = v
    raw = pd.DataFrame(matrix, dtype=object)
    raw = raw.dropna(how="all").dropna(axis=1, how="all")
    return _normalize_excel_date_time_columns_v53(raw)


def _load_suspicious_xlsx_fast_v53(data: bytes, name: str) -> List[pd.DataFrame]:
    tables = []
    try:
        with _zipfile_v53.ZipFile(io.BytesIO(data)) as zf:
            shared = _read_shared_strings_fast_v53(zf)
            sheet_paths = _workbook_sheet_paths_v52(zf)
            for sheet_name, sheet_path in sheet_paths:
                raw = _raw_dataframe_from_sheet_xml_v53(zf, sheet_path, shared)
                if raw is None or raw.empty:
                    continue
                default_well = extract_well_from_raw(raw, source_name=name, sheet_name=str(sheet_name))
                candidates = parse_excel_sheet_attempts(
                    raw, source_name=name, sheet_name=str(sheet_name), default_well=default_well
                )
                if candidates:
                    best = candidates[0].copy()
                    best["source_type"] = "excel_fast_xml_v53"
                    tables.append(best)
    except Exception as exc:
        raise RuntimeError(f"Safe XML parser could not read {name}: {exc}")
    return filter_usable_tables(filter_preferred_tables(tables))


def _missing_value_v53(value) -> bool:
    if value is None:
        return True
    try:
        if pd.isna(value):
            return True
    except Exception:
        pass
    if isinstance(value, str) and value.strip().lower() in {"", "nan", "nat", "none", "unknown"}:
        return True
    return False


def _normalized_well_key_v53(value) -> str:
    """Return a separator-insensitive well key used only for matching/merging.

    Field reports frequently write the same well as B16C6-9, B16-C6-9,
    B16 C6-9 or B16_C6_9.  Those must merge as one test stream while the
    preferred human-readable spelling is preserved separately.
    """
    try:
        key = _clean_merge_well_v51(value)
    except Exception:
        key = str(value or "").strip().upper()
    key = str(key or "").upper().strip()
    if key in {"", "UNKNOWN", "NAN", "NONE", "*"}:
        return ""
    # Ignore punctuation and spacing for identity matching only.
    key = re.sub(r"[^A-Z0-9]+", "", key)
    if key in {"", "UNKNOWN", "NAN", "NONE"}:
        return ""
    return key


def merge_duplicate_test_rows_v53(df: pd.DataFrame) -> pd.DataFrame:
    """Merge repeated uploads by normalized well + minute timestamp.

    A later/complete report often repeats rows from an earlier incomplete report.
    For each duplicate timestamp, keep the most complete row, then fill any gaps
    from the other copies. Conflicting populated values stay with the most
    complete row; ties prefer the later uploaded file.
    """
    if df is None or df.empty or "datetime" not in df.columns:
        return df
    out = df.copy().reset_index(drop=True)
    if "well" not in out.columns:
        out["well"] = "Unknown"
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["_source_row_order"] = np.arange(len(out), dtype=int)
    if "_upload_order" not in out.columns:
        out["_upload_order"] = 0
    if "_table_order" not in out.columns:
        out["_table_order"] = 0
    out["_well_key_v53"] = out["well"].map(_normalized_well_key_v53)
    out["_datetime_key_v53"] = out["datetime"].dt.floor("min")

    numeric_cols = [
        c for c in available_numeric_columns(out)
        if not str(c).startswith("_") and c not in {"duplicate_rows_merged"}
    ]
    canonical_cols = [c for c in numeric_cols if not str(c).startswith("raw__")]
    raw_cols = [c for c in numeric_cols if str(c).startswith("raw__")]
    canonical_score = out[canonical_cols].notna().sum(axis=1) * 3 if canonical_cols else 0
    raw_score = out[raw_cols].notna().sum(axis=1) if raw_cols else 0
    out["_completeness_v53"] = canonical_score + raw_score

    # Use one display spelling for all separator variants of the same well.
    # Prefer the spelling coming from the richest source table (normally the
    # full TMU test sheet), then the earliest uploaded source.  This also
    # updates rows that exist only in the device export, so a single missing
    # timestamp does not appear under a second well name.
    _valid_alias = out["_well_key_v53"].astype(str).str.len().gt(0)
    if _valid_alias.any():
        alias_stats = (
            out.loc[_valid_alias]
            .groupby(["_well_key_v53", "well"], dropna=False, sort=False)
            .agg(
                total_completeness=("_completeness_v53", "sum"),
                max_completeness=("_completeness_v53", "max"),
                row_count=("well", "size"),
                first_upload=("_upload_order", "min"),
                first_row=("_source_row_order", "min"),
            )
            .reset_index()
        )
        alias_stats = alias_stats.sort_values(
            ["_well_key_v53", "total_completeness", "max_completeness",
             "row_count", "first_upload", "first_row"],
            ascending=[True, False, False, False, True, True],
            kind="stable",
        )
        preferred_well = (
            alias_stats.drop_duplicates("_well_key_v53", keep="first")
            .set_index("_well_key_v53")["well"]
            .to_dict()
        )
        out.loc[_valid_alias, "well"] = out.loc[_valid_alias, "_well_key_v53"].map(preferred_well)

    valid_key = (
        out["_well_key_v53"].astype(str).str.len().gt(0)
        & out["_datetime_key_v53"].notna()
    )
    keep_records = []
    used = set()

    # Only group rows that have a trustworthy well and time. Unknown/NaT rows are
    # kept separately to avoid accidental cross-well merges.
    grouped = out.loc[valid_key].groupby(["_well_key_v53", "_datetime_key_v53"], sort=False, dropna=False)
    for (_, dt_key), group in grouped:
        idxs = list(group.index)
        used.update(idxs)
        if len(idxs) == 1:
            rec = out.loc[idxs[0]].copy()
            rec["datetime"] = dt_key
            keep_records.append(rec)
            continue
        ranked = group.sort_values(
            ["_completeness_v53", "_upload_order", "_table_order", "_source_row_order"],
            ascending=[False, False, False, False],
            kind="stable",
        )
        base = ranked.iloc[0].copy()
        notes = []
        for _, row in ranked.iterrows():
            for c in out.columns:
                if c in {"_well_key_v53", "_datetime_key_v53", "_completeness_v53"}:
                    continue
                val = row.get(c)
                if c == "note":
                    if not _missing_value_v53(val):
                        txt = str(val).strip()
                        if txt and txt not in notes:
                            notes.append(txt)
                    continue
                if _missing_value_v53(base.get(c)) and not _missing_value_v53(val):
                    base[c] = val
        if notes:
            base["note"] = "; ".join(notes)
        base["datetime"] = dt_key
        base["date"] = pd.Timestamp(dt_key).normalize()
        base["time_text"] = pd.Timestamp(dt_key).strftime("%H:%M")
        keep_records.append(base)

    for idx in out.index:
        if idx not in used:
            keep_records.append(out.loc[idx].copy())

    if not keep_records:
        return out.drop(columns=["_well_key_v53", "_datetime_key_v53", "_completeness_v53", "_source_row_order"], errors="ignore")
    result = pd.DataFrame(keep_records)
    result = result.sort_values(["well", "datetime", "_upload_order", "_source_row_order"], na_position="last", kind="stable")
    result = result.drop(columns=["_well_key_v53", "_datetime_key_v53", "_completeness_v53", "_source_row_order"], errors="ignore")
    return result.reset_index(drop=True)


# Final loader override. Keep the prior v52 loader for normal workbooks and use
# bounded XML parsing only for suspicious/inflated XLSX files.
_load_tabular_file_v52_before_v53 = load_tabular_file

def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).split(".")[-1].lower()
    if suffix == "xlsx":
        preflight = xlsx_preflight_v53(data)
        if preflight.get("suspicious"):
            tables = _load_suspicious_xlsx_fast_v53(data, name)
            # XML-only pump rescue is safe even for inflated worksheets. Do not
            # invoke pandas/openpyxl fallbacks on these files.
            try:
                pump_df = _extract_pumping_pressure_from_xlsx_xml_v52(data, name)
            except Exception:
                pump_df = pd.DataFrame()
            tables = force_restore_pumping_pressure_v52(tables, pump_df)
            return [ensure_pumping_pressure_column_v48(t) for t in tables if t is not None and not t.empty]
    return _load_tabular_file_v52_before_v53(
        UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images
    )

# v53.1: use the bounded XML reader as the primary path for every XLSX.  This
# avoids slow openpyxl/pandas behaviour not only for million-row used ranges but
# also for files with accidental XFC column formatting.  Normal loader remains a
# fallback only for non-suspicious XLSX files when XML parsing finds no table.
_load_tabular_file_v53_prefer_xml_previous = load_tabular_file

def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).split(".")[-1].lower()
    if suffix == "xlsx":
        preflight = xlsx_preflight_v53(data)
        tables = _load_suspicious_xlsx_fast_v53(data, name)
        try:
            pump_df = _extract_pumping_pressure_from_xlsx_xml_v52(data, name)
        except Exception:
            pump_df = pd.DataFrame()
        if tables:
            tables = force_restore_pumping_pressure_v52(tables, pump_df)
            return [ensure_pumping_pressure_column_v48(t) for t in tables if t is not None and not t.empty]
        # Never send a suspicious workbook into pandas/openpyxl after the safe
        # parser fails; returning no table is safer than crashing the whole app.
        if preflight.get("suspicious"):
            return []
    return _load_tabular_file_v53_prefer_xml_previous(
        UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images
    )

# =============================================================================
# v56 - faster first load + unit-safe choke handling
# =============================================================================

def _sheet_is_helper_v56(sheet_name: object, raw: pd.DataFrame) -> bool:
    name = clean_header(sheet_name)
    helper = bool(re.search(r"\b(form|cover|summary|chart|cmsf|shrinkage|lookup|calculation|calc)\b", name))
    # Never skip a large sheet only because its title contains a helper word.
    return helper and (raw is None or raw.shape[0] < 20)


def _extract_pumping_pressure_from_raw_v56(raw: pd.DataFrame, name: str, sheet_name: str,
                                            default_well: str) -> pd.DataFrame:
    """Extract hidden/far-right Pumping Pressure from an already-read raw sheet.

    v52 re-opened and rescanned the XLSX XML after the main parse. On files with
    inflated used ranges that doubled the load time. v56 reuses the same raw
    dataframe, so Pumping Pressure rescue adds only a small header scan.
    """
    if raw is None or raw.empty:
        return pd.DataFrame()
    scan_rows = min(50, len(raw))
    candidates = []
    for r in range(scan_rows):
        for c in range(raw.shape[1]):
            val = raw.iat[r, c]
            if val is None or not re.search(r"pump|pumping|circulation", clean_header(val), flags=re.I):
                continue
            nearby = []
            for rr in range(max(0, r - 2), min(len(raw), r + 3)):
                for cc in range(max(0, c - 1), min(raw.shape[1], c + 2)):
                    nearby.append(str(raw.iat[rr, cc] or ""))
            if _looks_like_pumping_pressure_header_v51(val, " ".join(nearby)):
                candidates.append((r, c))
    if not candidates:
        return pd.DataFrame()

    well = _clean_merge_well_v51(default_well or _well_from_sheet_or_file_v52(sheet_name, name))
    rows = []
    for header_r, c in candidates:
        for start_r in (header_r + 1, header_r + 2, header_r + 3):
            if start_r >= len(raw):
                continue
            idx = raw.index[start_r:]
            vals = _to_number_series_v51(pd.Series(raw.iloc[start_r:, c], index=idx))
            if vals.notna().sum() < 2:
                continue
            dt = _date_time_from_raw_v51(raw, idx)
            dt = pd.to_datetime(dt, errors="coerce")
            valid = vals.notna() & dt.notna()
            if valid.sum() < 2:
                continue
            for row_idx in vals[valid].index:
                dtv = pd.Timestamp(dt.loc[row_idx])
                rows.append({
                    "source": name,
                    "sheet": str(sheet_name),
                    "well": well,
                    "datetime": dtv,
                    "date": dtv.normalize(),
                    "time_text": dtv.strftime("%H:%M"),
                    "pumping_pressure_psi": float(vals.loc[row_idx]),
                    "source_type": "excel_raw_reused_pumping_v56",
                })
            break
    if not rows:
        return pd.DataFrame()
    out = pd.DataFrame(rows)
    out["well_key_v51"] = out["well"].map(_clean_merge_well_v51)
    return out.drop_duplicates(subset=["well_key_v51", "datetime"], keep="last")


def _parse_raw_sheet_fast_v56(raw: pd.DataFrame, source_name: str, sheet_name: str,
                              default_well: str) -> pd.DataFrame:
    """Use one strong parse for familiar TMU sheets; run all fallbacks only if needed."""
    try:
        primary = standardize_dataframe(
            table_from_raw(raw), source_name=source_name, sheet_name=sheet_name,
            default_well=default_well,
        )
    except Exception:
        primary = pd.DataFrame()

    if primary is not None and not primary.empty and is_valid_timeseries(primary):
        canonical = [
            c for c in available_numeric_columns(primary)
            if not str(c).startswith("raw__")
        ]
        valid_dt = pd.to_datetime(primary.get("datetime"), errors="coerce").notna().sum() if "datetime" in primary.columns else 0
        if len(canonical) >= 4 and valid_dt >= 2:
            return primary

    candidates = parse_excel_sheet_attempts(
        raw, source_name=source_name, sheet_name=sheet_name, default_well=default_well
    )
    return candidates[0].copy() if candidates else pd.DataFrame()


def _load_xlsx_fast_v56(data: bytes, name: str):
    tables = []
    pump_frames = []
    try:
        with _zipfile_v53.ZipFile(io.BytesIO(data)) as zf:
            shared = _read_shared_strings_fast_v53(zf)
            for sheet_name, sheet_path in _workbook_sheet_paths_v52(zf):
                raw = _raw_dataframe_from_sheet_xml_v53(zf, sheet_path, shared)
                if raw is None or raw.empty or _sheet_is_helper_v56(sheet_name, raw):
                    continue
                default_well = extract_well_from_raw(raw, source_name=name, sheet_name=str(sheet_name))
                pump = _extract_pumping_pressure_from_raw_v56(
                    raw, name=name, sheet_name=str(sheet_name), default_well=default_well
                )
                if pump is not None and not pump.empty:
                    pump_frames.append(pump)
                best = _parse_raw_sheet_fast_v56(
                    raw, source_name=name, sheet_name=str(sheet_name), default_well=default_well
                )
                if best is not None and not best.empty:
                    best = best.copy()
                    best["source_type"] = "excel_fast_xml_v56"
                    tables.append(best)
    except Exception as exc:
        raise RuntimeError(f"Safe XML parser could not read {name}: {exc}")

    tables = filter_usable_tables(filter_preferred_tables(tables))
    pump_df = pd.concat(pump_frames, ignore_index=True, sort=False) if pump_frames else pd.DataFrame()
    if not pump_df.empty:
        pump_df["datetime"] = pd.to_datetime(pump_df["datetime"], errors="coerce")
        pump_df["pumping_pressure_psi"] = pd.to_numeric(pump_df["pumping_pressure_psi"], errors="coerce")
        pump_df = pump_df.dropna(subset=["datetime", "pumping_pressure_psi"])
        pump_df = pump_df.drop_duplicates(subset=["well_key_v51", "datetime"], keep="last")
    return tables, pump_df


_load_tabular_file_v55_final_before_v56 = load_tabular_file


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).split(".")[-1].lower()
    if suffix == "xlsx":
        tables, pump_df = _load_xlsx_fast_v56(data, name)
        if tables:
            if pump_df is not None and not pump_df.empty:
                tables = force_restore_pumping_pressure_v52(tables, pump_df)
            return [
                ensure_pumping_pressure_column_v48(t)
                for t in tables if t is not None and not t.empty
            ]
        # Suspicious workbooks must never fall through to a memory-heavy reader.
        if xlsx_preflight_v53(data).get("suspicious"):
            return []
    return _load_tabular_file_v55_final_before_v56(
        UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images
    )

# =============================================================================
# v61 - simple device-export Excel support
# =============================================================================
PARSER_BUILD_ID_V61 = "v61-device-export-simple-timeseries-20260622"

DEVICE_TIMESERIES_FIELDS_V61 = [
    "pump_intake_pressure_psi", "pump_discharge_pressure_psi",
    "motor_current_amp", "motor_ama_amp", "pump_freq_hz", "drive_freq_hz",
    "intake_temp_c", "intake_temp_f", "motor_temp_c", "motor_temp_f",
    "motor_load_pct", "vibration_x", "vibration_y", "vibration_z",
]


def _simple_header_v61(value: object) -> str:
    """Normalize short device headers such as RUN FREQ., Pi and Tm."""
    text = safe_text(value).lower().replace("&", " and ")
    text = re.sub(r"[^a-z0-9%/]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _device_field_from_header_v61(header: object, values: Optional[pd.Series] = None) -> Optional[str]:
    h = _simple_header_v61(header)
    compact = h.replace(" ", "")
    if not h:
        return None

    # Files exported from this dashboard prefix unmapped/device channels with
    # "Raw:".  Treat that as metadata, not part of the engineering parameter
    # name, so a downloaded CSV can be uploaded again without losing channels.
    h_core = re.sub(r"^(?:raw|source|original|detected)\s+", "", h).strip()
    compact_core = h_core.replace(" ", "")
    # Remove common unit suffixes for semantic matching.
    compact_no_unit = re.sub(r"(?:psi|psig|bar|amp|amps|a|hz|degf|degc|fahrenheit|celsius)+$", "", compact_core)

    if re.search(r"\bmotor\s+current\b", h_core) or compact_no_unit in {"motorcurrent", "runcurrent"}:
        return "motor_current_amp"
    if re.search(r"\b(?:pump\s+)?intake\s+pressure\b", h_core) or compact_no_unit in {"pi", "pint", "pintake", "intakepressure", "pumpintakepressure"}:
        return "pump_intake_pressure_psi"
    if re.search(r"\b(?:pump\s+)?discharge\s+pressure\b", h_core) or compact_no_unit in {"pd", "pdis", "pdischarge", "dischargepressure", "pumpdischargepressure"}:
        return "pump_discharge_pressure_psi"
    if re.search(r"\b(?:run|pump|operating)\s+freq(?:uency)?\b", h_core):
        return "pump_freq_hz"
    if re.search(r"\bmotor\s+temp(?:erature)?\b", h_core):
        if re.search(r"(?:°?f|fahrenheit)\b", h_core):
            return "motor_temp_f"
        if re.search(r"(?:°?c|celsius)\b", h_core):
            return "motor_temp_c"
        nums = pd.to_numeric(values, errors="coerce").dropna() if values is not None else pd.Series(dtype=float)
        return "motor_temp_f" if (not nums.empty and float(nums.median()) > 200.0) else "motor_temp_c"
    if h in {"date", "reading date", "sample date"}:
        return "date"
    if h in {"time", "reading time", "sample time"}:
        return "time"
    if h in {"datetime", "date time", "timestamp", "time stamp"}:
        return "datetime"
    if h in {"well", "well name", "well no", "well number"}:
        return "well"

    # ESP/device-export abbreviations.
    if compact in {"ama", "motorama"}:
        return "motor_ama_amp"
    if compact in {"pi", "p/i", "pint", "pintake", "intakepressure", "pumpintakepressure"}:
        return "pump_intake_pressure_psi"
    if compact in {"pd", "p/d", "pdis", "pdischarge", "dischargepressure", "pumpdischargepressure"}:
        return "pump_discharge_pressure_psi"
    if compact in {"runfreq", "runfrequency", "pumpfreq", "pumpfrequency", "freq", "frequency", "operatingfreq", "operatingfrequency"}:
        return "pump_freq_hz"
    if compact in {"drivefreq", "drivefrequency", "vsdfreq", "vfdfreq", "vfdfrequency"}:
        return "drive_freq_hz"
    if compact in {"amp", "amps", "current", "motorcurrent", "runcurrent"}:
        return "motor_current_amp"
    if compact in {"vx", "vibx", "vibrationx"}:
        return "vibration_x"
    if compact in {"vy", "viby", "vibrationy"}:
        return "vibration_y"
    if compact in {"vz", "vibz", "vibrationz"}:
        return "vibration_z"
    if compact in {"motorload", "loadpct", "load%"}:
        return "motor_load_pct"

    # Tm/Ti are frequently exported without a unit. Infer only when the header
    # does not prove the unit. Values above 200 are operationally much more
    # plausible as degF than degC for ESP motor/intake temperature.
    if compact in {"tm", "motortemp", "motortemperature", "motorwindingtemp"}:
        if " f" in f" {h}" or "fahrenheit" in h:
            return "motor_temp_f"
        if " c" in f" {h}" or "celsius" in h:
            return "motor_temp_c"
        nums = pd.to_numeric(values, errors="coerce").dropna() if values is not None else pd.Series(dtype=float)
        return "motor_temp_f" if (not nums.empty and float(nums.median()) > 200.0) else "motor_temp_c"
    if compact in {"ti", "intaketemp", "intaketemperature", "pumpintaketemp"}:
        if " f" in f" {h}" or "fahrenheit" in h:
            return "intake_temp_f"
        if " c" in f" {h}" or "celsius" in h:
            return "intake_temp_c"
        nums = pd.to_numeric(values, errors="coerce").dropna() if values is not None else pd.Series(dtype=float)
        return "intake_temp_f" if (not nums.empty and float(nums.median()) > 200.0) else "intake_temp_c"
    return None


def _device_well_from_filename_v61(source_name: str, default_well: Optional[str] = None) -> str:
    """Prefer the real well token and ignore date-range text in export names."""
    stem = Path(str(source_name or "")).stem
    before_export = re.split(r"(?i)device[_ -]*export|report20\d{2}", stem, maxsplit=1)[0]
    # Example: BED_16 C6-9_device_export... -> C6-9
    matches = re.findall(r"\b([A-Za-z]{1,4}\d+[A-Za-z]*\s*[-_]\s*\d+[A-Za-z]?)\b", before_export)
    if matches:
        return re.sub(r"\s*[-_]\s*", "-", matches[-1]).upper()
    if default_well:
        d = str(default_well).strip()
        # Reject false guesses formed from report timestamps, e.g. TM2026-06.
        if d and not re.search(r"(?i)(?:19|20)\d{2}[-_/]\d{1,2}", d):
            return d
    guessed = guess_well_from_name(before_export)
    return guessed or "Unknown"


def _find_device_header_row_v61(raw: pd.DataFrame) -> Optional[int]:
    if raw is None or raw.empty:
        return None
    best = None
    best_score = -1
    for r in range(min(20, len(raw))):
        headers = [_simple_header_v61(v) for v in raw.iloc[r].tolist()]
        has_date = any(h in {"date", "reading date", "sample date", "datetime", "date time", "timestamp"} for h in headers)
        has_time = any(h in {"time", "reading time", "sample time", "datetime", "date time", "timestamp"} for h in headers)
        device_hits = sum(
            1 for c, h in enumerate(headers)
            if _device_field_from_header_v61(h, raw.iloc[r + 1:, c] if r + 1 < len(raw) else None)
            in DEVICE_TIMESERIES_FIELDS_V61
        )
        score = (4 if has_date else 0) + (4 if has_time else 0) + device_hits * 2
        if has_date and has_time and device_hits >= 2 and score > best_score:
            best, best_score = r, score
    return best


def _parse_simple_device_export_v61(raw: pd.DataFrame, source_name: str, sheet_name: str,
                                    default_well: Optional[str] = None) -> pd.DataFrame:
    """Parse compact Date/Time + ESP device-export worksheets.

    Supports real Excel date/time cells, strings, and mixed time cells where
    midnight values are represented as 1900-01-01 datetimes.
    """
    header_row = _find_device_header_row_v61(raw)
    if header_row is None:
        return pd.DataFrame()

    work = raw.iloc[header_row:].copy().dropna(how="all").dropna(axis=1, how="all")
    if len(work) < 3:
        return pd.DataFrame()
    headers = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(work.iloc[0].tolist())])
    data = work.iloc[1:].copy()
    data.columns = headers[: data.shape[1]]
    data = data.dropna(how="all")
    if data.empty:
        return pd.DataFrame()

    out = pd.DataFrame(index=data.index)
    out["source"] = source_name
    out["sheet"] = sheet_name
    out["well"] = _device_well_from_filename_v61(source_name, default_well)
    out["source_type"] = "excel_device_export_v61"

    date_col = time_col = datetime_col = None
    used = set()
    for col in data.columns:
        target = _device_field_from_header_v61(col, data[col])
        if target == "date":
            date_col = col
            used.add(col)
        elif target == "time":
            time_col = col
            used.add(col)
        elif target == "datetime":
            datetime_col = col
            used.add(col)
        elif target in DEVICE_TIMESERIES_FIELDS_V61:
            vals = data[col].map(extract_number).astype(float)
            if vals.notna().sum() >= 2:
                out[target] = vals
                used.add(col)

    # Preserve additional numeric device channels so the mapping panel can teach
    # their names later instead of rejecting the whole file.
    for col in data.columns:
        if col in used:
            continue
        vals = data[col].map(extract_number).astype(float)
        if vals.notna().sum() >= 2:
            out[raw_output_column_name(col, out.columns)] = vals

    parsed_time = None
    if datetime_col is not None:
        out["datetime"] = parse_datetime_series(data[datetime_col])
    else:
        dates = parse_date_series(data[date_col]) if date_col is not None else None
        parsed_time = parse_time_series(data[time_col]) if time_col is not None else None
        out["date"] = dates if dates is not None else pd.NaT
        out["datetime"] = combine_date_time(dates, data[time_col] if time_col is not None else None, None)
        # Some device exports leave the previous date on the 00:00 row and only
        # advance the Date column at 00:30. Detect that midnight rollover before
        # sorting so the sequence remains 21-Jun 23:30, 22-Jun 00:00, 00:30.
        if parsed_time is not None:
            out = adjust_datetime_rollover(out, parsed_time)

    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    valid_year = out["datetime"].dt.year.between(1900, 2100)
    out.loc[~valid_year.fillna(False), "datetime"] = pd.NaT
    out = out.loc[out["datetime"].notna()].copy()
    if out.empty:
        return pd.DataFrame()
    out["date"] = out["datetime"].dt.floor("D")
    out["time_text"] = out["datetime"].dt.strftime("%H:%M")

    numeric_cols = [c for c in out.columns if c in DEVICE_TIMESERIES_FIELDS_V61 or str(c).startswith("raw__")]
    useful = out[numeric_cols].notna().any(axis=1) if numeric_cols else pd.Series(False, index=out.index)
    out = out.loc[useful].copy()
    if out.empty:
        return pd.DataFrame()
    return out.drop_duplicates(subset=["well", "datetime"], keep="last").sort_values("datetime").reset_index(drop=True)


def is_usable_device_timeseries_v61(df: pd.DataFrame) -> bool:
    if df is None or df.empty or "datetime" not in df.columns:
        return False
    valid_dt = pd.to_datetime(df["datetime"], errors="coerce").notna().sum()
    if valid_dt < 2:
        return False
    recognized = 0
    for field in DEVICE_TIMESERIES_FIELDS_V61:
        if field in df.columns and pd.to_numeric(df[field], errors="coerce").notna().sum() >= 2:
            recognized += 1
    if recognized >= 2:
        return True
    # Safe generic fallback: only device-export rows with at least two numeric raw
    # channels are accepted. This avoids admitting random workbook summary sheets.
    source_type = df.get("source_type", pd.Series("", index=df.index)).astype(str)
    raw_cols = [c for c in df.columns if str(c).startswith("raw__") and pd.to_numeric(df[c], errors="coerce").notna().sum() >= 2]
    return source_type.str.contains("device_export", case=False, na=False).any() and len(raw_cols) >= 2


# Teach the normal alias engine about punctuation variants such as RUN FREQ.
_best_canonical_name_before_v61 = best_canonical_name

def best_canonical_name(column_name: str) -> Optional[str]:
    target = _device_field_from_header_v61(column_name)
    if target in DEVICE_TIMESERIES_FIELDS_V61:
        return target
    return _best_canonical_name_before_v61(column_name)


# Try the dedicated parser before the broader multi-template parser.
_parse_raw_sheet_fast_before_v61 = _parse_raw_sheet_fast_v56

def _parse_raw_sheet_fast_v56(raw: pd.DataFrame, source_name: str, sheet_name: str,
                              default_well: str) -> pd.DataFrame:
    device = _parse_simple_device_export_v61(
        raw, source_name=source_name, sheet_name=sheet_name, default_well=default_well
    )
    if is_usable_device_timeseries_v61(device):
        return device
    return _parse_raw_sheet_fast_before_v61(
        raw, source_name=source_name, sheet_name=sheet_name, default_well=default_well
    )


# The v46 production-only safety filter incorrectly removed ESP-only history
# tables. Keep its protection for normal files, but explicitly accept validated
# device-export time series.
_filter_usable_tables_before_v61 = filter_usable_tables

def filter_usable_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    accepted: List[pd.DataFrame] = []
    normal: List[pd.DataFrame] = []
    for table in tables or []:
        if table is None or table.empty:
            continue
        if is_usable_device_timeseries_v61(table):
            tt = table.copy()
            dt = pd.to_datetime(tt["datetime"], errors="coerce")
            useful = pd.Series(False, index=tt.index)
            for field in DEVICE_TIMESERIES_FIELDS_V61:
                if field in tt.columns:
                    useful |= pd.to_numeric(tt[field], errors="coerce").notna()
            for col in [c for c in tt.columns if str(c).startswith("raw__")]:
                useful |= pd.to_numeric(tt[col], errors="coerce").notna()
            tt = tt.loc[dt.notna() & useful].copy()
            if not tt.empty:
                accepted.append(tt)
        else:
            normal.append(table)
    accepted.extend(_filter_usable_tables_before_v61(normal))
    return accepted


# =============================================================================
# v62 - robust CSV encoding fallback
# =============================================================================
PARSER_BUILD_ID_V62 = "v62-robust-csv-encoding-20260622"
PARSER_BUILD_ID_V63 = "v63-merge-device-test-well-aliases-20260622"
PARSER_BUILD_ID_V64 = "v64-scientific-notation-numeric-fix-20260622"


def _decode_csv_bytes_v62(data: bytes) -> tuple[str, str]:
    """Decode CSVs saved by Excel in UTF-8, UTF-8 BOM, Windows-1252 or Latin-1.

    Excel may save the degree symbol as byte 0xB0 in Windows-1252. Reading that
    file as UTF-8 raises UnicodeDecodeError. We try strict encodings in a safe
    order and return normalized Unicode text.
    """
    if data is None:
        return "", "utf-8"
    if not isinstance(data, (bytes, bytearray)):
        data = bytes(data)
    raw = bytes(data)
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            text = raw.decode(encoding, errors="strict")
            # Remove embedded NULs that sometimes appear in UTF-16-like exports.
            text = text.replace("\x00", "")
            return text, encoding
        except UnicodeDecodeError:
            continue
    # Latin-1 can decode every byte, so this is only a defensive fallback.
    return raw.decode("latin-1", errors="replace").replace("\x00", ""), "latin-1"


def _csv_delimiter_v62(text: str) -> str:
    """Choose a likely delimiter without depending on locale-specific Excel defaults."""
    sample_lines = [line for line in str(text).splitlines()[:10] if line.strip()]
    sample = "\n".join(sample_lines)
    if not sample:
        return ","
    candidates = [",", ";", "\t", "|"]
    counts = {sep: sum(line.count(sep) for line in sample_lines) for sep in candidates}
    return max(candidates, key=lambda sep: counts[sep]) if max(counts.values()) > 0 else ","


def _load_csv_robust_v62(data: bytes, name: str) -> List[pd.DataFrame]:
    text, _encoding = _decode_csv_bytes_v62(data)
    if not text.strip():
        return []
    sep = _csv_delimiter_v62(text)
    tables: List[pd.DataFrame] = []

    # Try header-based parsing first. Keep all fields as flexible objects until
    # the application's own standardization logic converts them.
    try:
        header_df = pd.read_csv(io.StringIO(text), sep=sep, dtype=object)
        header_df = header_df.loc[:, ~header_df.columns.astype(str).str.match(r"^Unnamed")]
        if not header_df.empty:
            default_well = extract_well_from_raw(header_df, source_name=name, sheet_name="CSV")

            # Dedicated device parser accepts DateTime + ESP/device channels.
            raw_device = pd.concat(
                [pd.DataFrame([list(header_df.columns)]), header_df.reset_index(drop=True)],
                ignore_index=True,
            )
            device = _parse_simple_device_export_v61(
                raw_device, source_name=name, sheet_name="CSV", default_well=default_well
            )
            if is_usable_device_timeseries_v61(device):
                tables.append(device)

            # Standard production-test parser remains available for normal CSVs.
            std = standardize_dataframe(
                header_df, source_name=name, sheet_name="CSV", default_well=default_well
            )
            if is_valid_timeseries(std):
                tables.append(std)
    except Exception:
        pass

    if not tables:
        try:
            raw = pd.read_csv(io.StringIO(text), sep=sep, header=None, dtype=object)
            default_well = extract_well_from_raw(raw, source_name=name, sheet_name="CSV")
            device = _parse_simple_device_export_v61(
                raw, source_name=name, sheet_name="CSV", default_well=default_well
            )
            if is_usable_device_timeseries_v61(device):
                tables.append(device)
            table = table_from_raw(raw)
            std = standardize_dataframe(
                table, source_name=name, sheet_name="CSV", default_well=default_well
            )
            if is_valid_timeseries(std):
                tables.append(std)
        except Exception:
            pass

    # Avoid returning the same interpretation twice.
    unique: List[pd.DataFrame] = []
    signatures = set()
    for table in tables:
        if table is None or table.empty:
            continue
        dt = pd.to_datetime(table.get("datetime"), errors="coerce") if "datetime" in table.columns else pd.Series(dtype="datetime64[ns]")
        sig = (
            str(table.get("well", pd.Series([""])).iloc[0]) if len(table) else "",
            len(table),
            str(dt.min()) if not dt.empty else "",
            str(dt.max()) if not dt.empty else "",
            tuple(sorted(str(c) for c in table.columns)),
        )
        if sig not in signatures:
            signatures.add(sig)
            unique.append(table)
    return filter_usable_tables(unique)


_load_tabular_file_before_v62 = load_tabular_file


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 1000) -> List[pd.DataFrame]:
    data, name = _uploaded_bytes_v49(uploaded_file)
    suffix = str(name).rsplit(".", 1)[-1].lower() if "." in str(name) else ""
    if suffix == "csv":
        return _load_csv_robust_v62(data, name)
    return _load_tabular_file_before_v62(
        UploadedBytes(data, name), parse_images=parse_images, max_ocr_images=max_ocr_images
    )


# -----------------------------------------------------------------------------
# v65: robust pasted WhatsApp multi-report splitting + qualitative gas status
# -----------------------------------------------------------------------------
PARSER_BUILD_ID_V65 = "v65-whatsapp-multi-report-low-gas-20260622"


def _clean_pasted_whatsapp_text_v65(text: object) -> str:
    """Normalize pasted WhatsApp formatting without joining separate reports."""
    s = str(text or "")
    s = s.replace("\u200e", "").replace("\u200f", "").replace("\ufeff", "")
    s = s.replace("\xa0", " ")
    # Remove WhatsApp markdown characters only. Newlines are deliberately kept
    # because every TMU report is line-oriented.
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"\r\n?", "\n", s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n[ \t]+", "\n", s)
    return s.strip()


def split_messages(text: str) -> List[str]:
    """Split pasted text into complete TMU reports, one row per report.

    The older zero-width lookahead matched both ``PICO TMU-04`` and the nested
    ``TMU-04`` substring. That produced extra one-word ``PICO`` chunks and could
    behave inconsistently when several reports were pasted together. This
    version uses non-overlapping full-line report headers.
    """
    s = _clean_pasted_whatsapp_text_v65(text)
    if not s:
        return []

    header_re = re.compile(
        r"(?im)^\s*(?:PICO\s*T\s*MU|PICO\s*TMU|TMU)\s*[- ]?\s*\d+\b[^\n]*"
    )
    starts = [m.start() for m in header_re.finditer(s)]

    # Some reports omit the TMU header but repeat a Date line. Use Date as a
    # fallback only when at least two report starts are present.
    if not starts:
        date_re = re.compile(r"(?im)^\s*date\s*[:=@-]?\s*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")
        starts = [m.start() for m in date_re.finditer(s)]

    if len(starts) <= 1:
        return [s]

    chunks: List[str] = []
    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(s)
        chunk = s[start:end].strip()
        if chunk:
            chunks.append(chunk)
    return chunks


def _gas_status_from_text_v65(text: str) -> Optional[str]:
    """Return a non-numeric gas-rate status without inventing a gas value."""
    m = re.search(r"(?im)^\s*gas\s*rate\s*[:=@-]?\s*(.+?)\s*$", str(text or ""))
    if not m:
        return None
    raw = re.sub(r"[*_`~]+", "", m.group(1)).strip()
    norm = normalize_text(raw)
    if re.search(r"\blow\s*gas\b", norm):
        return "Low gas"
    if re.search(r"\b(no|zero|nil)\s*gas\b", norm):
        return "No gas"
    if re.search(r"\btrace\s*gas\b", norm):
        return "Trace gas"
    if re.search(r"\bn/?a\b|not\s+available", norm):
        return "Not available"
    return None


_parse_tmu_message_before_v65 = parse_tmu_message


def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    cleaned = _clean_pasted_whatsapp_text_v65(message)
    row = _parse_tmu_message_before_v65(cleaned, source_name=source_name)

    # Preserve qualitative gas descriptions separately. Never convert "Low gas"
    # into zero because zero would look like a measured shutdown value.
    gas_status = _gas_status_from_text_v65(cleaned)
    if gas_status:
        row["gas_rate_status"] = gas_status
        if pd.isna(pd.to_numeric(pd.Series([row.get("gas_rate_mmscfd")]), errors="coerce").iloc[0]):
            row["gas_rate_mmscfd"] = np.nan
        row["note"] = append_note(row.get("note"), f"Gas rate: {gas_status}")
    return row


def parse_many_tmu_messages(text: str, source_name: str = "WhatsApp_Text") -> pd.DataFrame:
    """Parse all pasted TMU reports independently and keep their own timestamps."""
    rows: List[Dict[str, object]] = []
    for message_index, chunk in enumerate(split_messages(text), start=1):
        row = parse_tmu_message(chunk, source_name=source_name)
        numeric_fields = [
            "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi",
            "sep_p_psi", "pumping_pressure_psi", "gas_rate_mmscfd", "bsw_pct",
            "salinity_kppm", "choke_pct", "choke_size_64", "h2s_ppm", "co2_mole_pct",
        ]
        valid_numeric_count = 0
        for field in numeric_fields:
            if field in row:
                value = pd.to_numeric(pd.Series([row.get(field)]), errors="coerce").iloc[0]
                if pd.notna(value):
                    valid_numeric_count += 1
        # A normal pasted report has many numeric values. Also accept a report
        # with a qualitative gas status as long as another real numeric reading exists.
        if valid_numeric_count >= 1 and pd.notna(row.get("datetime", pd.NaT)):
            row["message_index"] = message_index
            rows.append(row)

    if not rows:
        return pd.DataFrame()

    df = safe_object_columns(pd.DataFrame(rows))
    for col in ["source", "sheet", "well", "datetime", "date", "time_text"]:
        if col not in df.columns:
            df[col] = np.nan

    # Remove only truly repeated copies; never collapse different report times.
    subset = [c for c in ["well", "datetime", "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi", "sep_p_psi"] if c in df.columns]
    if subset:
        df = df.drop_duplicates(subset=subset, keep="last")
    return df.sort_values([c for c in ["well", "datetime"] if c in df.columns]).reset_index(drop=True)


_parse_whatsapp_plain_or_export_text_before_v65 = parse_whatsapp_plain_or_export_text


def parse_whatsapp_plain_or_export_text(text: str, source_name="WhatsApp_Text") -> pd.DataFrame:
    """Use export-chat parsing when applicable, otherwise robust pasted parsing."""
    export_df = parse_whatsapp_export_text(text, source_name=source_name)
    if export_df is not None and not export_df.empty:
        return export_df
    df = parse_many_tmu_messages(text, source_name=source_name)
    if df is not None and not df.empty:
        df = safe_object_columns(df)
        df["source_type"] = "pasted_whatsapp_text"
        df["link_status"] = np.where(
            df["well"].apply(clean_well_name_value).ne("Unknown"),
            "text_confirmed_by_well",
            "text_needs_well_review",
        )
    return df


def _repair_measurement_datetime_sequence_v67(out: pd.DataFrame) -> pd.DataFrame:
    """Repair measurement timestamps in original worksheet order.

    Fixes two common field-sheet defects without rewriting valid multi-day data:
    1. one Date cell differs from equal previous/next Date cells;
    2. Date is copied down at midnight while Time resets from evening to morning.
    """
    if out is None or out.empty or "datetime" not in out.columns:
        return out

    result = out.copy()
    if "date" in result.columns:
        dates = pd.to_datetime(result["date"], errors="coerce").dt.normalize().tolist()
    else:
        dates = pd.to_datetime(result["datetime"], errors="coerce").dt.normalize().tolist()

    time_values = []
    if "time" in result.columns:
        parsed = parse_time_series(result["time"])
        time_values = parsed.tolist()
    else:
        parsed_dt = pd.to_datetime(result["datetime"], errors="coerce")
        time_values = [
            (pd.Timestamp("1900-01-01") + pd.Timedelta(
                hours=ts.hour, minutes=ts.minute, seconds=ts.second
            )) if pd.notna(ts) else pd.NaT
            for ts in parsed_dt
        ]

    # Correct an isolated calendar-date typo surrounded by equal dates.
    for pos in range(1, len(dates) - 1):
        prev_d, cur_d, next_d = dates[pos - 1], dates[pos], dates[pos + 1]
        if pd.notna(prev_d) and pd.notna(next_d) and prev_d == next_d:
            if pd.notna(cur_d) and cur_d != prev_d:
                dates[pos] = prev_d

    repaired = []
    previous = pd.NaT
    for d, t in zip(dates, time_values):
        if pd.isna(d) or pd.isna(t):
            repaired.append(pd.NaT)
            continue
        current = pd.Timestamp(d).normalize() + pd.Timedelta(
            hours=t.hour, minutes=t.minute, seconds=t.second
        )
        if pd.notna(previous):
            previous = pd.Timestamp(previous)
            previous_seconds = previous.hour * 3600 + previous.minute * 60 + previous.second
            current_seconds = current.hour * 3600 + current.minute * 60 + current.second
            # Infer rollover only for a substantial clock reset and only when
            # the Date cell failed to advance. Small out-of-order edits remain
            # untouched and explicit Date advances are always trusted.
            if current.date() <= previous.date() and previous_seconds - current_seconds > 6 * 3600:
                while current <= previous:
                    current += pd.Timedelta(days=1)
        repaired.append(current)
        previous = current

    result["datetime"] = pd.to_datetime(repaired, errors="coerce")
    result["date"] = result["datetime"].dt.floor("D")
    result["time_text"] = result["datetime"].dt.strftime("%H:%M")
    return result


# =============================================================================
# v67 - cumulative Excel time / operational-note rollover fix
# =============================================================================
PARSER_BUILD_ID_V67 = "v67-cumulative-excel-time-sequence-fix-20260623"
PARSER_BUILD_ID = PARSER_BUILD_ID_V67


def _measurement_row_mask_v67(out: pd.DataFrame) -> pd.Series:
    """Return rows that contain at least one real engineering measurement.

    Operational notes in TMU sheets often have Date/Time values but no readings.
    Their clock times may be out of chronological order (for example 10:30,
    then a note recorded as 04:00).  Those note rows must never drive midnight
    rollover for the production-test rows that follow.
    """
    mask = pd.Series(False, index=out.index, dtype=bool)
    metadata = set(BASE_NON_PLOT_COLS) | {
        "source_type", "link_status", "review_required", "message_index",
        "source_priority", "source_row", "source_group", "data_quality_note",
        "gas_rate_status", "test_id", "test_sequence",
    }
    for col in out.columns:
        if col in metadata or str(col).startswith("_"):
            continue
        numeric = pd.to_numeric(out[col], errors="coerce")
        if numeric.notna().any():
            mask |= numeric.notna()
    return mask


def adjust_datetime_rollover(out: pd.DataFrame, parsed_time: Optional[pd.Series]) -> pd.DataFrame:
    """Repair midnight rollover using measurement rows only.

    The Date column is authoritative when it advances.  A day is added only
    when consecutive *measurement* rows keep the same Date while clock time
    moves backwards by a meaningful amount.  Note/event rows are ignored and
    cannot shift an entire test by one or more days.
    """
    if out is None or out.empty or parsed_time is None:
        return out
    if "date" not in out.columns or "datetime" not in out.columns:
        return out
    if parsed_time.notna().sum() < 2:
        return out

    result = out.copy()
    dates = pd.to_datetime(result["date"], errors="coerce")
    times = parsed_time.reindex(result.index)
    adjusted = pd.to_datetime(result["datetime"], errors="coerce").copy()
    measurement_rows = _measurement_row_mask_v67(result)

    day_offset = 0
    previous_time = None
    previous_base_date = None

    for idx in result.index:
        d = dates.loc[idx]
        t = times.loc[idx]
        if pd.isna(d) or pd.isna(t):
            continue

        base_date = pd.Timestamp(d).normalize()

        # Preserve a reasonable datetime for note rows, but do not let their
        # non-chronological clock entries affect rollover state.
        if not bool(measurement_rows.loc[idx]):
            adjusted.loc[idx] = base_date + pd.Timedelta(
                hours=t.hour, minutes=t.minute, seconds=t.second
            )
            continue

        # An explicit Date advance already represents the new calendar day.
        if previous_base_date is not None and base_date > previous_base_date:
            day_offset = 0
        elif previous_base_date is not None and base_date < previous_base_date:
            # Do not carry a stale offset into a new/repeated table section.
            day_offset = 0

        current_clock = (t.hour, t.minute, t.second)
        previous_clock = (
            (previous_time.hour, previous_time.minute, previous_time.second)
            if previous_time is not None else None
        )

        # Only infer midnight when Date did not advance and the time drop is
        # large enough to be a real rollover, not a small out-of-order edit.
        if (
            previous_clock is not None
            and previous_base_date is not None
            and base_date == previous_base_date
        ):
            previous_seconds = previous_clock[0] * 3600 + previous_clock[1] * 60 + previous_clock[2]
            current_seconds = current_clock[0] * 3600 + current_clock[1] * 60 + current_clock[2]
            if previous_seconds - current_seconds > 6 * 3600:
                day_offset += 1

        adjusted.loc[idx] = base_date + pd.Timedelta(days=day_offset) + pd.Timedelta(
            hours=t.hour, minutes=t.minute, seconds=t.second
        )
        previous_time = t
        previous_base_date = base_date

    result["datetime"] = pd.to_datetime(adjusted, errors="coerce")
    result["date"] = result["datetime"].dt.floor("D")
    result["time_text"] = result["datetime"].dt.strftime("%H:%M")
    return result


def _combine_date_time_v49(date_series, time_series):
    """Combine Date + Time without double-counting cumulative Excel days.

    Some workbooks store Time as 1.041667, 1.5, 2.020833, etc.  The integer
    portion is an elapsed/copy artefact while the Date column already advances.
    When Date is constant, the integer portion can still represent rollover, so
    it is applied relative to the first numeric time day—not as an absolute day.
    """
    dates = parse_date_series(date_series)
    times = _parse_time_series_v49(time_series)

    raw_day_parts = []
    for raw in list(time_series):
        part = None
        try:
            if isinstance(raw, (int, float, np.number)) and not isinstance(raw, bool):
                value = float(raw)
            elif re.fullmatch(r"[-+]?\d+(?:\.\d+)?", str(raw).strip()):
                value = float(str(raw).strip())
            else:
                value = np.nan
            if np.isfinite(value) and 0 <= value < 100:
                part = int(math.floor(value + 1e-9))
        except Exception:
            part = None
        raw_day_parts.append(part)

    valid_dates = pd.to_datetime(dates, errors="coerce").dropna().dt.normalize()
    date_column_advances = valid_dates.nunique() >= 2
    baseline_part = next((part for part in raw_day_parts if part is not None), 0)

    combined = []
    for d, t, raw_part in zip(dates, times, raw_day_parts):
        if pd.isna(d):
            combined.append(pd.NaT)
            continue
        extra_days = 0
        if not date_column_advances and raw_part is not None:
            extra_days = max(0, raw_part - baseline_part)
        if pd.notna(t):
            combined.append(
                pd.Timestamp(d).normalize()
                + pd.Timedelta(days=extra_days, hours=t.hour, minutes=t.minute, seconds=t.second)
            )
        else:
            combined.append(pd.Timestamp(d).normalize() + pd.Timedelta(days=extra_days))
    return pd.Series(combined, index=date_series.index, dtype="datetime64[ns]")

# v68 package marker. The unified parser is primary; this module remains the
# fallback for OCR, PDF, DOCX, WhatsApp ZIP and rare legacy workbook formats.
PARSER_BUILD_ID_V68 = "v68-corpus-driven-robust-ingestion-20260623"
PARSER_BUILD_ID = PARSER_BUILD_ID_V68

# =============================================================================
# v70 final OCR override: rectified CTU/HMI screen OCR + per-field audit data
# =============================================================================
# This override is intentionally appended at the end of the legacy module so all
# older ZIP/image loaders resolve this implementation at runtime.

CTU_OCR_BUILD_ID_V70 = "v70-rectified-screen-consensus-ocr-20260624"

CTU_SCREEN_SIZE_V70 = (1200, 750)
CTU_VALUE_ROIS_V70 = {
    # x1, y1, x2, y2 on the rectified 1200 x 750 ALL DATA display.
    "ctu_weight_lbf": (160, 92, 665, 192),
    "ctu_lt_weight_lbf": (670, 92, 1170, 192),
    "ctu_wellhead_pressure_psi": (165, 232, 665, 332),
    "ctu_circulation_pressure_psi": (670, 232, 1170, 332),
    "ctu_reel_depth_ft": (165, 357, 665, 458),
    "ctu_reel_speed_ftmin": (670, 357, 1170, 458),
    "ctu_fluid_rate_bpm": (165, 480, 665, 572),
    "ctu_n2_rate_scfm": (670, 480, 1170, 572),
    "ctu_fluid_total_bbl": (165, 602, 665, 697),
    "ctu_n2_total_scf": (670, 602, 1170, 697),
}

CTU_FIELD_RANGES_V70 = {
    "ctu_weight_lbf": (0.0, 250000.0),
    "ctu_lt_weight_lbf": (-5000.0, 5000.0),
    "ctu_wellhead_pressure_psi": (0.0, 10000.0),
    "ctu_circulation_pressure_psi": (0.0, 10000.0),
    "ctu_reel_depth_ft": (0.0, 60000.0),
    "ctu_reel_speed_ftmin": (-2000.0, 2000.0),
    "ctu_fluid_rate_bpm": (0.0, 500.0),
    "ctu_n2_rate_scfm": (0.0, 10000.0),
    "ctu_fluid_total_bbl": (0.0, 1000000.0),
    "ctu_n2_total_scf": (0.0, 2000000000.0),
}

CTU_INTEGER_FIELDS_V70 = {
    "ctu_weight_lbf", "ctu_lt_weight_lbf", "ctu_n2_rate_scfm", "ctu_n2_total_scf"
}

CTU_TARGET_ROIS_V70 = {
    "ctu_weight_lbf": (280, 100, 620, 190),
    "ctu_lt_weight_lbf": (760, 100, 1140, 190),
    "ctu_wellhead_pressure_psi": (280, 240, 620, 330),
    "ctu_circulation_pressure_psi": (760, 240, 1140, 330),
    "ctu_reel_depth_ft": (280, 365, 620, 455),
    "ctu_reel_speed_ftmin": (760, 365, 1140, 455),
    "ctu_fluid_rate_bpm": (280, 490, 620, 570),
    "ctu_n2_rate_scfm": (760, 490, 1140, 570),
    "ctu_fluid_total_bbl": (280, 610, 620, 690),
    "ctu_n2_total_scf": (760, 610, 1140, 690),
}


def _order_quad_points_v70(points):
    import numpy as _np
    pts = _np.asarray(points, dtype="float32").reshape(4, 2)
    total = pts.sum(axis=1)
    diff = _np.diff(pts, axis=1).reshape(-1)
    return _np.array([
        pts[_np.argmin(total)],   # top-left
        pts[_np.argmin(diff)],    # top-right
        pts[_np.argmax(total)],   # bottom-right
        pts[_np.argmax(diff)],    # bottom-left
    ], dtype="float32")


def _rectify_ctu_screen_v70(pil_image):
    """Find the HMI display rectangle and warp it to a stable 1200 x 750 canvas.

    The old OCR used ROIs relative to the full camera photo. That failed whenever
    the HMI occupied a different part of the photo. v70 first detects the display
    quadrilateral, then applies the fixed value-box ROIs to the rectified screen.
    """
    try:
        import cv2
        import numpy as _np
    except Exception:
        return pil_image, {"screen_rectified": False, "screen_detection_score": 0.0}

    rgb = _np.array(pil_image.convert("RGB"))
    bgr = cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)
    height, width = bgr.shape[:2]
    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(gray, 30, 120)
    edges = cv2.morphologyEx(
        edges, cv2.MORPH_CLOSE, _np.ones((9, 9), _np.uint8), iterations=2
    )
    contours, _ = cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)

    candidates = []
    image_area = float(max(1, width * height))
    for contour in contours:
        area_ratio = cv2.contourArea(contour) / image_area
        if not 0.14 <= area_ratio <= 0.85:
            continue
        x, y, w, h = cv2.boundingRect(contour)
        aspect = w / max(h, 1)
        if not 1.38 <= aspect <= 1.88:
            continue
        perimeter = cv2.arcLength(contour, True)
        quad = None
        for epsilon in (0.01, 0.015, 0.02, 0.025, 0.03, 0.04, 0.05):
            approx = cv2.approxPolyDP(contour, epsilon * perimeter, True)
            if len(approx) == 4:
                quad = approx.reshape(4, 2)
                break
        if quad is None:
            continue
        touches_edge = (
            x <= 4 or y <= 4 or x + w >= width - 4 or y + h >= height - 4
        )
        # ALL DATA displays are close to 1.60:1. The area target is deliberately
        # soft because some photos contain the inner display and others the bezel.
        score = abs(aspect - 1.60) + 0.10 * abs(area_ratio - 0.35)
        if touches_edge:
            score += 0.50
        candidates.append((score, quad, area_ratio, aspect))

    if not candidates:
        # Conservative center crop fallback. It is still safer than applying the
        # value ROIs directly to the entire camera photograph.
        x1, x2 = int(width * 0.08), int(width * 0.92)
        y1, y2 = int(height * 0.08), int(height * 0.92)
        crop = bgr[y1:y2, x1:x2]
        crop = cv2.resize(crop, CTU_SCREEN_SIZE_V70, interpolation=cv2.INTER_CUBIC)
        from PIL import Image as _Image
        return _Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB)), {
            "screen_rectified": False,
            "screen_detection_score": 0.0,
            "screen_detection_method": "center_crop_fallback",
        }

    candidates.sort(key=lambda item: item[0])
    score, quad, area_ratio, aspect = candidates[0]
    ordered = _order_quad_points_v70(quad)
    target_w, target_h = CTU_SCREEN_SIZE_V70
    destination = _np.array(
        [[0, 0], [target_w - 1, 0], [target_w - 1, target_h - 1], [0, target_h - 1]],
        dtype="float32",
    )
    matrix = cv2.getPerspectiveTransform(ordered, destination)
    warped = cv2.warpPerspective(bgr, matrix, (target_w, target_h))
    from PIL import Image as _Image
    return _Image.fromarray(cv2.cvtColor(warped, cv2.COLOR_BGR2RGB)), {
        "screen_rectified": True,
        "screen_detection_score": float(max(0.0, 1.0 - min(score, 1.0))),
        "screen_detection_method": "quadrilateral_perspective",
        "screen_area_ratio": float(area_ratio),
        "screen_aspect_ratio": float(aspect),
    }


def _clean_ocr_numeric_text_v70(text):
    value = str(text or "").strip()
    if not value or not re.search(r"\d", value):
        # Do not turn label letters such as I/l/O into fake numbers. Character
        # repair is used only when the OCR token already contains a real digit.
        return []
    value = value.replace(",", ".")
    value = value.translate(str.maketrans({"O": "0", "o": "0", "I": "1", "l": "1", "|": "1"}))
    # Keep possible punctuation inside the token but remove OCR decoration.
    value = re.sub(r"[^0-9.\-+]", "", value)
    value = re.sub(r"\.{2,}", ".", value)
    matches = re.findall(r"[-+]?\d+(?:\.\d+)?", value)
    output = []
    for match in matches:
        try:
            output.append(float(match))
        except Exception:
            continue
    return output


def _normalise_ctu_candidate_v70(field, value, raw_text=""):
    try:
        number = float(value)
    except Exception:
        return []
    if not np.isfinite(number):
        return []
    low, high = CTU_FIELD_RANGES_V70[field]
    # Most displayed values are nonnegative. Reel speed and line-tension weight
    # can legitimately be negative, so preserve their signs.
    if field not in {"ctu_reel_speed_ftmin", "ctu_lt_weight_lbf"}:
        number = abs(number)

    candidates = [(number, 0.0)]
    # Restore a lost decimal only when OCR returned an integer-like token. When a
    # real decimal point is already present (for example 0.01 reel speed), adding
    # 0.001/0.0001 alternatives causes the consensus median to collapse to zero.
    raw_has_decimal = "." in str(raw_text or "")
    if not raw_has_decimal or not (low <= number <= high):
        for divisor, penalty in ((10.0, 0.35), (100.0, 0.45), (1000.0, 0.65)):
            transformed = number / divisor
            if low <= transformed <= high:
                candidates.append((transformed, penalty))

    result = []
    seen = set()
    for candidate, penalty in candidates:
        if not low <= candidate <= high:
            continue
        rounded_key = round(candidate, 6)
        if rounded_key in seen:
            continue
        seen.add(rounded_key)
        # Avoid unreasonable decimal-restoration alternatives for integer counters.
        if field in CTU_INTEGER_FIELDS_V70 and penalty > 0 and abs(candidate - round(candidate)) > 1e-6:
            continue
        result.append((float(candidate), float(penalty)))
    return result


def _field_for_ocr_box_v70(left, top, width, height):
    center_x = float(left) + float(width) / 2.0
    center_y = float(top) + float(height) / 2.0
    for field, (x1, y1, x2, y2) in CTU_VALUE_ROIS_V70.items():
        if x1 <= center_x <= x2 and y1 <= center_y <= y2:
            return field
    return None


def _full_screen_candidates_v70(rectified_image):
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return {field: [] for field in CTU_VALUE_ROIS_V70}

    bgr = cv2.cvtColor(_np.array(rectified_image.convert("RGB")), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8)).apply(gray)
    median = cv2.medianBlur(gray, 3)
    otsu = cv2.threshold(clahe, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
    variants = {
        "raw": bgr,
        "median": median,
        "otsu": otsu,
    }
    output = {field: [] for field in CTU_VALUE_ROIS_V70}
    for variant_name, image in variants.items():
        try:
            data = pytesseract.image_to_data(
                image,
                config="--oem 3 --psm 11",
                output_type=pytesseract.Output.DATAFRAME,
            )
        except Exception:
            continue
        if data is None or data.empty:
            continue
        data = data.dropna(subset=["text"])
        for _, item in data.iterrows():
            field = _field_for_ocr_box_v70(
                item.get("left", 0), item.get("top", 0), item.get("width", 0), item.get("height", 0)
            )
            if not field:
                continue
            raw_text = str(item.get("text", "") or "")
            confidence = max(0.0, min(100.0, float(item.get("conf", 0) or 0))) / 100.0
            for raw_number in _clean_ocr_numeric_text_v70(raw_text):
                for number, transform_penalty in _normalise_ctu_candidate_v70(field, raw_number, raw_text):
                    output[field].append({
                        "value": number,
                        "confidence": confidence,
                        "raw_text": raw_text,
                        "variant": variant_name,
                        "penalty": transform_penalty,
                        "targeted": False,
                    })
    return output


def _adaptive_joined_screen_candidates_v77(rectified_image):
    """Recover split/glare-obscured digits with one adaptive full-screen pass.

    Tesseract often returns ``693`` and ``99.`` as separate tokens, or ``10149``
    and ``6`` for the reel depth.  Joining the tokens inside each known value
    box and applying the display's fixed decimal precision recovers the full
    engineering value without using a generic OCR guess.
    """
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return {field: [] for field in CTU_VALUE_ROIS_V70}

    rgb = _np.array(rectified_image.convert("RGB"))
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
    adaptive = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 9
    )
    try:
        data = pytesseract.image_to_data(
            adaptive,
            config="--oem 3 --psm 11 -c tessedit_char_whitelist=-0123456789.",
            output_type=pytesseract.Output.DATAFRAME,
        )
    except Exception:
        return {field: [] for field in CTU_VALUE_ROIS_V70}

    output = {field: [] for field in CTU_VALUE_ROIS_V70}
    if data is None or data.empty:
        return output
    tokens = {field: [] for field in CTU_VALUE_ROIS_V70}
    for _, item in data.dropna(subset=["text"]).iterrows():
        field = _field_for_ocr_box_v70(
            item.get("left", 0), item.get("top", 0), item.get("width", 0), item.get("height", 0)
        )
        text = str(item.get("text", "") or "").strip()
        if not field or not re.search(r"\d", text):
            continue
        tokens[field].append((float(item.get("left", 0)), text, float(item.get("conf", 0) or 0)))

    decimal_places = {
        "ctu_circulation_pressure_psi": 2,
        "ctu_reel_depth_ft": 1,
        "ctu_reel_speed_ftmin": 2,
        "ctu_fluid_rate_bpm": 2,
        "ctu_fluid_total_bbl": 1,
    }
    supported = {
        "ctu_weight_lbf", "ctu_circulation_pressure_psi", "ctu_reel_depth_ft",
        "ctu_reel_speed_ftmin", "ctu_fluid_rate_bpm", "ctu_fluid_total_bbl",
        "ctu_n2_total_scf",
    }
    for field in supported:
        parts = sorted(tokens.get(field, []), key=lambda item: item[0])
        if not parts:
            continue
        digits = "".join(re.sub(r"\D", "", part[1]) for part in parts)
        if not digits:
            continue
        sign = -1.0 if any("-" in part[1] for part in parts) and field == "ctu_reel_speed_ftmin" else 1.0
        if field in decimal_places:
            places = decimal_places[field]
            if len(digits) <= places:
                continue
            value = sign * (float(int(digits)) / (10.0 ** places))
            raw_text = " | ".join(part[1] for part in parts) + f" -> {value:g}"
        else:
            value = sign * float(int(digits))
            raw_text = " | ".join(part[1] for part in parts)
        low, high = CTU_FIELD_RANGES_V70[field]
        if not low <= value <= high:
            continue
        confidences = [max(0.0, min(100.0, part[2])) / 100.0 for part in parts]
        confidence = max(0.72, float(_np.mean(confidences)) if confidences else 0.72)
        output[field].append({
            "value": float(value),
            "confidence": confidence,
            "raw_text": raw_text,
            "variant": "adaptive_joined_v77",
            "penalty": 0.0,
            "targeted": True,
            "digit_count": len(digits),
        })
    return output


def _targeted_special_candidates_v77(rectified_image, field):
    """Small field-specific fallbacks for glare, signs and isolated zeros."""
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return []

    rgb = _np.array(rectified_image.convert("RGB"))
    x1, y1, x2, y2 = CTU_VALUE_ROIS_V70[field]
    crop = rgb[max(0, y1):min(rgb.shape[0], y2), max(0, x1):min(rgb.shape[1], x2)]
    if crop.size == 0:
        return []
    output = []

    def _read_candidates(image, variant_prefix, psms=(8, 13), confidence=0.94):
        local = []
        enlarged = cv2.resize(image, None, fx=3.0, fy=3.0, interpolation=cv2.INTER_CUBIC)
        for psm in psms:
            try:
                text = pytesseract.image_to_string(
                    enlarged,
                    config=f"--oem 3 --psm {psm} -c tessedit_char_whitelist=-0123456789.",
                ).strip()
            except Exception:
                continue
            for raw_number in _clean_ocr_numeric_text_v70(text):
                for number, penalty in _normalise_ctu_candidate_v70(field, raw_number, text):
                    local.append({
                        "value": number,
                        "confidence": confidence,
                        "raw_text": text,
                        "variant": f"{variant_prefix}_psm{psm}",
                        "penalty": penalty,
                        "targeted": True,
                        "digit_count": len(re.sub(r"\D", "", text)),
                    })
        return local

    bgr = cv2.cvtColor(crop, cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)

    if field == "ctu_wellhead_pressure_psi":
        blue, green, red = cv2.split(bgr)
        red_score = (red.astype(_np.int16) * 2 - green.astype(_np.int16) - blue.astype(_np.int16)).clip(0)
        red_score = cv2.normalize(red_score, None, 0, 255, cv2.NORM_MINMAX).astype(_np.uint8)
        output.extend(_read_candidates(red_score, "targeted_red_v77", confidence=0.97))

    elif field == "ctu_lt_weight_lbf":
        # The HMI displays a sign on the left and the magnitude on the right.
        # OCR them separately so "- 1" cannot become 7.
        width = crop.shape[1]
        right = gray[:, int(width * 0.35):]
        right = cv2.adaptiveThreshold(right, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 7)
        magnitudes = []
        enlarged_right = cv2.resize(right, None, fx=4.0, fy=4.0, interpolation=cv2.INTER_CUBIC)
        for psm in (8, 13):
            try:
                text = pytesseract.image_to_string(
                    enlarged_right,
                    config=f"--oem 3 --psm {psm} -c tessedit_char_whitelist=0123456789",
                ).strip()
            except Exception:
                continue
            match = re.search(r"\d+", text)
            if match:
                magnitudes.append((int(match.group(0)), text, psm))
        sign_negative = False
        left_red = cv2.split(bgr[:, :max(1, int(width * 0.45))])[2]
        try:
            sign_text = pytesseract.image_to_string(
                cv2.resize(left_red, None, fx=4.0, fy=4.0, interpolation=cv2.INTER_CUBIC),
                config="--oem 3 --psm 8 -c tessedit_char_whitelist=-",
            ).strip()
            sign_negative = "-" in sign_text
        except Exception:
            sign_text = ""
        # Tesseract may discard a standalone minus sign. Confirm it directly as
        # a short horizontal component in the middle-left of the value box,
        # excluding the long bottom border.
        if not sign_negative:
            sign_gray = gray[int(gray.shape[0] * 0.15):int(gray.shape[0] * 0.80), :max(1, int(width * 0.55))]
            for threshold in (160, 180):
                binary = cv2.threshold(sign_gray, threshold, 255, cv2.THRESH_BINARY_INV)[1]
                count, _, stats, _ = cv2.connectedComponentsWithStats(binary, 8)
                for component in range(1, count):
                    _, _, comp_w, comp_h, area = stats[component]
                    aspect = comp_w / max(comp_h, 1)
                    if 8 <= comp_w <= max(18, int(width * 0.18)) and comp_h <= 12 and aspect >= 2.5 and area >= 10:
                        sign_negative = True
                        sign_text = "-"
                        break
                if sign_negative:
                    break
        if magnitudes:
            counts = {}
            for magnitude, _, _ in magnitudes:
                counts[magnitude] = counts.get(magnitude, 0) + 1
            magnitude = max(counts, key=lambda key: (counts[key], key))
            value = -float(magnitude) if sign_negative else float(magnitude)
            output.append({
                "value": value,
                "confidence": 0.98 if counts[magnitude] >= 2 else 0.90,
                "raw_text": f"{sign_text} {magnitude}".strip(),
                "variant": "targeted_lt_signed_v77",
                "penalty": 0.0,
                "targeted": True,
                "digit_count": len(str(magnitude)),
            })

    elif field in {"ctu_n2_rate_scfm", "ctu_n2_total_scf"}:
        width = crop.shape[1]
        right = gray[:, int(width * 0.42):]
        right = cv2.adaptiveThreshold(right, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 7)
        output.extend(_read_candidates(right, "targeted_right_adaptive_v77", confidence=0.95))

    return output


def _refine_leading_counter_digit_v70(processed, text):
    """Repair a single confused leading digit in a long HMI counter.

    Whole-token OCR can confuse the leading 5 in a seven-segment-like counter
    with 9 (for example 547799 -> 947799). We locate the numeric token, crop only
    its first character, and require agreement from at least two independent page
    segmentation modes before replacing that one digit. The rest of the counter
    remains exactly as read by the whole-token pass. If agreement is absent, the
    original text is retained and the field remains review-required.
    """
    try:
        import cv2
        import pytesseract
    except Exception:
        return str(text or ""), False

    original = str(text or "").strip()
    digits = re.sub(r"\D", "", original)
    if len(digits) < 4:
        return original, False

    try:
        data = pytesseract.image_to_data(
            processed,
            config="--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789.",
            output_type=pytesseract.Output.DICT,
        )
    except Exception:
        return original, False

    best = None
    for index, token in enumerate(data.get("text", [])):
        token_text = str(token or "").strip()
        token_digits = re.sub(r"\D", "", token_text)
        if len(token_digits) < 4:
            continue
        candidate = (
            len(token_digits),
            float(data.get("conf", [0])[index] or 0),
            index,
            token_digits,
        )
        if best is None or candidate[:2] > best[:2]:
            best = candidate
    if best is None:
        return original, False

    _, _, index, token_digits = best
    try:
        left = int(data["left"][index])
        top = int(data["top"][index])
        width = int(data["width"][index])
        height = int(data["height"][index])
    except Exception:
        return original, False
    if width <= 0 or height <= 0:
        return original, False

    cell_width = width / max(len(token_digits), 1)
    x1 = max(0, int(left - 0.12 * cell_width))
    x2 = min(processed.shape[1], int(left + 1.18 * cell_width))
    y1 = max(0, int(top - 0.15 * height))
    y2 = min(processed.shape[0], int(top + 1.15 * height))
    first_cell = processed[y1:y2, x1:x2]
    if first_cell.size == 0:
        return original, False

    votes = []
    for psm in (10, 13, 8):
        try:
            read = pytesseract.image_to_string(
                first_cell,
                config=f"--oem 3 --psm {psm} -c tessedit_char_whitelist=0123456789",
            ).strip()
        except Exception:
            continue
        match = re.search(r"\d", read)
        if match:
            votes.append(match.group(0))
    if not votes:
        return original, False

    counts = {digit: votes.count(digit) for digit in set(votes)}
    voted_digit, vote_count = max(counts.items(), key=lambda item: item[1])
    if vote_count < 2 or voted_digit == token_digits[0]:
        return original, False

    refined_digits = voted_digit + token_digits[1:]
    return refined_digits, True


def _targeted_field_candidates_v70(rectified_image, field):
    """One field-specific OCR pass on the rectified value box."""
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return []

    bgr = cv2.cvtColor(_np.array(rectified_image.convert("RGB")), cv2.COLOR_RGB2BGR)
    x1, y1, x2, y2 = CTU_TARGET_ROIS_V70.get(field, CTU_VALUE_ROIS_V70[field])
    crop = bgr[max(0, y1):min(bgr.shape[0], y2), max(0, x1):min(bgr.shape[1], x2)]
    if crop.size == 0:
        return []

    psm = 7
    gray = cv2.cvtColor(crop, cv2.COLOR_BGR2GRAY)
    # Most large screen digits are best read from the unscaled Otsu image. This
    # preserves decimal points and avoids changing 6/8 at high interpolation.
    processed = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

    if field == "ctu_lt_weight_lbf":
        enlarged = cv2.resize(crop, None, fx=2.5, fy=2.5, interpolation=cv2.INTER_CUBIC)
        value_gray = cv2.cvtColor(enlarged, cv2.COLOR_BGR2GRAY)
        value_gray = cv2.createCLAHE(clipLimit=3.5, tileGridSize=(8, 8)).apply(value_gray)
        processed = cv2.threshold(value_gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        psm = 7
    elif field == "ctu_circulation_pressure_psi":
        enlarged = cv2.resize(crop, None, fx=2.5, fy=2.5, interpolation=cv2.INTER_CUBIC)
        blue, green, red = cv2.split(enlarged)
        blue_score = (blue.astype(_np.int16) * 2 - green.astype(_np.int16) - red.astype(_np.int16)).clip(0)
        # Keep the normalized blue-channel response grayscale. Otsu thresholding
        # changed the displayed 1363.61 to 1363.01 on a real field photo by
        # erasing the lower loop of the digit 6. The grayscale response reads all
        # three validation screens correctly (2188.97, 1363.61, 1362.09).
        processed = cv2.normalize(blue_score, None, 0, 255, cv2.NORM_MINMAX).astype(_np.uint8)
        psm = 7
    elif field in {"ctu_weight_lbf", "ctu_n2_rate_scfm"}:
        psm = 6
    elif field == "ctu_n2_total_scf":
        psm = 7

    try:
        text = pytesseract.image_to_string(
            processed,
            config=f"--oem 3 --psm {psm} -c tessedit_char_whitelist=-0123456789.",
        ).strip()
    except Exception:
        return []

    output = []
    if field == "ctu_n2_total_scf" and text:
        refined_text, changed = _refine_leading_counter_digit_v70(processed, text)
        if changed:
            for raw_number in _clean_ocr_numeric_text_v70(refined_text):
                for number, transform_penalty in _normalise_ctu_candidate_v70(field, raw_number, refined_text):
                    output.append({
                        "value": number,
                        "confidence": 0.97,
                        "raw_text": f"{text} -> {refined_text} (leading-digit consensus)",
                        "variant": "targeted_leading_digit_consensus",
                        "penalty": transform_penalty,
                        "targeted": True,
                        "digit_count": len(re.sub(r"\D", "", refined_text)),
                    })

    for raw_number in _clean_ocr_numeric_text_v70(text):
        for number, transform_penalty in _normalise_ctu_candidate_v70(field, raw_number, text):
            digit_count = len(re.sub(r"\D", "", text))
            output.append({
                "value": number,
                "confidence": min(0.94, 0.72 + 0.025 * min(digit_count, 8)) if text else 0.0,
                "raw_text": text,
                "variant": f"targeted_psm{psm}",
                "penalty": transform_penalty,
                "targeted": True,
                "digit_count": digit_count,
            })
    return output


def _values_close_v70(field, left, right):
    left = float(left)
    right = float(right)
    if field in CTU_INTEGER_FIELDS_V70:
        return abs(left - right) <= 0.5
    tolerance = max(0.02, 0.003 * max(abs(left), abs(right), 1.0))
    return abs(left - right) <= tolerance


def _select_ctu_candidate_v70(field, candidates):
    if not candidates:
        return np.nan, 0.0, "", "missing"

    # v77 audited field-specific reads. These are accepted only when the
    # dedicated preprocessing produced a value for the intended box.
    if field == "ctu_lt_weight_lbf":
        signed = [c for c in candidates if str(c.get("variant", "")) == "targeted_lt_signed_v77"]
        if signed:
            best = max(signed, key=lambda item: float(item.get("confidence", 0.0)))
            return (
                float(round(float(best["value"]))),
                float(best.get("confidence", 0.9)),
                str(best.get("raw_text", ""))[:250],
                "accepted" if float(best.get("confidence", 0.0)) >= 0.9 else "review_required",
            )

    if field in {"ctu_circulation_pressure_psi", "ctu_reel_depth_ft"}:
        joined = [c for c in candidates if str(c.get("variant", "")) == "adaptive_joined_v77"]
        if joined:
            best = max(joined, key=lambda item: (float(item.get("confidence", 0.0)), int(item.get("digit_count", 0))))
            value = float(best["value"])
            value = round(value, 2 if field == "ctu_circulation_pressure_psi" else 1)
            return value, min(0.99, float(best.get("confidence", 0.8)) + 0.08), str(best.get("raw_text", ""))[:250], "accepted"

    if field == "ctu_wellhead_pressure_psi":
        red = [c for c in candidates if str(c.get("variant", "")).startswith("targeted_red_v77")]
        if red:
            groups = {}
            for item in red:
                key = round(float(item["value"]), 2)
                groups.setdefault(key, []).append(item)
            value, group = max(groups.items(), key=lambda kv: (len(kv[1]), max(float(x.get("confidence", 0.0)) for x in kv[1])))
            confidence = max(float(x.get("confidence", 0.0)) for x in group)
            raw = " | ".join(sorted({str(x.get("raw_text", "")) for x in group if x.get("raw_text")}))
            return float(value), confidence, raw[:250], "accepted" if len(group) >= 2 else "review_required"

    # A leading-digit correction is accepted only after at least two independent
    # single-character OCR modes agree. Prefer that audited result over a larger
    # whole-token consensus that repeats the same glyph confusion.
    if field == "ctu_n2_total_scf":
        refined = [
            item for item in candidates
            if str(item.get("variant", "")) == "targeted_leading_digit_consensus"
        ]
        if refined:
            best = max(refined, key=lambda item: float(item.get("confidence", 0.0)))
            return (
                float(round(float(best["value"]))),
                min(0.99, float(best.get("confidence", 0.97))),
                str(best.get("raw_text", ""))[:250],
                "review_required",
            )

    groups = []
    for candidate in sorted(candidates, key=lambda item: float(item.get("value", 0.0))):
        placed = False
        for group in groups:
            center = float(np.median([item["value"] for item in group]))
            if _values_close_v70(field, center, candidate["value"]):
                group.append(candidate)
                placed = True
                break
        if not placed:
            groups.append([candidate])

    ranked = []
    for group in groups:
        values = [float(item["value"]) for item in group]
        base_conf = max(float(item.get("confidence", 0.0)) for item in group)
        variants = len({str(item.get("variant", "")) for item in group})
        targeted_bonus = 0.30 if any(item.get("targeted") for item in group) else 0.0
        penalty = min(float(item.get("penalty", 0.0)) for item in group)
        digit_count = max(
            int(item.get("digit_count", len(re.sub(r"\D", "", str(item.get("raw_text", ""))))))
            for item in group
        )
        precision_bonus = 0.0
        if field in {"ctu_wellhead_pressure_psi", "ctu_circulation_pressure_psi", "ctu_reel_depth_ft", "ctu_reel_speed_ftmin", "ctu_fluid_rate_bpm", "ctu_fluid_total_bbl"}:
            if any("." in str(item.get("raw_text", "")) and abs(float(item.get("value", 0.0)) - round(float(item.get("value", 0.0)))) > 1e-6 for item in group):
                precision_bonus = 0.18
        score = base_conf + 0.30 * min(variants, 4) + targeted_bonus + 0.035 * min(digit_count, 8) + precision_bonus - penalty
        value = float(np.median(values))
        if field in CTU_INTEGER_FIELDS_V70 and abs(value - round(value)) < 0.35:
            value = float(round(value))
            score += 0.10
        ranked.append((score, value, group))
    ranked.sort(key=lambda item: item[0], reverse=True)
    best_score, best_value, best_group = ranked[0]
    second_score = ranked[1][0] if len(ranked) > 1 else -1.0
    # A targeted decimal reading is preferred over a truncated integer from the
    # full-screen pass when both values are nearly identical (e.g. 12098.7 vs
    # 12098). This preserves the displayed decimal without inventing a new value.
    if field in {"ctu_wellhead_pressure_psi", "ctu_circulation_pressure_psi", "ctu_reel_depth_ft", "ctu_reel_speed_ftmin", "ctu_fluid_rate_bpm", "ctu_fluid_total_bbl"}:
        precise_targets = [
            item for item in candidates
            if item.get("targeted")
            and "." in str(item.get("raw_text", ""))
            and abs(float(item.get("value", 0.0)) - round(float(item.get("value", 0.0)))) > 1e-6
        ]
        if precise_targets:
            precise = max(precise_targets, key=lambda item: (float(item.get("confidence", 0.0)), len(re.sub(r"\D", "", str(item.get("raw_text", ""))))))
            precise_value = float(precise["value"])
            if abs(precise_value - best_value) <= max(1.0, 0.01 * max(abs(precise_value), abs(best_value), 1.0)):
                best_value = precise_value
                best_group = [precise]
                best_score = max(best_score, float(precise.get("confidence", 0.0)) + 0.75)
    confidence = max(0.0, min(1.0, 0.42 + 0.20 * best_score))
    status = "accepted"
    best_variants = len({str(item.get("variant", "")) for item in best_group})
    if confidence < 0.68 or best_variants < 2 or (second_score > 0 and best_score - second_score < 0.15):
        status = "review_required"
    if field in CTU_INTEGER_FIELDS_V70:
        best_value = float(round(best_value))
    elif field in {"ctu_wellhead_pressure_psi", "ctu_circulation_pressure_psi", "ctu_reel_speed_ftmin", "ctu_fluid_rate_bpm"}:
        best_value = float(round(best_value, 2))
    elif field in {"ctu_reel_depth_ft", "ctu_fluid_total_bbl"}:
        best_value = float(round(best_value, 1))
    raw = " | ".join(sorted({str(item.get("raw_text", "")) for item in best_group if item.get("raw_text")}))
    return best_value, confidence, raw[:250], status


def parse_ctu_all_data_screen_image(uploaded_file, source_name="Image_OCR") -> pd.DataFrame:
    """v70 CTU/HMI OCR for direct images and images inside WhatsApp ZIPs.

    It rectifies the display, combines four full-screen OCR passes, performs
    targeted fallbacks only where required, and exposes per-field confidence/raw
    text so no uncertain OCR number is silently trusted.
    """
    Image, _, _, _ = _try_import_ocr_libs()
    if Image is None:
        return pd.DataFrame()
    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        original = Image.open(uploaded_file).convert("RGB")
        rectified, screen_meta = _rectify_ctu_screen_v70(original)
        file_name = getattr(uploaded_file, "name", source_name)
        dt_from_name = parse_datetime_from_filename(file_name)

        all_candidates = _full_screen_candidates_v70(rectified)
        adaptive_joined = _adaptive_joined_screen_candidates_v77(rectified)
        for field, candidates in adaptive_joined.items():
            all_candidates.setdefault(field, []).extend(candidates)
        # Always reinforce the historically difficult boxes; for all other boxes,
        # use targeted OCR only when the full-screen candidate set is weak.
        difficult = {
            "ctu_lt_weight_lbf", "ctu_circulation_pressure_psi",
            "ctu_reel_speed_ftmin", "ctu_n2_rate_scfm", "ctu_n2_total_scf",
        }
        for field in CTU_VALUE_ROIS_V70:
            current = all_candidates.get(field, [])
            best_full_conf = max([float(item.get("confidence", 0.0)) for item in current] or [0.0])
            if field in difficult or len(current) == 0 or best_full_conf < 0.75:
                all_candidates[field].extend(_targeted_field_candidates_v70(rectified, field))
            all_candidates[field].extend(_targeted_special_candidates_v77(rectified, field))

        row = {
            "source": source_name,
            "sheet": "CTU_Image_OCR",
            "source_type": "ctu_image_ocr_v70",
            "ocr_template": "ctu_all_data_rectified_v70",
            "ocr_build_id": CTU_OCR_BUILD_ID_V70,
            "image_file": file_name,
            "well": "Unknown",
            "link_status": "ocr_manual_link_required",
            "review_required": True,
            **screen_meta,
        }
        if pd.notna(dt_from_name):
            row["datetime"] = dt_from_name
            row["date"] = pd.Timestamp(dt_from_name).floor("D")
            row["time_text"] = pd.Timestamp(dt_from_name).strftime("%H:%M")

        found = 0
        confidence_values = []
        low_confidence_fields = []
        for field in CTU_VALUE_ROIS_V70:
            value, confidence, raw_text, status = _select_ctu_candidate_v70(
                field, all_candidates.get(field, [])
            )
            # High-digit cumulative counters are always explicitly reviewed; a
            # single glyph error can change the value by hundreds of thousands.
            if field == "ctu_n2_total_scf" and pd.notna(value):
                status = "review_required"
            row[f"ocr_raw__{field}"] = raw_text
            row[f"ocr_conf__{field}"] = confidence
            row[f"ocr_status__{field}"] = status
            if pd.notna(value):
                row[field] = float(value)
                found += 1
                confidence_values.append(float(confidence))
                if status != "accepted":
                    low_confidence_fields.append(column_label(field))

        row["ocr_fields_found"] = found
        row["ocr_confidence"] = float(np.mean(confidence_values)) if confidence_values else 0.0
        row["ocr_low_confidence_fields"] = "; ".join(low_confidence_fields)
        row["ocr_status"] = (
            "parsed_review_required" if found >= 3 else "low_confidence_or_not_ctu_screen"
        )
        row["data_quality_note"] = append_note(
            row.get("data_quality_note"),
            "OCR values require field review before engineering use"
            + (f"; low confidence: {', '.join(low_confidence_fields)}" if low_confidence_fields else ""),
        )
        if found < 3:
            return pd.DataFrame()
        return pd.DataFrame([row])
    except Exception as exc:
        return pd.DataFrame([{
            "source": source_name,
            "sheet": "CTU_Image_OCR",
            "source_type": "ctu_image_ocr_v70",
            "image_file": getattr(uploaded_file, "name", source_name),
            "ocr_status": "ocr_failed",
            "data_quality_note": f"OCR failed: {exc}",
            "review_required": True,
            "well": "Unknown",
            "link_status": "ocr_manual_link_required",
        }])

# v70 accepts modern phone image exports as well as JPG/PNG.
try:
    IMAGE_SUFFIXES.add("webp")
except Exception:
    pass
