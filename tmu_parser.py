
import re
import warnings
import io
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import numpy as np
import pandas as pd


KEYWORDS = [
    "date", "time", "well", "choke", "whp", "w.h.p", "wellhead", "sep", "separator",
    "gas", "formation", "gross", "oil", "water", "bsw", "bs&w", "wc", "salinity",
    "h2s", "co2", "pump", "pumping", "pressure", "n2", "nitrogen", "ct", "rate",
]

UNIT_KEYWORDS = [
    "psi", "psig", "psia", "bbl", "bbl/d", "stb", "stb/d", "bpd", "mmscf",
    "mmscf/d", "scf", "ppm", "nacl", "mole", "deg api", "api", "%",
    "c", "f", "hh:mm", "d/mm", "gpm", "bpm",
]

PARENT_HEADERS = {
    "gas", "oil", "water", "gross rate", "production", "liquid rate",
    "liquid metering", "gas metering", "pressure and temperature measurements",
    "mpfm parameters", "fluid properties", "esp parameters", "well head",
    "separator", "oil or cond.", "water reading", "ratio", "tmu",
}

BASE_NON_PLOT_COLS = {
    "source", "sheet", "well", "date", "time", "time_text", "datetime", "note", "test_unit"
}

COLUMN_LABELS = {
    "choke_pct": "Choke (%)",
    "whp_psi": "WHP (psi)",
    "flp_psi": "FLP (psi)",
    "flt_c": "FLT (°C)",
    "sep_p_psi": "Separator Pressure (psi)",
    "gas_temp_c": "Gas Temp (°C)",
    "gas_temp_f": "Gas Temp (°F)",
    "gas_sg": "Gas Specific Gravity",
    "orifice_size_in": "Orifice Size (in)",
    "gas_rate_mmscfd": "Total Gas Rate (MMSCF/D)",
    "gas_formation_mmscfd": "Formation Gas Rate (MMSCF/D)",
    "gor_scf_bbl": "GOR (scf/bbl)",
    "h2s_ppm": "H2S (ppm)",
    "co2_mole_pct": "CO2 (mole %)",
    "oil_temp_c": "Oil Temp (°C)",
    "oil_temp_f": "Oil Temp (°F)",
    "oil_kf": "Oil K.F Factor",
    "oil_api": "Oil Gravity (API)",
    "oil_meter_increment_bbl": "Oil Meter Increment (bbl)",
    "oil_cmf": "Oil CMF",
    "oil_rate_stbd": "Oil Rate (STB/D)",
    "water_rate_bpd": "Water Rate (BBL/D)",
    "bsw_pct": "BS&W (%)",
    "salinity_kppm": "Salinity (K ppm NaCl)",
    "gross_rate_bpd": "Gross Rate (BBL/D)",
    "water_cum_bbl": "Water Cum (bbl)",
    "pumping_pressure_psi": "Pumping Pressure (psi)",
    "n2_rate_scfm": "N2 Rate (scfm)",
    "ct_pressure_psi": "CT Pressure (psi)",
    "ct_depth_m": "CT Depth (m)",
    "ct_running_speed_ftmin": "CT Running Speed (ft/min)",
    "ct_pipe_weight_lbf": "CT Pipe Weight (lbf)",
    "u2_pass_side_pump_pressure_psi": "U2 Pass-side Pump Pressure (psi)",
    "choke_size_64": "Choke Size (64ths)",
    "flow_press_psi": "Flow Pressure (psi)",
    "flow_temp_c": "Flow Temp (°C)",
    "wellhead_temp_c": "Wellhead Temp (°C)",
    "wellhead_temp_f": "Wellhead Temp (°F)",
    "sep_temp_c": "Separator Temp (°C)",
    "sep_temp_f": "Separator Temp (°F)",
    "mpfm_press_psig": "MPFM Pressure (psig)",
    "mpfm_temp_f": "MPFM Temp (°F)",
    "dp_mbar": "DP (mbar)",
    "qoil_s_stbd": "QOil(S) (STB/D)",
    "qwat_s_bpd": "QWat(S) (BBL/D)",
    "qgas_s_mmscfd": "QGas(S) (MMSCF/D)",
    "qoil_a_bpd": "QOil(A) (BBL/D)",
    "qwat_a_bpd": "QWat(A) (BBL/D)",
    "qgas_a_mmcfd": "QGas(A) (MMCF/D)",
    "wlr_s_pct": "WLR(S) (%)",
    "qgross_s_bpd": "QGross(S) (BBL/D)",
    "gor_s_scf_stb": "GOR(S) (SCF/STB)",
    "gvf_a_pct": "GVF(A) (%)",
    "flt_f": "FLT (°F)",
    "flow_temp_f": "Flow Temp (°F)",
    "flow_temp_c": "Flow Temp (°C)",
    "wellhead_temp_c": "Wellhead Temp (°C)",
    "wellhead_temp_f": "Wellhead Temp (°F)",
    "sep_temp_c": "Separator Temp (°C)",
    "sep_temp_f": "Separator Temp (°F)",
    "sep_dp_inh2o": "Separator DP (inH2O)",
    "gas_dp_inh2o": "Gas DP (inH2O)",
    "oil_cum_bbl": "Oil Cum (bbl)",
    "wat_cum_bbl": "Water Cum (bbl)",
    "gas_cum_mscf": "Gas Cum (MSCF)",
    "cgr_bbl_mmscf": "CGR (BBL/MMSCF)",
    "oil_sg": "Oil SG",
    "water_sg": "Water SG",
    "water_ph": "Water pH",
    "pump_freq_hz": "Pump Frequency (Hz)",
    "pump_intake_pressure_psi": "Pi / Intake Pressure (psi)",
    "pump_discharge_pressure_psi": "Pd / Discharge Pressure (psi)",
    "motor_current_amp": "Motor Current (A)",
    "motor_ama_amp": "AMA / Motor Current (A)",
    "intake_temp_c": "Intake Temperature (°C)",
    "intake_temp_f": "Intake Temperature (°F)",
    "motor_temp_c": "Motor Temperature (°C)",
    "motor_temp_f": "Motor Temperature (°F)",
    "motor_load_pct": "Motor Load (%)",
    "vibration_x": "Vibration X",
    "vibration_y": "Vibration Y",
    "vibration_z": "Vibration Z",
    "drive_freq_hz": "Drive Frequency (Hz)",
    "us_press_psi": "Upstream Pressure (psi)",
    "us_temp_c": "Upstream Temp (°C)",
    "ds_press_psi": "Downstream Pressure (psi)",
    "fw_pct": "FW (%)",
    "twc_pct": "TWC (%)",
    "liquid_volume_bbl": "Liquid Volume (bbl)",
    "gor_mmscf_bbl": "GOR (MMSCF/BBL)",
}


def normalize_text(x: object) -> str:
    if pd.isna(x):
        return ""
    s = str(x).strip().lower()
    s = s.replace("\n", " ").replace("\r", " ")
    s = re.sub(r"[\t_]+", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def safe_text(x: object) -> str:
    """Return a safe display/header string for any Excel/PDF cell value."""
    try:
        if pd.isna(x):
            return ""
    except Exception:
        pass
    try:
        s = str(x).strip()
    except Exception:
        s = ""
    if s.lower() in ["nan", "nat", "none"]:
        return ""
    return s


def safe_join(items, sep: str = " ") -> str:
    """Join mixed objects safely; prevents float-in-header join crashes."""
    return sep.join(safe_text(i) for i in items if safe_text(i))


def clean_header(s: object) -> str:
    s = normalize_text(s)
    s = s.replace("&", " and ")
    s = re.sub(r"[^a-z0-9/%.\- =]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def canonical_key(s: object) -> str:
    """Normalize a user/company header into a stable alias key.

    This is used by the Streamlit column-mapping UI.  Example:
    "W.H.P (psig)" and "WHP psig" become comparable keys.
    """
    c = clean_header(s)
    c = c.replace(" and ", " ")
    c = re.sub(r"[^a-z0-9]+", "_", c).strip("_")
    return c


def standard_column_options(include_meta: bool = False) -> Dict[str, str]:
    """Return canonical parser fields for the user-facing mapping UI."""
    labels = dict(COLUMN_LABELS)
    if include_meta:
        labels.update({
            "well": "Well Name",
            "date": "Date",
            "time": "Time",
            "datetime": "Date & Time",
            "note": "Note / Event",
        })
    return dict(sorted(labels.items(), key=lambda kv: kv[1].lower()))


def is_datetime_like(x: object) -> bool:
    import datetime as _dt
    return isinstance(x, (_dt.datetime, _dt.date, _dt.time, pd.Timestamp))


def is_numeric_like(x: object) -> bool:
    if pd.isna(x):
        return False
    if isinstance(x, (int, float, np.number)) and not isinstance(x, bool):
        return True
    s = str(x).strip().replace(",", "")
    if not s:
        return False
    return bool(re.fullmatch(r"[-+]?\d+(\.\d+)?", s))


def row_keywords_score(row: pd.Series) -> int:
    txt = safe_join([normalize_text(v) for v in row.tolist()])
    return sum(1 for kw in KEYWORDS if kw in txt)

def row_looks_like_data(row: pd.Series) -> bool:
    """Detect actual data/event rows so they are not swallowed as header rows."""
    vals = [v for v in row.tolist() if normalize_text(v)]
    if not vals:
        return False

    first = vals[0]
    first_txt = normalize_text(first)

    # Dates/timestamps in first column are almost always data rows, not header rows.
    if is_datetime_like(first):
        return True

    if re.match(r"^\d{1,2}[:.]\d{2}(:\d{2})?$", first_txt):
        # EXPRO/PDF rows that start with time.
        return True

    if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}(\s+\d{1,2}[:.]\d{2})?", first_txt):
        return True

    return False


def row_looks_like_units(row: pd.Series) -> bool:
    vals = [v for v in row.tolist() if normalize_text(v)]
    if not vals:
        return False
    txt = safe_join([normalize_text(v) for v in vals])
    unit_count = sum(1 for kw in UNIT_KEYWORDS if kw in txt)
    text_count = sum(not is_numeric_like(v) for v in vals)
    return unit_count >= 2 and text_count >= 2


def header_score(row: pd.Series) -> float:
    if row_looks_like_data(row):
        return -100.0

    kw = row_keywords_score(row)
    non_empty = int(row.notna().sum())
    textish = sum((not is_numeric_like(v)) and normalize_text(v) != "" for v in row.tolist())
    units = 6 if row_looks_like_units(row) else 0

    # Strongly favor rows that contain "date/time" plus operational keywords.
    row_txt = safe_join(normalize_text(v) for v in row.tolist())
    time_bonus = 12 if ("date" in row_txt and "time" in row_txt) or "hh:mm" in row_txt else 0

    return kw * 4 + min(textish, 25) + min(non_empty, 25) * 0.25 + units + time_bonus


def detect_header_row(raw_df: pd.DataFrame, max_scan_rows: int = 80) -> int:
    """Find the first row of a multi-row header block.

    Many well-test templates use broad parent rows above the actual column labels:
    e.g. PRODUCTION / GAS RATE / RATIO, or TMU / LAB DATA / Gas Metering.
    This scores candidate header blocks instead of a single row.
    """
    if raw_df.empty:
        return 0

    scan_rows = min(max_scan_rows, len(raw_df))
    best_idx = 0
    best_score = -1.0

    for i in range(scan_rows):
        if row_looks_like_data(raw_df.iloc[i]):
            continue

        # Do not start a header block from a single metadata/reference cell such as
        # "14.73 Psi and 60 °F Gas line Meter Run Size...". Parent merged headers
        # like "Gas" or "Oil" are still allowed if they match PARENT_HEADERS.
        row_vals_norm = [normalize_text(v) for v in raw_df.iloc[i].tolist() if normalize_text(v)]
        row_txt = safe_join(row_vals_norm)
        has_parent_cell = any(v in PARENT_HEADERS for v in row_vals_norm)
        has_date_time_marker = ("date" in row_txt and "time" in row_txt) or "hh:mm" in row_txt
        if len(row_vals_norm) < 2 and not has_parent_cell and not has_date_time_marker:
            continue

        row_score = header_score(raw_df.iloc[i])
        if row_score < 0:
            continue

        group_score = row_score
        row_joined_for_group = safe_join([normalize_text(v) for v in raw_df.iloc[i].tolist()])
        group_has_date_time = "date" in row_joined_for_group or "hh:mm" in row_joined_for_group
        header_rows = 1

        # Look ahead and include subheader/unit rows until a data row starts.
        for r in range(i + 1, min(len(raw_df), i + 6)):
            if row_looks_like_data(raw_df.iloc[r]):
                break
            if is_header_like_subrow(raw_df.iloc[r]) or row_looks_like_units(raw_df.iloc[r]) or header_score(raw_df.iloc[r]) > 8:
                group_score += max(header_score(raw_df.iloc[r]), 0) * 0.95
                header_rows += 1
                row_txt = safe_join([normalize_text(v) for v in raw_df.iloc[r].tolist()])
                if "date" in row_txt or "hh:mm" in row_txt:
                    group_has_date_time = True
            else:
                # allow one blank row inside metadata, but not inside a header block
                break

        if group_has_date_time:
            group_score += 15
        if header_rows >= 2:
            group_score += 8

        if group_score > best_score:
            best_score = group_score
            best_idx = i

    return best_idx if best_score >= 10 else 0


def is_header_like_subrow(row: pd.Series) -> bool:
    vals = [v for v in row.tolist() if normalize_text(v)]
    if not vals:
        return False

    if row_looks_like_data(row):
        return False

    txt = safe_join([normalize_text(v) for v in vals])
    kw_count = sum(1 for kw in KEYWORDS if kw in txt)
    unit_count = sum(1 for kw in UNIT_KEYWORDS if kw in txt)
    numeric_count = sum(is_numeric_like(v) and not is_datetime_like(v) for v in vals)
    text_count = sum(not is_numeric_like(v) for v in vals)

    return (kw_count + unit_count >= 2 and text_count >= 2 and numeric_count <= max(3, text_count + 1))


def fill_parent_groups(row_values: Iterable[object]) -> List[str]:
    """Forward-fill merged parent/group headers across blank cells.

    This handles rows such as:
      PRESSURE AND TEMPERATURE MEASUREMENTS | ... | PRODUCTION | GAS RATE
    where Excel merged cells appear as a value followed by blanks.
    """
    vals = [str(v).strip() if normalize_text(v) else "" for v in row_values]
    non_empty = sum(1 for v in vals if v)
    total = len(vals)

    normalized_vals = [normalize_text(v) for v in vals if normalize_text(v)]
    has_parent_cell = any(v in PARENT_HEADERS for v in normalized_vals)
    short_group_row = all(len(v) <= 60 for v in normalized_vals)

    should_fill = non_empty > 0 and short_group_row and (
        non_empty <= max(3, int(total * 0.55))
        or has_parent_cell
    )

    if not should_fill:
        return vals

    out = []
    current = ""
    for v in vals:
        if v:
            current = v
            out.append(v)
        else:
            out.append(current)
    return out


def build_combined_headers(raw_df: pd.DataFrame, header_row: int, max_header_rows: int = 6):
    header_rows = [header_row]

    for r in range(header_row + 1, min(len(raw_df), header_row + max_header_rows)):
        if row_looks_like_data(raw_df.iloc[r]):
            break
        if is_header_like_subrow(raw_df.iloc[r]) or row_looks_like_units(raw_df.iloc[r]) or header_score(raw_df.iloc[r]) > 8:
            header_rows.append(r)
        else:
            break

    header_matrix = []
    for r in header_rows:
        vals = fill_parent_groups(raw_df.iloc[r].tolist())
        header_matrix.append(vals)

    headers = []
    for c in range(raw_df.shape[1]):
        parts = []
        for row_vals in header_matrix:
            raw = row_vals[c] if c < len(row_vals) else ""
            val = normalize_text(raw)
            if not val or val == "nan" or is_datetime_like(raw):
                continue
            part = str(raw).strip()
            if part and part.lower() != "nan" and part not in parts:
                parts.append(part)

        # Safety: some Excel templates may pass numeric/float objects into
        # the combined header parts through merged or unit rows. Convert again
        # before joining to avoid: sequence item X: expected str instance, float found.
        parts = [str(p).strip() for p in parts if str(p).strip() and str(p).strip().lower() != "nan"]
        header = safe_join(parts).strip() or f"Column_{c + 1}"
        # compact repeated whitespace and obvious duplicate fragments
        header = re.sub(r"\s+", " ", str(header))
        headers.append(str(header))

    return make_unique(headers), header_rows




def make_unique(names: Iterable[object]) -> List[str]:
    seen: Dict[str, int] = {}
    out: List[str] = []

    for n in names:
        # Safety against float/numeric headers from Excel merged/unit rows.
        name = str(n).strip() if str(n).strip() else "unnamed"
        name = re.sub(r"\s+", " ", name)
        if name.lower() == "nan":
            name = "unnamed"
        if name in seen:
            seen[name] += 1
            name = f"{name}.{seen[name]}"
        else:
            seen[name] = 0
        out.append(str(name))

    return out




def table_from_raw(raw_df: pd.DataFrame) -> pd.DataFrame:
    raw_df = raw_df.dropna(how="all").dropna(axis=1, how="all")
    if raw_df.empty:
        return raw_df

    hdr = detect_header_row(raw_df)
    headers, header_rows = build_combined_headers(raw_df, hdr)
    data_start = max(header_rows) + 1

    df = raw_df.iloc[data_start:].copy()
    df.columns = [str(h) for h in headers[: df.shape[1]]]
    df = df.dropna(how="all")
    return df



def canonical_candidate_score(canon: str, column_name: str) -> int:
    c = clean_header(column_name)
    score = 0

    # Prefer exact operational rate/pressure fields over broader parent labels.
    if canon == "bsw_pct":
        if "bsw" in c or "bs and w" in c:
            score += 20
        if "wh" in c:
            score -= 8
    if canon == "oil_rate_stbd" and ("oil rate" in c or "oil q" in c or "oil or cond" in c):
        score += 20
    if canon == "water_rate_bpd" and ("water rate" in c or "water q" in c or "water reading" in c):
        score += 20
    if canon == "gross_rate_bpd" and ("gross" in c or "liq q" in c or "liquid rate" in c):
        score += 20
    if canon == "gas_rate_mmscfd" and ("gas rate" in c or "mmscf/d" in c):
        score += 20
    if canon == "gas_formation_mmscfd" and ("formation gas" in c or "gas formation" in c):
        score += 30
    if canon == "pumping_pressure_psi" and ("pump p" in c or "pumping p" in c or "pump pressure" in c):
        score += 25
    if canon == "gas_formation_mmscfd" and ("formation gas" in c or "gas formation" in c):
        score += 30
    if canon == "pumping_pressure_psi" and ("pump p" in c or "pumping p" in c or "pump pressure" in c):
        score += 25
    if canon == "sep_p_psi" and ("sep" in c or "separator" in c):
        score += 20
    if canon == "whp_psi" and ("well head" in c or "fthp" in c or "whp" in c or "u/s" in c):
        score += 20
    if canon == "datetime" and ("date" in c and "time" in c):
        score += 20

    # Penalize total/cumulative/event columns when looking for rates.
    if canon in {"oil_rate_stbd", "water_rate_bpd", "gross_rate_bpd", "gas_rate_mmscfd"}:
        if "cum" in c or "volume" in c:
            score -= 12

    return score



def canonical_candidate_score(canon: str, column_name: str) -> int:
    c = clean_header(column_name)
    score = 0

    if canon == "bsw_pct":
        if "bsw" in c or "bs and w" in c:
            score += 25
        if "wh" in c:
            score -= 8
    if canon == "oil_rate_stbd" and ("oil rate" in c or "oil q" in c or ("oil or condensate" in c and "rate" in c)):
        score += 25
    if canon == "water_rate_bpd" and ("water rate" in c or "water q" in c or ("water reading" in c and "rate" in c)):
        score += 25
    if canon == "gross_rate_bpd" and ("gross" in c or "liq q" in c or "liquid rate" in c):
        score += 20
    if canon == "gas_rate_mmscfd" and ("gas rate" in c or "mmscf/d" in c):
        score += 20
    if canon == "sep_p_psi" and ("sep" in c or "separator" in c or "gasp" in c):
        score += 20
    if canon == "whp_psi" and ("well head" in c or "fthp" in c or "whp" in c or "u/s" in c):
        score += 20
    if canon == "datetime" and ("date" in c and "time" in c):
        score += 20
    if canon == "pump_intake_pressure_psi" and re.search(r"\b(pi|intake)\b", c):
        score += 30
    if canon == "pump_discharge_pressure_psi" and re.search(r"\b(pd|discharge|disch)\b", c):
        score += 30
    if canon == "motor_current_amp" and ("amp" in c or "current" in c or re.fullmatch(r"cur\.?", c)):
        score += 30
    if canon == "pump_freq_hz" and ("freq" in c or "hz" in c):
        score += 30
    if canon == "salinity_kppm":
        if "salinity" in c:
            score += 35
        if "nacl" in c:
            score += 35
        if re.search(r"\bppm\b", c):
            score += 10
        if "water" in c:
            score += 5
        if any(term in c for term in ["api", "kf", "factor", "psig", "psi", "bbl/d", "bpd", "stb/d", "mmscf", "scf/stb", "h2o", "rate"]):
            score -= 40

    if canon in {"oil_rate_stbd", "water_rate_bpd", "gross_rate_bpd", "gas_rate_mmscfd"}:
        if "cum" in c or "volume" in c:
            score -= 14

    return score


def best_canonical_name(column_name: str) -> Optional[str]:
    c = clean_header(column_name)
    original = normalize_text(column_name)
    cpad = f" {c} "

    if c.startswith("calcul") or c in {"factor", "fb", "ftf", "fg", "fpv", "y2"}:
        return None

    # Date/time. Avoid broad fragments like "d m" because they match "liquid metering".
    if ("date" in c and "time" in c) or "hh:mm" in c or re.search(r"dd[/\-]?\s*mon[/\-]?\s*yy", c) or "[dd-mm-yy]" in original:
        return "datetime"
    if re.fullmatch(r"(date|test date|start date|dd/mm/yy|ddmmyy)", c) or (re.search(r"\bdate\b", c) and "time" not in c and "last update" not in c):
        return "date"
    if re.fullmatch(r"(time|test time|hour|hh:mm:ss?|hh mm)", c) or ("time" in c and ("hh:mm" in c or "hh mm" in c)):
        return "time"
    if c in {"hhmm", "hh mm", "time hh mm", "time hh:mm"}:
        return "time"

    if re.fullmatch(r"well( name| no\.?)?", c) or c in {"well_name", "wellname"}:
        return "well"
    if "event" in c or "remark" in c or "comment" in c or re.search(r"\bnote\b", c):
        return "note"

    # Choke.
    if "choke" in c:
        if "size" in c or "/64" in c or '64"' in c:
            return "choke_size_64"
        return "choke_pct"

    # Temperatures before pressures, because parent headers often contain "well head" or "separator".
    if re.search(r"\bflt\b", c) or "flow line temp" in c or "flowline temp" in c:
        return "flt_f" if ("deg f" in c or "°f" in c or re.search(r"\bf\b", c)) else "flt_c"
    has_temp_child = bool(re.search(r"\btemp\.?\b", c)) or "deg f" in c or "deg. f" in c or "deg c" in c or "deg. c" in c or "°f" in c or "°c" in c
    if "flow temp" in c or "d/s temp" in c or "ds temp" in c or "downstream temp" in c:
        return "flow_temp_f" if ("deg f" in c or "°f" in c) else "flow_temp_c"
    if "u/s temp" in c or "upstream temp" in c:
        return "us_temp_c"
    if "gast" in c or re.search(r"\bgas\s+(gas\s+)?t\b", c) or ("gas" in c and ("gas temp" in c or "gas temperature" in c)):
        if re.search(r"\bgas\s+(gas\s+)?t\s+c\b", c) or re.search(r"\bgas\s+temp\s+c\b", c):
            return "gas_temp_c"
        if re.search(r"\bgas\s+(gas\s+)?t\s+f\b", c) or re.search(r"\bgas\s+temp\s+f\b", c):
            return "gas_temp_f"
        return "gas_temp_f" if ("deg f" in c or "°f" in c or re.search(r"\bf\b", c)) else "gas_temp_c"
    if "oilt" in c or re.search(r"\boil\s+(oil\s+)?t\b", c) or ("oil" in c and ("oil temp" in c or "oil temperature" in c)):
        if re.search(r"\boil\s+(oil\s+)?t\s+c\b", c) or re.search(r"\boil\s+temp\s+c\b", c):
            return "oil_temp_c"
        if re.search(r"\boil\s+(oil\s+)?t\s+f\b", c) or re.search(r"\boil\s+temp\s+f\b", c):
            return "oil_temp_f"
        return "oil_temp_f" if ("deg f" in c or "°f" in c or re.search(r"\bf\b", c)) else "oil_temp_c"
    if "oil or condensate" in c and ("deg f" in c or "deg. f" in c or "deg c" in c or "deg. c" in c):
        return "oil_temp_f" if ("deg f" in c or "deg. f" in c) else "oil_temp_c"
    if "well head" in c and has_temp_child and "psig" not in c and "psi-g" not in c and "psi" not in c:
        return "wellhead_temp_f" if ("deg f" in c or "deg. f" in c or "°f" in c) else "wellhead_temp_c"
    if ("separator" in c or "sep" in c) and has_temp_child and "psig" not in c and "psi-g" not in c and "psi" not in c and "inh2o" not in c:
        return "sep_temp_f" if ("deg f" in c or "deg. f" in c or "°f" in c) else "sep_temp_c"

    # Gross rate can be polluted by merged parent rows that contain "salinity".
    if "gross rate" in c and ("bbl/d" in c or "bpd" in c):
        return "gross_rate_bpd"

    # Fluid properties before rates because parent words can include "oil".
    if "bs and w" in c or "bs&w" in original or "bsw" in c or "water cut" in c or "watercut" in c or re.fullmatch(r"wc %?", c):
        return "bsw_pct"
    if "fw" == c or re.fullmatch(r"fw %?", c) or " lab data fw " in f" {c} ":
        return "fw_pct"
    if "twc" in c:
        return "twc_pct"
    if ("salinity" in c or "nacl" in c or "salt" in c) and "gross rate" not in c:
        # Avoid false salinity mapping caused by merged-header forward-fill.
        # Example bad headers from wide TMU sheets: "api Salinity BBL/D",
        # "[psig] Salinity BBL/D", "Factor Salinity BBL/D". Those are
        # helper/pressure/rate columns that inherited the word Salinity from a
        # neighboring column; they must not be plotted as salinity.
        salinity_conflicts = [
            "api", "kf", "factor", "psig", "psi", "psia", "bbl/d", "bpd",
            "stb/d", "mmscf", "scf/stb", "h2o", "gross rate", "oil t",
            "air =1", "chart", "orifice", "gor", "rate",
        ]
        has_conflict = any(term in c for term in salinity_conflicts)
        has_strong_salinity_unit = ("nacl" in c) or re.search(r"\bsalinity\b.*\bppm\b", c)
        if has_conflict and not has_strong_salinity_unit:
            return None
        return "salinity_kppm"
    if ("oil" in c and "sg" in c):
        return "oil_sg"
    if ("water" in c and "sg" in c):
        return "water_sg"
    if ("water" in c and "ph" in c) or re.fullmatch(r"ph", c):
        return "water_ph"
    if ("gravity" in c and "api" in c) or "oil api" in c or "grav. deg api" in c:
        return "oil_api"
    if ("gravity" in c and "air" in c) or "gas sg" in c or "gas gravity" in c or "air = 1" in c:
        return "gas_sg"
    if "sp.gr" in c or "sp gr" in c or "specific gravity" in c:
        return "gas_sg"
    if re.search(r"\bh2s\b", c):
        return "h2s_ppm"
    if re.search(r"\bco2\b|co₂", c):
        return "co2_mole_pct"

    # MPFM/EXPRO explicit columns.
    if "qoil" in c and "(s" in c:
        return "qoil_s_stbd"
    if "qwat" in c and "(s" in c:
        return "qwat_s_bpd"
    if "qgas" in c and "(s" in c:
        return "qgas_s_mmscfd"
    if "qoil" in c and "(a" in c:
        return "qoil_a_bpd"
    if "qwat" in c and "(a" in c:
        return "qwat_a_bpd"
    if "qgas" in c and "(a" in c:
        return "qgas_a_mmcfd"
    if "qgross" in c:
        return "qgross_s_bpd"
    if re.search(r"\bwlr\b", c):
        return "wlr_s_pct"
    if "gvf" in c:
        return "gvf_a_pct"

    # Rates / cumulative volumes. Order matters.
    if "oil cum" in c or ("oil" in c and "cum" in c) or re.search(r"\bcum\.?\s*stb\b", c):
        return "oil_cum_bbl"
    if "wat cum" in c or "water cum" in c or "water cumulative" in c:
        return "wat_cum_bbl"
    if "gas cum" in c:
        return "gas_cum_mscf"
    if "liquid volume" in c or "lquid volume" in c:
        return "liquid_volume_bbl"

    # Formation gas must be checked before total gas because headers can contain
    # "formation gas rate", which also includes the phrase "gas rate".
    if "gas formation" in c or "formation gas" in c:
        return "gas_formation_mmscfd"

    # Gas rate first because some templates have polluted headers such as
    # "GAS WATER READING RATE MMSCF/D" after merged-cell extraction.
    if ("total gas rate" in c) or ("gas rate" in c) or ("mmscf/d" in c and "gas" in c) or re.search(r"\brate mmscf/d\b", c):
        return "gas_rate_mmscfd"

    if "oil rate" in c or "oil q" in c or ("oil or condensate" in c and "rate" in c) or (re.search(r"\brate\b", c) and "stb/d" in c and "water" not in c and "gas" not in c and "gross" not in c):
        return "oil_rate_stbd"
    if "water rate" in c or "water q" in c or "bwpd" in c or ("water reading" in c and "rate" in c and "mmscf" not in c):
        return "water_rate_bpd"
    if "gross rate" in c or ("gross" in c and ("bbl" in c or "bpd" in c)):
        return "gross_rate_bpd"
    if "liq q" in c:
        return "gross_rate_bpd"
    if "liquid rate" in c:
        if "oil" in c:
            return "oil_rate_stbd"
        if "water" in c:
            return "water_rate_bpd"
        return "gross_rate_bpd"
    if re.search(r"\brate\b", c) and "bbl/d" in c and "oil" not in c and "gas" not in c and "gross" not in c:
        return "water_rate_bpd"

    # ESP / artificial-lift short aliases.  These are intentionally handled
    # before generic pressure/temperature rules because headers like Pi/Pd/Amp/Freq
    # are short and may otherwise remain raw fallback columns.  The strict fullmatch
    # rules avoid misclassifying broad text.
    if re.fullmatch(r"(pi|p/i|p int|pint|pump intake p|pump intake pressure|intake pressure|intake p|pin|pi psi|pi psig)", c):
        return "pump_intake_pressure_psi"
    if re.fullmatch(r"(pd|p/d|p dis|pdis|pump discharge p|pump discharge pressure|discharge pressure|disch pressure|discharge p|pd psi|pd psig)", c):
        return "pump_discharge_pressure_psi"
    if re.fullmatch(r"(amp|amps|ampere|amperage|current|motor current|pump current|run current|cur|cur\.?|i)", c) or ("amp" in c and "temp" not in c):
        return "motor_current_amp"
    if re.fullmatch(r"(ama|a m a)", c):
        return "motor_ama_amp"
    if re.fullmatch(r"(freq|frequency|hz|run freq|run frequency|operating freq|operating frequency|pump freq|pump frequency)", c):
        return "pump_freq_hz"
    if re.fullmatch(r"(drive freq|drive frequency|vsd freq|vfd freq|vfd frequency|speed hz)", c):
        return "drive_freq_hz"
    if re.fullmatch(r"(ti|t/i|intake temp|intake temperature|pump intake temp|pump intake temperature|ti c|ti deg c)", c):
        return "intake_temp_c"
    if re.fullmatch(r"(ti f|ti deg f|intake temp f|intake temperature f)", c):
        return "intake_temp_f"
    if re.fullmatch(r"(tm|t/m|motor temp|motor temperature|motor winding temp|tm c|tm deg c)", c):
        return "motor_temp_c"
    if re.fullmatch(r"(tm f|tm deg f|motor temp f|motor temperature f)", c):
        return "motor_temp_f"
    if re.fullmatch(r"(motor load|load pct|load %|motor load %)", c):
        return "motor_load_pct"
    if re.fullmatch(r"(vx|vib x|vibration x|x vibration)", c):
        return "vibration_x"
    if re.fullmatch(r"(vy|vib y|vibration y|y vibration)", c):
        return "vibration_y"
    if re.fullmatch(r"(vz|vib z|vibration z|z vibration)", c):
        return "vibration_z"

    # Pressures / DP.
    if re.search(r"\bflp\b|flow line pressure|flowline pressure", c):
        return "flp_psi"
    if "flow press" in c or "d/s press" in c or "ds press" in c or "downstream" in c:
        return "flow_press_psi"
    if re.search(r"\bfthp\b|\bwhp\b|w\.?h\.?p|wellhead pressure|well head pressure", c):
        return "whp_psi"
    if ("well head" in c or "upstream" in c or "u/s" in c) and ("press" in c or "psi" in c):
        return "whp_psi"
    if ("sep" in c or "separator" in c or "gasp" in c or "gas p" in c) and ("dp" in c or "diff" in c or "inh2o" in c):
        return "sep_dp_inh2o"
    if "gasdp" in c:
        return "gas_dp_inh2o"
    if ("sep" in c or "separator" in c or "gasp" in c or "gas p" in c) and ("p" in c or "pressure" in c or "psig" in c or "psi" in c):
        return "sep_p_psi"
    if " dp " in f" {c} " or (re.search(r"\bdiff\b", c) and ("h2o" in c or "mbar" in c)):
        return "dp_mbar" if "mbar" in c else "sep_dp_inh2o"
    if re.search(r"\bpump\s*p\b|\bpumping\s*p\b|pumping pressure|pump pressure|circulation pressure", c):
        return "pumping_pressure_psi"
    if re.search(r"\bct pressure\b|ct press", c):
        return "ct_pressure_psi"

    # Ratios.
    if "cgr" in c:
        return "cgr_bbl_mmscf"
    if re.search(r"\bgor\b", c):
        if "mmscf/bbl" in c:
            return "gor_mmscf_bbl"
        return "gor_scf_bbl"

    # Other operational parameters.
    if "orifice" in c and "size" in c:
        return "orifice_size_in"
    if "mtr inc" in c or "meter inc" in c or "meter increment" in c:
        return "oil_meter_increment_bbl"
    if "cmf" in c and "oil" in c:
        return "oil_cmf"
    if "k .f" in c or c == "kf":
        return "oil_kf"
    if "n2" in c or "nitrogen" in c:
        return "n2_rate_scfm"
    if "ct depth" in c or ("depth" in c and "ct" in c):
        return "ct_depth_m"
    if "running speed" in c:
        return "ct_running_speed_ftmin"
    if "pipe weight" in c:
        return "ct_pipe_weight_lbf"
    if "u2 pass" in c and "pressure" in c:
        return "u2_pass_side_pump_pressure_psi"
    if "pump freq" in c or ("freq" in c and ("hz" in c or "pump" in c)):
        return "pump_freq_hz"

    return None

def extract_number(value: object) -> float:
    """Strict numeric extraction.

    This intentionally reads only numbers at the start of the value.
    It avoids pulling 0.750 from long event text such as
    "Installed orifice plate size 0.750".
    """
    import datetime as _dt

    if value is None or (isinstance(value, float) and np.isnan(value)):
        return np.nan
    if isinstance(value, _dt.time):
        return np.nan
    if isinstance(value, (_dt.datetime, _dt.date, pd.Timestamp)):
        return np.nan
    if isinstance(value, (int, float, np.number)) and not isinstance(value, bool):
        return float(value)

    s = str(value).strip().replace(",", "")
    if not s or re.fullmatch(r"n/?a|nan|null|-", s, flags=re.I):
        return np.nan

    m = re.match(r"^\s*[-+]?\d*\.?\d+", s)
    if not m:
        return np.nan

    try:
        return float(m.group(0))
    except ValueError:
        return np.nan


def clean_numeric_series(series: pd.Series, canonical_name: str) -> pd.Series:
    s = series.map(extract_number).astype(float)

    # Excel often stores choke as 1.0 for 100% in some TMU templates.
    if canonical_name == "choke_pct":
        valid = s.dropna()
        if not valid.empty and valid.median() <= 1.5:
            s = s.where(s > 1.5, s * 100.0)

    # BS&W can be stored either as percent (1 = 1%) or fraction (0.006 = 0.6%).
    # Convert only clear fraction-style values, but do not turn 1% into 100%.
    if canonical_name == "bsw_pct":
        valid = s.dropna()
        if not valid.empty and valid.max() <= 1.0 and valid.median() < 0.25:
            s = s * 100.0

    # Store salinity as K ppm because most plots/report templates use K ppm NaCl.
    if canonical_name == "salinity_kppm":
        med = s.dropna().median() if s.notna().any() else np.nan
        if pd.notna(med) and med > 1000:
            s = s / 1000.0

    return s


def parse_date_series(series: pd.Series) -> pd.Series:
    import datetime as _dt

    def one(x):
        if pd.isna(x):
            return pd.NaT
        if isinstance(x, pd.Timestamp):
            return pd.Timestamp(x.date())
        if isinstance(x, _dt.datetime):
            return pd.Timestamp(x.date())
        if isinstance(x, _dt.date):
            return pd.Timestamp(x)
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            return pd.to_datetime(x, errors="coerce", dayfirst=True)

    return series.map(one)


def parse_time_series(series: pd.Series) -> pd.Series:
    import datetime as _dt

    def one(x):
        if pd.isna(x):
            return pd.NaT

        if isinstance(x, pd.Timestamp):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, _dt.datetime):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, _dt.time):
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)

        if isinstance(x, (int, float, np.number)) and not isinstance(x, bool):
            if 0 <= float(x) < 1:
                seconds = int(round(float(x) * 24 * 3600))
                return pd.Timestamp("1900-01-01") + pd.Timedelta(seconds=seconds)

        s = str(x)
        m = re.search(r"(\d{1,2})[:.](\d{2})(?:[:.](\d{2}))?", s)
        if m:
            hh = int(m.group(1))
            mm = int(m.group(2))
            ss = int(m.group(3) or 0)
            return pd.Timestamp("1900-01-01") + pd.Timedelta(hours=hh, minutes=mm, seconds=ss)

        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            return pd.to_datetime(s, errors="coerce")

    return series.map(one)




def parse_datetime_series(series: pd.Series) -> pd.Series:
    """Parse mixed datetime values safely without flooding Streamlit logs with pandas warnings.

    ISO yyyy-mm-dd strings must not use dayfirst=True. Other field formats such
    as 10-06-2026 are parsed with dayfirst=True. Any unparseable values become NaT.
    """
    out = pd.Series(pd.NaT, index=series.index, dtype="datetime64[ns]")
    str_s = series.astype(str).str.strip()
    iso_mask = str_s.str.match(r"^\d{4}[-/]\d{1,2}[-/]\d{1,2}", na=False)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        if iso_mask.any():
            out.loc[iso_mask] = pd.to_datetime(series.loc[iso_mask], errors="coerce", dayfirst=False)
        if (~iso_mask).any():
            out.loc[~iso_mask] = pd.to_datetime(series.loc[~iso_mask], errors="coerce", dayfirst=True)
    return out

def combine_date_time(
    date_series: Optional[pd.Series] = None,
    time_series: Optional[pd.Series] = None,
    datetime_series: Optional[pd.Series] = None,
) -> pd.Series:
    """Build a real datetime column from any available date/time inputs.

    Important fix: Excel files often contain monthly/daily production history with a
    Date column only and no separate Time column. Older versions returned all NaT
    when time_series was missing, causing duplicate-removal to collapse the whole
    sheet into one row and the workbook was rejected as "no usable time-series".
    Date-only rows are now valid readings at midnight (00:00).
    """
    idx = None
    for s in [date_series, time_series, datetime_series]:
        if s is not None:
            idx = s.index
            break

    if idx is None:
        return pd.Series(dtype="datetime64[ns]")

    if datetime_series is not None:
        dt = parse_datetime_series(datetime_series)
    else:
        dt = pd.Series(pd.NaT, index=idx, dtype="datetime64[ns]")

    dates = parse_date_series(date_series) if date_series is not None else pd.Series(pd.NaT, index=idx, dtype="datetime64[ns]")
    times = parse_time_series(time_series) if time_series is not None else pd.Series(pd.NaT, index=idx, dtype="datetime64[ns]")

    out = []
    for d, t, existing in zip(dates, times, dt):
        if pd.notna(existing):
            out.append(pd.Timestamp(existing))
        elif pd.notna(d) and pd.notna(t):
            out.append(pd.Timestamp(d.date()) + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second))
        elif pd.notna(d):
            # Date-only production-history sheets are valid time series.
            out.append(pd.Timestamp(d.date()))
        elif pd.notna(t):
            out.append(pd.Timestamp("1900-01-01") + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second))
        else:
            out.append(pd.NaT)

    return pd.Series(out, index=idx, dtype="datetime64[ns]")





def adjust_datetime_rollover(out: pd.DataFrame, parsed_time: Optional[pd.Series]) -> pd.DataFrame:
    """Fix midnight rollover in templates where DATE is copied down but TIME resets from 23:30 to 00:00.

    Example: date column remains 2026-06-05, rows are 23:30 then 00:00.
    The 00:00 row is really 2026-06-06 00:00. We adjust in original row order before sorting.
    """
    if out is None or out.empty or parsed_time is None:
        return out
    if "date" not in out.columns or "datetime" not in out.columns:
        return out
    if parsed_time.notna().sum() < 2:
        return out

    dates = pd.to_datetime(out["date"], errors="coerce")
    times = parsed_time
    adjusted = out["datetime"].copy()

    day_offset = 0
    prev_valid_time = None
    prev_base_date = None

    for idx in out.index:
        d = dates.loc[idx]
        t = times.loc[idx]
        if pd.isna(d) or pd.isna(t):
            continue

        base_date = pd.Timestamp(d.date())

        # If the sheet explicitly advances the DATE column, reset offset.
        if prev_base_date is not None and base_date > prev_base_date:
            day_offset = 0

        # If DATE did not advance but TIME moved backwards, assume midnight rollover.
        if (
            prev_valid_time is not None
            and prev_base_date is not None
            and base_date == prev_base_date
            and (t.hour, t.minute, t.second) < (prev_valid_time.hour, prev_valid_time.minute, prev_valid_time.second)
        ):
            day_offset += 1

        adjusted.loc[idx] = base_date + pd.Timedelta(days=day_offset) + pd.Timedelta(
            hours=t.hour, minutes=t.minute, seconds=t.second
        )

        prev_valid_time = t
        prev_base_date = base_date

    out = out.copy()
    out["datetime"] = adjusted
    out["date"] = pd.to_datetime(out["datetime"], errors="coerce").dt.floor("D")
    out["time_text"] = pd.to_datetime(out["datetime"], errors="coerce").dt.strftime("%H:%M")
    return out


def normalize_well_name(value: object) -> Optional[str]:
    if value is None or pd.isna(value):
        return None

    s = str(value).strip()
    if not s:
        return None

    # Remove obvious label prefixes.
    s = re.sub(r"(?i)\bwell\s*(name|no\.?)?\s*[:=#-]*", "", s).strip()
    s = s.replace("#", "-")
    s = re.sub(r"\s*-\s*", "-", s)
    s = re.sub(r"\s+", " ", s).strip(" :-_")
    if not s:
        return None

    # Reject plain dates/numbers.
    if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", s):
        return None
    if re.fullmatch(r"\d+(\.\d+)?", s):
        return None

    # Normalize common forms.
    up = s.upper()

    # BAHGA 11 -> BAHGA-11
    up = re.sub(r"\b(BAHGA|ASSIL|KARAM|MAGD|ELMAGD|BED|SITRA|BAPETCO)\s+(\d)", r"\1-\2", up)

    # MAGD C 86-4 / ELMAGD 86-4 kept readable.
    up = re.sub(r"\bMAGD-C-", "MAGD C-", up)
    up = re.sub(r"\bELMAGD-", "ELMAGD-", up)

    # Remove duplicate spaces and spaces around hyphens.
    up = re.sub(r"\s*-\s*", "-", up)
    up = re.sub(r"\s+", " ", up).strip()

    # Keep not-too-long candidates only.
    if len(up) > 30:
        return None

    return up


def guess_well_from_name(name: str) -> Optional[str]:
    s = str(name)

    # Remove date-like parenthetical blocks to avoid returning 03-07 or 8-12 as a well.
    s_clean = re.sub(r"\([^)]*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}[^)]*\)", " ", s)
    s_clean = re.sub(r"\b\d{4}[-_]\d{2}[-_]\d{2}\b", " ", s_clean)

    patterns = [
        r"\b(BAHGA\s*#?\s*\d+[A-Z]?)\b",
        r"\b(MAGD\s*C\s*\d+\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b(ELMAGD\s*\d+\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b(BED\s*\d+[A-Z]?\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b(B\d+[A-Z]*\d*\s*[-_]\s*\d+[A-Z]?)\b",
        r"\b([A-Z][A-Z0-9]*\s*C?\s*\d+\s*[-_]\s*\d+[A-Z]?)\b",
    ]

    for pat in patterns:
        m = re.search(pat, s_clean, flags=re.I)
        if m:
            well = normalize_well_name(m.group(1))
            # Avoid false matches created by spreadsheet column labels plus dates, e.g. "Y2 06-06".
            if well and re.fullmatch(r"[XYZ]\d+[- ]\d{1,2}[-_/]\d{1,2}", well):
                continue
            if well:
                return well

    return None


def extract_well_from_raw(raw_df: pd.DataFrame, source_name: str = "", sheet_name: str = "") -> Optional[str]:
    """Extract well name from workbook metadata rows before falling back to filename."""
    try:
        max_r = min(len(raw_df), 40)
        max_c = min(raw_df.shape[1], 30)

        for r in range(max_r):
            row_vals = [raw_df.iloc[r, c] for c in range(max_c)]
            for c, val in enumerate(row_vals):
                txt = str(val).strip() if not pd.isna(val) else ""
                if not txt:
                    continue

                # Cell contains the whole value: "WELL: Bahga#9"
                m = re.search(r"(?i)\bwell\s*(?:name|no\.?)?\s*[:=#-]+\s*([A-Za-z0-9# \-_]+)", txt)
                if m:
                    well = normalize_well_name(m.group(1))
                    if well:
                        return well

                # Label in this cell, value in nearby cells to the right.
                if re.search(r"(?i)^\s*well\s*(name|no\.?)?\s*:?\s*$", txt):
                    for k in range(c + 1, min(c + 5, max_c)):
                        cand = normalize_well_name(raw_df.iloc[r, k])
                        if cand:
                            return cand

                # Some templates use "Well  :" in one cell and the next cell has value.
                if re.search(r"(?i)\bwell\b", txt) and len(txt) <= 20:
                    for k in range(c + 1, min(c + 5, max_c)):
                        cand = normalize_well_name(raw_df.iloc[r, k])
                        if cand and not cand.startswith("TEST"):
                            return cand

        # As a last metadata attempt, look for known well patterns in all early text.
        text_blob = safe_join(
            [raw_df.iloc[r, c] for r in range(max_r) for c in range(max_c)]
        )
        guessed = guess_well_from_name(text_blob)
        if guessed:
            return guessed

    except Exception:
        pass

    return guess_well_from_name(f"{source_name} {sheet_name}")



def raw_output_column_name(column_name: object, existing_columns: Iterable[str]) -> str:
    """Create a stable fallback name for numeric columns that do not match aliases yet."""
    label = str(column_name).strip() if str(column_name).strip() else "Column"
    safe = clean_header(label)
    safe = re.sub(r"[^a-z0-9]+", "_", safe).strip("_")
    if not safe:
        safe = "column"
    safe = safe[:80]
    base = f"raw__{safe}"
    name = base
    i = 2
    existing = set(existing_columns)
    while name in existing:
        name = f"{base}_{i}"
        i += 1
    return name


def is_useful_raw_numeric_column(column_name: object, series: pd.Series) -> bool:
    """Decide whether an unmapped column should still appear as a user-selectable raw column."""
    header = clean_header(column_name)
    if not header or header.startswith("unnamed") or header.startswith("column_"):
        return False
    if header in {"date", "time", "datetime", "well", "note", "event", "comments", "remarks"}:
        return False
    if "event" in header or "remark" in header or "comment" in header:
        return False

    # Avoid hidden calculation/helper columns that may be present far to the right
    # of operational Excel templates.
    helper_terms = [
        "calcul", "calc", "factor", "ftf", "fg", "fb", "fpv", "y2",
        "h2o", "psia", "base condition", "intermediate", "correction",
    ]
    if any(term in header for term in helper_terms):
        return False

    nums = series.map(extract_number).astype(float)
    return nums.notna().any()


def standardize_dataframe(
    df: pd.DataFrame,
    source_name: str = "",
    sheet_name: str = "",
    default_well: Optional[str] = None,
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    # Drop fully empty rows, but do NOT blindly drop fully empty columns here.
    # Some templates have a real field column whose values are all N/A, such as
    # Salinity = N/A. If we drop that exact Salinity column first, a later
    # helper/calculation column with a polluted merged header like
    # "api Salinity BBL/D" can be selected by mistake and plotted as salinity.
    df = df.copy().dropna(how="all")
    if df.empty:
        return pd.DataFrame()

    keep_cols = []
    for col in df.columns:
        if not df[col].isna().all() or best_canonical_name(str(col)) is not None:
            keep_cols.append(col)
    df = df.loc[:, keep_cols]
    if df.empty:
        return pd.DataFrame()

    # Choose the best source column for each canonical field.
    best_map: Dict[str, tuple] = {}
    for col in df.columns:
        canon = best_canonical_name(str(col))
        if not canon:
            continue
        score = canonical_candidate_score(canon, str(col))
        if canon not in best_map or score > best_map[canon][1]:
            best_map[canon] = (col, score)

    mapping = [(canon, col_score[0]) for canon, col_score in best_map.items()]

    out = pd.DataFrame(index=df.index)
    out["source"] = source_name
    out["sheet"] = sheet_name

    mapped_source_cols = set()
    for canon, col in mapping:
        mapped_source_cols.add(col)
        if canon in ["well", "date", "time", "datetime", "note"]:
            out[canon] = df[col]
        else:
            out[canon] = clean_numeric_series(df[col], canon)

    # Fallback: keep numeric columns not yet recognized by aliases only when the
    # table is mostly unknown. If a TMU template already produced many canonical
    # fields (WHP, FLP, Sep P, Gas rate, Oil rate, etc.), keeping every extra
    # helper/calculation/unit column creates confusing labels such as
    # "Raw: Column" or "Raw: Psig" and can even cause the Excel loader to choose
    # a lower-quality unit-row parse. Unknown templates still get raw columns.
    mapped_numeric_count = sum(
        1 for canon, _ in mapping
        if canon not in {"well", "date", "time", "datetime", "note"}
    )
    keep_raw_fallback = mapped_numeric_count < 6

    for col in df.columns:
        if col in mapped_source_cols:
            continue
        if not keep_raw_fallback:
            continue
        if not is_useful_raw_numeric_column(col, df[col]):
            continue
        raw_name = raw_output_column_name(col, out.columns)
        out[raw_name] = df[col].map(extract_number).astype(float)

    # Parse date/time/datetime robustly.
    if "datetime" in out.columns:
        out["datetime"] = parse_datetime_series(out["datetime"])

    if "date" in out.columns:
        out["date"] = parse_date_series(out["date"])
        out["date"] = out["date"].ffill().bfill()

    parsed_time = None
    if "time" in out.columns:
        parsed_time = parse_time_series(out["time"])
        out["time_text"] = parsed_time.dt.strftime("%H:%M")
        out.loc[parsed_time.isna(), "time_text"] = ""

    if ("datetime" not in out.columns or out["datetime"].isna().all()) and ("date" in out.columns or "time" in out.columns):
        out["datetime"] = combine_date_time(out.get("date"), out.get("time"), out.get("datetime"))

    # Derive date/time_text from datetime when a combined Date & Time column exists.
    if "datetime" in out.columns and out["datetime"].notna().any():
        if "date" not in out.columns:
            out["date"] = out["datetime"].dt.floor("D")
        if "time_text" not in out.columns:
            out["time_text"] = out["datetime"].dt.strftime("%H:%M")

    # Correct midnight rollover before sorting.
    if parsed_time is not None:
        out = adjust_datetime_rollover(out, parsed_time)

    # Well name: source metadata > sheet/filename fallback.
    if "well" in out.columns:
        out["well"] = out["well"].map(lambda x: normalize_well_name(x) or "")
    else:
        out["well"] = ""

    if out["well"].astype(str).str.strip().eq("").all():
        well = default_well or guess_well_from_name(f"{source_name} {sheet_name}") or "Unknown"
        out["well"] = well
    else:
        fallback = default_well or guess_well_from_name(f"{source_name} {sheet_name}") or "Unknown"
        out["well"] = out["well"].replace("", np.nan).ffill().bfill().fillna(fallback)

    numeric_cols = [c for c in out.columns if c not in BASE_NON_PLOT_COLS]
    canonical_numeric_cols = [c for c in numeric_cols if not str(c).startswith("raw__")]
    row_filter_cols = canonical_numeric_cols if canonical_numeric_cols else numeric_cols

    if row_filter_cols:
        useful = out[row_filter_cols].notna().any(axis=1)

        # For tables with a separate TIME column, require a valid time.
        # This removes event rows and rows such as Final Average that contain numbers but are not readings.
        if parsed_time is not None and parsed_time.notna().sum() >= 2:
            useful = useful & parsed_time.notna()
        elif "datetime" in out.columns and out["datetime"].notna().sum() >= 2:
            useful = useful & out["datetime"].notna()

        # Drop pure average/summary/total rows even if the date column was forward-filled.
        raw_text_for_row = df.apply(lambda r: safe_join(list(r.values), " | "), axis=1).str.lower()
        useful = useful & ~raw_text_for_row.str.contains(
            r"final\s+average|average|avg|summary|grand\s+total|^total",
            regex=True,
            na=False,
        )

        out = out.loc[useful].copy()

    if out.empty:
        return pd.DataFrame()

    # Remove duplicate readings at the same well/datetime. This also cleans old dashboard-exported CSVs
    # that may contain a duplicated Final Average row with a copied timestamp.
    if "datetime" in out.columns and "well" in out.columns and out["datetime"].notna().any():
        # Do not let all-NaT datetimes collapse a valid table to one row.
        out = out.drop_duplicates(subset=["well", "datetime"], keep="first")

    out["well"] = out["well"].astype(str).str.strip().replace("", "Unknown")
    sort_cols = ["well"] + (["datetime"] if "datetime" in out.columns else [])
    out = out.sort_values(sort_cols, na_position="last").reset_index(drop=True)

    return out


def dataframe_quality_score(df: pd.DataFrame) -> int:
    """Score parsed tables so the Excel loader keeps the real operational table.

    The old scoring over-valued the number of numeric columns. In wide TMU
    spreadsheets, a bad unit-row parse can create dozens of raw columns such as
    raw__psig/raw__column and beat the proper multi-row header parse. This scorer
    now strongly prefers:
      - recognized canonical TMU columns,
      - real calendar datetimes (not 1900 time-only placeholders),
      - date + time together,
      - fewer raw fallback columns.
    """
    if df is None or df.empty:
        return -10_000

    numeric_cols = available_numeric_columns(df)
    if not numeric_cols:
        return -10_000

    canonical_cols = [c for c in numeric_cols if not str(c).startswith("raw__")]
    raw_cols = [c for c in numeric_cols if str(c).startswith("raw__")]

    score = 0
    score += min(len(df), 500) * 4
    score += len(canonical_cols) * 300
    score += len(raw_cols) * 5

    # Penalize raw-only parses. They are useful as a fallback, but should not win
    # over a table where columns have been identified as WHP, FLP, gas rate, etc.
    if not canonical_cols:
        score -= 700
    if raw_cols and len(raw_cols) > max(5, len(canonical_cols) * 2):
        score -= (len(raw_cols) - max(5, len(canonical_cols) * 2)) * 20

    if "datetime" in df.columns:
        dt = pd.to_datetime(df["datetime"], errors="coerce")
        valid_dt = dt.dropna()
        score += int(valid_dt.size) * 12
        real_dt_count = int((valid_dt.dt.year > 1970).sum()) if not valid_dt.empty else 0
        score += real_dt_count * 45
        # Time-only fallback creates dates around 1900. It can be plotted, but it
        # must never beat a parse with real file dates.
        if valid_dt.size and real_dt_count == 0:
            score -= 900

    if "date" in df.columns:
        dates = pd.to_datetime(df["date"], errors="coerce").dropna()
        real_date_count = int((dates.dt.year > 1970).sum()) if not dates.empty else 0
        score += real_date_count * 25

    if "time_text" in df.columns:
        score += int(df["time_text"].astype(str).str.strip().ne("").sum()) * 5

    # Prefer the common field-test columns that users expect to see by name.
    priority_cols = {
        "choke_pct", "whp_psi", "flp_psi", "flt_c", "sep_p_psi",
        "gas_rate_mmscfd", "oil_rate_stbd", "water_rate_bpd", "gross_rate_bpd",
        "bsw_pct", "salinity_kppm",
    }
    score += sum(1 for c in priority_cols if c in df.columns) * 150

    return int(score)

def dataframe_key(df: pd.DataFrame) -> tuple:
    """Return a loose de-duplication key for repeated parsing attempts."""
    if df is None or df.empty:
        return (0, "", "")
    sheet = str(df["sheet"].iloc[0]) if "sheet" in df.columns and len(df) else ""
    well = str(df["well"].iloc[0]) if "well" in df.columns and len(df) else ""
    dt_min = ""
    dt_max = ""
    if "datetime" in df.columns:
        dt = pd.to_datetime(df["datetime"], errors="coerce").dropna()
        if not dt.empty:
            dt_min = str(dt.min())
            dt_max = str(dt.max())
    return (len(df), sheet, well, dt_min, dt_max)


def parse_datetime_value_count(series: pd.Series) -> int:
    """Count values that contain both date and time information."""
    parsed = parse_datetime_series(series)
    raw = series.astype(str).str.strip()
    raw_has_date_and_time = raw.str.contains(r"\d{1,4}[-/]\d{1,2}[-/]\d{1,4}.*\d{1,2}[:.]\d{2}|\d{1,2}[:.]\d{2}.*\d{1,4}[-/]\d{1,2}[-/]\d{1,4}", regex=True, na=False)

    import datetime as _dt
    object_has_time = series.map(
        lambda x: isinstance(x, (pd.Timestamp, _dt.datetime)) and (x.hour != 0 or x.minute != 0 or x.second != 0)
        if not pd.isna(x) else False
    )
    return int((parsed.notna() & (raw_has_date_and_time | object_has_time)).sum())


def parse_date_value_count(series: pd.Series) -> int:
    return int(parse_date_series(series).notna().sum())


def parse_time_value_count(series: pd.Series) -> int:
    parsed = parse_time_series(series)
    # Treat pure dates parsed as midnight as weak time evidence unless the raw value had ':' or was an Excel time fraction.
    raw = series.astype(str).str.strip()
    raw_time_like = raw.str.contains(r"\d{1,2}[:.]\d{2}", regex=True, na=False)
    excel_fraction_time = series.map(lambda x: isinstance(x, (int, float, np.number)) and not isinstance(x, bool) and 0 <= float(x) < 1 if not pd.isna(x) else False)
    return int((parsed.notna() & (raw_time_like | excel_fraction_time)).sum())


def best_value_column(raw_df: pd.DataFrame, kind: str, exclude: Optional[set] = None) -> Optional[object]:
    """Detect date/time/datetime columns from cell values, not only from headers."""
    exclude = exclude or set()
    best_col = None
    best_score = 0
    for col in raw_df.columns:
        if col in exclude:
            continue
        s = raw_df[col].dropna()
        if s.empty:
            continue
        sample = s.head(80)
        if kind == "datetime":
            score = parse_datetime_value_count(sample)
        elif kind == "date":
            score = parse_date_value_count(sample)
        else:
            score = parse_time_value_count(sample)
        # Require at least two useful values so metadata cells are not mistaken for a series.
        if score >= 2 and score > best_score:
            best_score = score
            best_col = col
    return best_col


def standardize_loose_timeseries(
    raw_df: pd.DataFrame,
    source_name: str = "",
    sheet_name: str = "",
    default_well: Optional[str] = None,
) -> pd.DataFrame:
    """Last-resort Excel parser for sheets whose headers are blank/merged/not recognized.

    It detects date/time columns from the actual values and keeps every numeric column as
    a raw plotting series. This prevents valid field Excel sheets from being rejected just
    because their header text does not match the alias list yet.
    """
    if raw_df is None or raw_df.empty:
        return pd.DataFrame()

    df = raw_df.copy().dropna(how="all").dropna(axis=1, how="all")
    if df.empty:
        return pd.DataFrame()

    # If the first non-empty row is mostly text, use it as labels for nicer raw column names.
    first_row = df.iloc[0]
    text_cells = sum((not is_numeric_like(v)) and bool(normalize_text(v)) for v in first_row.tolist())
    non_empty = sum(bool(normalize_text(v)) for v in first_row.tolist())
    first_row_is_data = row_looks_like_data(first_row) or (
        parse_date_value_count(pd.Series(first_row.tolist())) >= 1
        and parse_time_value_count(pd.Series(first_row.tolist())) >= 1
        and sum(is_numeric_like(v) for v in first_row.tolist()) >= 1
    )
    if not first_row_is_data and non_empty >= 2 and text_cells >= max(2, non_empty // 2):
        labels = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(first_row.tolist())])
        df = df.iloc[1:].copy()
        df.columns = labels[: df.shape[1]]
    else:
        df.columns = [f"Column_{i + 1}" for i in range(df.shape[1])]

    df = df.dropna(how="all")
    if df.empty:
        return pd.DataFrame()

    datetime_col = best_value_column(df, "datetime")
    used = set([datetime_col]) if datetime_col is not None else set()
    date_col = None if datetime_col is not None else best_value_column(df, "date", used)
    if date_col is not None:
        used.add(date_col)
    time_col = best_value_column(df, "time", used)
    if time_col is not None:
        used.add(time_col)

    out = pd.DataFrame(index=df.index)
    out["source"] = source_name
    out["sheet"] = sheet_name
    out["well"] = default_well or guess_well_from_name(f"{source_name} {sheet_name}") or "Unknown"

    if datetime_col is not None:
        out["datetime"] = parse_datetime_series(df[datetime_col])
    if date_col is not None:
        out["date"] = parse_date_series(df[date_col]).ffill().bfill()
    if time_col is not None:
        parsed_time = parse_time_series(df[time_col])
        out["time_text"] = parsed_time.dt.strftime("%H:%M")
        out.loc[parsed_time.isna(), "time_text"] = ""
    else:
        parsed_time = None

    if ("datetime" not in out.columns or out["datetime"].isna().all()) and (date_col is not None or time_col is not None):
        out["datetime"] = combine_date_time(out.get("date"), df[time_col] if time_col is not None else None, out.get("datetime"))

    if "datetime" in out.columns and out["datetime"].notna().any():
        if "date" not in out.columns:
            out["date"] = out["datetime"].dt.floor("D")
        if "time_text" not in out.columns:
            out["time_text"] = out["datetime"].dt.strftime("%H:%M")

    # Keep all numeric columns except detected date/time columns. Use header labels when possible.
    for col in df.columns:
        if col in used:
            continue
        nums = df[col].map(extract_number).astype(float)
        if nums.notna().sum() < 2:
            continue
        raw_name = raw_output_column_name(col, out.columns)
        out[raw_name] = nums

    numeric_cols = [c for c in out.columns if c not in BASE_NON_PLOT_COLS]
    if not numeric_cols:
        return pd.DataFrame()

    useful = out[numeric_cols].notna().any(axis=1)
    if "datetime" in out.columns and out["datetime"].notna().sum() >= 2:
        useful &= out["datetime"].notna()
    elif "time_text" in out.columns and out["time_text"].astype(str).str.strip().ne("").sum() >= 2:
        useful &= out["time_text"].astype(str).str.strip().ne("")

    out = out.loc[useful].copy()
    if out.empty:
        return pd.DataFrame()

    if "datetime" in out.columns and "well" in out.columns and out["datetime"].notna().any():
        # Do not let all-NaT datetimes collapse a valid table to one row.
        out = out.drop_duplicates(subset=["well", "datetime"], keep="first")

    sort_cols = ["well"] + (["datetime"] if "datetime" in out.columns else [])
    return out.sort_values(sort_cols, na_position="last").reset_index(drop=True)


def parse_excel_sheet_attempts(raw: pd.DataFrame, source_name: str, sheet_name: str, default_well: Optional[str]) -> List[pd.DataFrame]:
    """Try several Excel interpretations and return valid time-series candidates."""
    attempts: List[pd.DataFrame] = []

    def add_candidate(candidate: pd.DataFrame):
        if candidate is not None and not candidate.empty and is_valid_timeseries(candidate):
            attempts.append(candidate)

    # Existing smart multi-row header parser.
    table = table_from_raw(raw)
    add_candidate(standardize_dataframe(table, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # Common case: Excel already has a simple first-row header. Try it explicitly.
    df0 = raw.copy().dropna(how="all").dropna(axis=1, how="all")
    if not df0.empty and len(df0) >= 2:
        headers = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(df0.iloc[0].tolist())])
        direct = df0.iloc[1:].copy()
        direct.columns = headers[: direct.shape[1]]
        add_candidate(standardize_dataframe(direct, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # Scan possible header rows. This helps when title/metadata rows are above the actual table.
    max_scan = min(60, len(raw))
    for header_row in range(max_scan):
        row = raw.iloc[header_row]
        if row_looks_like_data(row):
            continue
        score = header_score(row)
        if score < 4:
            continue
        candidate_raw = raw.iloc[header_row:].dropna(how="all").dropna(axis=1, how="all")
        if candidate_raw.empty or len(candidate_raw) < 3:
            continue
        headers = make_unique([safe_text(v) or f"Column_{i + 1}" for i, v in enumerate(candidate_raw.iloc[0].tolist())])
        candidate = candidate_raw.iloc[1:].copy()
        candidate.columns = headers[: candidate.shape[1]]
        add_candidate(standardize_dataframe(candidate, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # Last-resort value-based parser. This catches headerless or unrecognized templates.
    add_candidate(standardize_loose_timeseries(raw, source_name=source_name, sheet_name=sheet_name, default_well=default_well))

    # De-duplicate and keep the strongest parse for this sheet.
    unique: Dict[tuple, pd.DataFrame] = {}
    for cand in attempts:
        key = dataframe_key(cand)
        if key not in unique or dataframe_quality_score(cand) > dataframe_quality_score(unique[key]):
            unique[key] = cand

    return sorted(unique.values(), key=dataframe_quality_score, reverse=True)


def is_valid_timeseries(df: pd.DataFrame) -> bool:
    """Return True only for real plottable time-series tables.

    A table must have at least one numeric reading column and repeated date/time
    evidence.  This prevents non-data text/PDF files from being accepted just
    because a few numbers were found somewhere in the document.
    """
    if df is None or df.empty:
        return False
    numeric_cols = available_numeric_columns(df)
    if len(numeric_cols) == 0:
        return False
    if "datetime" in df.columns and pd.to_datetime(df["datetime"], errors="coerce").notna().sum() >= 2:
        return True
    if "time_text" in df.columns and df["time_text"].astype(str).str.strip().ne("").sum() >= 2:
        return True
    if "date" in df.columns and pd.to_datetime(df["date"], errors="coerce").notna().sum() >= 2:
        return True
    return False


def is_usable_single_message_table(df: pd.DataFrame) -> bool:
    """Allow one-row WhatsApp/TMU messages while rejecting random text/PDF pages.

    WhatsApp reports are often one reading only, so they cannot pass the stricter
    repeated-time-series test.  They still must contain at least one numeric
    operational field and at least one usable date/time marker.
    """
    if df is None or df.empty:
        return False
    if not available_numeric_columns(df):
        return False
    has_dt = "datetime" in df.columns and pd.to_datetime(df["datetime"], errors="coerce").notna().any()
    has_date = "date" in df.columns and pd.to_datetime(df["date"], errors="coerce").notna().any()
    has_time = "time_text" in df.columns and df["time_text"].astype(str).str.strip().ne("").any()
    return bool(has_dt or has_date or has_time)


def filter_usable_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """Final safety filter for all upload types."""
    return [t for t in tables if is_valid_timeseries(t) or is_usable_single_message_table(t)]



def parse_expro_mpfm_text(text: str, source_name: str = "EXPRO_MPFM_PDF") -> pd.DataFrame:
    """Parse EXPRO MPFM PDF Data & Events text rows.

    Keeps only rows where the time is followed by the full numeric MPFM reading set.
    Skips event/comment rows such as 'BS&W is ...', 'bypassed the meter', etc.
    """
    if not text or "Data & Events" not in text or "QOil" not in text or "QWat" not in text:
        return pd.DataFrame()

    m = re.search(r"Well\s+(?:No\s+)?([A-Z0-9]+\s*-\s*\d+)", text, flags=re.I)
    if m:
        well = re.sub(r"\s*-\s*", "-", m.group(1).strip().upper())
    else:
        well = guess_well_from_name(source_name) or "Unknown"

    current_date = pd.NaT
    rows = []
    expro_cols = [
        "choke_size_64", "whp_psi", "flow_press_psi", "flow_temp_c",
        "mpfm_press_psig", "mpfm_temp_f", "dp_mbar",
        "qoil_s_stbd", "qwat_s_bpd", "qgas_s_mmscfd",
        "qoil_a_bpd", "qwat_a_bpd", "qgas_a_mmcfd",
        "wlr_s_pct", "qgross_s_bpd",
        "oil_sg", "water_sg", "water_ph", "salinity_kppm",
        "gas_sg", "co2_mole_pct", "h2s_ppm", "gor_s_scf_stb", "gvf_a_pct", "pump_freq_hz",
    ]

    date_pat = re.compile(r"^\s*(\d{1,2}/\d{1,2}/\d{2,4})\s*$")
    time_pat = re.compile(r"^\s*(\d{1,2}:\d{2}:\d{2})\s+(.*)$")

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        dm = date_pat.match(line)
        if dm:
            current_date = pd.to_datetime(dm.group(1), errors="coerce", dayfirst=True)
            continue

        tm = time_pat.match(line)
        if not tm:
            continue

        time_txt, rest = tm.group(1), tm.group(2).strip()
        if not re.match(r"^[+-]?\d", rest):
            continue

        nums = re.findall(r"[-+]?\d*\.\d+|[-+]?\d+", rest.replace(",", ""))
        if len(nums) < 25:
            continue

        vals = [float(x) for x in nums[:25]]
        row = dict(zip(expro_cols, vals))
        row["source"] = source_name
        row["sheet"] = "EXPRO_MPFM_Data_Events"
        row["well"] = well
        row["date"] = current_date
        row["time"] = time_txt
        row["time_text"] = time_txt[:5]

        t = parse_time_series(pd.Series([time_txt])).iloc[0]
        if pd.notna(current_date) and pd.notna(t):
            row["datetime"] = pd.Timestamp(current_date.date()) + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
        else:
            row["datetime"] = pd.NaT

        row["oil_rate_stbd"] = row["qoil_s_stbd"]
        row["water_rate_bpd"] = row["qwat_s_bpd"]
        row["gas_rate_mmscfd"] = row["qgas_s_mmscfd"]
        row["gross_rate_bpd"] = row["qgross_s_bpd"]
        row["bsw_pct"] = row["wlr_s_pct"]
        row["gor_scf_bbl"] = row["gor_s_scf_stb"]

        if pd.notna(row["salinity_kppm"]) and row["salinity_kppm"] > 1000:
            row["salinity_kppm"] = row["salinity_kppm"] / 1000.0

        rows.append(row)

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    df = df.drop_duplicates(subset=["well", "datetime", "source"], keep="first")
    df = df.sort_values(["well", "datetime"], na_position="last").reset_index(drop=True)
    return df



def filter_preferred_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """When a workbook has both summary/form sheets and detailed time-series sheets,
    keep the detailed time-series tables by default.
    """
    if len(tables) <= 1:
        return tables

    detailed = []
    for df in tables:
        if df is None or df.empty or len(df) < 3:
            continue
        dt_unique = df["datetime"].nunique(dropna=True) if "datetime" in df.columns else 0
        time_unique = df["time_text"].nunique(dropna=True) if "time_text" in df.columns else 0
        sheet_name = str(df["sheet"].iloc[0]).lower() if "sheet" in df.columns and len(df) else ""
        if max(dt_unique, time_unique) >= 3 and sheet_name not in {"form", "summary", "test summary"}:
            detailed.append(df)

    return detailed if detailed else tables

def load_tabular_file_base(uploaded_file) -> List[pd.DataFrame]:
    name = uploaded_file.name
    suffix = name.split(".")[-1].lower()
    tables: List[pd.DataFrame] = []

    if suffix in ["xlsx", "xls"]:
        xls = pd.ExcelFile(uploaded_file)
        for sheet in xls.sheet_names:
            raw = pd.read_excel(xls, sheet_name=sheet, header=None)
            default_well = extract_well_from_raw(raw, source_name=name, sheet_name=sheet)
            sheet_candidates = parse_excel_sheet_attempts(
                raw,
                source_name=name,
                sheet_name=sheet,
                default_well=default_well,
            )
            if sheet_candidates:
                # Keep only the best interpretation for each sheet to avoid duplicate plots.
                tables.append(sheet_candidates[0])

        tables = filter_preferred_tables(tables)

    elif suffix == "csv":
        # First try normal CSV headers. This is important for CSVs exported from this dashboard.
        uploaded_file.seek(0)
        try:
            header_df = pd.read_csv(uploaded_file)
            header_df = header_df.loc[:, ~header_df.columns.astype(str).str.match(r"^Unnamed")]
            known_header_hits = sum(1 for c in header_df.columns if best_canonical_name(str(c)) is not None)
            if known_header_hits >= 3:
                default_well = extract_well_from_raw(header_df, source_name=name, sheet_name="CSV")
                std = standardize_dataframe(header_df, source_name=name, sheet_name="CSV", default_well=default_well)
                if is_valid_timeseries(std):
                    tables.append(std)
        except Exception:
            pass

        if not tables:
            uploaded_file.seek(0)
            raw = pd.read_csv(uploaded_file, header=None)
            default_well = extract_well_from_raw(raw, source_name=name, sheet_name="CSV")
            table = table_from_raw(raw)
            std = standardize_dataframe(table, source_name=name, sheet_name="CSV", default_well=default_well)
            if is_valid_timeseries(std):
                tables.append(std)

    elif suffix == "txt":
        text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        msg_rows = parse_many_tmu_messages(text, source_name=name)
        if is_usable_single_message_table(msg_rows):
            tables.append(msg_rows)

    elif suffix == "docx":
        try:
            from docx import Document

            doc = Document(uploaded_file)
            text = "\n".join(safe_text(p.text) for p in doc.paragraphs)
            msg_rows = parse_many_tmu_messages(text, source_name=name)
            if is_usable_single_message_table(msg_rows):
                tables.append(msg_rows)

            for i, t in enumerate(doc.tables):
                rows = [[cell.text for cell in row.cells] for row in t.rows]
                raw = pd.DataFrame(rows)
                default_well = extract_well_from_raw(raw, source_name=name, sheet_name=f"DOCX_Table_{i + 1}")
                table = table_from_raw(raw)
                std = standardize_dataframe(table, source_name=name, sheet_name=f"DOCX_Table_{i + 1}", default_well=default_well)
                if is_valid_timeseries(std):
                    tables.append(std)
        except Exception as e:
            raise RuntimeError(f"Could not read DOCX file {name}: {e}")

    elif suffix == "pdf":
        try:
            import pdfplumber

            with pdfplumber.open(uploaded_file) as pdf:
                full_text = []
                for pi, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    full_text.append(text)

                    for ti, table in enumerate(page.extract_tables() or []):
                        raw = pd.DataFrame(table)
                        default_well = extract_well_from_raw(raw, source_name=name, sheet_name=f"PDF_Page_{pi + 1}_Table_{ti + 1}")
                        table_df = table_from_raw(raw)
                        std = standardize_dataframe(table_df, source_name=name, sheet_name=f"PDF_Page_{pi + 1}_Table_{ti + 1}", default_well=default_well)
                        if is_valid_timeseries(std):
                            tables.append(std)

                all_text = "\n".join(safe_text(x) for x in full_text)

                # Special support for EXPRO MPFM PDF reports where the Data & Events
                # table is embedded as text rather than a clean extractable table.
                # If this detailed parser succeeds, keep it and discard summary/event-only tables.
                expro_rows = parse_expro_mpfm_text(all_text, source_name=name)
                if not expro_rows.empty:
                    tables = [expro_rows]
                else:
                    msg_rows = parse_many_tmu_messages(all_text, source_name=name)
                    if is_usable_single_message_table(msg_rows):
                        tables.append(msg_rows)
                    tables = filter_preferred_tables(tables)
        except Exception as e:
            raise RuntimeError(f"Could not read PDF file {name}: {e}")

    return filter_usable_tables(tables)


def value_by_patterns(text: str, patterns: List[str]) -> Optional[str]:
    for p in patterns:
        pattern = rf"{p}\s*[:=@-]?\s*([^\n\r]+)"
        m = re.search(pattern, text, flags=re.I)
        if m:
            value = m.group(1).strip()
            value = re.split(
                r"\s{2,}(?=(choke|w\.?h\.?p|sep|gas|gross|oil|water|bs|salinity|h2s|co2|note|pumping)\b)",
                value,
                flags=re.I,
            )[0].strip()
            return value
    return None


def parse_salinity_to_kppm(value: object) -> float:
    if value is None:
        return np.nan

    s = str(value).replace(",", "")
    num = extract_number(s)
    if pd.isna(num):
        return np.nan
    if re.search(r"\bk\b|kppm|k ppm", s, flags=re.I):
        return num
    if num > 1000:
        return num / 1000.0
    return num


def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    text = message.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)

    row: Dict[str, object] = {"source": source_name, "sheet": "WhatsApp_Text"}

    tmu = re.search(r"\b(PICO\s*T\s*MU[-\s]*\d+|PICO\s*TMU[-\s]*\d+|TMU[-\s]*\d+)\b", text, flags=re.I)
    if tmu:
        row["test_unit"] = re.sub(r"\s+", " ", tmu.group(1).upper()).replace("T MU", "TMU")

    date_v = value_by_patterns(text, [r"\bdate\b"])
    time_v = value_by_patterns(text, [r"\btime\s*@?", r"\btime\b"])
    well_v = value_by_patterns(text, [r"\bwell\s*name\b", r"\bwell\b"])

    if well_v:
        row["well"] = well_v.split()[0].strip()
    else:
        row["well"] = guess_well_from_name(text) or "Unknown"

    if date_v:
        row["date"] = pd.to_datetime(date_v, errors="coerce", dayfirst=True)
    if time_v:
        row["time"] = time_v

    fields = {
        "choke_pct": [r"\bchoke\b"],
        "whp_psi": [r"\bw\.?\s*h\.?\s*p\.?\b", r"\bwellhead pressure\b"],
        "sep_p_psi": [r"\bsep\.?\s*p\.?\b", r"\bseparator pressure\b"],
        "gas_rate_mmscfd": [r"\bgas rate\b"],
        "gas_formation_mmscfd": [r"\bgas formation\b", r"\bformation gas\b"],
        "gross_rate_bpd": [r"\bgross rate\b", r"\bgross\b"],
        "oil_rate_stbd": [r"\boil rate\b"],
        "water_rate_bpd": [r"\bwater rate\b"],
        "bsw_pct": [r"\bbs\s*&\s*w\b", r"\bbsw\b", r"\bwc\b", r"\bwater cut\b"],
        "salinity_kppm": [r"\bsalinity\b"],
        "h2s_ppm": [r"\bh2s\b"],
        "co2_mole_pct": [r"\bco2\b", r"\bco₂\b"],
        "water_cum_bbl": [r"\bwater cum\b", r"\bwater cumulative\b"],
        "pumping_pressure_psi": [r"\bpumping\s*p\b", r"\bpumping pressure\b", r"\bpump pressure\b"],
        "n2_rate_scfm": [r"\bn2 standard rate\b", r"\bn2 rate\b", r"\bnitrogen rate\b"],
        "note": [r"\bnote\b"],
    }

    for canon, pats in fields.items():
        raw_val = value_by_patterns(text, pats)
        if raw_val is None:
            continue
        if canon == "note":
            row[canon] = raw_val.strip()
        elif canon == "salinity_kppm":
            row[canon] = parse_salinity_to_kppm(raw_val)
        else:
            row[canon] = extract_number(raw_val)

    if "date" in row and "time" in row:
        d = pd.to_datetime(row["date"], errors="coerce", dayfirst=True)
        tser = parse_time_series(pd.Series([row["time"]]))
        if pd.notna(d) and pd.notna(tser.iloc[0]):
            t = tser.iloc[0]
            row["datetime"] = pd.Timestamp(d.date()) + pd.Timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
            row["time_text"] = f"{t.hour:02d}:{t.minute:02d}"

    return row


def split_messages(text: str) -> List[str]:
    markers = list(re.finditer(r"(?=PICO\s*T?\s*MU|TMU[-\s]*\d+|Date\s*:)", text, flags=re.I))
    if len(markers) <= 1:
        return [text]

    chunks = []
    for i, m in enumerate(markers):
        start = m.start()
        end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

    return chunks


def parse_many_tmu_messages(text: str, source_name: str = "WhatsApp_Text") -> pd.DataFrame:
    rows = []
    for chunk in split_messages(text):
        row = parse_tmu_message(chunk, source_name=source_name)
        has_rate = any(
            k in row
            for k in [
                "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi",
                "pumping_pressure_psi", "gas_rate_mmscfd", "bsw_pct",
            ]
        )
        if has_rate:
            rows.append(row)

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    for col in ["source", "sheet", "well", "datetime", "date", "time_text"]:
        if col not in df.columns:
            df[col] = np.nan

    return df


def apply_fill_method(df: pd.DataFrame, numeric_cols: List[str], method: str) -> pd.DataFrame:
    if df.empty or not numeric_cols or method == "No fill":
        return df

    out = df.sort_values(["well", "datetime"], na_position="last").copy()

    if method == "Forward fill":
        out[numeric_cols] = out.groupby("well", dropna=False)[numeric_cols].ffill()

    elif method == "Forward + backward fill":
        out[numeric_cols] = out.groupby("well", dropna=False)[numeric_cols].ffill().bfill()

    elif method == "Linear interpolation by row":
        out[numeric_cols] = out.groupby("well", dropna=False)[numeric_cols].transform(
            lambda x: x.interpolate(limit_direction="both")
        )

    elif method == "Time interpolation":
        pieces = []
        for _, g in out.groupby("well", dropna=False):
            g = g.copy()
            if "datetime" in g.columns and g["datetime"].notna().sum() >= 2:
                g = g.set_index("datetime")
                g[numeric_cols] = g[numeric_cols].interpolate(method="time", limit_direction="both")
                g = g.reset_index()
            else:
                g[numeric_cols] = g[numeric_cols].interpolate(limit_direction="both")
            pieces.append(g)
        out = pd.concat(pieces, ignore_index=True)

    return out


def apply_user_column_mappings(df: pd.DataFrame, mapping: Optional[Dict[str, str]] = None) -> pd.DataFrame:
    """Apply user-confirmed column mappings after auto parsing.

    mapping keys may be actual DataFrame column names (for example raw__amp) or
    normalized alias keys from canonical_key(). Values are canonical parser names
    such as motor_current_amp, pump_intake_pressure_psi, or the special value
    __drop__ to hide a column.

    If target already exists, non-null values from the source fill gaps only; this
    prevents a manual alias from overwriting a stronger auto-detected column.
    """
    if df is None or df.empty or not mapping:
        return df

    out = df.copy()
    options = standard_column_options(include_meta=False)
    drop_cols = []

    for col in list(out.columns):
        key_candidates = [str(col), canonical_key(col)]
        target = None
        for key in key_candidates:
            if key in mapping:
                target = mapping.get(key)
                break
        if not target or target in {"__keep__", "Keep as-is"}:
            continue
        if target == "__drop__":
            drop_cols.append(col)
            continue
        if target not in options:
            continue
        if col in BASE_NON_PLOT_COLS:
            continue

        vals = clean_numeric_series(out[col], target)
        if target in out.columns:
            out[target] = pd.to_numeric(out[target], errors="coerce").combine_first(vals)
        else:
            out[target] = vals
        if col != target:
            drop_cols.append(col)

    if drop_cols:
        out = out.drop(columns=[c for c in drop_cols if c in out.columns], errors="ignore")
    return out


def available_numeric_columns(df: pd.DataFrame) -> List[str]:
    return [
        c
        for c in df.columns
        if c not in BASE_NON_PLOT_COLS
        and pd.api.types.is_numeric_dtype(df[c])
        and df[c].notna().any()
    ]


def column_label(column_name: str) -> str:
    if str(column_name).startswith("raw__"):
        label = str(column_name)[5:].replace("_", " ").strip()
        return f"Raw: {label.title()}"
    return COLUMN_LABELS.get(column_name, column_name)

# -----------------------------------------------------------------------------
# v43 additions: WhatsApp export ZIP + CTU image OCR + safe test segmentation
# -----------------------------------------------------------------------------

# Keep CTU / OCR metadata out of normal numeric detection and plotting unless the
# value columns themselves are selected.
BASE_NON_PLOT_COLS.update({
    "source_type", "ocr_template", "ocr_fields_found", "ocr_status", "ocr_confidence",
    "image_file", "attachment_name", "chat_sender", "chat_datetime", "message_index",
    "test_id", "test_start", "test_end", "test_sequence", "link_status",
    "suggested_well", "suggested_test_id", "suggested_link_reason", "suggested_link_gap_hours",
    "review_required", "caption_text", "whatsapp_message_body", "source_member",
})

COLUMN_LABELS.update({
    "ctu_weight_lbf": "CTU Weight (LBF)",
    "ctu_lt_weight_lbf": "CTU Lt Weight (LBF)",
    "ctu_wellhead_pressure_psi": "CTU Wellhead Pressure (psi)",
    "ctu_circulation_pressure_psi": "CTU Circulation Pressure (psi)",
    "ctu_reel_depth_ft": "CTU Reel Depth (ft)",
    "ctu_reel_speed_ftmin": "CTU Reel Speed (ft/min)",
    "ctu_fluid_rate_bpm": "CTU Fluid Rate (bpm)",
    "ctu_n2_rate_scfm": "CTU N2 Flow (scf/min)",
    "ctu_fluid_total_bbl": "CTU Fluid Total (bbl)",
    "ctu_n2_total_scf": "CTU N2 Total (scf)",
})

for _kw in [
    "circulation", "reel", "lt weight", "fluid total", "n2 total", "flare test",
    "wellhead pressure", "all data", "ctu", "coil tubing", "coiled tubing",
]:
    if _kw not in KEYWORDS:
        KEYWORDS.append(_kw)

IMAGE_SUFFIXES = {"jpg", "jpeg", "png"}
DATA_SUFFIXES = {"xlsx", "xls", "csv", "docx", "pdf", "txt"}


class UploadedBytes(io.BytesIO):
    """BytesIO object that behaves like a Streamlit UploadedFile for recursive parsers."""

    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


def append_note(old_note, new_note):
    old = "" if old_note is None or pd.isna(old_note) else str(old_note).strip()
    new = str(new_note or "").strip()
    if not new:
        return old
    if not old:
        return new
    if new.lower() in old.lower():
        return old
    return f"{old}; {new}"


def parse_datetime_from_filename(name: str):
    """Parse WhatsApp media filenames such as PHOTO-2026-06-12-20-35-05.jpg."""
    name = str(name or "")
    patterns = [
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(\d{2})[-_](\d{2})[-_](20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
    ]
    for p in patterns:
        m = re.search(p, name)
        if not m:
            continue
        nums = list(map(int, m.groups()))
        try:
            if len(nums) == 6 and str(nums[0]).startswith("20"):
                y, mo, d, hh, mm, ss = nums
            elif len(nums) == 5:
                y, mo, d, hh, mm = nums
                ss = 0
            else:
                d, mo, y, hh, mm, ss = nums
            return pd.Timestamp(year=y, month=mo, day=d, hour=hh, minute=mm, second=ss)
        except Exception:
            pass
    return pd.NaT


def parse_attachment_reference(body: str) -> str:
    """Return attached filename mentioned in exported WhatsApp text, if any."""
    text = str(body or "")
    patterns = [
        r"<attached:\s*([^>]+)>",
        r"([^\s<>]+\.(?:jpg|jpeg|png|xlsx|xls|csv|pdf|docx))",
    ]
    for p in patterns:
        m = re.search(p, text, flags=re.I)
        if m:
            return Path(m.group(1).strip()).name
    return ""


def _try_import_ocr_libs():
    try:
        from PIL import Image, ImageOps, ImageEnhance
        import pytesseract
        return Image, ImageOps, ImageEnhance, pytesseract
    except Exception:
        return None, None, None, None


def _numbers_from_ocr_text(txt: str) -> List[str]:
    return re.findall(r"-?\d+(?:\.\d+)?", str(txt or "").replace(" ", ""))


def _best_number_from_ocr_text(txt: str):
    nums = _numbers_from_ocr_text(txt)
    if not nums:
        return np.nan
    # Prefer the longest candidate; this avoids choosing a stray unit digit.
    nums = sorted(nums, key=lambda x: len(x.replace(".", "").replace("-", "")), reverse=True)
    return extract_number(nums[0])


def _ocr_numeric_region_pil(crop_img):
    """OCR a numeric value from a CTU/PICO HMI ROI using multiple simple preprocesses."""
    Image, ImageOps, ImageEnhance, pytesseract = _try_import_ocr_libs()
    if Image is None:
        return np.nan, 0.0, "ocr_dependency_missing"

    best_val = np.nan
    best_score = -1
    best_text = ""

    try:
        gray = crop_img.convert("L")
        gray = ImageOps.autocontrast(gray)
    except Exception:
        return np.nan, 0.0, "ocr_image_error"

    variants = []
    for invert in [False, True]:
        g0 = ImageOps.invert(gray) if invert else gray
        for contrast in [2.0, 3.0, 5.0]:
            g = ImageEnhance.Contrast(g0).enhance(contrast)
            g = ImageEnhance.Sharpness(g).enhance(2.0)
            g = g.resize((max(90, g.width * 4), max(40, g.height * 4)))
            variants.append(g)
            for thr in [110, 140, 170, 200]:
                variants.append(g.point(lambda p, t=thr: 255 if p > t else 0))

    for img_try in variants:
        try:
            txt = pytesseract.image_to_string(
                img_try,
                config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.",
            )
        except Exception:
            continue

        nums = _numbers_from_ocr_text(txt)
        if not nums:
            continue
        val = _best_number_from_ocr_text(txt)
        digit_count = max(len(n.replace(".", "").replace("-", "")) for n in nums)
        score = digit_count + (2 if any("." in n for n in nums) else 0)
        if pd.notna(val) and score > best_score:
            best_score = score
            best_val = val
            best_text = str(txt).strip()

    confidence = min(1.0, max(0.0, best_score / 8.0)) if best_score >= 0 else 0.0
    return best_val, confidence, best_text


# ROIs are value boxes only, relative to the full image. They fit the CTU "ALL DATA"
# screen style in the user's sample; all rows remain review_required unless safely linked.


def _ocr_numeric_region_cv2(crop_img):
    """Optional OpenCV OCR path that improves low-contrast colored HMI digits."""
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return np.nan, 0.0, "cv2_or_tesseract_missing"

    try:
        arr = _np.array(crop_img.convert("RGB"))
        bgr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
        bgr = cv2.resize(bgr, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
        lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        gray = clahe.apply(l)
        variants = [gray, 255 - gray]
        for block in [31, 51, 71]:
            variants.append(cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block, 3))
            variants.append(cv2.adaptiveThreshold(255 - gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block, 3))

        best_val = np.nan
        best_score = -1
        best_text = ""
        for v in variants:
            try:
                txt = pytesseract.image_to_string(
                    v,
                    config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.",
                )
            except Exception:
                continue
            nums = _numbers_from_ocr_text(txt)
            if not nums:
                continue
            val = _best_number_from_ocr_text(txt)
            digit_count = max(len(n.replace(".", "").replace("-", "")) for n in nums)
            score = digit_count + (2 if any("." in n for n in nums) else 0)
            if pd.notna(val) and score > best_score:
                best_score = score
                best_val = val
                best_text = str(txt).strip()
        confidence = min(1.0, max(0.0, best_score / 8.0)) if best_score >= 0 else 0.0
        return best_val, confidence, best_text
    except Exception:
        return np.nan, 0.0, "cv2_ocr_error"


def _ocr_numeric_region(crop_img):
    """Use PIL and OpenCV OCR paths, then choose the higher-confidence numeric value."""
    pil_val, pil_conf, pil_text = _ocr_numeric_region_pil(crop_img)
    cv_val, cv_conf, cv_text = _ocr_numeric_region_cv2(crop_img)
    if pd.notna(cv_val) and (pd.isna(pil_val) or cv_conf >= pil_conf):
        return cv_val, cv_conf, cv_text
    return pil_val, pil_conf, pil_text

CTU_ALL_DATA_ROIS = {
    "ctu_weight_lbf": (0.26, 0.13, 0.55, 0.30),
    "ctu_lt_weight_lbf": (0.66, 0.13, 0.92, 0.30),
    "ctu_wellhead_pressure_psi": (0.31, 0.31, 0.55, 0.47),
    "ctu_circulation_pressure_psi": (0.64, 0.31, 0.92, 0.47),
    "ctu_reel_depth_ft": (0.31, 0.49, 0.55, 0.66),
    "ctu_reel_speed_ftmin": (0.62, 0.49, 0.92, 0.66),
    "ctu_fluid_rate_bpm": (0.31, 0.66, 0.55, 0.80),
    "ctu_n2_rate_scfm": (0.64, 0.66, 0.92, 0.80),
    "ctu_fluid_total_bbl": (0.31, 0.81, 0.55, 0.95),
    "ctu_n2_total_scf": (0.64, 0.81, 0.92, 0.95),
}


def parse_ctu_all_data_screen_image(uploaded_file, source_name="Image_OCR") -> pd.DataFrame:
    """Parse CTU/PICO ALL DATA screen photos as auxiliary OCR rows.

    Safety rule: OCR rows are never treated as confirmed well-test readings by
    themselves. If the image does not explicitly carry a well/date, the row is
    marked review_required and remains Well=Unknown until the app/user approves
    linking. This avoids silently assigning CTU data to the wrong test.
    """
    Image, _, _, _ = _try_import_ocr_libs()
    if Image is None:
        return pd.DataFrame()

    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        img = Image.open(uploaded_file).convert("RGB")
        w, h = img.size
        file_name = getattr(uploaded_file, "name", source_name)
        dt_from_name = parse_datetime_from_filename(file_name)

        row = {
            "source": source_name,
            "sheet": "CTU_Image_OCR",
            "source_type": "ctu_image_ocr",
            "ocr_template": "ctu_all_data",
            "image_file": file_name,
            "well": "Unknown",
            "link_status": "unlinked_needs_review",
            "review_required": True,
        }
        if pd.notna(dt_from_name):
            row["datetime"] = dt_from_name
            row["date"] = pd.Timestamp(dt_from_name.date())
            row["time_text"] = pd.Timestamp(dt_from_name).strftime("%H:%M")

        fields_found = 0
        confidences = []
        for field, (x1, y1, x2, y2) in CTU_ALL_DATA_ROIS.items():
            crop = img.crop((int(x1 * w), int(y1 * h), int(x2 * w), int(y2 * h)))
            val, conf, _raw = _ocr_numeric_region(crop)
            if pd.notna(val):
                row[field] = float(val)
                fields_found += 1
                confidences.append(float(conf))

        row["ocr_fields_found"] = fields_found
        row["ocr_confidence"] = float(np.mean(confidences)) if confidences else 0.0
        row["ocr_status"] = "parsed_review_required" if fields_found >= 2 else "low_confidence_or_not_ctu_screen"

        if fields_found < 1:
            return pd.DataFrame()
        return pd.DataFrame([row])
    except Exception:
        return pd.DataFrame()


WHATSAPP_SYSTEM_PATTERNS = [
    r"messages and calls are end-to-end encrypted",
    r"joined using this group",
    r"changed the subject",
    r"changed the group description",
    r"changed this group's icon",
    r"this message was deleted",
    r"missed voice call",
    r"missed video call",
    r"created group",
]


def is_system_or_noise_message(body: str) -> bool:
    b = normalize_text(body)
    if not b:
        return True
    for p in WHATSAPP_SYSTEM_PATTERNS:
        if re.search(p, b, flags=re.I):
            return True
    # Very short acknowledgement messages are not useful data.
    if b in {"ok", "okay", "thanks", "thank you", "done", "تمام", "اوكي"}:
        return True
    return False


def score_tmu_body(body: str) -> int:
    b = normalize_text(body)
    if is_system_or_noise_message(b):
        return 0
    score = 0
    for kw in [
        "tmu", "well", "date", "time", "choke", "w.h.p", "whp", "sep",
        "gas rate", "gross rate", "oil rate", "water rate", "bsw", "bs&w",
        "salinity", "h2s", "co2", "pumping", "n2", "flare test", "all data",
        "circulation pressure", "reel depth", "wellhead pressure",
    ]:
        if kw in b:
            score += 1
    score += min(5, len(re.findall(r"[-+]?\d+(?:\.\d+)?", body)))
    return score


def parse_whatsapp_export_messages(text: str) -> List[Dict[str, object]]:
    """Parse WhatsApp exported _chat.txt into message dicts.

    Handles common Android/iOS formats:
      12/06/2026, 00:30 - Sender: body
      [12/06/2026, 00:30:05] Sender: body
    Multi-line messages are preserved.
    """
    lines = str(text or "").replace("\ufeff", "").splitlines()
    messages = []
    current = None

    patterns = [
        r"^(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{2,4}),\s*(?P<time>\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\s*-\s*(?P<sender>[^:]+):\s*(?P<body>.*)$",
        r"^\[(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{2,4}),\s*(?P<time>\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\]\s*(?P<sender>[^:]+):\s*(?P<body>.*)$",
        r"^(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\s+(?P<time>\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\s*-\s*(?P<sender>[^:]+):\s*(?P<body>.*)$",
    ]

    for line in lines:
        line = line.rstrip()
        if not line.strip():
            continue

        matched = False
        for p in patterns:
            m = re.match(p, line)
            if m:
                if current:
                    messages.append(current)
                current = m.groupdict()
                current["sender"] = str(current.get("sender", "")).strip()
                current["body"] = str(current.get("body", "")).strip()
                matched = True
                break

        if not matched and current:
            current["body"] = str(current.get("body", "")) + "\n" + line.strip()

    if current:
        messages.append(current)

    for i, m in enumerate(messages):
        m["message_index"] = i
        m["datetime"] = pd.to_datetime(
            f"{m.get('date', '')} {m.get('time', '')}",
            errors="coerce",
            dayfirst=True,
        )
        m["attachment_name"] = parse_attachment_reference(m.get("body", ""))

    return messages


def parse_whatsapp_export_text(text: str, source_name="WhatsApp_Export") -> pd.DataFrame:
    """Parse exported WhatsApp chat text into TMU production rows.

    Important: this parser uses well name and explicit TMU keywords. It does not
    assume one chat equals one test.
    """
    messages = parse_whatsapp_export_messages(text)
    if not messages:
        return pd.DataFrame()

    rows = []
    for m in messages:
        body = str(m.get("body", ""))
        if score_tmu_body(body) < 4:
            continue

        row = parse_tmu_message(body, source_name=source_name)
        # If a message is only a continuation line, parse_tmu_message may not get
        # its own date/time. Use WhatsApp timestamp as fallback only.
        msg_dt = m.get("datetime", pd.NaT)
        if ("datetime" not in row or pd.isna(row.get("datetime", pd.NaT))) and pd.notna(msg_dt):
            row["datetime"] = msg_dt
            row["date"] = pd.Timestamp(msg_dt).floor("D")
            row["time_text"] = pd.Timestamp(msg_dt).strftime("%H:%M")

        row["source_type"] = "whatsapp_export_text"
        row["chat_sender"] = m.get("sender", "")
        row["chat_datetime"] = msg_dt
        row["message_index"] = m.get("message_index", np.nan)
        row["attachment_name"] = m.get("attachment_name", "")
        row["whatsapp_message_body"] = body[:500]
        row["link_status"] = "text_confirmed_by_well" if row.get("well") and str(row.get("well")).lower() != "unknown" else "text_needs_well_review"

        has_useful = any(
            k in row for k in [
                "gross_rate_bpd", "oil_rate_stbd", "water_rate_bpd", "whp_psi",
                "sep_p_psi", "pumping_pressure_psi", "gas_rate_mmscfd", "bsw_pct",
            ]
        )
        if has_useful:
            rows.append(row)

    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def parse_whatsapp_plain_or_export_text(text: str, source_name="WhatsApp_Text") -> pd.DataFrame:
    """Try exported-chat parsing first; fallback to pasted/block TMU parser."""
    export_df = parse_whatsapp_export_text(text, source_name=source_name)
    if not export_df.empty:
        return export_df
    df = parse_many_tmu_messages(text, source_name=source_name)
    if not df.empty:
        df["source_type"] = "pasted_whatsapp_text"
        df["link_status"] = "text_confirmed_by_well"
    return df


def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    """Assign stable test_id by well name + time gap.

    Rule requested by user:
      - Same well continues the same test until the time gap exceeds gap_hours.
      - A different well is always a different test stream.
      - Unknown/unlinked OCR rows are NOT silently assigned to a well/test.
    """
    if df is None or df.empty:
        return df

    out = df.copy()
    if "well" not in out.columns:
        out["well"] = "Unknown"
    if "datetime" not in out.columns:
        out["datetime"] = pd.NaT

    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["well"] = out["well"].fillna("Unknown").astype(str).str.strip().replace("", "Unknown")
    out["test_id"] = out.get("test_id", pd.Series([np.nan] * len(out), index=out.index))
    out["test_sequence"] = out.get("test_sequence", pd.Series([np.nan] * len(out), index=out.index))

    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}

    for well, g in sortable.groupby("well", dropna=False):
        well_txt = str(well or "Unknown").strip() or "Unknown"
        if well_txt.lower() == "unknown":
            # Keep unknown rows separate/unlinked. They may be reviewed in the app.
            for i in g.index:
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well", np.nan)
            continue

        seq = 0
        last_dt = pd.NaT
        current_id = None
        current_start = pd.NaT
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                # Same well but no time: keep as separate manual-review row.
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, seq)
                continue

            new_test = (
                current_id is None
                or pd.isna(last_dt)
                or (pd.Timestamp(dt) - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            )
            if new_test:
                seq += 1
                current_start = pd.Timestamp(dt)
                current_id = f"{well_txt}_{current_start.strftime('%Y%m%d_%H%M')}"

            assigned[i] = (current_id, seq)
            last_dt = pd.Timestamp(dt)

    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = tid
        out.at[i, "test_sequence"] = seq

    # Add test start/end for confirmed test rows.
    if "test_id" in out.columns and "datetime" in out.columns:
        valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
        if not valid.empty:
            starts = valid.groupby("test_id")["datetime"].min()
            ends = valid.groupby("test_id")["datetime"].max()
            out["test_start"] = out["test_id"].map(starts)
            out["test_end"] = out["test_id"].map(ends)

    return out


def suggest_links_for_ocr_rows(df: pd.DataFrame, max_gap_hours: float = 3.0) -> pd.DataFrame:
    """Suggest, but do not apply, well/test links for CTU image rows.

    A suggestion is created only when there is one nearest confirmed non-OCR row
    within max_gap_hours. The row remains review_required; app/user must approve
    before it becomes part of the test plot.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    for c in ["source_type", "well", "datetime", "test_id", "suggested_well", "suggested_test_id", "suggested_link_reason", "suggested_link_gap_hours"]:
        if c not in out.columns:
            out[c] = np.nan
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")

    anchor_mask = (
        out["datetime"].notna()
        & out["well"].notna()
        & (out["well"].astype(str).str.strip().str.lower() != "unknown")
        & (~out["source_type"].astype(str).str.contains("ocr", case=False, na=False))
    )
    anchors = out[anchor_mask].copy()
    if anchors.empty:
        return out

    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    for i, row in out[ocr_mask].iterrows():
        dt = row.get("datetime", pd.NaT)
        if pd.isna(dt):
            continue
        deltas = (anchors["datetime"] - pd.Timestamp(dt)).abs()
        if deltas.empty:
            continue
        nearest_idx = deltas.idxmin()
        gap_h = float(deltas.loc[nearest_idx].total_seconds() / 3600.0)
        if gap_h <= float(max_gap_hours):
            out.at[i, "suggested_well"] = anchors.at[nearest_idx, "well"]
            out.at[i, "suggested_test_id"] = anchors.at[nearest_idx, "test_id"]
            out.at[i, "suggested_link_gap_hours"] = round(gap_h, 3)
            out.at[i, "suggested_link_reason"] = f"Nearest confirmed text/Excel row within {gap_h:.2f} hr; review before approving."
            out.at[i, "link_status"] = "suggested_needs_user_approval"
            out.at[i, "review_required"] = True
    return out


def approve_suggested_ocr_links(df: pd.DataFrame) -> pd.DataFrame:
    """Apply suggested OCR links after explicit user approval in Streamlit."""
    if df is None or df.empty:
        return df
    out = df.copy()
    required = {"suggested_well", "suggested_test_id", "source_type"}
    if not required.issubset(set(out.columns)):
        return out
    mask = (
        out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
        & out["suggested_well"].notna()
        & out["suggested_test_id"].notna()
    )
    out.loc[mask, "well"] = out.loc[mask, "suggested_well"]
    out.loc[mask, "test_id"] = out.loc[mask, "suggested_test_id"]
    out.loc[mask, "link_status"] = "ocr_link_approved"
    out.loc[mask, "review_required"] = False
    return out


def load_tabular_file(uploaded_file, parse_images: bool = True, max_ocr_images: int = 100) -> List[pd.DataFrame]:
    """v43 upload router.

    Supports existing Excel/CSV/DOCX/PDF/TXT parsing plus:
      - WhatsApp exported ZIP bundles (_chat.txt + attachments + images)
      - Direct CTU screen image OCR (JPG/JPEG/PNG)

    OCR safety: CTU image rows are parsed as auxiliary rows and left unlinked until
    reviewed/approved in the app. No nearest-well fill is applied automatically.
    """
    name = getattr(uploaded_file, "name", "uploaded")
    suffix = Path(str(name)).suffix.lower().lstrip(".")

    if suffix in IMAGE_SUFFIXES:
        if not parse_images:
            return []
        return filter_usable_tables([parse_ctu_all_data_screen_image(uploaded_file, source_name=name)])

    if suffix == "zip":
        try:
            raw = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
            tables: List[pd.DataFrame] = []
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                members = [m for m in zf.namelist() if not m.endswith("/") and not Path(m).name.startswith("._")]

                # Build chat-message index first so OCR/image rows can get timestamps and suggestions.
                all_messages: List[Dict[str, object]] = []
                chat_names = []
                for member in members:
                    member_name = Path(member).name
                    if Path(member_name).suffix.lower().lstrip(".") == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        text = zf.read(member).decode("utf-8", errors="ignore")
                        chat_names.append(member_name)
                        all_messages.extend(parse_whatsapp_export_messages(text))
                        df = parse_whatsapp_export_text(text, source_name=f"{name}:{member_name}")
                        if not df.empty:
                            df["source_member"] = member
                            tables.append(df)

                attachment_context = {}
                for m in all_messages:
                    att = str(m.get("attachment_name", "") or "").strip()
                    if not att:
                        continue
                    attachment_context[Path(att).name] = m

                # Parse non-chat attachments and OCR images.
                for member in members:
                    member_name = Path(member).name
                    ext = Path(member_name).suffix.lower().lstrip(".")
                    if not member_name or ext not in (DATA_SUFFIXES | IMAGE_SUFFIXES):
                        continue
                    if ext == "txt" and ("chat" in member_name.lower() or "_chat" in member_name.lower()):
                        continue

                    sub_file = UploadedBytes(zf.read(member), member_name)
                    if ext in IMAGE_SUFFIXES:
                        if not parse_images or max_ocr_images <= 0:
                            continue
                        max_ocr_images -= 1
                    sub_tables = load_tabular_file(sub_file, parse_images=parse_images, max_ocr_images=max_ocr_images)
                    ctx = attachment_context.get(member_name, {})
                    for t in sub_tables or []:
                        if t is None or t.empty:
                            continue
                        t = t.copy()
                        t["attachment_name"] = member_name
                        t["source_member"] = member
                        if ctx:
                            t["chat_sender"] = ctx.get("sender", "")
                            t["chat_datetime"] = ctx.get("datetime", pd.NaT)
                            t["message_index"] = ctx.get("message_index", np.nan)
                            if "datetime" not in t.columns or pd.to_datetime(t["datetime"], errors="coerce").isna().all():
                                if pd.notna(ctx.get("datetime", pd.NaT)):
                                    t["datetime"] = ctx.get("datetime")
                                    t["date"] = pd.Timestamp(ctx.get("datetime")).floor("D")
                                    t["time_text"] = pd.Timestamp(ctx.get("datetime")).strftime("%H:%M")
                            if ext in IMAGE_SUFFIXES:
                                t["caption_text"] = str(ctx.get("body", ""))[:500]
                        tables.append(t)

            if not tables:
                return []
            merged = pd.concat(tables, ignore_index=True, sort=False)
            merged = assign_test_ids(merged, gap_hours=12.0)
            # No nearest-time OCR linking here. CTU/OCR image rows remain unlinked
            # unless the user manually selects the correct Well/Test in Streamlit.
            return filter_usable_tables([merged])
        except Exception as e:
            raise RuntimeError(f"Could not read WhatsApp ZIP {name}: {e}")

    if suffix == "txt":
        try:
            text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        except Exception:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        df = parse_whatsapp_plain_or_export_text(text, source_name=name)
        if not df.empty:
            df = assign_test_ids(df, gap_hours=12.0)
            return filter_usable_tables([df])
        return []

    tables = load_tabular_file_base(uploaded_file)
    out_tables = []
    for t in tables or []:
        if t is not None and not t.empty:
            t = assign_test_ids(t, gap_hours=12.0)
            out_tables.append(t)
    return filter_usable_tables(out_tables)

# v43.1 safety refinements: allow direct OCR rows even if no timestamp; keep them unlinked.
def parse_datetime_from_filename(name: str):
    """Parse common WhatsApp media filenames, including 'WhatsApp Image 2026-06-13 at 15.29.01.jpeg'."""
    name = str(name or "")
    patterns = [
        r"(20\d{2})[-_](\d{2})[-_](\d{2})\s+at\s+(\d{1,2})[.:\-_](\d{2})[.:\-_](\d{2})",
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
        r"(\d{2})[-_](\d{2})[-_](20\d{2})[-_](\d{2})[-_](\d{2})[-_](\d{2})",
    ]
    for p in patterns:
        m = re.search(p, name, flags=re.I)
        if not m:
            continue
        nums = list(map(int, m.groups()))
        try:
            if len(nums) == 6 and str(nums[0]).startswith("20"):
                y, mo, d, hh, mm, ss = nums
            elif len(nums) == 5:
                y, mo, d, hh, mm = nums
                ss = 0
            else:
                d, mo, y, hh, mm, ss = nums
            return pd.Timestamp(year=y, month=mo, day=d, hour=hh, minute=mm, second=ss)
        except Exception:
            pass
    return pd.NaT


def is_usable_ocr_table(df: pd.DataFrame) -> bool:
    if df is None or df.empty:
        return False
    if "source_type" not in df.columns:
        return False
    if not df["source_type"].astype(str).str.contains("ocr", case=False, na=False).any():
        return False
    return any(c in df.columns and pd.to_numeric(df[c], errors="coerce").notna().any() for c in COLUMN_LABELS if str(c).startswith("ctu_"))


def filter_usable_tables(tables: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """Final safety filter for all upload types, including direct CTU OCR rows."""
    return [t for t in tables if is_valid_timeseries(t) or is_usable_single_message_table(t) or is_usable_ocr_table(t)]

# v43.2 optimized OCR overrides: fewer OCR passes for Streamlit performance.
def _ocr_numeric_region_pil(crop_img):
    Image, ImageOps, ImageEnhance, pytesseract = _try_import_ocr_libs()
    if Image is None:
        return np.nan, 0.0, "ocr_dependency_missing"
    try:
        gray = ImageOps.autocontrast(crop_img.convert("L"))
        g = ImageEnhance.Contrast(gray).enhance(3.0)
        g = ImageEnhance.Sharpness(g).enhance(2.0)
        g = g.resize((max(90, g.width * 4), max(40, g.height * 4)))
        variants = [g, ImageOps.invert(g), g.point(lambda p: 255 if p > 150 else 0)]
        best = (-1, np.nan, "")
        for im in variants:
            try:
                txt = pytesseract.image_to_string(im, config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.")
            except Exception:
                continue
            nums = _numbers_from_ocr_text(txt)
            if not nums:
                continue
            val = _best_number_from_ocr_text(txt)
            score = max(len(n.replace(".", "").replace("-", "")) for n in nums) + (2 if any("." in n for n in nums) else 0)
            if pd.notna(val) and score > best[0]:
                best = (score, val, str(txt).strip())
        conf = min(1.0, max(0.0, best[0] / 8.0)) if best[0] >= 0 else 0.0
        return best[1], conf, best[2]
    except Exception:
        return np.nan, 0.0, "ocr_image_error"


def _ocr_numeric_region_cv2(crop_img):
    try:
        import cv2
        import numpy as _np
        import pytesseract
    except Exception:
        return np.nan, 0.0, "cv2_or_tesseract_missing"
    try:
        arr = _np.array(crop_img.convert("RGB"))
        bgr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
        bgr = cv2.resize(bgr, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
        lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
        l, _, _ = cv2.split(lab)
        gray = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8)).apply(l)
        variants = [gray, 255 - gray, cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 51, 3)]
        best = (-1, np.nan, "")
        for im in variants:
            try:
                txt = pytesseract.image_to_string(im, config="--oem 3 --psm 7 -c tessedit_char_whitelist=-0123456789.")
            except Exception:
                continue
            nums = _numbers_from_ocr_text(txt)
            if not nums:
                continue
            val = _best_number_from_ocr_text(txt)
            score = max(len(n.replace(".", "").replace("-", "")) for n in nums) + (2 if any("." in n for n in nums) else 0)
            if pd.notna(val) and score > best[0]:
                best = (score, val, str(txt).strip())
        conf = min(1.0, max(0.0, best[0] / 8.0)) if best[0] >= 0 else 0.0
        return best[1], conf, best[2]
    except Exception:
        return np.nan, 0.0, "cv2_ocr_error"


def _ocr_numeric_region(crop_img):
    pil_val, pil_conf, pil_text = _ocr_numeric_region_pil(crop_img)
    cv_val, cv_conf, cv_text = _ocr_numeric_region_cv2(crop_img)
    if pd.notna(cv_val) and (pd.isna(pil_val) or cv_conf >= pil_conf):
        return cv_val, cv_conf, cv_text
    return pil_val, pil_conf, pil_text

# v43.3 dtype-safe test ID assignment override.
def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    out = df.copy()
    if "well" not in out.columns:
        out["well"] = "Unknown"
    if "datetime" not in out.columns:
        out["datetime"] = pd.NaT
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["well"] = out["well"].fillna("Unknown").astype(str).str.strip().replace("", "Unknown")
    out["test_id"] = out["test_id"].astype("object") if "test_id" in out.columns else pd.Series([None] * len(out), index=out.index, dtype="object")
    out["test_sequence"] = pd.to_numeric(out["test_sequence"], errors="coerce") if "test_sequence" in out.columns else pd.Series([np.nan] * len(out), index=out.index, dtype="float")
    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}
    for well, g in sortable.groupby("well", dropna=False):
        well_txt = str(well or "Unknown").strip() or "Unknown"
        if well_txt.lower() == "unknown":
            for i in g.index:
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well", np.nan)
            continue
        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            new_test = current_id is None or pd.isna(last_dt) or (pd.Timestamp(dt) - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{pd.Timestamp(dt).strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = pd.Timestamp(dt)
    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = tid
        out.at[i, "test_sequence"] = seq
    if "test_id" in out.columns and "datetime" in out.columns:
        valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
        if not valid.empty:
            starts = valid.groupby("test_id")["datetime"].min()
            ends = valid.groupby("test_id")["datetime"].max()
            out["test_start"] = out["test_id"].map(starts)
            out["test_end"] = out["test_id"].map(ends)
    return out

# v43.4 preserve old TMU parser but add operation-note detection such as Flare test.
_parse_tmu_message_base_v43 = parse_tmu_message

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v43(message, source_name=source_name)
    text = str(message or "")
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row


# -----------------------------------------------------------------------------
# v44 safety patch: no nearest-time CTU linking + robust WhatsApp bold well names
# -----------------------------------------------------------------------------
def clean_well_name_value(value: object) -> str:
    """Clean WhatsApp/Excel well names without inventing a well.

    Examples:
      '*S8-58*' -> 'S8-58'
      '*'       -> 'Unknown'
    """
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-"}:
        return "Unknown"
    # Prefer the first token/candidate containing at least one digit, which is how
    # almost all field well names appear: S8-58, B3C18-7, A-C83-1, BED15-33.
    token_match = re.search(r"[A-Za-z0-9]+(?:[-/][A-Za-z0-9]+)*", s)
    if token_match:
        candidate = token_match.group(0).strip("-_/ .")
        if candidate and re.search(r"[A-Za-z0-9]", candidate):
            return candidate
    return "Unknown"


def suggest_links_for_ocr_rows(df: pd.DataFrame, max_gap_hours: float = 0.0) -> pd.DataFrame:
    """Disabled in v44 by design.

    CTU/OCR rows are not linked or suggested by nearest time because that can
    attach image data to the wrong test. The app asks the user to manually choose
    the Well/Test ID after reviewing the image OCR numbers.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    if "source_type" in out.columns:
        ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
        for c in ["link_status", "review_required"]:
            if c not in out.columns:
                out[c] = pd.Series([None] * len(out), index=out.index, dtype="object")
        out.loc[ocr_mask, "link_status"] = out.loc[ocr_mask, "link_status"].fillna("ocr_manual_link_required")
        out.loc[ocr_mask, "review_required"] = True
    return out


def approve_suggested_ocr_links(df: pd.DataFrame) -> pd.DataFrame:
    """No-op kept for backward compatibility with older app.py builds."""
    return df


_parse_tmu_message_base_v44 = parse_tmu_message

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v44(message, source_name=source_name)
    row["well"] = clean_well_name_value(row.get("well", "Unknown"))
    text = str(message or "")
    if row.get("well", "Unknown") == "Unknown":
        guessed = guess_well_from_name(re.sub(r"[*_`~]+", "", text))
        row["well"] = clean_well_name_value(guessed)
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row


def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    """Assign test IDs by well name and time gap only.

    - Different well names are always separate streams.
    - Same well continues until the inactive gap exceeds gap_hours.
    - OCR rows with Unknown well stay Unlinked until manual review.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    for c in ["source", "sheet", "source_type", "well", "datetime", "test_id", "test_sequence", "link_status", "review_required"]:
        if c not in out.columns:
            default = None if c in {"well", "test_id", "link_status", "source_type", "source", "sheet"} else np.nan
            out[c] = pd.Series([default] * len(out), index=out.index, dtype="object")
    out["well"] = out["well"].apply(clean_well_name_value).astype("object")
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["test_id"] = out["test_id"].astype("object")
    out["link_status"] = out["link_status"].astype("object")
    out["review_required"] = out["review_required"].astype("object")

    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}
    for well, g in sortable.groupby("well", dropna=False):
        well_txt = clean_well_name_value(well)
        if well_txt == "Unknown":
            for i, row in g.iterrows():
                is_ocr = str(row.get("source_type", "")).lower().find("ocr") >= 0
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well" if is_ocr else "Unknown_Well_Unlinked", np.nan)
            continue
        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            dt = pd.Timestamp(dt)
            new_test = current_id is None or pd.isna(last_dt) or (dt - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{dt.strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = dt

    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = str(tid)
        out.at[i, "test_sequence"] = seq

    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    unknown_ocr = ocr_mask & ((out["well"].astype(str).str.lower() == "unknown") | out["test_id"].astype(str).str.startswith("Unlinked"))
    out.loc[unknown_ocr, "link_status"] = out.loc[unknown_ocr, "link_status"].fillna("ocr_manual_link_required")
    out.loc[unknown_ocr, "review_required"] = True

    valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
    if not valid.empty:
        starts = valid.groupby("test_id")["datetime"].min()
        ends = valid.groupby("test_id")["datetime"].max()
        out["test_start"] = out["test_id"].map(starts)
        out["test_end"] = out["test_id"].map(ends)
    return out


# v44.1 extra well-name cleanup for WhatsApp bold/plain first-line well names.
_guess_well_from_name_base_v441 = guess_well_from_name

def guess_well_from_name(name: str) -> Optional[str]:
    base = _guess_well_from_name_base_v441(name)
    if base:
        return clean_well_name_value(base) if 'clean_well_name_value' in globals() else base
    s = re.sub(r"[*_`~]+", "", str(name or ""))
    extra_patterns = [
        r"\b(OB\s*[-_ ]\s*\d+[A-Z]?)\b",          # Obaiyed OB-69
        r"\b(S\d+\s*[-_ ]\s*\d+[A-Z]?)\b",        # S8-58
        r"\b([A-Z]{1,4}\d+[A-Z]*\s*[-_ ]\s*\d+[A-Z]?)\b",
    ]
    for pat in extra_patterns:
        m = re.search(pat, s, flags=re.I)
        if m:
            return clean_well_name_value(m.group(1)) if 'clean_well_name_value' in globals() else normalize_well_name(m.group(1))
    return None

# Override clean_well_name_value to reject non-well words such as 'to' and '*'.
def clean_well_name_value(value: object) -> str:
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-", "to"}:
        return "Unknown"
    # Candidate must contain at least one digit to avoid false wells from normal words.
    m = re.search(r"[A-Za-z0-9]*\d[A-Za-z0-9]*(?:[-/][A-Za-z0-9]+)*", s)
    if m:
        candidate = m.group(0).strip("-_/ .")
        if candidate and re.search(r"\d", candidate):
            return candidate.upper() if re.match(r"^[A-Za-z]{1,6}[-/ ]?\d", candidate) else candidate
    return "Unknown"

# Re-override parser and test-id assignment so the final clean_well_name_value is used.
_parse_tmu_message_base_v441 = _parse_tmu_message_base_v44

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v441(message, source_name=source_name)
    text = str(message or "")
    cleaned_guess = clean_well_name_value(row.get("well", "Unknown"))
    if cleaned_guess == "Unknown":
        guessed = guess_well_from_name(text)
        cleaned_guess = clean_well_name_value(guessed)
    row["well"] = cleaned_guess
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row

# v44.2 final well cleaning: keep prefixes like OB-69, S8-58, B3C18-7, A-C83-1.
def clean_well_name_value(value: object) -> str:
    s = safe_text(value)
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[*_`~]+", "", s)
    s = re.sub(r"(?i)\bwell\s*name\b|\bwell\b", "", s)
    s = re.sub(r"^[\s:=@\-]+", "", s).strip()
    s = re.split(r"[\n\r,;]|\s{2,}", s)[0].strip()
    s = s.strip(" .:;=()[]{}<>|'\"")
    if not s or s.lower() in {"nan", "nat", "none", "unknown", "null", "-", "to"}:
        return "Unknown"
    # Find all candidate well-looking tokens that contain digits and optional letter prefixes.
    pat = r"\b(?:[A-Za-z]{1,12}[-/ ]*)?[A-Za-z0-9]*\d[A-Za-z0-9]*(?:[-/][A-Za-z0-9]+)*\b"
    candidates = []
    for m in re.finditer(pat, s):
        cand = re.sub(r"\s+", "", m.group(0).strip("-_/ ."))
        if not cand or not re.search(r"\d", cand):
            continue
        if re.fullmatch(r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4}", cand):
            continue
        candidates.append(cand)
    if candidates:
        # Prefer the last candidate because strings like 'Obaiyed OB-69' contain
        # a field name first and the actual well token at the end.
        return candidates[-1].upper()
    return "Unknown"

_parse_tmu_message_base_v442 = _parse_tmu_message_base_v441

def parse_tmu_message(message: str, source_name: str = "WhatsApp_Text") -> Dict[str, object]:
    row = _parse_tmu_message_base_v442(message, source_name=source_name)
    text = str(message or "")
    cleaned_guess = clean_well_name_value(row.get("well", "Unknown"))
    if cleaned_guess == "Unknown":
        guessed = guess_well_from_name(text)
        cleaned_guess = clean_well_name_value(guessed)
    row["well"] = cleaned_guess
    if re.search(r"\bflare\s+test\b", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "Flare test")
    if re.search(r"\bctu\b|coiled\s+tubing|coil\s+tubing", text, flags=re.I):
        row["note"] = append_note(row.get("note"), "CTU operation")
    return row

# Re-override assign_test_ids to use v44.2 clean_well_name_value.
def assign_test_ids(df: pd.DataFrame, gap_hours: float = 12.0) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    out = df.copy()
    for c in ["source", "sheet", "source_type", "well", "datetime", "test_id", "test_sequence", "link_status", "review_required"]:
        if c not in out.columns:
            default = None if c in {"well", "test_id", "link_status", "source_type", "source", "sheet"} else np.nan
            out[c] = pd.Series([default] * len(out), index=out.index, dtype="object")
    out["well"] = out["well"].apply(clean_well_name_value).astype("object")
    out["datetime"] = pd.to_datetime(out["datetime"], errors="coerce")
    out["test_id"] = out["test_id"].astype("object")
    out["link_status"] = out["link_status"].astype("object")
    out["review_required"] = out["review_required"].astype("object")
    sortable = out.sort_values(["well", "datetime", "source"], na_position="last").copy()
    assigned = {}
    for well, g in sortable.groupby("well", dropna=False):
        well_txt = clean_well_name_value(well)
        if well_txt == "Unknown":
            for i, row in g.iterrows():
                is_ocr = str(row.get("source_type", "")).lower().find("ocr") >= 0
                assigned[i] = ("Unlinked_OCR_or_Unknown_Well" if is_ocr else "Unknown_Well_Unlinked", np.nan)
            continue
        seq = 0
        last_dt = pd.NaT
        current_id = None
        for i, row in g.iterrows():
            dt = row.get("datetime", pd.NaT)
            if pd.isna(dt):
                seq += 1
                current_id = f"{well_txt}_T{seq:02d}_NoTime"
                assigned[i] = (current_id, float(seq))
                continue
            dt = pd.Timestamp(dt)
            new_test = current_id is None or pd.isna(last_dt) or (dt - pd.Timestamp(last_dt) > pd.Timedelta(hours=float(gap_hours)))
            if new_test:
                seq += 1
                current_id = f"{well_txt}_{dt.strftime('%Y%m%d_%H%M')}"
            assigned[i] = (current_id, float(seq))
            last_dt = dt
    for i, (tid, seq) in assigned.items():
        out.at[i, "test_id"] = str(tid)
        out.at[i, "test_sequence"] = seq
    ocr_mask = out["source_type"].astype(str).str.contains("ocr", case=False, na=False)
    unknown_ocr = ocr_mask & ((out["well"].astype(str).str.lower() == "unknown") | out["test_id"].astype(str).str.startswith("Unlinked"))
    out.loc[unknown_ocr, "link_status"] = out.loc[unknown_ocr, "link_status"].fillna("ocr_manual_link_required")
    out.loc[unknown_ocr, "review_required"] = True
    valid = out[~out["test_id"].astype(str).str.startswith("Unlinked") & out["datetime"].notna()].copy()
    if not valid.empty:
        starts = valid.groupby("test_id")["datetime"].min()
        ends = valid.groupby("test_id")["datetime"].max()
        out["test_start"] = out["test_id"].map(starts)
        out["test_end"] = out["test_id"].map(ends)
    return out
